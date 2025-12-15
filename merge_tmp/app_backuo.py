from flask import Flask, request, jsonify, send_file, redirect, url_for,send_from_directory,make_response,after_this_request
from db import get_connection
import os, tempfile, time,base64 ,datetime as dt,hashlib, random,difflib
import win32com.client as win32
import pythoncom
import re
import io
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx2pdf import convert
from io import BytesIO
import   zipfile, shutil,  unicodedata, copy
import xml.etree.ElementTree as ET
from docx import Document as DocxDocument
import pythoncom
from datetime import datetime
from pathlib import Path
# ↑ cerca de otros imports
import json
import traceback
import pythoncom
import subprocess, socket, requests, traceback, shutil, json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf import PdfReader
from shutil import copy2
from flask_cors import CORS, cross_origin
from difflib import SequenceMatcher
import mammoth
from docx import Document
import re as _re
from copy import deepcopy
...
try:
    from docxcompose.composer import Composer
except Exception:
    Composer = None

try:
    from docx import Document as _DocxDocument
except Exception:
    _DocxDocument = None

try:
    from pypdf import PdfReader as _PdfReader
except Exception:
    _PdfReader = None

import win32com.client as win32
from docx.enum.text import WD_COLOR_INDEX  # fallback para highlight



# (opcional, ayuda con rutas largas/espacios)
def _short_path(p: str) -> str:
    """Devuelve ruta en 8.3 si es posible; si no, la ruta absoluta normal."""
    try:
        import win32api
        return win32api.GetShortPathName(os.path.abspath(p))
    except Exception:
        return os.path.abspath(p)

def _short83(p: str) -> str:
    """Alias más usado en el código (sin guion bajo intermedio)."""
    return _short_path(p)
# rutas absolutas
BACKEND_DIR = os.path.abspath(os.path.dirname(__file__))
PROYECTO_DIR     = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DESCARGAS_DIR = os.path.join(PROYECTO_DIR, "descargas")
STATIC_DIR     = os.path.join(PROYECTO_DIR, "static")     # asegúrate que existe
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(os.path.join(STATIC_DIR, "previews"), exist_ok=True)
os.makedirs(DESCARGAS_DIR, exist_ok=True)

app = Flask(__name__, static_folder=STATIC_DIR, static_url_path="/static")
CORS(app, resources={r"/*": {"origins": "*"}})
NS = {
    "w":  "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r":  "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m":  "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a":  "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic":"http://schemas.openxmlformats.org/drawingml/2006/picture",
    "v":  "urn:schemas-microsoft-com:vml",
    "o":  "urn:schemas-microsoft-com:office:office",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}
for p, u in NS.items():
    ET.register_namespace(p, u)
# rutas absolutas
BACKEND_DIR = os.path.abspath(os.path.dirname(__file__))

app.config["DESCARGAS_FOLDER"]  = DESCARGAS_DIR
app.config['UPLOAD_FOLDER'] = os.path.join(BACKEND_DIR, "uploads")

EXAM_DIR = os.path.join(BACKEND_DIR , "uploads", "examenes")
app.config["UPLOADS_EXAM_DIR"] = EXAM_DIR
os.makedirs(EXAM_DIR, exist_ok=True)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DESCARGAS_FOLDER'], exist_ok=True)


PREVIEWS_DIR = os.path.join(STATIC_DIR, "previews")
os.makedirs(PREVIEWS_DIR, exist_ok=True)
print("PREVIEWS_DIR    =>", PREVIEWS_DIR)
# (si quieres verificar en consola)


print("DESCARGAS_FOLDER =>", app.config["DESCARGAS_FOLDER"])
print("STATIC_DIR       =>", STATIC_DIR)


app.config['PREGUNTAS_DIR'] = os.path.join(os.path.dirname(__file__), 'temas_archivos')
os.makedirs(app.config['PREGUNTAS_DIR'], exist_ok=True)

# cerca de los imports superiores
class DocxVacioError(Exception):
    def __init__(self, paths): self.paths = paths

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GRUPOS_OUT_DIR = os.path.join(app.config['DESCARGAS_FOLDER'], "grupos")
os.makedirs(GRUPOS_OUT_DIR, exist_ok=True)
DATA_DIR      = os.path.join(BASE_DIR, "data")

UPLOADS_DIR   = os.path.join(DATA_DIR, "uploads")
OUTPUTS_DIR   = os.path.join(DATA_DIR, "outputs")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(OUTPUTS_DIR, exist_ok=True)


# --- CONFIG UPLOADS BANCO ---
BANCO_DIR = os.path.join(app.root_path, "uploads", "banco")
BANCO_PREG_DIR = os.path.join(BANCO_DIR, "preguntas")
BANCO_SOL_DIR  = os.path.join(BANCO_DIR, "solucionarios")
os.makedirs(BANCO_PREG_DIR, exist_ok=True)
os.makedirs(BANCO_SOL_DIR, exist_ok=True)

def _row_to_dict_list(cur):
    cols = [c[0] for c in cur.description]
    return [dict(zip(cols, r)) for r in cur.fetchall()]


@app.route("/probar-conexion")
def probar_conexion():
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT DATABASE()")
        nombre_bd = cursor.fetchone()[0]
        cursor.close()
        conn.close()
        return jsonify({"conexion": "ok", "base_datos": nombre_bd})
    except Exception as e:
        return jsonify({"conexion": "error", "detalle": str(e)})


# =============== MODO DE CORRECCIÓN ===============
# True  -> comportamiento clásico (como tu código antiguo, corrige "igual de bien")
# False -> heurísticas anti-falsos positivos (tu código nuevo)
USE_CLASSIC_LT = True

# --- Utilidades y constantes ---
TOKEN_RX = re.compile(r'(\s+|[^\wÁÉÍÓÚáéíóúÑñ]+)', flags=re.UNICODE)

# Marcas para preview
MARK_WORD_DELETE  = '---'
MARK_SPACE_DELETE = '--'

app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024
# ====== Languagetool ======
LT_PORT = int(os.environ.get("LT_PORT", "8010"))

def _java_cmd():
    emb = os.path.join(BASE_DIR, "jre", "bin", "java.exe" if os.name=="nt" else "java")
    return emb if os.path.exists(emb) else "java"

LT_DIR = os.environ.get("LT_DIR", "").strip() or None
_LT_CANDIDATES = [
    os.path.join(BASE_DIR, "libs", "LanguageTool"),
    os.path.join(BASE_DIR, "LanguageTool"),
    os.path.join(BASE_DIR, "languagetool"),
    os.path.join(BASE_DIR, "languagetools"),
    BASE_DIR,
]

def _resolve_lt_dir() -> str:
    forced = os.environ.get("LT_DIR", "").strip()
    if forced and os.path.isdir(forced):
        for r, _d, files in os.walk(forced):
            for f in files:
                name = f.lower()
                if name.endswith("server.jar") and "languagetool" in name:
                    print("[LT] LT_DIR (env) ->", r); return r
    for root in _LT_CANDIDATES:
        if os.path.isdir(root):
            for r, _d, files in os.walk(root):
                for f in files:
                    name = f.lower()
                    if name.endswith("server.jar") and "languagetool" in name:
                        print("[LT] LT_DIR detectado ->", r); return r
    raise RuntimeError("No se encontró LanguageTool (server.jar).")

def _find_lt_jar() -> str:
    if not LT_DIR or not os.path.isdir(LT_DIR):
        raise RuntimeError(f"No se encontró la carpeta LanguageTool: {LT_DIR}")
    for r, _d, files in os.walk(LT_DIR):
        for f in files:
            name = f.lower()
            if name.endswith("server.jar") and "languagetool" in name:
                return os.path.join(r, f)
    raise RuntimeError(f"No se encontró el jar del servidor dentro de: {LT_DIR}")

def lt_is_running(host="127.0.0.1", port=LT_PORT) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.25)
        try:
            s.connect((host, port)); return True
        except Exception:
            return False

def _detect_ngrams_dir() -> str | None:
    forced = os.environ.get("NGRAMS_DIR", "").strip()
    if forced and os.path.isdir(forced):
        if os.path.isdir(os.path.join(forced, "1grams")) and os.path.isdir(os.path.join(forced, "2grams")):
            print("[LT] NGRAMS_DIR (env) ->", forced); return forced
    candidates = []
    if LT_DIR:
        candidates += [os.path.join(LT_DIR, "ngrams", "es"), os.path.join(LT_DIR, "ngrams")]
        parent = os.path.dirname(LT_DIR)
        candidates += [os.path.join(parent, "languagetools", "ngrams", "es"),
                       os.path.join(parent, "languagetools", "ngrams")]
    candidates += [
        os.path.join(BASE_DIR, "ngrams", "es"),
        os.path.join(BASE_DIR, "ngrams"),
        os.path.join(BASE_DIR, "languagetools", "ngrams", "es"),
        os.path.join(BASE_DIR, "languagetools", "ngrams"),
    ]
    for c in candidates:
        if os.path.isdir(os.path.join(c, "1grams")) and os.path.isdir(os.path.join(c, "2grams")):
            print("[LT] NGRAMS detectado ->", c); return c
    print("[LT] NGRAMS no detectado (se intentará iniciar sin languageModel)")
    return None

_LT_PROC = None
def _spawn_lt(args):
    creationflags = 0x08000000 if os.name == "nt" else 0
    return subprocess.Popen(args, cwd=LT_DIR, stdout=subprocess.DEVNULL, stderr=subprocess.STDOUT,
                            creationflags=creationflags)

def lt_start_server():
    global _LT_PROC, LT_DIR
    if lt_is_running(): return
    if not LT_DIR: LT_DIR = _resolve_lt_dir()
    lt_jar = _find_lt_jar(); java = _java_cmd()
    base_args = [java, "-Xms256m", "-Xmx2048m", "-jar", lt_jar, "-p", str(LT_PORT)]
    custom_rules = os.path.join(BASE_DIR, "libs", "LanguageTool", "custom-rules-es.xml")
    if os.path.isfile(custom_rules): base_args += ["--rules", custom_rules]
    ngrams_dir = _detect_ngrams_dir()
    try_args = [base_args + (["--languageModel", ngrams_dir] if ngrams_dir else []), base_args]
    for args in try_args:
        print("[LT] cmd:", " ".join(args))
        try:
            _LT_PROC = _spawn_lt(args)
            for _ in range(100):
                if lt_is_running():
                    print("[LT] Servidor iniciado en", LT_PORT); return
                time.sleep(0.1)
        except Exception as e:
            print("[LT] Error al iniciar:", e)
    raise RuntimeError("No se pudo iniciar LanguageTool local.")

# ============ Helpers de resaltado anaranjado ============
def _shade_run_orange_exact(run):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFA500')  # naranja
    for old in rPr.findall(qn('w:shd')):
        rPr.remove(old)
    rPr.append(shd)
    try:
        run.font.highlight_color = WD_COLOR_INDEX.NONE
    except Exception:
        pass

def contiene_omml(run):
    r = run._r
    if r.xpath('.//m:oMath | .//m:oMathPara'):
        return True
    if r.xpath('.//w:drawing | .//w:pict | .//w:object | .//w:fldChar'):
        return True
    return False

def _split_run_keep_style(paragraph, run_idx, cut_pos):
    run = paragraph.runs[run_idx]
    txt = run.text or ""
    if cut_pos <= 0 or cut_pos >= len(txt):
        return run_idx
    left_txt, right_txt = txt[:cut_pos], txt[cut_pos:]
    new_r = deepcopy(run._r)
    run._r.addnext(new_r)
    from docx.text.run import Run
    right_run = Run(new_r, run._parent)
    run.text = left_txt
    right_run.text = right_txt
    return run_idx + 1

def _char_index_map(paragraph):
    index_map = []
    for i, r in enumerate(paragraph.runs):
        if contiene_omml(r): continue
        t = r.text or ""
        for k in range(len(t)):
            index_map.append((i, k))
    return index_map

def _highlight_range_by_char(paragraph, start_char, end_char):
    if start_char >= end_char: return
    index_map = _char_index_map(paragraph)
    n = len(index_map)
    if n == 0: return
    start_char = max(0, min(start_char, n))
    end_char   = max(0, min(end_char,   n))
    if start_char >= end_char: return

    start_run, start_off = index_map[start_char]
    end_run,   end_off   = index_map[end_char - 1]

    if start_off < len(paragraph.runs[start_run].text or ""):
        start_run = _split_run_keep_style(paragraph, start_run, start_off)

    index_map = _char_index_map(paragraph)
    end_char = max(0, min(end_char, len(index_map)))
    end_run, end_off = index_map[end_char - 1]
    end_txt_len = len(paragraph.runs[end_run].text or "")
    cut_pos = end_off + 1
    if cut_pos < end_txt_len:
        _ = _split_run_keep_style(paragraph, end_run, cut_pos)

    index_map = _char_index_map(paragraph)
    s_run = index_map[start_char][0]
    e_run = index_map[end_char - 1][0]

    for i in range(s_run, e_run + 1):
        r = paragraph.runs[i]
        if not contiene_omml(r) and (r.text or ""):
            _shade_run_orange_exact(r)

# --- Ayudas para decidir agrupación ---
_WORD_RE = re.compile(r'[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]')

def _is_word_char(ch: str) -> bool:
    return bool(_WORD_RE.match(ch))

def _is_single_word(s: str) -> bool:
    s = s.strip()
    if not s: return False
    if re.search(r'\s', s): return False
    return all(_is_word_char(c) for c in s)

def _is_two_or_more_words(s: str) -> bool:
    s = s.strip()
    if not s: return False
    if re.search(r'[^\wÁÉÍÓÚÜÑáéíóúüñ\s]', s):  # signos -> no agrupar
        return False
    tokens = [t for t in re.split(r'\s+', s) if t]
    if len(tokens) < 2: return False
    return all(all(_is_word_char(c) for c in t) for t in tokens)

def _highlight_tokens_in_range(paragraph, base_start, frag):
    if not frag: return
    visible = "".join(r.text or "" for r in paragraph.runs if not contiene_omml(r))
    start_global = base_start
    end_global   = base_start + len(frag)
    n_visible    = len(visible)
    if start_global < 0: start_global = 0
    if end_global   > n_visible: end_global = n_visible
    if start_global >= end_global: return
    frag_has_word = any(_is_word_char(c) for c in frag)
    if not frag_has_word:
        _highlight_range_by_char(paragraph, start_global, end_global)
        return
    def is_word_char_local(ch: str) -> bool:
        return _is_word_char(ch)
    for m in re.finditer(r'\S+', frag, flags=re.UNICODE):
        a_loc, b_loc = m.span()
        a = start_global + a_loc
        b = start_global + b_loc
        L = a
        R = b
        while L > 0 and is_word_char_local(visible[L - 1]): L -= 1
        while R < n_visible and is_word_char_local(visible[R]): R += 1
        if L < R:
            _highlight_range_by_char(paragraph, L, R)

def _highlight_replacement_smart(paragraph, base_start, oldfrag, newfrag):
    if oldfrag and _is_single_word(oldfrag) and _is_two_or_more_words(newfrag):
        _highlight_range_by_char(paragraph, base_start, base_start + len(newfrag))
    else:
        _highlight_tokens_in_range(paragraph, base_start, newfrag)

# ================================================================
def insertar_marcas_eliminacion(texto_original: str, texto_corregido: str) -> str:
    orig_toks = tokenize_preservando(texto_original)
    corr_toks = tokenize_preservando(texto_corregido)
    def _is_space(tok: str) -> bool: return bool(tok and tok.isspace())
    def _is_word(tok: str) -> bool:  return es_token_palabra(tok) and not _is_space(tok)
    salida: list[str] = []
    pending_space_after_marker = False
    def _append_token(tok: str, out: list[str]):
        nonlocal pending_space_after_marker
        if pending_space_after_marker:
            if not _is_space(tok): out.append(" ")
            pending_space_after_marker = False
        out.append(tok)
    def _emit_marker_for_deleted(deleted_tokens: list[str], out: list[str]) -> bool:
        nonlocal pending_space_after_marker
        has_word  = any(_is_word(t) for t in deleted_tokens)
        has_space = any(_is_space(t) for t in deleted_tokens)
        if has_word:
            if not out or not _is_space(out[-1]): out.append(" ")
            out.append(MARK_WORD_DELETE); pending_space_after_marker = True; return True
        elif has_space:
            if not out or out[-1] != MARK_SPACE_DELETE: out.append(MARK_SPACE_DELETE)
        return False
    sm = SequenceMatcher(a=orig_toks, b=corr_toks, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ("equal", "insert"):
            for tok in corr_toks[j1:j2]: _append_token(tok, salida)
        elif tag == "delete":
            _emit_marker_for_deleted(orig_toks[i1:i2], salida)
        elif tag == "replace":
            seg_o = orig_toks[i1:i2]; seg_c = corr_toks[j1:j2]
            tmp_out: list[str] = []
            inner = SequenceMatcher(a=seg_o, b=seg_c, autojunk=False)
            emitted_delete = False
            for t2, a1, a2, b1, b2 in inner.get_opcodes():
                if t2 in ("equal", "insert", "replace"):
                    for tok in seg_c[b1:b2]: _append_token(tok, tmp_out)
                if t2 == "delete":
                    if _emit_marker_for_deleted(seg_o[a1:a2], tmp_out): emitted_delete = True
            if not emitted_delete:
                words_o = sum(1 for t in seg_o if es_token_palabra(t) and not t.isspace())
                words_c = sum(1 for t in seg_c if es_token_palabra(t) and not t.isspace())
                if words_o > words_c:
                    left_idx = 0
                    while left_idx < len(tmp_out) and (tmp_out[left_idx].isspace() if tmp_out[left_idx] else False): left_idx += 1
                    while left_idx < len(tmp_out) and not (tmp_out[left_idx].isspace() if tmp_out[left_idx] else False): left_idx += 1
                    while left_idx < len(tmp_out) and (tmp_out[left_idx].isspace() if tmp_out[left_idx] else False): left_idx += 1
                    if left_idx == 0 or not (tmp_out[left_idx-1].isspace() if tmp_out[left_idx-1] else False):
                        tmp_out.insert(left_idx, " "); left_idx += 1
                    tmp_out.insert(left_idx, MARK_WORD_DELETE); left_idx += 1
                    if left_idx >= len(tmp_out) or not (tmp_out[left_idx].isspace() if tmp_out[left_idx] else False):
                        tmp_out.insert(left_idx, " ")
            for tok in tmp_out: _append_token(tok, salida)
    return "".join(salida)

# --- Sesión HTTP reutilizable para LT
LT_HTTP = requests.Session()
try:
    from requests.adapters import HTTPAdapter
    LT_HTTP.mount("http://", HTTPAdapter(pool_connections=4, pool_maxsize=4, max_retries=1))
except Exception:
    pass

LT_SOFT_CHUNK = 20000

def _lt_request(texto: str, lang: str, use_picky=True, use_variant=True) -> dict:
    if not lt_is_running(): lt_start_server()
    url = f"http://127.0.0.1:{LT_PORT}/v2/check"
    data = {"text": texto, "language": lang}
    if use_picky:   data["level"] = "picky"
    if use_variant: data["preferredVariants"] = "es-ES"
    try:
        resp = LT_HTTP.post(url, data=data, timeout=120)
    except requests.exceptions.ConnectionError:
        lt_start_server(); resp = LT_HTTP.post(url, data=data, timeout=120)
    if resp.status_code == 400: return {"_status": 400, "_body": resp.text}
    resp.raise_for_status(); out = resp.json(); out["_status"] = resp.status_code; return out

def lt_check_smart(texto: str, lang="es"):
    r1 = _lt_request(texto, lang, use_picky=True, use_variant=True)
    if r1.get("_status") != 400: return r1
    too_long = ("too long" in (r1.get("_body","")).lower()) or (len(texto) > LT_SOFT_CHUNK*1.2)
    r2 = _lt_request(texto, lang, use_picky=False, use_variant=False)
    if r2.get("_status") != 400 and not too_long: return r2
    matches_all = []; offset_base = 0; n = len(texto); start = 0
    while start < n:
        end = min(start + LT_SOFT_CHUNK, n)
        cut = texto.rfind("\n\n", start, end)
        if cut == -1 or cut <= start + 1000: cut = end
        chunk = texto[start:cut]
        r = _lt_request(chunk, lang, use_picky=False, use_variant=False)
        if r.get("_status") == 400:
            r = _lt_request(chunk, lang, use_picky=False, use_variant=True)
            if r.get("_status") == 400:
                raise RuntimeError(f"LT 400 en chunk {start}: {r.get('_body','')}")
        for m in r.get("matches", []):
            m2 = dict(m); m2["offset"] = m.get("offset", 0) + offset_base; matches_all.append(m2)
        offset_base += len(chunk); start = cut
    return {"matches": matches_all}

# ===== util doc/pdf =====
def extraer_texto_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def extraer_texto_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)

def crear_docx_desde_texto(texto: str) -> Document:
    doc = Document()
    for linea in texto.split("\n"):
        doc.add_paragraph(linea)
    return doc

def generar_diff_html(original: str, corregido: str) -> str:
    diff  = difflib.HtmlDiff(wrapcolumn=90)
    tabla = diff.make_table(original.splitlines(), corregido.splitlines(),
                            "Original", "Corregido", context=True, numlines=2)
    estilo = """
    <style>
      table.diff {font-family:monospace; font-size:13px; border:1px solid #ccc; width:100%}
      .diff_header {background:#f8f9fa} .diff_add {background:#d4edda}
      .diff_chg {background:#fff3cd} .diff_sub {background:#f8d7da}
      td,th {padding:4px 6px; border:1px solid #e9ecef; vertical-align:top}
    </style>"""
    return estilo + tabla

def now_mysql(): return dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def normalize_ocr_noise(t: str) -> str:
    if not t: return t
    t = (t.replace("\u00AD","").replace("\u200B","").replace("\u200C","")
           .replace("\u200D","").replace("\u2060",""))
    t = (t.replace("\u00A0"," ").replace("\u202F"," "))
    t = re.sub(r'[\u2000-\u200A]',' ', t)
    t = re.sub(r'(\w)-\n(\w)', r'\1\2', t)
    t = t.replace("\r\n","\n").replace("\r","\n")
    t = re.sub(r'[ \t]{2,}',' ', t)
    return t

def _strip_accents(s: str) -> str:
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

def _same_casing(src: str, dst: str) -> str:
    if src.isupper(): return dst.upper()
    if src.istitle(): return dst[:1].upper() + dst[1:]
    if src.islower(): return dst.lower()
    return dst

def _prev_token(texto: str, offset: int) -> str:
    izq = texto[:offset]
    m = _re.search(r'([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)\s*$', izq)
    return m.group(1) if m else ""

def _choose_best_suggestion(original: str, suggestions: list, prev: str) -> str | None:
    cand = [s.get("value","") for s in suggestions if s.get("value")]
    if not cand: return None
    base = _strip_accents(original.lower())
    solo_acento = [c for c in cand if _strip_accents(c.lower()) == base and c != original]
    if solo_acento: return solo_acento[0]
    prev_l = prev.lower()
    if prev.istitle() or prev_l in {"el","él","ella"}:
        pref_3a = [c for c in cand if c.endswith(("ó","a")) and not c.endswith("é")]
        if pref_3a: return pref_3a[0]
    same_len = [c for c in cand if len(c) == len(original)]
    if same_len: return same_len[0]
    return cand[0]

_MATH_VARS_1L = {"n","x","y","z","m","k","i","j","t"}

# ====== NUEVO: Detector de siglas en MAYÚSCULAS ======
_ACRONYM_LETTERS_RX = re.compile(r'^[A-ZÁÉÍÓÚÜÑ]{2,12}$')
def is_upper_acronym(token: str) -> bool:
    if not token: return False
    t = token.strip()
    if t.startswith("(") and t.endswith(")"): t = t[1:-1].strip()
    t = (t.replace(".", "").replace("-", "").replace("–", "").replace("—", "").replace("/", ""))
    if len(t) < 2 or len(t) > 12 or not t.isalpha(): return False
    return bool(_ACRONYM_LETTERS_RX.fullmatch(t))

# ====== NUEVO: No tocar números / dígitos (incluye superíndices) ======
_SUPERSCRIPT_DIGITS = set("⁰¹²³⁴⁵⁶⁷⁸⁹")
def _has_any_digit(s: str) -> bool:
    return any(ch.isdigit() or ch in _SUPERSCRIPT_DIGITS for ch in s or "")

# ====== NUEVO: filtro de “parecido razonable” ======
def _edit_distance(a: str, b: str) -> int:
    la, lb = len(a), len(b)
    if la == 0: return lb
    if lb == 0: return la
    prev = list(range(lb + 1))
    cur = [0] * (lb + 1)
    for i, ca in enumerate(a, 1):
        cur[0] = i
        for j, cb in enumerate(b, 1):
            cost = 0 if ca == cb else 1
            cur[j] = min(prev[j] + 1, cur[j-1] + 1, prev[j-1] + cost)
        prev, cur = cur, prev
    return prev[-1]

def _looks_reasonable_replacement(src: str, dst: str) -> bool:
    if not src or not dst: return False
    if _strip_accents(src.lower()) == _strip_accents(dst.lower()) and src != dst:
        return True
    a = _strip_accents(src.lower())
    b = _strip_accents(dst.lower())
    if len(a) <= 3 or len(b) <= 3:
        return SequenceMatcher(None, a, b).ratio() >= 0.60
    sim = SequenceMatcher(None, a, b).ratio()
    ed  = _edit_distance(a, b)
    if sim < 0.80: return False
    max_ed = max(2, min(len(a), len(b)) // 4)
    if ed > max_ed: return False
    return True

# ====== NUEVO: spans de alternativas y utilidades ======
_RX_ALTERNATIVA = re.compile(r'^\s*[A-E]\)\s')  # PATRÓN pedido: Letra mayúscula A–E + ") " al inicio de línea

def detectar_spans_alternativas(texto: str):
    """Devuelve lista de (ini, fin) por cada línea que sea alternativa tipo 'A) '."""
    spans = []
    pos = 0
    for linea in texto.split("\n"):
        L = len(linea)
        if _RX_ALTERNATIVA.match(linea):
            spans.append((pos, pos + L))  # proteger toda la línea (sin el \n)
        pos += L + 1  # +1 por el salto de línea
    return spans

def intersecta_spans(o: int, ln: int, spans):
    """True si el rango [o, o+ln) toca algún (a,b) en spans."""
    if not spans or ln <= 0: return False
    r0, r1 = o, o + ln
    for a, b in spans:
        if not (r1 <= a or r0 >= b):
            return True
    return False

def restaurar_segmentos_protegidos(original: str, corregido: str, spans):
    """Copia desde 'original' los segmentos protegidos a 'corregido' (mismas posiciones)."""
    if not spans: return corregido
    out = []
    last = 0
    for a, b in sorted(spans):
        out.append(corregido[last:a])
        out.append(original[a:b])  # restaurar exactamente
        last = b
    out.append(corregido[last:])
    return "".join(out)

def apply_lt_corrections_smart(texto: str, matches: list, protected_spans=None) -> str:
    out = texto
    for m in sorted(matches, key=lambda x: x.get("offset",0), reverse=True):
        o = m.get("offset",0); ln = m.get("length",0)
        reps = m.get("replacements",[])
        if ln <= 0 or not reps: continue

        # NO tocar si cae en alternativa protegida
        if intersecta_spans(o, ln, protected_spans): 
            continue

        frag = out[o:o+ln]
        if len(frag) == 1 and frag.lower() in _MATH_VARS_1L: continue
        if _has_any_digit(frag): continue
        if is_upper_acronym(frag): continue
        prev = _prev_token(out, o)
        best = _choose_best_suggestion(frag, reps, prev)
        if best is None: continue
        if not _looks_reasonable_replacement(frag, best): continue
        best = _same_casing(frag, best)
        out = out[:o] + best + out[o+ln:]
    return out

# ===== PARCHE: modo clásico también ignora dígitos/siglas y exige parecido razonable
def apply_lt_corrections_classic(texto: str, matches: list, protected_spans=None) -> str:
    out = texto
    for m in sorted(matches, key=lambda x: x.get("offset",0), reverse=True):
        o = m.get("offset",0); ln = m.get("length",0)
        reps = m.get("replacements",[])
        if ln <= 0 or not reps:
            continue

        # NO tocar si cae en alternativa protegida
        if intersecta_spans(o, ln, protected_spans): 
            continue

        frag = out[o:o+ln]

        # Filtros de seguridad (mismos que smart)
        if len(frag) == 1 and frag.lower() in _MATH_VARS_1L:
            continue
        if _has_any_digit(frag):      # jamás tocar secuencias con dígitos
            continue
        if is_upper_acronym(frag):    # ni siglas mayúsculas
            continue

        prev = _prev_token(out, o)
        best = _choose_best_suggestion(frag, reps, prev)
        if best is None:
            continue

        if not _looks_reasonable_replacement(frag, best):
            continue

        best = _same_casing(frag, best)
        out = out[:o] + best + out[o+ln:]
    return out

def apply_lt_corrections(texto: str, matches: list, protected_spans=None) -> str:
    if USE_CLASSIC_LT:
        return apply_lt_corrections_classic(texto, matches, protected_spans=protected_spans)
    else:
        return apply_lt_corrections_smart(texto, matches, protected_spans=protected_spans)

# ====== REGLAS POST ======
POST_RULES = [
    (r'(?:(?<=^)|(?<=[\.\?\!]\s))((?:[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+|Él|El|Élla|Ella))\s+llegue\b', r'\1 llegó'),
    (r'\b([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+|Él|El|Élla|Ella)\s+llegue\b', r'\1 llegó'),
    (r'\bnumeros\b', 'números'),
    (r'\b(solo|sólo)\s+velos\b', r'\1 ve los'),
    (r'\benter(?:\s|[\u00AD\u200B\u200C\u200D\u2060])?s\b', 'enteros'),
]

def load_agreement_config():
    default = {
        "masc_nouns": ["número","vector","conjunto","ángulo","polígono","intervalo","resultado","valor","punto"],
        "fem_nouns":  ["fracción","suma","resta","potencia","matriz","media","mediana","varianza","raíz","medida"],
        "adjectives": ["entero","real","natural","primo","compuesto","racional","irracional",
                       "positivo","negativo","par","impar","mínimo","máximo","mayor","menor"]
    }
    path = os.path.join(DATA_DIR, "agreement_es.json")
    if os.path.isfile(path):
        try:
            with open(path,"r",encoding="utf-8") as f:
                cfg = json.load(f)
            for k in ["masc_nouns","fem_nouns","adjectives"]:
                if k not in cfg or not isinstance(cfg[k], list): cfg[k] = default[k]
            return cfg
        except Exception:
            return default
    else:
        try:
            with open(path,"w",encoding="utf-8") as f:
                json.dump(default,f,ensure_ascii=False,indent=2)
        except Exception:
            pass
        return default

_AGR_CFG   = load_agreement_config()
_MASC_NOUNS= set(_AGR_CFG["masc_nouns"])
_FEM_NOUNS = set(_AGR_CFG["fem_nouns"])
_ADJ       = set(_AGR_CFG["adjectives"])

def _agree_adj(noun: str, adj: str) -> str:
    def endcase(src, dst):
        if src.isupper(): return dst.upper()
        if src.istitle(): return dst[:1].upper() + dst[1:]
        return dst
    base = re.sub(r'(o|a|os|as)$', '', adj, flags=re.IGNORECASE)
    is_plural = noun.lower().endswith("s")
    noun_base = noun.lower().rstrip("s")
    if noun_base in _MASC_NOUNS: end = "os" if is_plural else "o"
    elif noun_base in _FEM_NOUNS: end = "as" if is_plural else "a"
    else: return adj
    fixed = base + end
    return endcase(adj, fixed)

def _fix_noun_adj_agreement(texto: str) -> str:
    if not texto: return texto
    def pluralize(w): return w + "s" if not w.endswith("s") else w
    masc = list(_MASC_NOUNS) + [pluralize(w) for w in _MASC_NOUNS]
    fem  = list(_FEM_NOUNS) + [pluralize(w) for w in _FEM_NOUNS]
    nouns_union = "|".join(sorted(set(masc + fem), key=len, reverse=True))
    adjs_union  = "|".join(sorted(_ADJ, key=len, reverse=True))
    patron = rf'\b({nouns_union})\s+(({adjs_union})(?:o|a|os|as)?)\b'
    def repl(m):
        noun = m.group(1); adj = m.group(2)
        base = re.sub(r'(o|a|os|as)$', '', adj, flags=re.IGNORECASE).lower()
        if base not in _ADJ: return m.group(0)
        return f"{noun} {_agree_adj(noun, adj)}"
    return re.sub(patron, repl, texto, flags=re.IGNORECASE)

def _fix_numero_entero_specific(t: str) -> str:
    t = re.sub(r'\bn[úu]mero\s+enter(?:a|as|os)\b', 'número entero', t, flags=re.IGNORECASE)
    t = re.sub(r'\bn[úu]meros\s+enter(?:o|a|as)\b', 'números enteros', t, flags=re.IGNORECASE)
    return t

_MATH_ADJ_NUMERO = {"entero","real","natural","primo","compuesto","racional",
                    "irracional","positivo","negativo","par","impar",
                    "mínimo","máximo","mayor","menor","cuadrado","cúbico"}

def _strip_gender_ending(w: str) -> str:
    wl = w.lower()
    for suf in ("os","as","o","a"):
        if wl.endswith(suf): return w[:-len(suf)]
    return w

def _fix_numero_fallback(texto: str) -> str:
    patron = rf'\b(n[uú]mero(?:s)?)\s+([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)\b'
    def repl(m):
        head = m.group(1)
        adj_full = m.group(2)
        lemma = _strip_gender_ending(adj_full).lower()
        if lemma not in _MATH_ADJ_NUMERO and (lemma + "o") not in _MATH_ADJ_NUMERO:
            return m.group(0)
        plural = head.lower().endswith("s")
        base = lemma if lemma in _MATH_ADJ_NUMERO else (lemma + "o")
        base = re.sub(r'enteroo$', 'entero', base)
        nuevo = base + ("os" if plural else "o")
        if adj_full.isupper(): nuevo = nuevo.upper()
        elif adj_full.istitle(): nuevo = nuevo[:1].upper() + nuevo[1:]
        return f"{head} {nuevo}"
    s = re.sub(patron, repl, texto, flags=re.IGNORECASE)
    s = re.sub(r'\benteroo(s)?\b', r'entero\1', s, flags=re.IGNORECASE)
    s = re.sub(r'\benteroos\b', 'enteros', s, flags=re.IGNORECASE)
    return s

# ------------------------
# Reglas de contexto + signos + mayúsculas
# ------------------------
_DIAS = ["lunes","martes","miércoles","jueves","viernes","sábado","domingo"]

def _lowercase_weekdays_mid_sentence(t: str) -> str:
    for d in _DIAS:
        pat = r'(?<!^)(?<![\.!\?\n]\s)'+d.capitalize()
        t = re.sub(pat, d, t)
    return t

def _context_fixes(t: str) -> str:
    if not t: return t
    t = re.sub(r'\b(Hay que)\s+de\s+', r'\1 ', t, flags=re.IGNORECASE)
    t = re.sub(r'(hasta[^.,;\n]*?)\s+avisa\b', r'\1, avisa', t, flags=re.IGNORECASE)
    t = re.sub(r'\bSe me olvidó\s+de\s+', 'Se me olvidó ', t, flags=re.IGNORECASE)
    t = _lowercase_weekdays_mid_sentence(t)
    return t

def _fix_del_contractions(t: str) -> str:
    return re.sub(
        r'\b([Dd])e\s+el\b',
        lambda m: ('D' if m.group(1) == 'D' else 'd') + 'el',
        t
    )

def _same_case_pair(src: str, dst_cap: str, dst_low: str) -> str:
    return dst_cap if src[:1].isupper() else dst_low

def _fix_aver_haber_context(t: str) -> str:
    def repl_haber_trigger(m):
        haber = m.group(1)
        trig  = m.group(2)
        return _same_case_pair(haber, "A ver", "a ver") + " " + trig
    t = re.sub(r'\b(Haber|haber)\s+(si|que|qué)\b', repl_haber_trigger, t)
    def repl_sentence_start(m):
        sep   = m.group(1)
        haber = m.group(2)
        after = m.group(3) or ""
        return f"{sep}{_same_case_pair(haber, 'A ver', 'a ver')}{after}"
    t = re.sub(r'(^|[\.!\?\n]\s)\b(Haber|haber)\b(\s*[,;:])', repl_sentence_start, t)
    return t

def _fix_cuyo_lado_agreement(t: str) -> str:
    def repl(m):
        lados = m.group(1)
        return 'cuyos lados' if lados.lower().endswith('s') else 'cuyo lado'
    t = re.sub(r'(?i)\bcuya(?:s)?\s+(lado|lados)\b', repl, t)
    t = re.sub(r'(?m)^(Cuya|Cuyas)\b', 'Cuyo', t)
    return t

# Prefijo de viñetas/numeración:  –  -  •   o bien  (1)  1)  1.
_BULLET_PREFIX_RX = re.compile(r'^\s*(?:[–\-•]\s*)?(?:\(?\d+[.)]\s*)?')

def _ensure_opening_mark(line: str) -> str:
    s = line
    def _do(close_mark, open_mark):
        nonlocal s
        if open_mark in s or close_mark not in s: return
        m = _BULLET_PREFIX_RX.match(s); pos = m.end() if m else 0
        s = s[:pos] + open_mark + s[pos:]
    _do("?", "¿"); _do("!", "¡")
    s = re.sub(r'\?{2,}\s*$', '?', s); s = re.sub(r'!{2,}\s*$', '!', s)
    m = _BULLET_PREFIX_RX.match(s); pos = m.end() if m else 0
    if s[pos:pos+2] == '¿¿': s = s[:pos] + '¿' + s[pos+2:]
    if s[pos:pos+2] == '¡¡': s = s[:pos] + '¡' + s[pos+2:]
    if '¿' in s and '?' not in s: s = s.rstrip() + '?'
    if '¡' in s and '!' not in s: s = s.rstrip() + '!'
    return s

def _add_opening_spanish_marks(t: str) -> str:
    return "\n".join(_ensure_opening_mark(ln) for ln in (t or "").split("\n"))

_LOWER_START_RX = re.compile(r'(^|(?<=[\.!\?]\s))([a-záéíóúñ])', flags=re.MULTILINE)
_NUM_BULLET_RX  = re.compile(r'^(\s*\d+\)\s*)([a-záéíóúñ])',      flags=re.MULTILINE)

def _capitalize_starts(t: str) -> str:
    t = _LOWER_START_RX.sub(lambda m: m.group(1) + m.group(2).upper(), t)
    t = _NUM_BULLET_RX.sub (lambda m: m.group(1) + m.group(2).upper(), t)
    return t

def post_correcciones(t: str) -> str:
    t = t or ""
    t = normalize_ocr_noise(t)
    for pat, repl in POST_RULES:
        t = re.sub(pat, repl, t, flags=re.IGNORECASE)
    t = _fix_numero_entero_specific(t)
    t = _fix_numero_fallback(t)
    t = _fix_noun_adj_agreement(t)
    t = _context_fixes(t)
    t = _fix_del_contractions(t)
    t = _fix_aver_haber_context(t)
    t = _fix_cuyo_lado_agreement(t)
    t = _add_opening_spanish_marks(t)
    t = _capitalize_starts(t)
    return t

# ====== DOCX → combinar con original preservando estilos (con opción de resaltado)
def tokenize_preservando(texto: str):
    if not texto: return []
    partes = TOKEN_RX.split(texto)
    return [t for t in partes if t is not None and t != ""]

def es_token_palabra(tok: str):
    return not TOKEN_RX.fullmatch(tok)

def reemplazo_en_runs_parciales(paragraph, old_txt: str, new_txt: str):
    if not old_txt: return False
    for run in paragraph.runs:
        if contiene_omml(run): continue
        t = run.text or ""
        idx = t.find(old_txt)
        if idx != -1:
            run.text = t[:idx] + new_txt + t[idx+len(old_txt):]
            return True
    return False

def reemplazo_en_runs_flexible(paragraph, old_txt: str, new_txt: str):
    if not old_txt: return False
    eligible = []
    for run in paragraph.runs:
        eligible.append((run, None if contiene_omml(run) else (run.text if run.text else "")))
    concat = []; index_map = []
    for idx, (run, txt) in enumerate(eligible):
        if txt is None: continue
        for k, ch in enumerate(txt):
            concat.append(ch); index_map.append((idx, k))
    full = "".join(concat); pos = full.find(old_txt)
    if pos == -1: return False
    start = pos; end = pos + len(old_txt)
    if start >= len(index_map) or end > len(index_map): return False
    first_run_idx = index_map[start][0]; last_run_idx  = index_map[end - 1][0]
    for idx in range(first_run_idx, last_run_idx + 1):
        _run, _txt = eligible[idx]
        if _txt is None: return False
    old_block = ""; per_run_text = []
    for idx in range(first_run_idx, last_run_idx + 1):
        _run, _txt = eligible[idx]; per_run_text.append(_txt); old_block += _txt
    acc_before = 0
    for idx in range(first_run_idx):
        _run, _txt = eligible[idx]
        if _txt is not None: acc_before += len(_txt)
    rel_start = start - acc_before; rel_end = rel_start + len(old_txt)
    new_block = old_block[:rel_start] + new_txt + old_block[rel_end:]
    lens = [len(x) for x in per_run_text]; rebuilt = []; cursor = 0
    for k, L in enumerate(lens):
        if k < len(lens) - 1:
            rebuilt.append(new_block[cursor: cursor + L]); cursor += L
        else:
            rebuilt.append(new_block[cursor:])
    for off, idx in enumerate(range(first_run_idx, last_run_idx + 1)):
        run, _txt = eligible[idx]; run.text = rebuilt[off]
    return True

def pares_reemplazo_palabra_a_palabra(texto_original: str, texto_corregido: str):
    orig_toks = tokenize_preservando(texto_original)
    corr_toks = tokenize_preservando(texto_corregido)
    sm = SequenceMatcher(a=orig_toks, b=corr_toks, autojunk=False)
    pares = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "replace":
            seg_o = orig_toks[i1:i2]; seg_c = corr_toks[j1:j2]
            if len(seg_o) == 1 and len(seg_c) == 1 and es_token_palabra(seg_o[0]) and es_token_palabra(seg_c[0]):
                if seg_o[0] != seg_c[0]: pares.append((seg_o[0], seg_c[0]))
            else:
                for so, sc in zip(seg_o, seg_c):
                    if es_token_palabra(so) and es_token_palabra(sc) and so != sc:
                        pares.append((so, sc))
    vistos = set(); dedup = []
    for a, b in pares:
        key = (a, b)
        if key not in vistos:
            dedup.append((a, b)); vistos.add(key)
    return dedup

def similitud_suave(a: str, b: str):
    return SequenceMatcher(None, a, b).ratio()

# ---------- Forzado de bloque “A ver” / “Por favor” ----------
def _force_block_highlight_phrases(paragraph, original_text: str, candidate_text: str):
    orig_l = (original_text or "").lower().replace("\u00ad","")
    cand   = candidate_text or ""
    need_aver  = any(w in orig_l for w in ["haber", "aver", "haver"]) and "a ver" in cand.lower()
    need_porf  = "porfavor" in orig_l and "por favor" in cand.lower()
    if not (need_aver or need_porf): return
    if need_aver:
        for m in re.finditer(r'\b[Aa]\s+ver\b', cand):
            a, b = m.span(); _highlight_range_by_char(paragraph, a, b)
    if need_porf:
        for m in re.finditer(r'\b[Pp]or\s+favor\b', cand):
            a, b = m.span(); _highlight_range_by_char(paragraph, a, b)
# -----------------------------------------------------------------

# ====== generar_docx_corregido (con anti-falsos positivos + highlight opcional) ======
def generar_docx_corregido(path_docx_original: str, texto_corregido: str, path_salida_docx: str,
                           highlight: bool = False, texto_original_para_highlight: str | None = None):
    from docx import Document as _Doc
    import unicodedata as _unic
    def nfc(s: str) -> str: return _unic.normalize("NFC", s or "")
    def simil(a: str, b: str) -> float: return SequenceMatcher(None, a, b).ratio()
    def parrafo_tiene_objetos(p) -> bool:
        r = p._p
        if r.xpath('.//m:oMath | .//m:oMathPara'): return True
        if r.xpath('.//w:drawing | .//w:pict | .//w:object | .//w:fldChar'): return True
        return False
    def _reemplazos_por_diff(original: str, candidato: str):
        rep = []; sm = SequenceMatcher(a=original, b=candidato, autojunk=False)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag != "replace": continue
            old = original[i1:i2]; new = candidato[j1:j2]
            if not old or not new: continue
            if max(len(old), len(new)) > 64: continue
            rep.append((old, new))
        vistos, out = set(), []
        for a, b in rep:
            k = (a, b)
            if k not in vistos: out.append((a, b)); vistos.add(k)
        return out
    def _reemplazo_en_runs_flexible(paragraph, old_txt: str, new_txt: str) -> bool:
        return reemplazo_en_runs_flexible(paragraph, old_txt, new_txt)
    def _reescribir_parrafo_preservando_runs(paragraph, new_text: str):
        runs = paragraph.runs; lens = [len(r.text or "") for r in runs]; total = sum(lens)
        if total == 0:
            for r in runs:
                if not contiene_omml(r): r.text = ""
            return
        cursor = 0
        for i, r in enumerate(runs):
            if contiene_omml(r): continue
            L = lens[i]
            if i < len(runs) - 1:
                r.text = new_text[cursor: cursor + L]; cursor += L
            else:
                r.text = new_text[cursor:]
    def _reemplazo_pequeno_valido(old: str, new: str) -> bool:
        if not old or not new: return False
        ratio_len = len(new) / max(1, len(old))
        if ratio_len < 0.40 or ratio_len > 1.60: return False
        if simil(old, new) < 0.60: return False
        return True

    # ---- Detectores y firmas de alternativas/viñetas ----
    _RX_LETTER_OPT = re.compile(r'^\s*([A-Ea-e])\)\s*')     # A) B) C) ...
    _RX_NUM_BULLET = re.compile(r'^\s*\(?(\d+)[\.)]\s*')    # (1) 1) 1. ...
    _RX_DIGITS     = re.compile(r'\d+')

    def _option_label(line: str) -> tuple[str, str] | None:
        if not line: return None
        m = _RX_LETTER_OPT.match(line)
        if m: return ('LETTER', m.group(1).upper())
        m = _RX_NUM_BULLET.match(line)
        if m: return ('NUM', m.group(1))
        return None

    def _same_option_label(a: str, b: str) -> bool:
        la = _option_label(a); lb = _option_label(b)
        if not la or not lb: return False
        return la == lb

    def _numbers_signature(s: str) -> tuple:
        return tuple(_RX_DIGITS.findall(s or ""))

    def _safe_option_rewrite_ok(original_line: str, candidate_line: str) -> bool:
        """
        Permite reescritura SOLO si:
        - misma etiqueta (ya se filtró),
        - misma firma numérica (mismos números y orden),
        - y similitud alta (>= 0.92).
        """
        if _numbers_signature(original_line) != _numbers_signature(candidate_line):
            return False
        sim = SequenceMatcher(None, original_line, candidate_line).ratio()
        return sim >= 0.92

    doc = _Doc(path_docx_original)
    parrafos_docx = doc.paragraphs
    parrafos_corregidos = [nfc(s) for s in texto_corregido.splitlines()]
    usados = set()

    for i, p in enumerate(parrafos_docx):
        original = p.text or ""; original_nfc = nfc(original)

        # --- Selección de candidato robusta y estable ---
        candidatos = []
        if i < len(parrafos_corregidos): candidatos.append((i, parrafos_corregidos[i]))
        for j in range(max(0, i-2), min(len(parrafos_corregidos), i+3)):
            if j != i: candidatos.append((j, parrafos_corregidos[j]))

        # Si es alternativa/viñeta, exige misma etiqueta y prioriza mismo índice.
        label_orig = _option_label(original_nfc)
        if label_orig:
            candidatos = [(j, c) for (j, c) in candidatos if _same_option_label(original_nfc, c)]
            if not candidatos:
                candidatos = [(i, original_nfc)]  # no tocar
            else:
                mismos_idx = [(j, c) for (j, c) in candidatos if j == i]
                if mismos_idx: candidatos = mismos_idx

        mejor_idx, mejor_score = None, -1.0
        for j, cand in candidatos:
            score = simil(original_nfc, cand) - 0.03 * abs(i - j)
            if score > mejor_score:
                mejor_score, mejor_idx = score, j

        if mejor_idx is None:
            continue

        candidato = parrafos_corregidos[mejor_idx]

        # Umbrales base
        score_crudo = simil(original_nfc, candidato)
        umbral = 0.55 if mejor_idx == i else 0.80
        if score_crudo < umbral:
            usados.add(mejor_idx); continue

        # Si cambió el índice y hay dígitos, no tocar
        if mejor_idx != i and any(ch.isdigit() for ch in original_nfc):
            usados.add(mejor_idx); continue

        # Diferencia de longitud controlada
        max_diff = 0.45 if mejor_idx == i else 0.30
        if len(original_nfc) > 0 and abs(len(candidato) - len(original_nfc)) / max(1, len(original_nfc)) > max_diff:
            usados.add(mejor_idx); continue

        # Reglas estrictas para alternativas/viñetas
        if label_orig:
            # Debe mantener misma firma numérica y ser casi idéntico
            if not _safe_option_rewrite_ok(original_nfc, candidato):
                usados.add(mejor_idx); continue

        tiene_obj = parrafo_tiene_objetos(p)

        if not tiene_obj:
            _reescribir_parrafo_preservando_runs(p, candidato)
            if highlight:
                try:
                    sm  = SequenceMatcher(a=original_nfc, b=candidato, autojunk=False)
                    ops = list(sm.get_opcodes())
                    def _has_old(op): return op[2] > op[1]
                    k = 0
                    while k < len(ops):
                        tag, i1, i2, j1, j2 = ops[k]
                        if k == 0 or ops[k-1][4] != j1:
                            q = k + 1; jL, jR = j1, j2; iL, iR = (i1, i2) if _has_old(ops[k]) else (None, None)
                            while q < len(ops) and ops[q][3] == jR:
                                _, qi1, qi2, qj1, qj2 = ops[q]
                                jR = qj2
                                if qi2 > qi1:
                                    if iL is None: iL, iR = qi1, qi2
                                    else: iL = min(iL, qi1); iR = max(iR, qi2)
                                q += 1
                            oldfrag = original_nfc[iL:iR] if iL is not None else ""
                            newfrag = candidato[jL:jR]
                            if oldfrag and _is_single_word(oldfrag) and _is_two_or_more_words(newfrag):
                                _highlight_range_by_char(p, jL, jR); k = q; continue
                        tag, i1, i2, j1, j2 = ops[k]
                        if tag == "insert":
                            if original_nfc.strip():
                                frag = candidato[j1:j2]; _highlight_tokens_in_range(p, j1, frag)
                        elif tag == "replace":
                            oldfrag = original_nfc[i1:i2]; newfrag = candidato[j1:j2]
                            _highlight_replacement_smart(p, j1, oldfrag, newfrag)
                        k += 1
                    _force_block_highlight_phrases(p, original_nfc, candidato)
                    pat = r'(?:' + re.escape(MARK_WORD_DELETE) + r'|' + re.escape(MARK_SPACE_DELETE) + r')'
                    for m in re.finditer(pat, candidato):
                        a, b = m.span(); _highlight_range_by_char(p, a, b)
                except Exception:
                    pass
        else:
            rep_palabra = pares_reemplazo_palabra_a_palabra(original_nfc, candidato)
            rep_diff    = _reemplazos_por_diff(original_nfc, candidato)
            vistos, reemplazos = set(), []
            for par in rep_palabra + rep_diff:
                if par not in vistos: reemplazos.append(par); vistos.add(par)
            for old, new in reemplazos:
                old_n, new_n = nfc(old), nfc(new)
                if not _reemplazo_pequeno_valido(old_n, new_n): continue
                changed = False
                if reemplazo_en_runs_parciales(p, old_n, new_n): changed = True
                elif _reemplazo_en_runs_flexible(p, old_n, new_n): changed = True
                if highlight and changed:
                    concat = "".join(r.text or "" for r in p.runs if not contiene_omml(r))
                    start = 0
                    while True:
                        pos = concat.find(new_n, start)
                        if pos == -1: break
                        if _is_single_word(old_n) and _is_two_or_more_words(new_n):
                            _highlight_range_by_char(p, pos, pos + len(new_n))
                        else:
                            _highlight_tokens_in_range(p, pos, new_n)
                        start = pos + len(new_n)
            if highlight:
                concat = "".join(r.text or "" for r in p.runs if not contiene_omml(r))
                _force_block_highlight_phrases(p, original_nfc, concat)
                pat = r'(?:' + re.escape(MARK_WORD_DELETE) + r'|' + re.escape(MARK_SPACE_DELETE) + r')'
                for m in re.finditer(pat, concat):
                    a, b = m.span(); _highlight_range_by_char(p, a, b)
        usados.add(mejor_idx)
    doc.save(path_salida_docx)

# ====== HELPERS DESCARGAS / CACHE ======
def _save_into_descargas(src_path: str, target_name: str) -> str:
    os.makedirs(DESCARGAS_DIR, exist_ok=True)
    dst = os.path.join(DESCARGAS_DIR, target_name); copy2(src_path, dst); return dst

def _sha1_file(path: str, block=1024*1024) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        while True:
            b = f.read(block)
            if not b: break
            h.update(b)
    return h.hexdigest()

# ====== Exportar vistas ======
def generar_pdf(ruta_word: str) -> str:
    # Usamos ruta corta 8.3 por si hay espacios o rutas largas
    ruta_word_abs   = os.path.abspath(ruta_word)
    ruta_word_short = _short83(ruta_word_abs)

    tmp = tempfile.NamedTemporaryFile(prefix="ltpdf_", suffix=".pdf", delete=False)
    ruta_pdf = tmp.name
    tmp.close()

    wdExportFormatPDF         = 17
    wdExportOptimizeForPrint  = 0
    wdExportAllDocument       = 0
    wdExportDocumentContent   = 0
    wdExportCreateNoBookmarks = 0

    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        doc = word.Documents.Open(ruta_word_short, ReadOnly=True)
        doc.ExportAsFixedFormat(
            OutputFileName=ruta_pdf,
            ExportFormat=wdExportFormatPDF,
            OpenAfterExport=False,
            OptimizeFor=wdExportOptimizeForPrint,
            Range=wdExportAllDocument,
            From=1,
            To=1,
            Item=wdExportDocumentContent,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=wdExportCreateNoBookmarks,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
        doc.Close(False)

        for _ in range(20):
            if os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 0:
                break
            time.sleep(0.1)

        if not os.path.exists(ruta_pdf) or os.path.getsize(ruta_pdf) == 0:
            raise RuntimeError("Word no generó el PDF (archivo vacío).")

        return ruta_pdf
    finally:
        try:
            word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()


# === NUEVO: generar PDF estable para previsualización (original/corregido) ===
def generar_pdf_preview(ruta_docx: str, nombre_base: str | None = None) -> str:
    """
    Genera (o reutiliza) un PDF en DESCARGAS_DIR a partir de un DOCX,
    pensado para previsualización en el navegador (iframe).
    """
    ruta_docx = os.path.abspath(ruta_docx)
    os.makedirs(DESCARGAS_DIR, exist_ok=True)

    # Nombre base estable por contenido
    h = _sha1_file(ruta_docx)[:12]
    base = nombre_base or f"{h}_preview"
    dst_pdf = os.path.join(DESCARGAS_DIR, f"{base}.pdf")

    # Si ya existe y es válido, reutilizar
    if os.path.exists(dst_pdf) and os.path.getsize(dst_pdf) > 0:
        return dst_pdf

    # Generar PDF temporal vía Word y copiar a descargas
    tmp_pdf = generar_pdf(ruta_docx)
    copy2(tmp_pdf, dst_pdf)
    try:
        os.remove(tmp_pdf)
    except Exception:
        pass

    return dst_pdf

def _force_utf8_html(html_path: str):
    import re
    try:
        try:
            txt = open(html_path, "r", encoding="utf-8").read()
            decoded_as = "utf-8"
        except UnicodeDecodeError:
            txt = open(html_path, "r", encoding="cp1252", errors="strict").read()
            decoded_as = "cp1252"
        # Normalizar declaración charset a utf-8
        txt = re.sub(r'(?i)charset\s*=\s*["\']?[-\w]+["\']?', 'charset=utf-8', txt, count=1)
        txt = re.sub(r'(?i)<meta\s+charset=["\']?[-\w]+["\']?\s*/?>', '<meta charset="utf-8">', txt, count=1)
        if '<meta' not in txt.lower():
            txt = txt.replace('<head>', '<head><meta charset="utf-8">', 1)
        with open(html_path, "w", encoding="utf-8", newline="") as f: f.write(txt)
        print(f"[HTML] {os.path.basename(html_path)} normalizado a UTF-8 (desde {decoded_as}).")
    except Exception as e:
        print("[HTML] No se pudo forzar UTF-8:", e)

def _postprocess_word_html(html_path: str):
    try:
        with open(html_path, "r", encoding="utf-8", errors="ignore") as f: txt = f.read()
        txt = re.sub(r'<!--\s*\[if\s+gte\s+vml\s+1\s*\]>.*?<!\s*\[endif\]\s*-->', '', txt, flags=re.I | re.S)
        txt = re.sub(r'<!--\s*\[if\s*!vml\s*\]-->', '', txt, flags=re.I)
        txt = re.sub(r'<!--\s*<!\s*\[endif\]\s*-->', '', txt, flags=re.I)
        extra_css = "<style>img{max-width:none!important;height:auto;}body{margin:0;}</style>"
        if "</head>" in txt.lower(): txt = re.sub(r'</head>', extra_css + '</head>', txt, flags=re.I, count=1)
        else: txt = extra_css + txt
        with open(html_path, "w", encoding="utf-8", newline="") as f: f.write(txt)
    except Exception as e:
        print("[HTML] Post-proceso fallido:", e)

def generar_html_desde_docx(ruta_docx: str, nombre_base: str | None = None) -> str:
    wdFormatFilteredHTML = 10; msoEncodingUTF8 = 65001
    ruta_docx = os.path.abspath(ruta_docx)
    h = _sha1_file(ruta_docx)[:12]; base = nombre_base or f"{h}_preview"
    dst_html = os.path.join(DESCARGAS_DIR, f"{base}.htm")
    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application"); word.Visible = False; word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(ruta_docx, ReadOnly=True)
        try:
            doc.WebOptions.AllowPNG = True; doc.WebOptions.OptimizeForBrowser = True; doc.WebOptions.RelyOnCSS = True
        except Exception: pass
        doc.SaveAs2(FileName=os.path.abspath(dst_html), FileFormat=wdFormatFilteredHTML, Encoding=msoEncodingUTF8)
        doc.Close(False)
    finally:
        try: word.Quit()
        except Exception: pass
        pythoncom.CoUninitialize()
    _force_utf8_html(dst_html); _postprocess_word_html(dst_html); return dst_html

# ====== RUTAS ======
@app.route("/login", methods=["POST"])
@cross_origin()
def login():
    data = request.json; usuario = data["usuario"]; clave = data["clave"]
    conn = get_connection(); cursor = conn.cursor()
    cursor.execute("SELECT * FROM usuarios WHERE username=%s AND password=%s", (usuario, clave))
    user = cursor.fetchone(); cursor.close(); conn.close()
    if user: return jsonify({"status": "ok", "mensaje": "Login correcto"})
    return jsonify({"status": "error", "mensaje": "Credenciales inválidas"}), 401

@app.route("/lt/status")
def lt_status(): return jsonify({"running": lt_is_running(), "dir": LT_DIR})

@app.route("/lt/ensure")
def lt_ensure():
    try:
        lt_start_server()
        return jsonify({"ok": True, "running": lt_is_running(), "dir": LT_DIR, "port": LT_PORT})
    except Exception as e:
        traceback.print_exc(); return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/descargas/<path:nombre>", methods=["GET"])
def serve_descargas_inline(nombre):
    path = os.path.join(DESCARGAS_DIR, nombre)
    if not os.path.exists(path):
        return "No existe el archivo", 404

    # Si es PDF → mostrar en el iframe (inline)
    if nombre.lower().endswith(".pdf"):
        return send_file(path, mimetype="application/pdf")

    # Cualquier otro archivo (docx, etc.) → descarga
    return send_file(path, as_attachment=True)


@app.route("/api/descargar_pdf_corregido/<path:nombre_docx>", methods=["GET"])
def descargar_pdf_corregido(nombre_docx):
    try:
        docx_path = os.path.join(DESCARGAS_DIR, nombre_docx)
        if not os.path.exists(docx_path): return jsonify({"ok": False, "error": f"No existe DOCX: {nombre_docx}"}), 404
        pdf_tmp = generar_pdf(docx_path)
        base = os.path.splitext(os.path.basename(nombre_docx))[0]
        download_name = f"{base}.pdf"
        @after_this_request
        def _cleanup(response):
            try:
                if os.path.exists(pdf_tmp): os.remove(pdf_tmp)
            except Exception: pass
            return response
        return send_file(pdf_tmp, as_attachment=True, download_name=download_name,
                         mimetype="application/pdf", max_age=0)
    except Exception as e:
        traceback.print_exc(); return jsonify({"ok": False, "error": str(e)}), 500

# === Previsualización del original .docx como PDF (visor del navegador) ===
@app.route("/api/render_vista", methods=["POST"])
def render_vista():
    """
    Recibe un .docx, lo convierte a PDF y devuelve la URL para
    previsualizarlo en el iframe del corrector.
    """
    try:
        if "archivo" not in request.files:
            return jsonify({"ok": False, "error": "Falta 'archivo'"}), 400

        up = request.files["archivo"]
        if not up.filename:
            return jsonify({"ok": False, "error": "Nombre de archivo vacío"}), 400

        filename = secure_filename(up.filename)
        base, ext = os.path.splitext(filename)
        if ext.lower() != ".docx":
            return jsonify({"ok": False, "error": "Solo se admite .docx para la vista."}), 400

        # Aseguramos carpetas
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        os.makedirs(DESCARGAS_DIR, exist_ok=True)

        # Guardar DOCX original
        docx_path = os.path.join(UPLOAD_DIR, filename)
        up.save(docx_path)

        # Convertir a PDF con la función que ya usas en exportar_examen
        pdf_tmp = generar_pdf(docx_path)

        # Nombre estable para la vista
        pdf_name = f"{base}_vista.pdf"
        pdf_path = os.path.join(DESCARGAS_DIR, pdf_name)

        # Copiar el PDF temporal a la carpeta descargas
        copy2(pdf_tmp, pdf_path)
        try:
            os.remove(pdf_tmp)
        except Exception:
            pass

        # Devolvemos la URL que el frontend usará en el iframe
        return jsonify({"ok": True, "html_url": f"/descargas/{pdf_name}"})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500

# === Render de DOCX corregido (ahora como PDF para preview) ===
@app.route("/api/render_docx_guardado/<path:nombre_docx>")
def render_docx_guardado(nombre_docx):
    # Si piden el limpio, intentamos mostrar el de PREVIEW con resaltados
    if nombre_docx.endswith("_corregido_limpio.docx"):
        candidato_preview = nombre_docx.replace("_corregido_limpio.docx", "_corregido.docx")
        preview_path = os.path.join(DESCARGAS_DIR, candidato_preview)
        if os.path.exists(preview_path):
            nombre_docx = candidato_preview

    docx_path = os.path.join(DESCARGAS_DIR, nombre_docx)
    if not os.path.exists(docx_path): return jsonify({"ok": False, "error": "No existe DOCX"}), 404
    try:
        base_in = os.path.splitext(nombre_docx)[0]

        # Generar (o reutilizar) PDF de preview desde el DOCX corregido
        pdf_path = generar_pdf_preview(
            docx_path,
            nombre_base=f"{base_in}_preview"
        )
        pdf_name = os.path.basename(pdf_path)

        # Seguimos usando 'html_url' para compatibilidad con el frontend
        return jsonify({"ok": True, "html_url": f"/descargas/{pdf_name}"})
    except Exception as e:
        traceback.print_exc(); return jsonify({"ok": False, "error": str(e)}), 500

# ====== CORRECCIÓN (con preview resaltado y descarga limpia) ======
@app.route('/api/corregir_archivo', methods=['POST'])
def corregir_archivo():
    try:
        if 'archivo' not in request.files: return jsonify({"ok": False, "error": "Falta 'archivo'"}), 400
        up = request.files['archivo']
        if not up.filename: return jsonify({"ok": False, "error": "Nombre de archivo vacío"}), 400

        idioma = (request.form.get('idioma') or 'es').strip()
        modo   = (request.form.get('modo')   or 'corregir').strip().lower()

        nombre_archivo = secure_filename(up.filename)
        path_in = os.path.join(UPLOAD_DIR, nombre_archivo)
        up.save(path_in)

        if not nombre_archivo.lower().endswith('.docx'):
            return jsonify({"ok": False, "error": "Ahora solo se admite .docx."}), 400

        texto = extraer_texto_docx(path_in)
        texto = normalize_ocr_noise(texto)

        # NUEVO: detectar y proteger spans de alternativas tipo "A) "..."E) "
        spans_alternativas = detectar_spans_alternativas(texto)

        if modo == "preview":
            return jsonify({
                "ok": True, "texto_original": texto, "texto_corregido": "",
                "corregido_html_inline": "", "total_alertas": 0, "descargas": {}
            })

        try:
            resp = lt_check_smart(texto, lang=idioma)
            matches = resp.get("matches", [])
        except Exception as e:
            traceback.print_exc()
            return jsonify({"ok": False, "error": f"LanguageTool: {e}"}), 500

        # Aplicar correcciones SIN tocar alternativas
        texto_corregido_base    = apply_lt_corrections(texto, matches, protected_spans=spans_alternativas)
        texto_corregido_puro    = post_correcciones(texto_corregido_base)

        # Por si alguna regla "post" tocara algo dentro del span, restauramos exactamente:
        texto_corregido_puro    = restaurar_segmentos_protegidos(texto, texto_corregido_puro, spans_alternativas)

        # Marcado de eliminaciones (se genera desde el puro para que el diff ignore alternativas)
        texto_corregido_marcado = insertar_marcas_eliminacion(texto, texto_corregido_puro)

        base = os.path.splitext(os.path.basename(nombre_archivo))[0]
        out_docx_preview = os.path.join(DESCARGAS_DIR, f"{base}_corregido.docx")
        out_docx_clean   = os.path.join(DESCARGAS_DIR, f"{base}_corregido_limpio.docx")

        generar_docx_corregido(path_in, texto_corregido_marcado, out_docx_preview,
                               highlight=True, texto_original_para_highlight=texto)
        generar_docx_corregido(path_in, texto_corregido_puro, out_docx_clean, highlight=False)

        corregido_html_inline = (texto_corregido_marcado
                                 .replace("&", "&amp;").replace("<", "&lt;")
                                 .replace(">", "&gt;").replace("\n", "<br/>"))

        return jsonify({
            "ok": True,
            "total_alertas": len(matches),
            "nombre_archivo_base": base,
            "descargas": { "docx": f"/descargas/{os.path.basename(out_docx_clean)}" },
            "texto_original": texto[:50000],
            "texto_corregido": texto_corregido_marcado[:50000],
            "diff_html_inline": "",
            "corregido_html_inline": corregido_html_inline
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"ok": False, "error": str(e)}), 500


#
#fin corrector ortografico
@app.route("/api/examenes")
def obtener_examenes():
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT idexamenes, nombre, numero, institucion, anio FROM examenes")
        examenes = cursor.fetchall()
        cursor.close()
        conn.close()
        return jsonify(examenes)
    except Exception as e:
        return jsonify({"status": "error", "detalle": str(e)}), 500


@app.route('/api/examenes/<int:idexamen>', methods=['DELETE'])
def eliminar_examen(idexamen):
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM examenes WHERE idexamenes = %s", (idexamen,))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'mensaje': 'Examen eliminado correctamente'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/importar_examen', methods=['POST'])

def importar_examen():
    if 'archivo' not in request.files:
        return jsonify({'error': 'No se envió archivo'}), 400

    archivo = request.files['archivo']
    if archivo.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400

    # ✅ Usar nombre original para extraer datos
    nombre_original = archivo.filename
    nombre_sin_extension = nombre_original.replace(".docx", "")

    # ✅ Expresión regular para extraer datos
    patron = r"examen\s+([a-zA-ZáéíóúÁÉÍÓÚñÑ\s]+)\s+(I{1,2})\s+([a-zA-ZáéíóúÁÉÍÓÚñÑ]+)\s+(\d{4})"
    match = re.search(patron, nombre_sin_extension, re.IGNORECASE)

    if not match:
        return jsonify({'error': 'Nombre de archivo no tiene formato válido'}), 400

    # ✅ Extraer valores
    nombre = f"Examen {match.group(1).strip().title()}"
    numero = match.group(2)
    institucion = match.group(3).upper()
    anio = int(match.group(4))

    # ✅ Guardar archivo de forma segura
    filename = secure_filename(nombre_original)
    ruta_archivo = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    archivo.save(ruta_archivo)

    try:
        conn = get_connection()
        cursor = conn.cursor()
        # Verificar si ya existe un archivo con el mismo nombre
        cursor.execute("SELECT COUNT(*) FROM examenes WHERE archivo_nombre = %s", (filename,))
        existe = cursor.fetchone()[0]

        if existe > 0:
            cursor.close()
            conn.close()
            return jsonify({'error': 'Ya se ha importado un examen con ese nombre de archivo'}), 400



        sql = """
            INSERT INTO examenes (nombre, numero, institucion, anio, archivo_nombre, archivo_ruta)
            VALUES (%s, %s, %s, %s, %s, %s)
        """
        cursor.execute(sql, (nombre, numero, institucion, anio, filename, ruta_archivo))
        conn.commit()
        cursor.close()
        conn.close()
        return jsonify({'exito': True, 'mensaje': 'Examen importado correctamente'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    


@app.route("/api/exportar_examen/<int:idexamen>")
def exportar_examen(idexamen):
    formato = (request.args.get("formato") or "pdf").lower()
    if formato not in ("pdf", "word"):
        return jsonify({"error": "formato inválido (pdf|word)"}), 400

    # 1) Obtener ruta del DOCX desde la BD
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute(
            "SELECT archivo_ruta, archivo_nombre FROM examenes WHERE idexamenes=%s",
            (idexamen,)
        )
        row = cur.fetchone()
        cur.close(); conn.close()
    except Exception as e:
        return jsonify({"error": f"DB error: {e}"}), 500

    if not row:
        return jsonify({"error": "No existe el examen"}), 404

    ruta_docx = os.path.abspath(row["archivo_ruta"])
    if not os.path.isfile(ruta_docx):
        return jsonify({"error": f"No existe el archivo en disco: {ruta_docx}"}), 500

    # 2) Entregar DOCX
    if formato == "word":
        return send_file(
            ruta_docx,
            as_attachment=True,
            download_name=row["archivo_nombre"],
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            max_age=0,
        )

    # 3) Convertir a PDF con tu función generar_pdf(ruta_word)
    try:
        ruta_pdf = generar_pdf(ruta_docx)  # <- tu función existente
        # asegura ruta absoluta
        ruta_pdf = os.path.abspath(ruta_pdf)

        # enviar y borrar el temporal al finalizar
        resp = send_file(
            ruta_pdf,
            as_attachment=True,
            download_name=os.path.splitext(row["archivo_nombre"])[0] + ".pdf",
            mimetype="application/pdf",
            max_age=0,
        )

        @resp.call_on_close
        def _cleanup():
            try:
                if os.path.exists(ruta_pdf):
                    os.remove(ruta_pdf)
            except Exception:
                pass

        return resp

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": f"Fallo al convertir a PDF: {e}"}), 500
def generar_pdf(ruta_word: str) -> str:
    ruta_word = os.path.abspath(ruta_word)

    # carpeta temporal única donde dejaremos el PDF
    tmpdir = tempfile.mkdtemp(prefix="pdf_")
    nombre_pdf = os.path.splitext(os.path.basename(ruta_word))[0] + ".pdf"
    ruta_pdf = os.path.join(tmpdir, nombre_pdf)

    # Constantes Word
    wdExportFormatPDF = 17
    wdExportOptimizeForPrint = 0
    wdExportAllDocument = 0
    wdExportDocumentContent = 0
    wdExportCreateNoBookmarks = 0

    # Requerido para hilos que crean objetos COM
    pythoncom.CoInitialize()

    # Inicia Word en un proceso aislado y silencioso
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        # Abrir solo lectura
        doc = word.Documents.Open(ruta_word, ReadOnly=True)

        # Exportar como PDF (fiel al diseño impreso)
        doc.ExportAsFixedFormat(
            OutputFileName=ruta_pdf,
            ExportFormat=wdExportFormatPDF,
            OpenAfterExport=False,
            OptimizeFor=wdExportOptimizeForPrint,
            Range=wdExportAllDocument,
            From=1, To=1,  # ignorados cuando Range = All
            Item=wdExportDocumentContent,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=wdExportCreateNoBookmarks,
            DocStructureTags=True,      # accesibilidad/estructura
            BitmapMissingFonts=True,    # si falta una fuente, rasteriza
            UseISO19005_1=False         # pon True si quieres PDF/A-1b
        )
        doc.Close(False)

        # A veces Word tarda un instante en cerrar el handle al archivo
        for _ in range(10):
            if os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 0:
                break
            time.sleep(0.1)

        if not os.path.exists(ruta_pdf) or os.path.getsize(ruta_pdf) == 0:
            raise RuntimeError("Word no generó el PDF (archivo vacío).")

        return ruta_pdf

    finally:
        try:
            word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()
#
#fin de corrector ortofgrafico
#
@app.route("/api/examen_nombre/<int:idexamen>")
def examen_nombre(idexamen):
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute(
            "SELECT archivo_nombre FROM examenes WHERE idexamenes=%s",
            (idexamen,)
        )
        row = cur.fetchone()
        cur.close(); conn.close()
        if not row:
            return jsonify({"error": "No existe el examen"}), 404
        return jsonify({"archivo_nombre": row["archivo_nombre"]})
    except Exception as e:
        return jsonify({"error": f"DB error: {e}"}), 500
    
# ---------------------------
# CRUD TEMAS (cursos)
# ---------------------------

@app.route("/api/temas", methods=["GET"])
def temas_listar():
    """?all=1 para incluir inactivos; por defecto solo activos."""
    include_all = request.args.get("all") == "1"
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)

        # Opción con LEFT JOIN + COUNT(idpreguntas) (robusta y eficiente)
        if include_all:
            cur.execute("""
                SELECT t.id, t.nombre, t.activo,
                       COALESCE(COUNT(p.idpreguntas), 0) AS n_preguntas
                FROM temas t
                LEFT JOIN preguntas p ON p.tema_id = t.id
                GROUP BY t.id, t.nombre, t.activo
                ORDER BY t.nombre
            """)
        else:
            cur.execute("""
                SELECT t.id, t.nombre, t.activo,
                       COALESCE(COUNT(p.idpreguntas), 0) AS n_preguntas
                FROM temas t
                LEFT JOIN preguntas p ON p.tema_id = t.id
                WHERE t.activo = 1
                GROUP BY t.id, t.nombre, t.activo
                ORDER BY t.nombre
            """)

        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/api/temas", methods=["POST"])
def temas_crear():
    data = request.get_json(silent=True) or {}
    nombre = (data.get("nombre") or "").strip()
    if not nombre:
        return jsonify({"error": "El nombre es requerido"}), 400
    if len(nombre) > 100:
        return jsonify({"error": "Máximo 100 caracteres"}), 400
    try:
        conn = get_connection(); cur = conn.cursor()
        # Evitar duplicados (case-insensitive)
        cur.execute("SELECT id FROM temas WHERE LOWER(nombre)=LOWER(%s)", (nombre,))
        if cur.fetchone():
            cur.close(); conn.close()
            return jsonify({"error": "Ya existe un curso con ese nombre"}), 409
        cur.execute("INSERT INTO temas (nombre, activo) VALUES (%s, 1)", (nombre,))
        conn.commit()
        nuevo_id = cur.lastrowid
        cur.close(); conn.close()
        return jsonify({"exito": True, "id": nuevo_id, "nombre": nombre, "activo": 1}), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/temas/<int:tema_id>", methods=["PUT"])
def temas_editar(tema_id):
    data = request.get_json(silent=True) or {}
    nombre = (data.get("nombre") or "").strip()
    if not nombre:
        return jsonify({"error": "El nombre es requerido"}), 400
    try:
        conn = get_connection(); cur = conn.cursor()
        # Chequeo de duplicado
        cur.execute("SELECT id FROM temas WHERE LOWER(nombre)=LOWER(%s) AND id<>%s", (nombre, tema_id))
        if cur.fetchone():
            cur.close(); conn.close()
            return jsonify({"error": "Ya existe otro curso con ese nombre"}), 409
        cur.execute("UPDATE temas SET nombre=%s WHERE id=%s", (nombre, tema_id))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/temas/<int:tema_id>/toggle", methods=["PATCH"])
def temas_toggle(tema_id):
    """Habilitar/Deshabilitar (soft-delete)."""
    try:
        conn = get_connection(); cur = conn.cursor()
        cur.execute("UPDATE temas SET activo=1-activo WHERE id=%s", (tema_id,))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/temas/<int:tema_id>", methods=["DELETE"])
def temas_eliminar(tema_id):
    """
    Eliminar definitivo.
    - por defecto: si hay preguntas, bloquea.
    - si ?force=1: borra también sus preguntas (¡peligroso!).
    """
    force = request.args.get("force") == "1"
    try:
        conn = get_connection(); cur = conn.cursor()
        # ¿tiene preguntas?
        cur.execute("SELECT COUNT(*) FROM preguntas WHERE tema_id=%s", (tema_id,))
        n = cur.fetchone()[0]

        if n > 0 and not force:
            cur.close(); conn.close()
            return jsonify({"error": f"El tema tiene {n} preguntas. Deshabilítalo o elimina con force=1."}), 409

        if force:
            cur.execute("DELETE FROM preguntas WHERE tema_id=%s", (tema_id,))
        cur.execute("DELETE FROM temas WHERE id=%s", (tema_id,))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    

# ---------------------------
# CRUD preguntas
# ---------------------------
import os, re, zipfile, shutil, tempfile, unicodedata, copy
import xml.etree.ElementTree as ET

# -----------------------------------------------------------------------------------
# Helpers DOCX / XML
# -----------------------------------------------------------------------------------
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{%s}" % NS_W
ns_doc = {"w": NS_W}

def paragraph_text(p):
    """
    Reconstruye el texto real de un <w:p> sin meter espacios falsos entre runs,
    respetando xml:space='preserve' y considerando <w:tab>/<w:br> como espacio.
    """
    parts = []
    for node in p.iter():
        if node.tag == W + "t":
            txt = node.text or ""
            if node.get(W + "space") == "preserve":
                parts.append(txt)
            else:
                parts.append(txt.strip())
        elif node.tag in (W + "tab", W + "br"):
            parts.append(" ")
    return "".join(parts)

def _norm(s: str) -> str:
    if not s:
        return ""
    s = s.lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")  # quita acentos
    s = re.sub(r"\s+", " ", s)
    return s



def _norm_tema(s: str) -> str:
    """
    Normaliza nombres de temas para comparación:
    - reutiliza _norm (minúsculas, sin tildes, espacios colapsados)
    - luego pasa a MAYÚSCULAS para clave canónica
    """
    return _norm(s).upper()

def _slug(s: str) -> str:
    t = _norm(s)
    t = t.replace(" ", "_")
    return re.sub(r"[^a-z0-9_]+", "", t) or "tema"

def is_centered_bold_heading(p):
    """Heurística: párrafo centrado/negrita y corto ⇒ probable encabezado."""
    pPr = p.find("w:pPr", ns_doc)
    centered = False
    if pPr is not None:
        jc = pPr.find("w:jc", ns_doc)
        centered = (jc is not None and jc.get(W + "val") == "center")
    bold = any((r.find("w:rPr", ns_doc) is not None and r.find("w:rPr/w:b", ns_doc) is not None)
               for r in p.findall("w:r", ns_doc))
    txt = _norm(paragraph_text(p))
    return (centered or bold) and len(txt) <= 30 and re.fullmatch(r"[a-z\s]+", txt or "") is not None

def _reempacar_docx(work_dir: str, elementos_xml, destino_docx: str):
    tmp = tempfile.mkdtemp(prefix="docx_")
    shutil.copytree(work_dir, tmp, dirs_exist_ok=True)

    doc_path = os.path.join(tmp, 'word', 'document.xml')
    tree = ET.parse(doc_path)
    root = tree.getroot()
    body = root.find(W + 'body')

    # 1) sectPr original (igual que ya tenías)
    sectPr = body.find(W + 'sectPr')
    if sectPr is None and len(list(body)) > 0:
        last = list(body)[-1]
        if last.tag == W + 'p':
            pPr = last.find(W + 'pPr')
            if pPr is not None:
                sectPr = pPr.find(W + 'sectPr')
    if sectPr is None:
        sectPr = ET.Element(W + 'sectPr')

    # 2) limpiar body
    for ch in list(body):
        body.remove(ch)

    # 3) añadir SOLO contenido saneado
    for el in elementos_xml:
        frag = copy.deepcopy(el)
        _sanear_fragmento(frag)
        body.append(frag)

    # 4) remate válido: p vacío + sectPr final
    body.append(ET.Element(W + 'p'))
    body.append(copy.deepcopy(sectPr))

    tree.write(doc_path, xml_declaration=True, encoding='UTF-8', method='xml')

    with zipfile.ZipFile(destino_docx, 'w') as z:
        for base, _, files in os.walk(tmp):
            for f in files:
                p = os.path.join(base, f)
                z.write(p, os.path.relpath(p, tmp))
    shutil.rmtree(tmp, ignore_errors=True)


def reparar_docx_inplace(path_in: str) -> None:
    pythoncom.CoInitialize()
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(path_in, ReadOnly=False, OpenAndRepair=True,
                                  ConfirmConversions=False, AddToRecentFiles=False, Visible=False)
        # Guardar sobre el mismo archivo
        wdFormatXMLDocument = 12
        doc.SaveAs(path_in, FileFormat=wdFormatXMLDocument)
        doc.Close(False)
    finally:
        try: word.Quit()
        except: pass
        pythoncom.CoUninitialize()

# -----------------------------------------------------------------------------------
# CRUD PREGUNTAS / PARTIR EXAMEN POR TEMAS (ROBUSTO)
# -----------------------------------------------------------------------------------
@app.route("/api/examenes/<int:idexamen>/partir_y_guardar", methods=["POST"])
def partir_y_guardar(idexamen):
    overwrite = (request.args.get("overwrite") == "1")

    # ====== utilidades/constantes del "código firme" ======
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns_doc = {"w": NS_W}

    LETTER_FMTS = {"upperLetter", "lowerLetter"}
    ROMAN_FMTS  = {"upperRoman", "lowerRoman"}
    BULLET_FMTS = {"bullet"}
    QUESTION_FMTS = {"decimal", "decimalZero"}
    NON_QUESTION_TOPLEVEL_FMTS = LETTER_FMTS | ROMAN_FMTS | BULLET_FMTS

    def normalize(s: str) -> str:
        s = (s or "").lower().strip()
        s = unicodedata.normalize("NFD", s)
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        s = re.sub(r"\s+", " ", s)
        return s

    def slugify(s: str) -> str:
        s = normalize(s)
        s = re.sub(r"[^a-z0-9]+", "_", s)
        s = re.sub(r"_+", "_", s).strip("_")
        return s or "tema"

    def paragraph_text(p) -> str:
        parts = []
        for node in p.iter():
            if node.tag == W + "t":
                txt = node.text or ""
                if node.get(W + "space") == "preserve":
                    parts.append(txt)
                else:
                    parts.append(txt.strip())
            elif node.tag in (W + "tab", W + "br"):
                parts.append(" ")
        return "".join(parts)

    def is_centered_bold_heading(p) -> bool:
        pPr = p.find("w:pPr", ns_doc)
        centered = False
        if pPr is not None:
            jc = pPr.find("w:jc", ns_doc)
            centered = (jc is not None and jc.get(W + "val") == "center")
        bold = any((r.find("w:rPr", ns_doc) is not None and r.find("w:rPr/w:b", ns_doc) is not None)
                   for r in p.findall("w:r", ns_doc))
        txt = normalize(paragraph_text(p))
        return (centered or bold) and len(txt) <= 40 and re.fullmatch(r"[a-z\s]+", txt or "") is not None

    def is_question_start(plain_text: str, numPr, fmt: str, ilvl: str) -> bool:
        if re.match(r"^\s*\(?[1-9]\d{0,2}\)?[.)]?(?:\s+|(?=[^\s]))", plain_text):
            return True
        if numPr is None or ilvl != "0" or fmt in NON_QUESTION_TOPLEVEL_FMTS:
            return False
        return (fmt in QUESTION_FMTS) or (fmt == "")

    def match_topic_name(line_norm: str, CANON_TOPICS: dict) -> str | None:
        cleaned = re.sub(r"^[\s\-\–\—\:;,\.\(\)\[\]\{\}]+|[\s\-\–\—\:;,\.\(\)\[\]\{\}]+$", "", line_norm).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        return CANON_TOPICS.get(cleaned, None)

    def heading_block_topic(elems, start, CANON_TOPICS):
        texts = []
        j = start
        while j < len(elems) and elems[j].tag == W + "p":
            p = elems[j]
            if p.find(".//w:numPr", ns_doc) is not None:
                break
            tnorm = normalize(paragraph_text(p))
            if not tnorm:
                break
            if not is_centered_bold_heading(p):
                break
            texts.append(tnorm)
            merged = normalize(" ".join(texts))
            exact = match_topic_name(merged, CANON_TOPICS)
            if exact:
                return exact, (j - start + 1)
            j += 1
            if len(texts) >= 3:
                break
        return None, 0
    def looks_like_title(p, line_norm: str) -> bool:
            """
            Título si: (centrado/negrita) o (sin numeración y corto).
            Aumentamos el largo permitido a 80 para tolerar colas como '... del Perú y del mundo'.
            """
            if is_centered_bold_heading(p):
                return True
            if p.find(".//w:numPr", ns_doc) is None:
                if len(line_norm) <= 80 and re.fullmatch(r"[a-z0-9\s]+", line_norm or ""):
                    return True
            return False
    def soft_topic_match(line_norm: str, CANON_TOPICS: dict) -> str | None:
            """
            Soft-match simétrico por tokens (más tolerante con temas cortos como 'geografia').
            Acepta si:
            - cobertura_tema ≥ 0.5 y cobertura_linea ≥ 0.35, o
            - Jaccard ≥ 0.5
            Además: fallback por prefijo (la línea empieza con el nombre del tema).
            """
            def toks(s): 
                return [w for w in re.split(r"[^a-z0-9]+", normalize(s)) if w]

            lt = set(toks(line_norm))
            if not lt:
                return None

            best, best_score = None, 0.0
            for canon_norm, original in CANON_TOPICS.items():
                tt = set(toks(canon_norm))
                if not tt:
                    continue
                inter = len(tt & lt)
                if inter == 0:
                    continue

                topic_cov = inter / len(tt)
                line_cov  = inter / len(lt)
                jaccard   = inter / len(tt | lt)
                score = (topic_cov + line_cov + jaccard) / 3.0

                ok = ((topic_cov >= 0.5 and line_cov >= 0.35) or (jaccard >= 0.5))
                if ok and score > best_score:
                    best, best_score = original, score

            if best:
                return best

            # Prefijo: 'geografia del peru...' debe mapear a 'GEOGRAFÍA'
            for canon_norm, original in CANON_TOPICS.items():
                if line_norm.startswith(canon_norm + " "):
                    return original

            return None

    def detectar_tema(elem, texto_norm, plain_text, numPr_elem, fmt, ilvl, CANON_TOPICS, TOPIC_TRIGGERS):
        """
        Orden:
          1) Coincidencia EXACTA de línea con el nombre de tema.
          2) Trigger al inicio (tolerante): start<=2, before_words<=2, after_words<=8.
             Si además es inicio de pregunta, acepta sin mirar longitudes.
          3) Si 'parece título' (looks_like_title), aplicar soft_topic_match.
        """
        # 1) Exacto
        mt = match_topic_name(texto_norm, CANON_TOPICS)
        if mt:
            return mt

        # Limpia prefijos (I., 1), A)) y colas de puntuación
        t = re.sub(r"^(?:[ivxlcdm]+|[0-9]+|[a-z])[\.\)]\s*", "", texto_norm)
        t = re.sub(r"[\s\-\–\—\:;,\.\(\)\[\]\{\}]+$", "", t)

        # 2) Trigger al inicio (más laxo)
        for tema, triggers in TOPIC_TRIGGERS.items():
            for trig in triggers:
                m = re.search(rf"(^|\W){re.escape(trig)}(\W|$)", t)
                if not m:
                    continue
                start = m.start(0); end = m.end(0)
                before_words = len(t[:start].strip().split())
                after_words  = len(t[end:].strip().split())

                # Si el mismo párrafo es inicio de pregunta, aceptamos directo
                if start <= 2 and is_question_start(plain_text, numPr_elem, fmt, ilvl):
                    return tema

                # Título puro pero con cola algo más larga (p.ej. '... del perú y del mundo')
                if start <= 2 and before_words <= 2 and after_words <= 8:
                    return tema

        # 3) Soft por tokens si 'parece título' (no solo centrado/negrita)
        if looks_like_title(elem, texto_norm):
            st = soft_topic_match(texto_norm, CANON_TOPICS)
            if st:
                return st

        return None

    # ====== 1) Traer examen y TEMAS desde BD ======
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("SELECT archivo_ruta, archivo_nombre FROM examenes WHERE idexamenes=%s", (idexamen,))
        exam = cur.fetchone()
        if not exam:
            cur.close(); conn.close()
            return jsonify({"error": "Examen no existe"}), 404
        ruta_docx = os.path.abspath(exam["archivo_ruta"])
        if not os.path.isfile(ruta_docx):
            cur.close(); conn.close()
            return jsonify({"error": f"Archivo DOCX no encontrado: {ruta_docx}"}), 500

        cur.execute("SELECT id, nombre FROM temas")
        temas_rows = cur.fetchall()
        cur.close(); conn.close()
    except Exception as e:
        return jsonify({"error": f"DB error: {e}"}), 500

    # ⚠️ TEMAS solo desde la BD
    TEMAS = [r["nombre"] for r in temas_rows]
    CANON_TOPICS = { normalize(t): t for t in TEMAS }  # norm->original
    temas_bd_map = { normalize(r["nombre"]): (r["id"], r["nombre"]) for r in temas_rows }

    # Base de sinónimos del código firme (solo se aplican si el tema existe en BD)
    TOPIC_TRIGGERS_BASE = {
        "RAZONAMIENTO MATEMÁTICO": ["razonamiento matematico","raz matematico","razonamiento de matematica"],
        "RAZONAMIENTO VERBAL": ["razonamiento verbal","raz verbal"],
        "COMUNICACIÓN": ["comunicacion","capacidad comunicativa"],
        "ÁLGEBRA": ["algebra"],
        "GEOMETRÍA": ["geometria"],
        "TRIGONOMETRÍA": ["trigonometria"],
        "FÍSICA": ["fisica"],
        "QUÍMICA": ["quimica"],
        "ECOLOGÍA Y MEDIO AMBIENTE": ["ecologia y medio ambiente"],
        "BIOLOGÍA": ["biologia"],
        "ZOOLOGÍA": ["zoologia"],
        "ECONOMÍA": ["economia"],
        "HISTORIA DEL PERÚ EN EL CONTEXTO MUNDIAL": ["historia del peru en el contexto mundial","historia del peru"],
        "EDUCACIÓN CÍVICA": ["educacion civica","civica"],
        "GEOGRAFÍA DEL PERÚ Y DEL MUNDO": ["geografia del peru y del mundo","geografia del peru","geografia"],
        "ARITMÉTICA": ["aritmetica"],
        "MATEMÁTICA II": ["matematica ii","matematica 2"],
        "COMPETENCIA LINGÜÍSTICA": ["competencia linguistica","competencia linguistica comunicativa"],
    }
    # Triggers efectivos: solo para temas que existan en BD; si no hay base, usa su propio nombre
    TOPIC_TRIGGERS = {
        name: TOPIC_TRIGGERS_BASE.get(name, [normalize(name)])
        for name in TEMAS
    }

    base_examen_dir = os.path.join(app.config['PREGUNTAS_DIR'], f"examen_{idexamen}")
    os.makedirs(base_examen_dir, exist_ok=True)

    # ====== 2) Limpiar si overwrite ======
    if overwrite:
        try:
            conn = get_connection(); cur = conn.cursor()
            cur.execute("DELETE FROM preguntas WHERE examenes_idexamenes=%s", (idexamen,))
            conn.commit()
            cur.close(); conn.close()
        except Exception as e:
            return jsonify({"error": f"No se pudo limpiar preguntas previas: {e}"}), 500
        if os.path.isdir(base_examen_dir):
            shutil.rmtree(base_examen_dir, ignore_errors=True)
        os.makedirs(base_examen_dir, exist_ok=True)

    # ====== 3) Unzip DOCX y numbering ======
    tmp_root = tempfile.mkdtemp(prefix="split_")
    work = os.path.join(tmp_root, "work"); os.mkdir(work)
    with zipfile.ZipFile(ruta_docx, "r") as z:
        z.extractall(work)

    num_fmt_map = {}
    numbering_xml_path = os.path.join(work, "word", "numbering.xml")
    if os.path.exists(numbering_xml_path):
        tree_num = ET.parse(numbering_xml_path)
        root_num = tree_num.getroot()
        for num in root_num.findall(".//w:num", ns_doc):
            numId = num.attrib.get(W + "numId")
            abs_el = num.find("./w:abstractNumId", ns_doc)
            if abs_el is None: continue
            abs_id = abs_el.attrib.get(W + "val")
            abstract = root_num.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns_doc)
            if abstract is None: continue
            for lvl in abstract.findall("./w:lvl", ns_doc):
                ilvl = lvl.attrib.get(W + "ilvl")
                nf = lvl.find("./w:numFmt", ns_doc)
                if nf is not None:
                    num_fmt_map[(numId, ilvl)] = nf.attrib.get(W + "val")

    # ====== 4) Partir como en el código firme ======
    document_xml_path = os.path.join(work, "word", "document.xml")
    tree_doc = ET.parse(document_xml_path)
    root_doc = tree_doc.getroot()
    body = root_doc.find(W + "body")
    elems = list(body)

    preguntas_por_tema = {t: [] for t in TEMAS}  # solo temas que existen en BD
    actual_pregunta, copiando = [], False
    tema_actual, tema_anterior = "", ""
    active_q_numId = None

    i = 0
    while i < len(elems):
        elem = elems[i]
        if elem.tag != W + "p":
            i += 1
            continue

        plain_text = paragraph_text(elem)
        texto_norm = normalize(plain_text)

        # 0) título repartido
        topic_from_block, span = heading_block_topic(elems, i, CANON_TOPICS)
        if topic_from_block:
            if copiando and tema_anterior in preguntas_por_tema and actual_pregunta:
                preguntas_por_tema[tema_anterior].append(actual_pregunta)
            tema_actual = topic_from_block
            tema_anterior = tema_actual
            copiando = False
            actual_pregunta = []
            active_q_numId = None
            i += span
            continue

        # 1) num info
        numPr = elem.find(".//w:numPr", ns_doc)
        ilvl_val = numPr.find("./w:ilvl", ns_doc).get(W+"val") if (numPr is not None and numPr.find("./w:ilvl", ns_doc) is not None) else ""
        numId_val = numPr.find("./w:numId", ns_doc).get(W+"val") if (numPr is not None and numPr.find("./w:numId", ns_doc) is not None) else ""
        fmt = num_fmt_map.get((numId_val, ilvl_val), "")

        # 2) detección de tema
        posible_tema = detectar_tema(elem, texto_norm, plain_text, numPr, fmt, ilvl_val, CANON_TOPICS, TOPIC_TRIGGERS)
        if posible_tema:
            if copiando and tema_anterior in preguntas_por_tema and actual_pregunta:
                preguntas_por_tema[tema_anterior].append(actual_pregunta)
            # ⚠️ Solo aceptamos si existe en BD
            if posible_tema in preguntas_por_tema:
                tema_actual = posible_tema
                tema_anterior = tema_actual
                copiando = False
                actual_pregunta = []
                active_q_numId = None
                if not is_question_start(plain_text, numPr, fmt, ilvl_val):
                    i += 1
                    continue
            else:
                # tema no registrado en BD → ignorar este encabezado
                i += 1
                continue

        # 3) inicio de pregunta
        qstart = is_question_start(plain_text, numPr, fmt, ilvl_val)
        if qstart and numPr is not None and ilvl_val == "0" and fmt not in NON_QUESTION_TOPLEVEL_FMTS:
            if active_q_numId is None:
                active_q_numId = numId_val
        if qstart and active_q_numId and numPr is not None and numId_val != active_q_numId:
            if fmt in NON_QUESTION_TOPLEVEL_FMTS:
                qstart = False

        if tema_actual in preguntas_por_tema:
            if not copiando and qstart:
                copiando = True
                actual_pregunta = [elem]
            elif copiando and qstart:
                preguntas_por_tema[tema_actual].append(actual_pregunta)
                actual_pregunta = [elem]
            elif copiando:
                actual_pregunta.append(elem)

        i += 1

    if copiando and tema_actual in preguntas_por_tema and actual_pregunta:
        preguntas_por_tema[tema_actual].append(actual_pregunta)

    # ====== 5) Guardar .docx y registrar (sin crear temas) ======
    insertados, resumen = 0, {}
    try:
        conn = get_connection(); cur = conn.cursor()
        for tema_nombre, preguntas in preguntas_por_tema.items():
            if not preguntas:
                continue
            tema_key = normalize(tema_nombre)
            if tema_key not in temas_bd_map:
                # seguridad extra (no debería ocurrir porque TEMAS viene de BD)
                continue

            tema_id, tema_nombre_bd = temas_bd_map[tema_key]
            out_dir = os.path.join(base_examen_dir, slugify(tema_nombre_bd))
            os.makedirs(out_dir, exist_ok=True)

            for idx, contenido in enumerate(preguntas, start=1):
                nombre_archivo = f"{slugify(tema_nombre_bd)}_pregunta_{idx}.docx"
                ruta_archivo = os.path.abspath(os.path.join(out_dir, nombre_archivo))
                _reempacar_docx(work, contenido, ruta_archivo)
                try:
                    reparar_docx_inplace(ruta_archivo)   # ← limpia cualquier rastro dañino
                except Exception:
                    pass  # continúa aunque no se pueda reparar


                cur.execute("""
                    SELECT 1 FROM preguntas
                    WHERE examenes_idexamenes=%s AND tema_id=%s AND numero_p=%s
                """, (idexamen, tema_id, idx))
                if not cur.fetchone():
                    cur.execute("""
                        INSERT INTO preguntas (examenes_idexamenes, tema_id, numero_p, archivo_nombre, archivo_ruta)
                        VALUES (%s, %s, %s, %s, %s)
                    """, (idexamen, tema_id, idx, nombre_archivo, ruta_archivo))
                    insertados += 1

            resumen[tema_nombre_bd] = len(preguntas)

        conn.commit()
        cur.close(); conn.close()
    except Exception as e:
        try:
            cur.close(); conn.close()
        except Exception:
            pass
        shutil.rmtree(tmp_root, ignore_errors=True)
        return jsonify({"error": f"Fallo guardando en BD: {e}"}), 500

    shutil.rmtree(tmp_root, ignore_errors=True)
    return jsonify({
        "ok": True,
        "examen": idexamen,
        "carpeta_base": base_examen_dir,
        "preguntas_insertadas": insertados,
        "por_tema": resumen
    })


# ========== TEMAS DEL EXAMEN (con conteo de preguntas ya partidas) ==========
# === TEMAS DE UN EXAMEN (para el modal "Buscar") ===
# === TEMAS DE UN EXAMEN (para el modal "Buscar") ===
# === TEMAS DE UN EXAMEN (para el modal "Buscar") ===
def _temas_de_examen_impl(idexamen: int):
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT  t.id, t.nombre, t.activo,
                    COALESCE(COUNT(p.idpreguntas), 0) AS n_preguntas
            FROM temas t
            LEFT JOIN preguntas p
                   ON p.tema_id = t.id
                  AND p.examenes_idexamenes = %s
            GROUP BY t.id, t.nombre, t.activo
            ORDER BY t.nombre;
        """, (idexamen,))
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Un único handler que normaliza el id (venga como string o int)
def temas_de_examen(idexamen):
    try:
        idexamen = int(idexamen)
    except (TypeError, ValueError):
        return jsonify({"error": "id de examen inválido"}), 400
    return _temas_de_examen_impl(idexamen)

# Registra TODAS las variantes para evitar 404 por detalles de ruta
app.url_map.strict_slashes = False  # ya lo tienes, lo dejo por claridad
try:
    app.add_url_rule(
        "/api/examenes/<int:idexamen>/temas",
        endpoint="temas_de_examen_int",
        view_func=temas_de_examen,
        methods=["GET"],
    )
    app.add_url_rule(
        "/api/examenes/<idexamen>/temas",
        endpoint="temas_de_examen_str",
        view_func=temas_de_examen,
        methods=["GET"],
    )
    app.add_url_rule(
        "/api/examenes/<int:idexamen>/temas/",
        endpoint="temas_de_examen_int_slash",
        view_func=temas_de_examen,
        methods=["GET"],
    )
    app.add_url_rule(
        "/api/examenes/<idexamen>/temas/",
        endpoint="temas_de_examen_str_slash",
        view_func=temas_de_examen,
        methods=["GET"],
    )
except AssertionError:
    # Si ya existían, ignora el error de duplicado.
    pass
# ⭐ alias: si alguien hace GET /api/examenes/<id> redirige a /temas
@app.route("/api/examenes/<int:idexamen>", methods=["GET"])
def examenes_alias_temas(idexamen):
    return redirect(url_for("temas_de_examen", idexamen=idexamen), code=308)


# === LISTAR PREGUNTAS (filtrable por examen y/o tema) ===
@app.route("/api/preguntas", methods=["GET"])
def preguntas_listar():
    examen_id = request.args.get("examen", type=int)
    tema_id   = request.args.get("tema", type=int)

    where = []
    params = []
    if examen_id is not None:
        where.append("examenes_idexamenes=%s"); params.append(examen_id)
    if tema_id is not None:
        where.append("tema_id=%s"); params.append(tema_id)

    sql = """
        SELECT idpreguntas, examenes_idexamenes, tema_id,
               numero_p, archivo_nombre, archivo_ruta
        FROM preguntas
    """
    if where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY numero_p"

    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute(sql, tuple(params))
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ---------------------------
# CRUD GRUPOS (A, B, C, ES, etc.)
# ---------------------------

def _sanitize_clave(clave: str) -> str:
    return (clave or "").strip().upper()[:5]

@app.route("/api/grupos", methods=["GET"])
def grupos_listar():
    include_all = request.args.get("all") == "1"
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        sql = """
            SELECT g.idgrupo, g.clave, g.nombre, g.activo, g.fecha_creacion,
                   COALESCE(SUM(gt.cantidad), 0) AS total_preguntas
            FROM grupos g
            LEFT JOIN grupo_tema gt ON gt.grupos_idgrupo = g.idgrupo
        """
        if not include_all:
            sql += " WHERE g.activo=1"
        sql += """
            GROUP BY g.idgrupo, g.clave, g.nombre, g.activo, g.fecha_creacion
            ORDER BY g.clave
        """
        cur.execute(sql)
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _sanitize_clave(clave: str) -> str:
    return (clave or "").strip().upper()[:5]

@app.route("/api/grupos", methods=["POST"])
def grupos_crear():
    data = request.get_json(force=True) or {}
    clave  = _sanitize_clave(data.get("clave", ""))
    nombre = (data.get("nombre") or "").strip()
    cuotas = data.get("cuotas") or []  # [{tema_id, cantidad}, ...]

    if not clave:
        return jsonify({"error": "La 'clave' es requerida"}), 400
    if len(nombre) > 100:
        return jsonify({"error": "El 'nombre' admite máximo 100 caracteres"}), 400

    try:
        conn = get_connection(); cur = conn.cursor()
        cur.execute("SELECT 1 FROM grupos WHERE UPPER(clave)=UPPER(%s)", (clave,))
        if cur.fetchone():
            cur.close(); conn.close()
            return jsonify({"error": f"Ya existe un grupo con clave '{clave}'"}), 409

        cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES (%s,%s,1)", (clave, nombre))
        idgrupo = cur.lastrowid

        for c in cuotas:
            tema_id = int(c["tema_id"]); cant = int(c["cantidad"])
            if cant <= 0:  # opcional: ignora/valida
                continue
            cur.execute("""
                INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad)
                VALUES (%s,%s,%s)
                ON DUPLICATE KEY UPDATE cantidad=VALUES(cantidad)
            """, (idgrupo, tema_id, cant))

        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True, "idgrupo": idgrupo})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/grupos/<int:idgrupo>", methods=["PUT"])
def grupos_editar(idgrupo):
    """
    Body JSON (campos opcionales):
    { "clave": "B", "nombre": "Humanidades", "activo": 1 }
    """
    data = request.get_json(silent=True) or {}
    clave  = data.get("clave", None)
    nombre = data.get("nombre", None)
    activo = data.get("activo", None)

    updates = []
    params  = []

    if clave is not None:
        clave = _sanitize_clave(clave)
        if not clave:
            return jsonify({"error": "La 'clave' no puede ser vacía"}), 400
        updates.append("clave=%s"); params.append(clave)

    if nombre is not None:
        nombre = (nombre or "").strip()
        if len(nombre) > 100:
            return jsonify({"error": "El 'nombre' admite máximo 100 caracteres"}), 400
        updates.append("nombre=%s"); params.append(nombre)

    if activo is not None:
        try:
            activo = 1 if int(activo) else 0
        except Exception:
            return jsonify({"error": "El campo 'activo' debe ser 0 o 1"}), 400
        updates.append("activo=%s"); params.append(activo)

    if not updates:
        return jsonify({"error": "No hay campos para actualizar"}), 400

    try:
        conn = get_connection(); cur = conn.cursor()

        # Chequeo de duplicado de clave si se cambia
        if "clave=%s" in updates:
            cur.execute("SELECT idgrupo FROM grupos WHERE UPPER(clave)=UPPER(%s) AND idgrupo<>%s", (clave, idgrupo))
            if cur.fetchone():
                cur.close(); conn.close()
                return jsonify({"error": f"Ya existe otro grupo con clave '{clave}'"}), 409

        sql = f"UPDATE grupos SET {', '.join(updates)} WHERE idgrupo=%s"
        params.append(idgrupo)
        cur.execute(sql, tuple(params))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/grupos/<int:idgrupo>/toggle", methods=["PATCH"])
def grupos_toggle(idgrupo):
    """Activa/Desactiva un grupo (soft)."""
    try:
        conn = get_connection(); cur = conn.cursor()
        cur.execute("UPDATE grupos SET activo = 1 - activo WHERE idgrupo=%s", (idgrupo,))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/grupos/<int:idgrupo>", methods=["DELETE"])
def grupos_eliminar(idgrupo):
    """
    Borrado definitivo.
    - Bloquea si hay cuotas en grupo_tema o exámenes generados.
    - ?force=1 elimina primero sus cuotas (NO borra exámenes generados).
    """
    force = request.args.get("force") == "1"
    try:
        conn = get_connection(); cur = conn.cursor()

        cur.execute("SELECT COUNT(*) FROM grupo_tema WHERE grupos_idgrupo=%s", (idgrupo,))
        n_gt = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM examen_generado WHERE grupos_idgrupo=%s", (idgrupo,))
        n_gen = cur.fetchone()[0]

        if (n_gt > 0 or n_gen > 0) and not force:
            cur.close(); conn.close()
            return jsonify({"error": f"No se puede borrar: cuotas={n_gt}, generados={n_gen}. Desactívalo o usa force=1."}), 409

        if force and n_gt > 0:
            cur.execute("DELETE FROM grupo_tema WHERE grupos_idgrupo=%s", (idgrupo,))

        # Por seguridad NO borramos examen_generado. Si existe FK, el delete fallará.
        cur.execute("DELETE FROM grupos WHERE idgrupo=%s", (idgrupo,))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"exito": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# =======================
# EXAMEN ↔ GRUPOS (lecturas para UI)
# =======================
# Ver cuotas asignadas al grupo (solo las asignadas)
@app.route("/api/grupos/<clave>/cuotas", methods=["GET"])
def grupos_cuotas_get(clave):
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT t.id AS tema_id, t.nombre AS tema, gt.cantidad
            FROM grupo_tema gt
            JOIN grupos g ON g.idgrupo = gt.grupos_idgrupo
            JOIN temas  t ON t.id     = gt.tema_id
            WHERE g.clave = %s
            ORDER BY t.nombre
        """, (clave.upper(),))
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Reemplazar cuotas del grupo (PUT con la lista completa)
# GET: cuotas existentes (solo las que tienen cantidad)
# GET: cuotas existentes (solo las que tienen cantidad)
@app.route("/api/grupos/<int:idgrupo>/cuotas", methods=["GET"])
def grupo_cuotas_get(idgrupo):
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT gt.tema_id, t.nombre AS tema, gt.cantidad
            FROM grupo_tema gt
            JOIN temas t ON t.id = gt.tema_id
            WHERE gt.grupos_idgrupo = %s
            ORDER BY t.nombre
        """, (idgrupo,))
        rows = cur.fetchall()
        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# PUT: reemplaza todas las cuotas del grupo
@app.route("/api/grupos/<int:idgrupo>/cuotas", methods=["PUT"])
def grupo_cuotas_put(idgrupo):
    data = request.get_json(force=True) or {}
    cuotas = data.get("cuotas", [])
    try:
        conn = get_connection(); cur = conn.cursor()

        tema_ids = [int(c["tema_id"]) for c in cuotas if "tema_id" in c]
        # Borra las que ya no están
        if tema_ids:
            cur.execute(
                "DELETE FROM grupo_tema WHERE grupos_idgrupo=%s AND tema_id NOT IN (%s)" %
                ("%s", ",".join(["%s"]*len(tema_ids))),
                (idgrupo, *tema_ids)
            )
        else:
            cur.execute("DELETE FROM grupo_tema WHERE grupos_idgrupo=%s", (idgrupo,))

        # Upsert de cada cuota (asegúrate de tener UNIQUE (grupos_idgrupo, tema_id))
        for c in cuotas:
            cur.execute("""
                INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad)
                VALUES (%s, %s, %s)
                ON DUPLICATE KEY UPDATE cantidad=VALUES(cantidad)
            """, (idgrupo, int(c["tema_id"]), int(c["cantidad"])))

        conn.commit()
        cur.close(); conn.close()
        return jsonify({"status": "ok"})
    except Exception as e:
        try: conn.rollback()
        except: pass
        return jsonify({"error": str(e)}), 500




# =======================
# GENERAR DOC (WORD/PDF) SEGÚN GRUPO
# =======================
# --- Implementación común ---
def _grupos_generar_doc_impl(idgrupo: int, formato: str):
    formato = (formato or "word").lower()
    if formato not in ("word", "pdf"):
        return jsonify({"error": "formato inválido (word|pdf)"}), 400

    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)

        # Datos del grupo
        cur.execute("SELECT idgrupo, clave, nombre FROM grupos WHERE idgrupo=%s", (idgrupo,))
        g = cur.fetchone()
        if not g:
            cur.close(); conn.close()
            return jsonify({"error":"Grupo no existe"}), 404
        clave = g["clave"] or f"G{idgrupo}"

        # Cuotas
        cur.execute("""
            SELECT gt.tema_id, t.nombre AS tema, gt.cantidad
            FROM grupo_tema gt
            JOIN temas t ON t.id = gt.tema_id
            WHERE gt.grupos_idgrupo = %s
            ORDER BY t.nombre
        """, (idgrupo,))
        cuotas = cur.fetchall()
        if not cuotas:
            cur.close(); conn.close()
            return jsonify({"error":"El grupo no tiene cuotas configuradas."}), 400

        total_requeridas = sum(max(0, int(c["cantidad"])) for c in cuotas)
        if total_requeridas <= 0:
            cur.close(); conn.close()
            return jsonify({"error": "Todas las cuotas están en 0. Asigna cantidades > 0."}), 400

        tema_ids = [c["tema_id"] for c in cuotas if int(c["cantidad"]) > 0]
        if not tema_ids:
            cur.close(); conn.close()
            return jsonify({"error":"No hay cuotas con cantidad > 0."}), 400

        # Disponibilidad
        cur.execute(
            "SELECT tema_id, COUNT(*) AS disponibles "
            "FROM preguntas "
            "WHERE tema_id IN (" + ",".join(["%s"]*len(tema_ids)) + ") "
            "GROUP BY tema_id", tuple(tema_ids)
        )
        disp = cur.fetchall()
        mapa_disp = {int(d["tema_id"]): int(d["disponibles"]) for d in disp}
        faltantes = []
        for c in cuotas:
            req = max(0, int(c["cantidad"]))
            if req and mapa_disp.get(int(c["tema_id"]), 0) < req:
                faltantes.append({
                    "tema_id": int(c["tema_id"]),
                    "tema": c["tema"],
                    "requeridas": req,
                    "disponibles": mapa_disp.get(int(c["tema_id"]), 0)
                })
        if faltantes:
            cur.close(); conn.close()
            return jsonify({"ok": False, "error": "No hay preguntas suficientes para crear este examen.", "faltantes": faltantes}), 409

        # Modo diagnóstico
        if request.args.get("debug") == "1":
            cur.close(); conn.close()
            return jsonify({
                "grupo_id": idgrupo,
                "cuotas": cuotas,
                "disponibilidad_por_tema": mapa_disp,
                "total_requeridas": total_requeridas
            })

        # Selección aleatoria
        # --- Selección aleatoria agrupada por tema (respeta el orden de cuotas) ---
        grouped = []  # [(tema_nombre, [ruta_docx, ...]), ...]
        for c in cuotas:
            tema_nombre = c["tema"]
            cantidad = int(c["cantidad"])
            if cantidad <= 0:
                continue

            cur.execute("""
                SELECT archivo_ruta
                FROM preguntas
                WHERE tema_id = %s
                ORDER BY RAND()
                LIMIT %s
            """, (c["tema_id"], cantidad))
            filas = cur.fetchall()
            if len(filas) < cantidad:
                cur.close(); conn.close()
                return jsonify({"ok": False, "error": "Stock insuficiente inesperado."}), 409

            paths = [os.path.abspath(f["archivo_ruta"]) for f in filas]
            grouped.append((tema_nombre, paths))

            # justo después de construir 'grouped'
        vacios = [(t, len(fs)) for (t, fs) in grouped if not fs]
        if vacios:
            return jsonify({"ok": False, "error": "No se encontraron preguntas para algunos temas",
                            "detalles": [{"tema": t, "encontradas": n} for t, n in vacios]}), 409

                # --- Unir y normalizar (igual que ya tienes) ---
               # --- Unir y normalizar ---
        ts = time.strftime("%Y-%m-%d %H-%M")
        # timestamp legible
        friendly = f"Examen del grupo {clave} - {ts}"
        # NOMBRE AMIGABLE

        base_name = f"{friendly}.docx"


        destino_docx = os.path.join(app.config['DESCARGAS_FOLDER'], base_name)

        # Puedes forzar COM con ?com=1 (Windows). Por defecto usamos el camino Python (estable).
        usar_com = (request.args.get("com") == "1") and _com_disponible()
        if not usar_com and Composer is None and _com_disponible():
            usar_com = True 
        malos = []
        if usar_com:
            # COM (Word) + títulos
            destino_docx, _, malos = _merge_grouped_with_headings_wordcom(grouped, destino_docx)
        else:
            # Python (docxcompose) + títulos
            destino_docx, _, malos = _merge_grouped_with_headings(grouped, destino_docx)
            # Normaliza numeración decimal como en tu Colab
            try:
                _post_merge_fix_numbering(destino_docx)    # fuerza numId=1 ilvl=0 para 'decimal'
            except Exception:
                pass
            # Secciones continuas y limpieza opcional
            try:
                _hacer_secciones_continuas(destino_docx)   # cambia nextPage->continuous y quita <w:br type="page"/>
            except Exception:
                pass
       # ... después de producir destino_docx (con COM o con python) ...
        try:
            # 🔁 Convierte cualquier viñeta a numeración 1., 2., 3., …
            bullets_to_numbers_docx(destino_docx)
        except Exception as _e:
            # no abortes la generación por esto: lo dejamos como advertencia silenciosa
           print("bullet->decimal warning:", _e)



        # Reparación/compactado con Word si está disponible (opcional pero útil tras editar XML)
        try:
            reparar_docx_inplace(destino_docx)
        except Exception:
            pass
        # === Vista previa HTML (Word -> Web Page, Filtered) ===
      #  try:
            # Carpeta única por examen dentro de /static/previews
       #  ===   base_safe = re.sub(r'[^A-Za-z0-9_-]+', '_', os.path.splitext(base_name)[0])
       #     out_dir = os.path.join(PREVIEWS_DIR, base_safe)
        #    html_abs, html_rel = docx_a_html_filtrado(destino_docx, out_dir)  # <- tu función
        #    preview_url = html_rel  # ej: /static/previews/Examen_A_2025-11-02_21-53/Examen_A_2025-11-02_21-53.html
        #except Exception as _e:
        #    print("[preview] fallo generando HTML:", _e)
         #   preview_url = None
     
        # --- Guardar y devolver JSON (sin descargar) ---
        # --- Guardar y devolver JSON (sin descargar) ---
          # --- Guardar y devolver JSON (sin descargar) ---
        advertencias = []  # o usa los “malos” que devuelven los merges

        
        # nombres finales en /descargas
        pdf_name  = f"{friendly}.pdf"
        pdf_final = os.path.join(app.config['DESCARGAS_FOLDER'], pdf_name)

        # aseguramos la carpeta
        os.makedirs(app.config['DESCARGAS_FOLDER'], exist_ok=True)

        ruta_rel_pdf  = None
        ruta_pdf_dl   = None
        ruta_pdf_inline = None

            # 🔧 Normaliza el DOCX con Word para evitar “archivo corrompido”
        try:
            tmp_norm = os.path.join(app.config['DESCARGAS_FOLDER'], "_tmp_norm.docx")
            resave_docx_formatted(destino_docx, tmp_norm)
            shutil.move(tmp_norm, destino_docx)  # reescribe el original normalizado
        except Exception as _e:
            print("[normalize] aviso:", _e)

       # 1) convertir DOCX -> PDF directamente al destino final
        try:
            pdf_generado = docx_a_pdf(destino_docx, pdf_final)
            print("[PDF grupos] OK:", pdf_generado, "size=",
                os.path.getsize(pdf_generado) if os.path.exists(pdf_generado) else 0)

            ruta_rel_pdf    = f"/api/descargas/{pdf_name}"         # descarga
            ruta_pdf_dl     = ruta_rel_pdf
            ruta_pdf_inline = f"/descargas/{pdf_name}"             # preview en <iframe>
        except Exception as e:
            print("[PDF grupos] ERROR al convertir:", repr(e))
            ruta_rel_pdf = ruta_pdf_dl = ruta_pdf_inline = None     # forzar fallback a HTML


        # 2) rutas de descarga del DOCX (esto ya lo tenías)
        ruta_docx_dl = f"/api/descargas/{base_name}"

        # 3) decidir qué mostrar en el iframe
        if ruta_pdf_inline:
            # tenemos PDF bueno → forzamos PDF
            preview_url  = ruta_pdf_inline
            preview_kind = "pdf"
        else:
            # no hubo PDF → generamos HTML como siempre
            html_abs = generar_html_desde_docx(destino_docx, nombre_base=f"{friendly}_preview")
            preview_url  = f"/descargas/{os.path.basename(html_abs)}" if html_abs else None
            preview_kind = "html"

        # 4) armar respuesta
        result = {
            "ok": True,
            "formato": formato,
            "archivo_docx": base_name,
            "ruta_rel": ruta_docx_dl,          # descarga DOCX
            "archivo_pdf": pdf_name if ruta_rel_pdf else None,
            "ruta_rel_pdf": ruta_rel_pdf,      # descarga PDF (si existe)
            "preview_url": preview_url,        # lo que va al iframe
            "preview_kind": preview_kind,      # "pdf" o "html"
            "warnings": [{"path": p, "motivo": m} for (p, m) in advertencias]
        }
        # =========================
        # ⛑️ PLAN B: aplanar numeración si ?flat=1
        # (convierte toda la numeración a texto y elimina numbering.xml)
        # =========================
        if request.args.get("flat") == "1":
            try:
                aplanar_listas_a_texto(destino_docx)
                reparar_docx_inplace(destino_docx)  # opcional
            except Exception:
                pass

        # Si pidieron PDF, convertir AHORA (después del aplanado si aplica)
        if formato == "pdf":
            result.update({
                "archivo_pdf": pdf_name,
                "ruta_rel_pdf": f"/api/descargas/{pdf_name}"
                # NO cambiar preview_url: ya apunta al PDF inline
            })


        return jsonify(result)



    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# Acepta ID numérico y también GET para compatibilidad
@app.route("/api/grupos/<int:idgrupo>/generar_doc", methods=["GET", "POST"])
def grupos_generar_doc_por_id(idgrupo: int):
    return _grupos_generar_doc_impl(idgrupo, request.args.get("formato"))

# NUEVA: acepta clave tipo "A", "B", "ES" y también GET
@app.route("/api/grupos/<clave>/generar_doc", methods=["GET", "POST"])
def grupos_generar_doc_por_clave(clave: str):
    try:
        conn = get_connection(); cur = conn.cursor()
        cur.execute("SELECT idgrupo FROM grupos WHERE UPPER(clave)=UPPER(%s)", (clave,))
        row = cur.fetchone()
        cur.close(); conn.close()
        if not row:
            return jsonify({"error": f"Grupo con clave '{clave}' no existe"}), 404
        idgrupo = int(row[0])
        return _grupos_generar_doc_impl(idgrupo, request.args.get("formato"))
    except Exception as e:
        return jsonify({"error": str(e)}), 500






@app.get("/api/descargas/<path:nombre>")
def descargar_archivo(nombre):
    # seguridad básica: evita subir directorios
    nombre = os.path.normpath(nombre).replace("\\", "/")
    if "/" in nombre:
        # solo se acepta el nombre dentro de descargas
        return jsonify({"error": "nombre inválido"}), 400

    ruta = os.path.join(app.config['DESCARGAS_FOLDER'], nombre)
    if not os.path.isfile(ruta):
        return jsonify({"error": "archivo no existe"}), 404

    # ahora sí, descarga al cliente
    return send_file(
        ruta,
        as_attachment=True,
        download_name=nombre,
        max_age=0,
        mimetype=("application/pdf" if nombre.lower().endswith(".pdf")
                  else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    )
def _tiene_texto_o_contenido(p):
    import zipfile
    patterns = (
        b"<w:t",           # texto
        b"<w:tbl",         # tablas
        b"<w:drawing",     # imágenes/ecuaciones renderizadas
        b"<w:pict",        # VML antiguo
        b"<w:object",      # OLE
        b"<m:oMath",       # ecuaciones OMML
        b"<w:txbxContent", # texto en cuadros de texto
        b"<w:sdt",         # content controls
        b"<w:fldSimple",   # campos simples
        b"<w:instrText",   # instrucciones de campo
    )
    with zipfile.ZipFile(p, 'r') as z:
        # cuerpo + headers/footers por si el contenido quedó allí
        names = ['word/document.xml'] + \
                [n for n in z.namelist()
                 if n.startswith('word/header') or n.startswith('word/footer')]
        data = b''.join(z.read(n) for n in names if n in z.namelist())
    return any(pat in data for pat in patterns)

def _sanear_fragmento(el):
    """
    Elimina elementos que suelen quedar desbalanceados al recortar:
    bookmarks, rangos de comentarios, control de cambios, proofErr, y
    sectPr internos (conservaremos solo el sectPr final en el body).
    """
    BAD_TAGS = {
        W + "bookmarkStart",
        W + "bookmarkEnd",
        W + "commentRangeStart",
        W + "commentRangeEnd",
        W + "commentReference",
        W + "proofErr",
        W + "permStart",
        W + "permEnd",
        W + "moveFrom",
        W + "moveTo",
        W + "ins",
        W + "del",
        W + "smartTag",
        W + "sectPr",  # NO queremos sectPr dentro de párrafos copiados
    }
    # recorrer en profundidad y borrar las malas
    for bad in list(el.iter()):
        if bad.tag in BAD_TAGS:
            parent = bad.getparent() if hasattr(bad, "getparent") else None
            # ElementTree estándar no tiene getparent; hacemos workaround:
            if parent is None:
                # buscar padre manualmente
                for anc in el.iter():
                    for ch in list(anc):
                        if ch is bad:
                            anc.remove(ch)
                            break
            else:
                parent.remove(bad)
    # limpia atributos de control de cambios/rastros
    for node in el.iter():
        for attr in list(node.attrib.keys()):
            if any(k in attr for k in ("rsid",)):  # rsidR, rsidDel, etc.
                node.attrib.pop(attr, None)

def _post_merge_fix_numbering(docx_path: str):
    """
    Replica lo que hiciste en Colab:
    - Descomprime el .docx
    - Lee numbering.xml para mapear (numId, ilvl) -> numFmt
    - Reescribe en document.xml los w:numPr que sean 'decimal' a numId=1, ilvl=0
    - Reempaqueta el .docx
    """
    import os, zipfile, shutil
    import xml.etree.ElementTree as ET

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns = {"w": NS_W}

    tmp_dir = os.path.join(os.path.dirname(docx_path), "_tmp_fix_num")
    if os.path.exists(tmp_dir):
        shutil.rmtree(tmp_dir, ignore_errors=True)
    os.makedirs(tmp_dir, exist_ok=True)

    # 1) unzip
    with zipfile.ZipFile(docx_path, "r") as z:
        z.extractall(tmp_dir)

    # 2) mapear formatos desde numbering.xml
    num_fmt_map = {}
    numbering_xml_path = os.path.join(tmp_dir, "word", "numbering.xml")
    if os.path.exists(numbering_xml_path):
        tnum = ET.parse(numbering_xml_path)
        rnum = tnum.getroot()

        for num in rnum.findall(".//w:num", ns):
            numId = num.attrib.get(W + "numId")
            abs_el = num.find("w:abstractNumId", ns)
            if abs_el is None:
                continue
            abs_id = abs_el.attrib.get(W + "val")
            abstract = rnum.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns)
            if abstract is None:
                continue
            for lvl in abstract.findall("w:lvl", ns):
                ilvl = lvl.attrib.get(W + "ilvl")
                nfmt = lvl.find("w:numFmt", ns)
                if nfmt is not None:
                    num_fmt_map[(numId, ilvl)] = nfmt.attrib.get(W + "val")

    # 3) reescribir document.xml (solo decimal -> numId=1, ilvl=0)
    doc_xml_path = os.path.join(tmp_dir, "word", "document.xml")
    tdoc = ET.parse(doc_xml_path)
    rdoc = tdoc.getroot()

    for p in rdoc.findall(".//w:p", ns):
        numPr = p.find(".//w:numPr", ns)
        if numPr is None:
            continue
        numId = numPr.find("w:numId", ns)
        ilvl = numPr.find("w:ilvl", ns)
        numId_val = numId.attrib.get(W + "val") if numId is not None else "-"
        ilvl_val = ilvl.attrib.get(W + "val") if ilvl is not None else "-"

        fmt = num_fmt_map.get((numId_val, ilvl_val), "decimal")
        if fmt == "decimal":
            if numId is not None:
                numId.attrib[W + "val"] = "1"
            if ilvl is not None:
                ilvl.attrib[W + "val"] = "0"

    ET.register_namespace("w", NS_W)
    tdoc.write(doc_xml_path, encoding="utf-8", xml_declaration=True)

    # 4) rezip sobre el MISMO archivo
    tmp_docx = docx_path + ".tmp"
    with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(tmp_dir):
            for f in files:
                p = os.path.join(root, f)
                z.write(p, os.path.relpath(p, tmp_dir))

    shutil.move(tmp_docx, docx_path)
    shutil.rmtree(tmp_dir, ignore_errors=True)

def _hacer_secciones_continuas(docx_path: str):
    """
    Convierte los section breaks de 'nextPage' a 'continuous' y elimina <w:br w:type="page"/>.
    """
    import zipfile, io, xml.etree.ElementTree as ET
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns = {"w": NS_W}

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    if "word/document.xml" not in files:
        return

    root = ET.fromstring(files["word/document.xml"])

    # 1) sectPr -> type continuous
    for sectPr in root.findall(".//w:sectPr", ns):
        t = sectPr.find("./w:type", ns)
        if t is not None and t.get(W + "val") in ("nextPage", "page"):
            t.set(W + "val", "continuous")

    # 2) quita <w:br w:type="page"/>
    for br in root.findall(".//w:br", ns):
        if br.get(W + "type") == "page":
            parent = None
            # ET estándar no tiene getparent; buscamos a mano
            for p in root.iter():
                for ch in list(p):
                    if ch is br:
                        parent = p
                        break
            if parent is not None:
                parent.remove(br)

    doc_bytes = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    # reempacar DOCX
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, b in files.items():
            if n == "word/document.xml":
                zout.writestr(n, doc_bytes)
            else:
                zout.writestr(n, b)






def _tmp_heading_doc(texto: str) -> str:
    """Crea un DOCX temporal con un párrafo centrado y en negrita como título de sección."""
    from docx import Document as DocxDocument
    import tempfile, os
    doc = DocxDocument()
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    # centrado
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # un poquito de espacio abajo
    p.paragraph_format.space_after = 300  # 15pt aprox

    fd, path = tempfile.mkstemp(suffix=".docx", prefix="titulo_")
    os.close(fd)
    doc.save(path)
    return path


def _merge_grouped_with_headings(grouped, out_path):
    """
    grouped: [(tema_nombre, [docx1, docx2, ...]), ...]
    Une metiendo un doc temporal con el título antes de cada bloque.
    """


    # Documento maestro vacío
    maestro = DocxDocument()
    comp = Composer(maestro)

    tmp_titles = []
    try:
        for (tema, files) in grouped:
            # 1) insertar título como un pequeño DOCX
            tpath = _tmp_heading_doc(tema)
            tmp_titles.append(tpath)
            comp.append(DocxDocument(tpath))  # añade el título

            # 2) añadir todas las preguntas del tema
            for f in files:
                comp.append(DocxDocument(os.path.abspath(f)))

        comp.save(out_path)
        return out_path, [], []  # compatibles con tu interfaz
    finally:
        # limpiar temporales
        for p in tmp_titles:
            try:
                os.remove(p)
            except Exception:
                pass


def _merge_grouped_with_headings_wordcom(grouped, out_path):
    """
    grouped: [(tema, [paths_pregunta,...]), ...]
    - Crea un DOCX temporal de título por tema.
    - Inserta con Word/COM.
    - NO pone salto de página después del título; SÍ después de cada pregunta.
    """
    tmp_titles = []
    flat = []  # [(path, is_title)]
    for (tema, files) in grouped:
        tpath = _tmp_heading_doc(tema)
        tmp_titles.append(tpath)
        flat.append((tpath, True))
        for f in files:
            flat.append((os.path.abspath(f), False))

    try:
        out, _, malos = _merge_with_word(flat, out_path)  # usa COM con flags
        # Si TODO lo que no es título falló, verás solo títulos → devuélvelo en JSON
        if malos:
            print("⚠️ Archivos que Word no pudo insertar:", malos[:5], "… total:", len(malos))
        return out, [], malos
    finally:
        for p in tmp_titles:
            try: os.remove(p)
            except: pass


def _merge_with_word(marked_paths, out_path):
    """
    marked_paths: [(path, is_title)]
    Inserta cada archivo con Word. PageBreak solo tras preguntas (is_title=False).
    Hace fallback copiando FormattedText si InsertFile falla.
    Filtra inexistentes y vacíos antes de llamar a Word.
    """
    import os, pythoncom
    import win32com.client as win32

    wdCollapseEnd = 0
    wdFormatXMLDocument = 12
    wdPageBreak = 7

    # Filtrado previo (evita rutas malas que 'matan' el merge)
    cleaned = []
    for p, is_title in marked_paths:
        if not p or not os.path.isfile(p):
            continue
        try:
            if _tiene_texto_o_contenido(p):  # ya la tienes definida
                cleaned.append((os.path.abspath(p), is_title))
        except Exception:
            # si no puedo leer el zip, lo salto
            continue

    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    malos = []
    try:
        doc_dest = word.Documents.Add()

        def end_range():
            r = doc_dest.Content
            r.Collapse(wdCollapseEnd)
            return r

        for p_abs, is_title in cleaned:
            try:
                r = end_range()
                # InsertFile suele preservar más formato que FormattedText
                r.InsertFile(p_abs)
            except Exception as e:
                # Fallback: abrir y volcar FormattedText (sin portapapeles)
                malos.append((p_abs, f"InsertFile: {e}"))
                try:
                    doc_src = word.Documents.Open(
                        p_abs, ReadOnly=True, ConfirmConversions=False,
                        AddToRecentFiles=False, Revert=False, Visible=False,
                        OpenAndRepair=True
                    )
                    src = doc_src.Content
                    dst = end_range()
                    dst.FormattedText = src.FormattedText
                    doc_src.Close(False)
                except Exception as e2:
                    malos.append((p_abs, f"FormattedText: {e2}"))
                    continue  # este archivo definitivamente no entró

            # 👇 Salto de página SOLO después de preguntas
            if not is_title:
                end_range().InsertBreak(wdPageBreak)

        try:
            doc_dest.SaveAs2(out_path, FileFormat=wdFormatXMLDocument)
        except Exception:
            doc_dest.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
        doc_dest.Close(False)
        return out_path, [], malos
    finally:
        try: word.Quit()
        except: pass
        pythoncom.CoUninitialize()




def aplanar_listas_a_texto(docx_path: str):
    """
    Elimina TODA numeración de Word (w:numPr) y numera las preguntas como texto:
    1. ..., 2. ..., 3. ... (continuo en todo el documento).
    Borra además numbering.xml y su relación/Content-Type para que Word no repare nada.
    """
    import re, zipfile, xml.etree.ElementTree as ET

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns = {"w": NS_W}

    # Un "inicio de pregunta" escrito como texto: 1)  / 1.  / (1)  / 1⌀ (pegado)
    RX_Q = re.compile(r"^\s*\(?\d{1,3}\)?[.)]?(?:\s+|(?=[^\s]))")

    def _para_text(p):
        return "".join((t.text or "") for t in p.findall(".//w:t", ns))

    def _strip_prefix_in_runs(p, n_chars):
        # quita n_chars del principio del párrafo, repartidos en los runs
        remaining = n_chars
        for r in p.findall(".//w:r", ns):
            t = r.find("./w:t", ns)
            if t is None or remaining <= 0:
                continue
            cur = t.text or ""
            if len(cur) <= remaining:
                t.text = ""
                remaining -= len(cur)
            else:
                t.text = cur[remaining:]
                break

    def _prepend_text(p, txt):
        # inserta un run con texto preservando espacios al inicio del párrafo
        r = ET.Element(W + "r")
        t = ET.SubElement(r, W + "t", {W + "space": "preserve"})
        t.text = txt
        # colócalo antes del primer hijo que no sea pPr
        children = list(p)
        idx = 0
        if children and children[0].tag == W + "pPr":
            idx = 1
        p.insert(idx, r)

    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    if "word/document.xml" not in files:
        return

    root = ET.fromstring(files["word/document.xml"])
    contador = 1
    changed_doc = False

    for p in root.findall(".//w:p", ns):
        # 1) Si tiene numeración de Word, elimínala
        pPr = p.find("./w:pPr", ns)
        if pPr is not None:
            numPr = pPr.find("./w:numPr", ns)
            if numPr is not None:
                pPr.remove(numPr)
                # marcarlo como "pregunta" (para numerarlo)
                txt = _para_text(p)
                # si empieza con número textual, lo quitamos; si no, nada
                m = RX_Q.match(txt)
                if m:
                    _strip_prefix_in_runs(p, m.end())
                _prepend_text(p, f"{contador}. ")
                contador += 1
                changed_doc = True
                continue

        # 2) Si no tenía lista pero empieza con número textual → numerar y limpiar prefijo
        txt = _para_text(p)
        m = RX_Q.match(txt)
        if m:
            _strip_prefix_in_runs(p, m.end())
            _prepend_text(p, f"{contador}. ")
            contador += 1
            changed_doc = True

    if changed_doc:
        files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    # 3) Elimina numbering.xml y su relación/override para que Word no "repairs"
    if "word/numbering.xml" in files:
        files.pop("word/numbering.xml", None)

        # limpia la relación
        rels_name = "word/_rels/document.xml.rels"
        if rels_name in files:
            rels_root = ET.fromstring(files[rels_name])
            REL = "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
            to_del = []
            for r in rels_root.findall(REL):
                if (r.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
                        and r.get("Target") == "numbering.xml"):
                    to_del.append(r)
            for r in to_del:
                rels_root.remove(r)
            files[rels_name] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)

        # limpia el override de content types
        ct_name = "[Content_Types].xml"
        if ct_name in files:
            Types = ET.fromstring(files[ct_name])
            ns_ct = "http://schemas.openxmlformats.org/package/2006/content-types"
            Override = "{%s}Override" % ns_ct
            to_del = []
            for el in list(Types):
                if el.tag == Override and el.get("PartName") == "/word/numbering.xml":
                    to_del.append(el)
            for el in to_del:
                Types.remove(el)
            files[ct_name] = ET.tostring(Types, encoding="utf-8", xml_declaration=True)

    # 4) Reempaqueta el DOCX
    _safe_rezip(docx_path, files)


def _com_disponible() -> bool:
    """True si estamos en Windows y win32com funciona."""
    try:
        return (os.name == "nt") and (win32 is not None)
    except Exception:
        return False



def _safe_rezip(docx_path: str, files: dict):
    import zipfile, os, tempfile, time, shutil

    dest_dir = os.path.dirname(os.path.abspath(docx_path))
    os.makedirs(dest_dir, exist_ok=True)

    # 1) temporal EN LA MISMA CARPETA DEL DESTINO (misma unidad)
    fd, tmp = tempfile.mkstemp(suffix=".docx", prefix="safe_", dir=dest_dir)
    os.close(fd)  # ¡clave en Windows!

    # 2) escribir el zip al temporal
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, b in files.items():
            n = n.replace("\\", "/")
            if isinstance(b, str):
                b = b.encode("utf-8")
            zout.writestr(n, b)

    # 3) reemplazo atómico con reintentos
    for _ in range(8):
        try:
            os.replace(tmp, docx_path)  # misma unidad ⇒ OK
            return
        except PermissionError:
            time.sleep(0.15)
        except OSError:
            # No debería ocurrir al estar en la misma carpeta,
            # pero por si acaso: copia y borra.
            try:
                shutil.copy2(tmp, docx_path)
                os.remove(tmp)
                return
            except Exception:
                time.sleep(0.15)

    # si todo falla, limpia y propaga
    try: os.remove(tmp)
    except: pass
    raise

def bullets_to_numbers_docx(docx_path: str):
    """
    Abre el .docx y convierte TODAS las definiciones de listas 'bullet'
    en listas decimales (1., 2., 3., …) en word/numbering.xml.
    No toca document.xml: al cambiar el abstractNum a decimal,
    todos los párrafos que usan ese numId pasan a numerados.
    """
    import zipfile, xml.etree.ElementTree as ET, io

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns = {"w": NS_W}

    # Leer archivos del docx
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    if "word/numbering.xml" not in files:
        # No hay listas en el documento
        return

    root = ET.fromstring(files["word/numbering.xml"])

    # Reescribir todos los niveles con bullet -> decimal y lvlText -> %N.
    for absN in root.findall(".//w:abstractNum", ns):
        for lvl in absN.findall("./w:lvl", ns):
            ilvl = lvl.get(W + "ilvl", "0")

            # <w:numFmt w:val="bullet"/>  →  decimal
            numFmt = lvl.find("./w:numFmt", ns)
            if numFmt is not None and numFmt.get(W + "val") == "bullet":
                numFmt.set(W + "val", "decimal")

            # Asegura <w:lvlText w:val="%{n}."/> (n = ilvl+1)
            lvlText = lvl.find("./w:lvlText", ns)
            desired = f"%{int(ilvl) + 1}."
            if lvlText is None:
                lvlText = ET.SubElement(lvl, W + "lvlText", {W + "val": desired})
            else:
                lvlText.set(W + "val", desired)

            # Opcional: inicio en 1 y alineación a la izquierda
            start = lvl.find("./w:start", ns)
            if start is None:
                ET.SubElement(lvl, W + "start", {W + "val": "1"})
            else:
                start.set(W + "val", "1")
            lvlJc = lvl.find("./w:lvlJc", ns)
            if lvlJc is None:
                ET.SubElement(lvl, W + "lvlJc", {W + "val": "left"})
            else:
                lvlJc.set(W + "val", "left")

    # Guardar de nuevo el DOCX con numbering.xml modificado
    new_numbering = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, b in files.items():
            if n == "word/numbering.xml":
                b = new_numbering
            zout.writestr(n, b)



def guardar_pdf(ruta_word: str) -> str:
    ruta_word = os.path.abspath(ruta_word)

    tmpdir = tempfile.mkdtemp(prefix="pdf_")
    base = os.path.splitext(os.path.basename(ruta_word))[0]
    ruta_docx_tmp = os.path.join(tmpdir, base + "_norm.docx")
    ruta_pdf = os.path.join(tmpdir, base + ".pdf")

    shutil.copy2(ruta_word, ruta_docx_tmp)

    # Word consts
    wdExportFormatPDF = 17
    wdExportOptimizeForPrint = 0
    wdExportAllDocument = 0
    wdExportDocumentContent = 0
    wdExportCreateNoBookmarks = 0
    wdFormatXMLDocument = 12

    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        doc = word.Documents.Open(
            ruta_docx_tmp,
            ReadOnly=False,
            ConfirmConversions=False,
            AddToRecentFiles=False,
            Visible=False,
            OpenAndRepair=True
        )
        # Normaliza
        doc.SaveAs(ruta_docx_tmp, FileFormat=wdFormatXMLDocument)

        # Actualiza campos (ítems numerados/TOC/refs) y repagina
        try:
            doc.Fields.Update()
        except Exception:
            pass
        try:
            doc.Repaginate()
        except Exception:
            pass

        # Exporta a PDF (impresión)
        doc.ExportAsFixedFormat(
            OutputFileName=ruta_pdf,
            ExportFormat=wdExportFormatPDF,
            OpenAfterExport=False,
            OptimizeFor=wdExportOptimizeForPrint,
            Range=wdExportAllDocument,
            From=1, To=1,
            Item=wdExportDocumentContent,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=wdExportCreateNoBookmarks,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False
        )
        doc.Close(False)

        for _ in range(20):
            if os.path.exists(ruta_pdf) and os.path.getsize(ruta_pdf) > 0:
                break
            time.sleep(0.1)

        if not os.path.exists(ruta_pdf) or os.path.getsize(ruta_pdf) == 0:
            raise RuntimeError("Word no generó el PDF (vacío).")

        return ruta_pdf

    finally:
        try:
            word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()

def _sha1_file(path: str, block=1024*1024) -> str:
    import hashlib
    h = hashlib.sha1()
    with open(path, "rb") as f:
        while True:
            b = f.read(block)
            if not b: break
            h.update(b)
    return h.hexdigest()

# --- Forzar UTF-8 + parches HTML Word ---
def _force_utf8_html(html_path: str):
    try:
        try:
            txt = open(html_path, "r", encoding="utf-8").read()
            decoded_as = "utf-8"
        except UnicodeDecodeError:
            txt = open(html_path, "r", encoding="cp1252", errors="strict").read()
            decoded_as = "cp1252"
        import re
        txt = re.sub(r'(?i)charset\s*=\s*[-\w]+', 'charset=utf-8', txt, count=1)
        txt = re.sub(r'(?i)<meta\s+charset=["\']?[-\w]+["\']?\s*/?>',
                     '<meta charset="utf-8">', txt, count=1)
        if '<meta' not in txt.lower():
            txt = txt.replace('<head>', '<head><meta charset="utf-8">', 1)
        with open(html_path, "w", encoding="utf-8", newline="") as f:
            f.write(txt)
        print(f"[HTML] {os.path.basename(html_path)} normalizado a UTF-8 (desde {decoded_as}).")
    except Exception as e:
        print("[HTML] No se pudo forzar UTF-8:", e)

def _postprocess_word_html(html_path: str):
    """
    Correcciones típicas del HTML de Word en navegadores + CSS base.
    """
    try:
        with open(html_path, "r", encoding="utf-8", errors="ignore") as f:
            txt = f.read()

        import re
        # Quitar comentarios condicionales VML de Word
        txt = re.sub(r'<!--\s*\[if\s+gte\s+vml\s+1\s*\]>.?<!\s\[endif\]\s*-->',
                     '', txt, flags=re.I | re.S)
        txt = re.sub(r'<!--\s*\[if\s*!vml\s*\]-->', '', txt, flags=re.I)
        txt = re.sub(r'<!--\s*<!\s*\[endif\]\s*-->', '', txt, flags=re.I)

        extra_css = """
<style>
img { max-width: none !important; height: auto; }
body { margin:0; }
</style>
"""
        if "</head>" in txt.lower():
            txt = re.sub(r'</head>', extra_css + '</head>', txt, flags=re.I, count=1)
        else:
            txt = extra_css + txt

        with open(html_path, "w", encoding="utf-8", newline="") as f:
            f.write(txt)
    except Exception as e:
        print("[HTML] Post-proceso fallido:", e)



def _wait_exists_nonzero(path, tries=80, delay=0.1):
    for _ in range(tries):
        try:
            if os.path.exists(path) and os.path.getsize(path) > 0:
                # prueba abrir en lectura para asegurar desbloqueo
                with open(path, "rb"):
                    pass
                return True
        except Exception:
            pass
        time.sleep(delay)
    return False


    
def _preview_with_mammoth(docx_path: str, nombre_base: str) -> str:
    """
    Convierte DOCX -> HTML usando Mammoth (sin Word).
    Devuelve la ruta absoluta del HTML generado en DESCARGAS_FOLDER.
    """
    import mammoth, base64, os

    dst_html = os.path.join(app.config['DESCARGAS_FOLDER'], f"{nombre_base}.htm")

    # Incrusta las imágenes como data URI para que se vean en el iframe
    def _img_inline(element):
        if element.content_type and element.read:
            data = element.read()
            b64 = base64.b64encode(data).decode("ascii")
            return {"src": f"data:{element.content_type};base64,{b64}"}
        return None

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(_img_inline)
        )
        html = result.value  # HTML

    # Asegura UTF-8 y un CSS mínimo
    extra_css = "<style>img{max-width:none;height:auto} body{margin:0}</style>"
    if "<head>" in html.lower():
        html = html.replace("<head>", "<head><meta charset='utf-8'>" + extra_css, 1)
    else:
        html = "<meta charset='utf-8'>" + extra_css + html

    with open(dst_html, "w", encoding="utf-8", newline="") as f:
        f.write(html)

    return dst_html

def _exportar_html_con_mammoth(docx_path: str, base_name: str) -> tuple[str, list]:
    """
    Fallback: DOCX -> HTML con Mammoth. Incrusta imágenes en data URI.
    """
    import base64, mammoth, re
    warnings = []
    safe_base = re.sub(r'[\\/:*?"<>|]+', "_", base_name).strip()
    html_dst  = os.path.join(PREVIEWS_DIR, f"{safe_base}.htm")

    def _img_inline(image):
        try:
            with image.open() as f:
                raw = f.read()
            mime = image.content_type or "image/png"
            b64  = base64.b64encode(raw).decode("ascii")
            return {"src": f"data:{mime};base64,{b64}"}
        except Exception as e:
            warnings.append(f"img_inline:{e}")
            return None

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f, convert_image=mammoth.images.img_element(_img_inline))
        html = result.value
        for m in result.messages:
            warnings.append(f"mammoth:{m}")

    extra = "<meta charset='utf-8'><style>img{max-width:none;height:auto}body{margin:0}</style>"
    if "<head>" in html.lower():
        html = html.replace("<head>", "<head>"+extra, 1)
    else:
        html = extra + html

    with open(html_dst, "w", encoding="utf-8", newline="") as fo:
        fo.write(html)

    return html_dst, warnings
def _exportar_html_con_word(docx_path: str, base_name: str) -> tuple[str, list]:
    """
    Exporta con Word a 'Web Page, Filtered' dentro de PREVIEWS_DIR.
    Devuelve (ruta_html, warnings). Mantiene imágenes/OMML mejor que Mammoth.
    """
    import os, re, shutil, tempfile
    import pythoncom
    import win32com.client as win32

    wdFormatFilteredHTML = 10
    wdFormatXMLDocument  = 12
    msoAutomationSecurityForceDisable = 3
    warnings: list[str] = []

    # Nombre HTML estable y carpeta de recursos "<base>_files"
    safe_base = re.sub(r'[\\/:*?"<>|]+', "_", base_name).strip()
    html_dst  = os.path.join(PREVIEWS_DIR, f"{safe_base}.htm")
    res_folder_name = f"{safe_base}_files"
    res_dst = os.path.join(PREVIEWS_DIR, res_folder_name)

    tmp_dir = tempfile.mkdtemp(prefix="prev_word_")
    try:
        # Copia temporal del DOCX (Word a veces bloquea el original)
        tmp_docx = os.path.join(tmp_dir, os.path.basename(docx_path))
        shutil.copy2(docx_path, tmp_docx)

        # Ruta 8.3 ayuda MUCHÍSIMO a Word en Windows
        try:
            import win32api
            src_open = win32api.GetShortPathName(tmp_docx)
        except Exception:
            src_open = tmp_docx

        pythoncom.CoInitialize()
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            try:
                word.AutomationSecurity = msoAutomationSecurityForceDisable
            except Exception:
                pass

            doc = None
            # Estrategia A: abrir con reparación
            try:
                doc = word.Documents.Open(
                    src_open,
                    ReadOnly=True,
                    ConfirmConversions=False,
                    AddToRecentFiles=False,
                    Visible=False,
                    OpenAndRepair=True,
                )
            except Exception:
                # Estrategia B: InsertFile en un documento en blanco
                try:
                    doc = word.Documents.Add()
                    rng = doc.Range(0, 0)
                    rng.InsertFile(src_open, ConfirmConversions=False, Link=False, Attachment=False)
                except Exception:
                    # Estrategia C: copiar FormattedText desde doc fuente
                    srcdoc = word.Documents.Open(
                        src_open,
                        ReadOnly=True,
                        ConfirmConversions=False,
                        AddToRecentFiles=False,
                        Visible=False,
                        OpenAndRepair=True,
                    )
                    doc = word.Documents.Add()
                    dst = doc.Range(0, 0)
                    dst.FormattedText = srcdoc.Content.FormattedText
                    srcdoc.Close(False)

            # Normaliza a DOCX "limpio" y exporta HTML filtrado directamente en PREVIEWS_DIR
            norm_docx = os.path.join(tmp_dir, "norm.docx")
            doc.SaveAs(norm_docx, FileFormat=wdFormatXMLDocument)
            try:
                doc.WebOptions.AllowPNG = True
                doc.WebOptions.OptimizeForBrowser = True
                doc.WebOptions.RelyOnCSS = True
            except Exception:
                pass
            doc.SaveAs(os.path.abspath(html_dst), FileFormat=wdFormatFilteredHTML)
            doc.Close(False)

        finally:
            try: word.Quit()
            except Exception: pass
            pythoncom.CoUninitialize()

        # Si Word creó la carpeta de recursos con otro nombre, la movemos a <safe_base>_files
        parent = os.path.dirname(html_dst)
        generadas = [d for d in os.listdir(parent) if d.endswith("_files")]
        if res_folder_name not in generadas and generadas:
            generadas.sort(key=lambda d: os.path.getmtime(os.path.join(parent, d)), reverse=True)
            origen = os.path.join(parent, generadas[0])
            if origen != res_dst:
                try:
                    if os.path.exists(res_dst):
                        shutil.rmtree(res_dst, ignore_errors=True)
                    shutil.move(origen, res_dst)
                except Exception as e:
                    warnings.append(f"move_files:{e}")

        # Reescribe referencias a recursos → /static/previews/<base>_files/...
        try:
            with open(html_dst, "r", encoding="utf-8", errors="ignore") as f:
                html = f.read()

            html = re.sub(
                r'(?i)(["\'])\.?/[^"\']+?_files/',
                rf'\1/static/previews/{res_folder_name}/',
                html,
            )
            # Fuerza UTF-8
            html = re.sub(r'(?i)charset\s*=\s*[-\w]+', 'charset=utf-8', html, count=1)
            if "<meta" not in html[:400].lower():
                html = html.replace("<head>", "<head><meta charset=\"utf-8\">", 1)

            with open(html_dst, "w", encoding="utf-8", newline="") as f:
                f.write(html)
        except Exception as e:
            warnings.append(f"rewrite_src:{e}")

        return html_dst, warnings

    finally:
        try: shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception: pass


# --- sirve /static desde la carpeta de proyecto ---
@app.route("/static/<path:subpath>")
def serve_static_project(subpath):
    return send_from_directory(STATIC_DIR, subpath)

    
def generar_html_desde_docx(ruta_docx: str, nombre_base: str | None = None) -> str:
    """
    Exporta una COPIA del DOCX a HTML (Filtered) con Word/COM.
    Si Word falla, usa Mammoth como fallback.
    Devuelve la ruta absoluta del .htm en DESCARGAS_FOLDER.
    """
    wdFormatFilteredHTML = 10
    wdFormatXMLDocument  = 12
    msoEncodingUTF8      = 65001
    msoAutomationSecurityForceDisable = 3  # desactiva avisos por macros/plantillas

    ruta_docx = os.path.abspath(ruta_docx)
    if not os.path.isfile(ruta_docx):
        raise FileNotFoundError(ruta_docx)

    os.makedirs(app.config['DESCARGAS_FOLDER'], exist_ok=True)
    base = nombre_base or (os.path.splitext(os.path.basename(ruta_docx))[0] + "_preview")
    dst_html = os.path.join(app.config['DESCARGAS_FOLDER'], f"{base}.htm")

    tmpdir = tempfile.mkdtemp(prefix="preview_")
    try:
        src_copy = os.path.join(tmpdir, os.path.basename(ruta_docx))
        shutil.copy2(ruta_docx, src_copy)
        if not _wait_exists_nonzero(src_copy):
            raise RuntimeError("Copia temporal no lista.")

        src_open = _short_path(src_copy)  # rutas 8.3 ayudan a Word
        norm_doc = os.path.join(tmpdir, "norm.docx")

        # ---------- Intento con Word ----------
        try:
            pythoncom.CoInitialize()
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            try:
                word.AutomationSecurity = msoAutomationSecurityForceDisable
            except Exception:
                pass

            try:
                doc = None
                # A) Open & Repair
                try:
                    doc = word.Documents.Open(
                        src_open,
                        ReadOnly=True,
                        ConfirmConversions=False,
                        AddToRecentFiles=False,
                        Visible=False,
                        OpenAndRepair=True
                    )
                except Exception:
                    # B) Blanco + Range.InsertFile(...)
                    try:
                        doc = word.Documents.Add()
                        rng = doc.Range(0, 0)
                        rng.InsertFile(src_open, ConfirmConversions=False, Link=False, Attachment=False)
                    except Exception:
                        # C) Abrir fuente y copiar FormattedText
                        srcdoc = word.Documents.Open(
                            src_open,
                            ReadOnly=True,
                            ConfirmConversions=False,
                            AddToRecentFiles=False,
                            Visible=False,
                            OpenAndRepair=True
                        )
                        doc = word.Documents.Add()
                        dst = doc.Range(0, 0)
                        dst.FormattedText = srcdoc.Content.FormattedText
                        srcdoc.Close(False)

                # Normaliza a DOCX y exporta HTML filtrado
                doc.SaveAs(norm_doc, FileFormat=wdFormatXMLDocument)
                try:
                    doc.WebOptions.AllowPNG = True
                    doc.WebOptions.OptimizeForBrowser = True
                    doc.WebOptions.RelyOnCSS = True
                except Exception:
                    pass
                doc.SaveAs2(FileName=os.path.abspath(dst_html),
                            FileFormat=wdFormatFilteredHTML,
                            Encoding=msoEncodingUTF8)
                doc.Close(False)
            finally:
                try: word.Quit()
                except Exception: pass
                pythoncom.CoUninitialize()

            if not _wait_exists_nonzero(dst_html):
                raise RuntimeError("Word no generó el HTML de vista previa.")

            _force_utf8_html(dst_html)
            _postprocess_word_html(dst_html)
            return dst_html

        except Exception:
            # ---------- Fallback Mammoth ----------
            try:
                return _preview_with_mammoth(ruta_docx, base)
            except Exception as e2:
                raise RuntimeError(f"No se pudo generar la vista previa (Word/Mammoth): {e2}") from e2

    finally:
        try: shutil.rmtree(tmpdir, ignore_errors=True)
        except Exception: pass


def docx_a_html_filtrado(docx_path, out_dir):
    """
    Convierte DOCX a HTML filtrado con Word.
    Devuelve (html_abs_path, html_rel_url) o (None, None) si falla.
    """
    os.makedirs(out_dir, exist_ok=True)

    # Nombre base sin espacios raros
    base = re.sub(r'[^A-Za-z0-9_-]+', '_', os.path.splitext(os.path.basename(docx_path))[0])
    html_name = f"{base}.html"
    html_out = os.path.join(out_dir, html_name)

    # Word: 10 = wdFormatFilteredHTML
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path, ReadOnly=True)
        doc.SaveAs2(html_out, FileFormat=10)
        doc.Close(False)
    except Exception as e:
        print("[HTML] Error exportando:", e)
        try:
            if word: word.Quit()
        except: pass
        pythoncom.CoUninitialize()
        return None, None
    finally:
        try:
            if word: word.Quit()
        except: pass
        pythoncom.CoUninitialize()

    # Word crea una carpeta “<base>_archivos” junto al HTML
    recursos_dir = os.path.join(out_dir, f"{base}_archivos")
    # (opcional) renómbrala a algo estable
    if os.path.isdir(recursos_dir):
        fijo = os.path.join(out_dir, "assets")
        if os.path.isdir(fijo):
            shutil.rmtree(fijo, ignore_errors=True)
        os.rename(recursos_dir, fijo)
        # reescribir rutas en el HTML para apuntar a /assets/
        try:
            with open(html_out, "r", encoding="utf-8", errors="ignore") as f:
                html = f.read()
            html = re.sub(rf'{re.escape(base)}_archivos/', 'assets/', html)
            # añade <base href="/static/previews/…/"> para que todo resuelva
            html = html.replace("<head>", f'<head><base href="./">', 1)
            with open(html_out, "w", encoding="utf-8") as f:
                f.write(html)
        except Exception as e:
            print("[HTML] No se pudo normalizar rutas:", e)

    # Devuelve ruta relativa para servirlo como /static/previews/...
    # Ajusta esto a tu layout real de carpetas
    rel = f"/static/previews/{os.path.basename(out_dir)}/{html_name}"
    return html_out, rel



@app.post("/api/pdf_from_docx")
def api_pdf_from_docx():
    """
    Convierte un DOCX que ya existe en /descargas a PDF.
    Body JSON: { "docx": "Nombre del archivo.docx" }  (también acepta "nombre")
    Devuelve: { ok, archivo_pdf, ruta_rel_pdf }
    """
    try:
        data = request.get_json(force=True) or {}
        docx_name = (data.get("docx") or data.get("nombre") or "").strip()
        if not docx_name or "/" in docx_name or "\\" in docx_name:
            return jsonify(ok=False, error="nombre DOCX inválido"), 400

        docx_path = os.path.join(app.config['DESCARGAS_FOLDER'], docx_name)
        if not os.path.isfile(docx_path):
            return jsonify(ok=False, error=f"No existe DOCX: {docx_name}"), 404

        pdf_name = os.path.splitext(docx_name)[0] + ".pdf"
        pdf_path = os.path.join(app.config['DESCARGAS_FOLDER'], pdf_name)

        # usa tu función robusta
        try:
            final_pdf = docx_a_pdf(docx_path, pdf_path)
        except Exception as e:
            return jsonify(ok=False, error=f"docx_a_pdf: {e}"), 500

        if not (os.path.exists(final_pdf) and os.path.getsize(final_pdf) > 0):
            return jsonify(ok=False, error="PDF no se generó (vacío)"), 500

        return jsonify(
            ok=True,
            archivo_pdf=pdf_name,
            ruta_rel_pdf=f"/api/descargas/{pdf_name}",
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify(ok=False, error=str(e)), 500



def to_pdf_insert_only(src, dst):
    """Abre un doc en blanco y hace InsertFile(src) -> ExportAsFixedFormat(dst)."""
    src = _short83(src)
    dst = _short83(dst)

    pythoncom.CoInitialize()
    w = win32.DispatchEx("Word.Application")
    w.Visible = False
    w.DisplayAlerts = 0
    try:
        d = w.Documents.Add()
        d.Range(0, 0).InsertFile(src, ConfirmConversions=False, Link=False, Attachment=False)
        # 1º intento
        try:
            d.ExportAsFixedFormat(OutputFileName=dst, ExportFormat=17, OpenAfterExport=False)
        except Exception:
            # 2º intento (algunos Word prefieren SaveAs2 con FileFormat=17)
            d.SaveAs2(dst, FileFormat=17)
        d.Close(False)
    finally:
        try: w.Quit()
        finally: pythoncom.CoUninitialize()

    # esperar a que el PDF exista y tenga tamaño
    for _ in range(50):
        if os.path.exists(dst) and os.path.getsize(dst) > 0:
            return dst
        time.sleep(0.1)
    raise RuntimeError("no pdf")

def resave_docx_formatted(src, dst_docx):
    src = _short83(src); dst_docx = _short83(dst_docx)
    pythoncom.CoInitialize()
    w = win32.DispatchEx("Word.Application"); w.Visible=False; w.DisplayAlerts=0
    try:
        # Abrir con reparación si se puede; si no, InsertFile
        try:
            s = w.Documents.Open(src, ReadOnly=True, ConfirmConversions=False, Visible=False, OpenAndRepair=True)
            d = w.Documents.Add()
            d.Range(0,0).FormattedText = s.Content.FormattedText
            s.Close(False)
        except Exception:
            d = w.Documents.Add()
            d.Range(0,0).InsertFile(src, ConfirmConversions=False, Link=False, Attachment=False)
        d.SaveAs2(dst_docx, FileFormat=12)  # DOCX normalizado
        d.Close(False)
    finally:
        try: w.Quit()
        finally: pythoncom.CoUninitialize()
    return dst_docx

def docx_a_pdf(docx_path: str, pdf_path: str) -> str:
    """
    DOCX -> PDF con 3 estrategias:
    A) Open&Repair + ExportAsFixedFormat
    B) Doc en blanco + InsertFile + ExportAsFixedFormat
    C) Copiar FormattedText + ExportAsFixedFormat
    Espera a que el PDF exista y tenga tamaño. Lanza RuntimeError si no puede.
    """
    import os, time, pythoncom
    import win32com.client as win32

    docx_path = os.path.abspath(docx_path)
    pdf_path  = os.path.abspath(pdf_path)
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)

    src = _short83(docx_path)
    dst = _short83(pdf_path)

    # Pequeña espera por si acabas de reescribir el DOCX
    for _ in range(20):
        if os.path.exists(docx_path) and os.path.getsize(docx_path) > 0:
            break
        time.sleep(0.1)

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
        except Exception:
            pass

        doc = None
        # A) Open&Repair
        try:
            doc = word.Documents.Open(
                src, ReadOnly=True, ConfirmConversions=False,
                AddToRecentFiles=False, Visible=False, OpenAndRepair=True
            )
        except Exception as e_open:
            # B) Blanco + InsertFile
            try:
                doc = word.Documents.Add()
                doc.Range(0, 0).InsertFile(src, ConfirmConversions=False, Link=False, Attachment=False)
            except Exception as e_ins:
                # C) Copiar FormattedText
                try:
                    s = word.Documents.Open(
                        src, ReadOnly=True, ConfirmConversions=False,
                        AddToRecentFiles=False, Visible=False, OpenAndRepair=True
                    )
                    doc = word.Documents.Add()
                    doc.Range(0, 0).FormattedText = s.Content.FormattedText
                    s.Close(False)
                except Exception as e_fmt:
                    raise RuntimeError(f"Open:{e_open} Insert:{e_ins} Fmt:{e_fmt}")

        # Exporta a PDF
        try:
            doc.ExportAsFixedFormat(
                OutputFileName=dst, ExportFormat=17,  # wdExportFormatPDF
                OpenAfterExport=False, OptimizeFor=0,  # Print
                Range=0, Item=0, IncludeDocProps=True, KeepIRM=True,
                CreateBookmarks=0, DocStructureTags=True,
                BitmapMissingFonts=True, UseISO19005_1=False
            )
        except Exception:
            # fallback raro: algunos Word aceptan SaveAs2 con FileFormat=17
            doc.SaveAs2(dst, FileFormat=17)
        finally:
            doc.Close(False)

        # Espera a que realmente exista y con tamaño > 0
        for _ in range(60):
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                return pdf_path
            time.sleep(0.1)

        # Intento adicional: normalizar DOCX y reintentar con InsertFile
        tmp_norm = os.path.join(os.path.dirname(pdf_path), "_norm.docx")
        resave_docx_formatted(docx_path, tmp_norm)
        to_pdf_insert_only(tmp_norm, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            return pdf_path

        raise RuntimeError("Word no generó el PDF o quedó vacío.")
    finally:
        try:
            if word: word.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()



# boton banco solucionario 

def _save_docx(file_storage, folder):
    filename = secure_filename(file_storage.filename)
    # nombre único para evitar colisiones
    ts = int(time.time() * 1000)
    name, ext = os.path.splitext(filename)
    final_name = f"{name}_{ts}{ext}"
    path = os.path.join(folder, final_name)
    file_storage.save(path)
    return final_name, path

def _insert_pregunta_from_doc(conn, tema_id, doc_name, doc_path):
    """
    Crea una nueva fila en `preguntas` para un DOCX del banco.

    - examenes_idexamenes = NULL  (es del banco, no de un examen concreto)
    - tema_id = tema del DOCX
    - numero_p = correlativo dentro del tema (solo para las del banco)
    - archivo_nombre / archivo_ruta = DOCX que se subió
    """
    cur = conn.cursor()

    # número siguiente solo entre las preguntas del banco (examenes_idexamenes IS NULL)
    cur.execute("""
        SELECT COALESCE(MAX(numero_p), 0) + 1
        FROM preguntas
        WHERE tema_id = %s AND examenes_idexamenes IS NULL
    """, (tema_id,))
    next_num = cur.fetchone()[0] or 1

    cur.execute("""
        INSERT INTO preguntas
            (examenes_idexamenes, tema_id, numero_p, archivo_nombre, archivo_ruta)
        VALUES (NULL, %s, %s, %s, %s)
    """, (tema_id, next_num, doc_name, doc_path))

    cur.close()



@app.route("/api/banco_preguntas", methods=["GET"])
def banco_listar():
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT td.id, td.tema_id, t.nombre AS tema_nombre,
                   td.doc_preguntas_nombre, td.doc_preguntas_ruta,
                   td.doc_sol_nombre, td.doc_sol_ruta,
                   td.fecha_creacion
            FROM tema_docs td
            INNER JOIN temas t ON t.id = td.tema_id
            ORDER BY t.nombre
        """)
        rows = cur.fetchall()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try: cur.close(); conn.close()
        except: pass

@app.route("/api/banco_preguntas", methods=["POST"])
def banco_importar_tema():
    """
    FormData:
      - tema_id
      - doc_preguntas (file .docx)

    Ahora SIEMPRE inserta una nueva fila en `tema_docs` para ese tema,
    aunque ya existan otros DOCX del mismo tema.
    Además crea una nueva fila en `preguntas`.
    """
    try:
        tema_id = request.form.get("tema_id", type=int)
        fdoc = request.files.get("doc_preguntas")

        if not tema_id or not fdoc:
            return jsonify({"error": "tema_id y doc_preguntas son obligatorios"}), 400

        doc_name, doc_path = _save_docx(fdoc, BANCO_PREG_DIR)

        conn = get_connection()
        cur = conn.cursor()

        # 🔹 ANTES se hacía SELECT por tema_id y UPDATE (reemplazaba)
        # AHORA: siempre INSERT, permite varios DOCX por el mismo tema
        cur.execute("""
            INSERT INTO tema_docs(tema_id, doc_preguntas_nombre, doc_preguntas_ruta)
            VALUES(%s, %s, %s)
        """, (tema_id, doc_name, doc_path))

        # Crear también la pregunta asociada en `preguntas`
        _insert_pregunta_from_doc(conn, tema_id, doc_name, doc_path)

        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        print("ERROR banco_importar_tema:", e)
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            cur.close()
            conn.close()
        except:
            pass

@app.route("/api/banco_preguntas/<int:id>/reemplazar/preguntas", methods=["POST"])
def banco_reemplazar_preguntas(id):
    fdoc = request.files.get("doc_preguntas")
    if not fdoc:
        return jsonify({"error": "doc_preguntas es obligatorio"}), 400

    doc_name, doc_path = _save_docx(fdoc, BANCO_PREG_DIR)

    conn = get_connection()
    cur = conn.cursor(dictionary=True)

    # 1) Obtener datos actuales del registro
    cur.execute("""
        SELECT tema_id, doc_preguntas_ruta
        FROM tema_docs
        WHERE id = %s
    """, (id,))
    row = cur.fetchone()
    if not row:
        cur.close()
        conn.close()
        return jsonify({"error": "Registro no encontrado"}), 404

    tema_id = row["tema_id"]
    old_ruta = row["doc_preguntas_ruta"]

    # 2) Actualizar tema_docs
    cur.execute("""
        UPDATE tema_docs
        SET doc_preguntas_nombre = %s,
            doc_preguntas_ruta   = %s
        WHERE id = %s
    """, (doc_name, doc_path, id))

    # 3) Intentar actualizar la fila correspondiente en `preguntas`
    cur2 = conn.cursor()
    cur2.execute("""
        UPDATE preguntas
        SET archivo_nombre = %s,
            archivo_ruta   = %s
        WHERE tema_id = %s
          AND examenes_idexamenes IS NULL
          AND archivo_ruta = %s
    """, (doc_name, doc_path, tema_id, old_ruta))

    if cur2.rowcount == 0:
        # Si no encontramos la fila por ruta (caso raro), creamos una nueva
        _insert_pregunta_from_doc(conn, tema_id, doc_name, doc_path)

    cur2.close()

    conn.commit()
    cur.close()
    conn.close()
    return jsonify({"ok": True})

@app.route("/api/banco_preguntas/solucionario", methods=["POST"])
def banco_agregar_solucionario():
    """
    FormData:
      - tema_id
      - doc_solucionario (file .docx)
    """
    try:
        tema_id = request.form.get("tema_id", type=int)
        fsol = request.files.get("doc_solucionario")

        if not tema_id or not fsol:
            return jsonify({"error":"tema_id y doc_solucionario son obligatorios"}), 400

        sol_name, sol_path = _save_docx(fsol, BANCO_SOL_DIR)

        conn = get_connection()
        cur = conn.cursor()

        cur.execute("SELECT id FROM tema_docs WHERE tema_id=%s", (tema_id,))
        exist = cur.fetchone()
        if not exist:
            return jsonify({"error":"Primero importa el tema (doc preguntas)."}), 400

        cur.execute("""
            UPDATE tema_docs
            SET doc_sol_nombre=%s, doc_sol_ruta=%s
            WHERE tema_id=%s
        """, (sol_name, sol_path, tema_id))

        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try: cur.close(); conn.close()
        except: pass

@app.route("/api/banco_preguntas/<int:id>/download/preguntas", methods=["GET"])
def banco_download_preguntas(id):
    conn = get_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT doc_preguntas_ruta, doc_preguntas_nombre FROM tema_docs WHERE id=%s", (id,))
    row = cur.fetchone()
    cur.close(); conn.close()
    if not row or not os.path.exists(row["doc_preguntas_ruta"]):
        return jsonify({"error":"Archivo no encontrado"}), 404
    return send_file(row["doc_preguntas_ruta"], as_attachment=True,
                     download_name=row["doc_preguntas_nombre"])

@app.route("/api/banco_preguntas/<int:id>/download/solucionario", methods=["GET"])
def banco_download_sol(id):
    conn = get_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT doc_sol_ruta, doc_sol_nombre FROM tema_docs WHERE id=%s", (id,))
    row = cur.fetchone()
    cur.close(); conn.close()
    if not row or not row["doc_sol_ruta"] or not os.path.exists(row["doc_sol_ruta"]):
        return jsonify({"error":"Solucionario no disponible"}), 404
    return send_file(row["doc_sol_ruta"], as_attachment=True,
                     download_name=row["doc_sol_nombre"])

@app.route("/api/banco_preguntas/<int:id>", methods=["PUT"])
def banco_editar(id):
    """
    JSON:
      - tema_id (opcional)
    """
    data = request.get_json(force=True)
    tema_id = data.get("tema_id")

    try:
        conn = get_connection()
        cur = conn.cursor()
        if tema_id:
            cur.execute("UPDATE tema_docs SET tema_id=%s WHERE id=%s", (tema_id, id))
        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        try: cur.close(); conn.close()
        except: pass



@app.route("/api/banco_preguntas/<int:id>/reemplazar/solucionario", methods=["POST"])
def banco_reemplazar_sol(id):
    fsol = request.files.get("doc_solucionario")
    if not fsol: return jsonify({"error":"doc_solucionario es obligatorio"}), 400
    sol_name, sol_path = _save_docx(fsol, BANCO_SOL_DIR)
    conn = get_connection(); cur = conn.cursor()
    cur.execute("""
        UPDATE tema_docs
        SET doc_sol_nombre=%s, doc_sol_ruta=%s
        WHERE id=%s
    """, (sol_name, sol_path, id))
    conn.commit(); cur.close(); conn.close()
    return jsonify({"ok": True})

@app.route("/api/banco_preguntas/<int:id>", methods=["DELETE"])
def banco_eliminar(id):
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)

        cur.execute("""
            SELECT tema_id, doc_preguntas_ruta, doc_sol_ruta
            FROM tema_docs
            WHERE id = %s
        """, (id,))
        row = cur.fetchone()
        if not row:
            return jsonify({"error": "Registro no existe"}), 404

        tema_id = row["tema_id"]
        doc_ruta = row["doc_preguntas_ruta"]
        sol_ruta = row["doc_sol_ruta"]

        # 1) Borrar solo este registro del banco
        cur.execute("DELETE FROM tema_docs WHERE id = %s", (id,))

        # 2) Borrar SOLO la pregunta asociada a este DOCX
        cur.execute("""
            DELETE FROM preguntas
            WHERE tema_id = %s
              AND examenes_idexamenes IS NULL
              AND archivo_ruta = %s
        """, (tema_id, doc_ruta))

        conn.commit()
        cur.close()
        conn.close()

        # 3) Eliminar archivos físicos
        for p in [doc_ruta, sol_ruta]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except:
                    pass

        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
from zipfile import ZipFile

@app.route("/api/banco_preguntas/<int:id>/download", methods=["GET"])
def banco_download_full(id):
    conn = get_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("""SELECT doc_preguntas_nombre, doc_preguntas_ruta,
                          doc_sol_nombre, doc_sol_ruta
                   FROM tema_docs WHERE id=%s""", (id,))
    row = cur.fetchone()
    cur.close(); conn.close()

    if not row:
        return jsonify({"error": "Registro no encontrado"}), 404

    if not os.path.exists(row["doc_preguntas_ruta"]):
        return jsonify({"error": "Archivo de preguntas no encontrado"}), 404

    # Crear ZIP en memoria
    zip_buffer = io.BytesIO()
    with ZipFile(zip_buffer, 'w') as zipf:
        zipf.write(row["doc_preguntas_ruta"], row["doc_preguntas_nombre"])
        if row["doc_sol_ruta"] and os.path.exists(row["doc_sol_ruta"]):
            zipf.write(row["doc_sol_ruta"], row["doc_sol_nombre"])

    zip_buffer.seek(0)

    nombre_zip = f"{row['doc_preguntas_nombre'].split('.')[0]}_paquete.zip"
    return send_file(zip_buffer, as_attachment=True,
                     download_name=nombre_zip,
                     mimetype="application/zip")
 





# ==============================
# MATRIZ (MySQL)
# ==============================

# POST /api/matriz  -> crea cabecera + detalle
# Body JSON:
#   { "nombre":"Matriz X",
#     "items":[ { "tema_id":1, "tema_nombre":"BIOLOGÍA", "cantidad":10 }, ... ] }
@app.route("/api/matriz", methods=["POST"])
def matriz_crear_db():
    data = request.get_json(force=True) or {}
    nombre = (data.get("nombre") or "Matriz").strip()
    items  = data.get("items") or []
    if not items:
        return jsonify({"error": "items vacío"}), 400

    # valida
    norm = []
    for it in items:
        try:
            tema_id = int(it.get("tema_id"))
            cantidad = int(it.get("cantidad") or 0)
        except Exception:
            return jsonify({"error": "items inválidos"}), 400
        if tema_id <= 0:
            return jsonify({"error": "tema_id inválido"}), 400
        norm.append((tema_id, max(0, cantidad)))

    try:
        conn = get_connection(); cur = conn.cursor()
        cur.execute("INSERT INTO matriz (nombre) VALUES (%s)", (nombre,))
        matriz_id = cur.lastrowid

        # inserta detalle
        for tema_id, cantidad in norm:
            cur.execute(
                "INSERT INTO matriz_detalle (matriz_id, tema_id, cantidad) VALUES (%s,%s,%s) "
                "ON DUPLICATE KEY UPDATE cantidad=VALUES(cantidad)",
                (matriz_id, tema_id, cantidad)
            )

        conn.commit()
        cur.close(); conn.close()
        # crea carpeta de uploads para los docx de esta matriz
        os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], "matrices", str(matriz_id)), exist_ok=True)
        return jsonify({"ok": True, "matriz_id": matriz_id})
    except Exception as e:
        try: conn.rollback()
        except: pass
        return jsonify({"error": str(e)}), 500


# GET /api/matriz?detail=1&search=texto  -> lista matrices
@app.route("/api/matriz", methods=["GET"])
def matriz_listar_db():
    detail = request.args.get("detail") in ("1","true","True")
    search = (request.args.get("search") or "").strip()

    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)

        base_sql = """
            SELECT m.id, m.nombre, m.fecha_creacion,
                   COUNT(md.id)                   AS n_items,
                   SUM(CASE WHEN md.archivo_ruta IS NOT NULL AND md.archivo_ruta <> '' THEN 1 ELSE 0 END) AS n_archivos_subidos
            FROM matriz m
            LEFT JOIN matriz_detalle md ON md.matriz_id = m.id
        """
        where = []; params = []
        if search:
            where.append("(m.nombre LIKE %s OR m.id = %s)")
            params.extend([f"%{search}%", search if search.isdigit() else 0])
        if where:
            base_sql += " WHERE " + " AND ".join(where)
        base_sql += " GROUP BY m.id ORDER BY m.id DESC"
        cur.execute(base_sql, tuple(params))
        rows = cur.fetchall()

        if detail and rows:
            ids = [r["id"] for r in rows]
            fmt = ",".join(["%s"]*len(ids))
            cur.execute(f"""
                SELECT md.id, md.matriz_id, md.tema_id, t.nombre AS tema_nombre, md.cantidad, md.archivo_ruta
                FROM matriz_detalle md
                JOIN temas t ON t.id = md.tema_id
                WHERE md.matriz_id IN ({fmt})
                ORDER BY md.matriz_id, t.nombre
            """, tuple(ids))
            dets = cur.fetchall()
            by_m = {}
            for d in dets:
                by_m.setdefault(d["matriz_id"], []).append(d)
            for r in rows:
                r["items"] = by_m.get(r["id"], [])

        cur.close(); conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# GET /api/matriz/<id>  -> cabecera + detalle
@app.route("/api/matriz/<int:matriz_id>", methods=["GET"])
def matriz_get_db(matriz_id:int):
    try:
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("SELECT id, nombre, fecha_creacion FROM matriz WHERE id=%s", (matriz_id,))
        head = cur.fetchone()
        if not head:
            cur.close(); conn.close()
            return jsonify({"error": "Matriz no existe"}), 404

        cur.execute("""
            SELECT md.id, md.tema_id, t.nombre AS tema_nombre, md.cantidad, md.archivo_ruta
            FROM matriz_detalle md
            JOIN temas t ON t.id = md.tema_id
            WHERE md.matriz_id = %s
            ORDER BY t.nombre
        """, (matriz_id,))
        items = cur.fetchall()
        cur.close(); conn.close()

        head["n_items"] = len(items)
        head["n_archivos_subidos"] = sum(1 for x in items if x["archivo_ruta"])
        head["items"] = items
        return jsonify(head)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# POST /api/matriz/<id>/upload  (form-data: file, tema_id, [cantidad])
@app.route("/api/matriz/<int:matriz_id>/upload", methods=["POST"])
def matriz_upload_db(matriz_id:int):
    f = request.files.get("file")
    tema_id = request.form.get("tema_id", type=int)
    cantidad = request.form.get("cantidad", type=int)

    if not f or not tema_id:
        return jsonify({"error": "Faltan parámetros"}), 400
    if os.path.splitext(f.filename.lower())[1] != ".docx":
        return jsonify({"error": "Solo .docx"}), 400

    # guarda archivo
    dest_dir = os.path.join(app.config['UPLOAD_FOLDER'], "matrices", str(matriz_id))
    os.makedirs(dest_dir, exist_ok=True)
    filename = f"tema_{tema_id}.docx"
    dest_path = os.path.abspath(os.path.join(dest_dir, filename))
    f.save(dest_path)

    try:
        conn = get_connection(); cur = conn.cursor()

        # verifica que el detalle exista
        cur.execute("SELECT id FROM matriz_detalle WHERE matriz_id=%s AND tema_id=%s", (matriz_id, tema_id))
        row = cur.fetchone()
        if not row:
            # si no existe, crea la fila con cantidad 0
            cur.execute(
                "INSERT INTO matriz_detalle (matriz_id, tema_id, cantidad, archivo_ruta) VALUES (%s,%s,%s,%s)",
                (matriz_id, tema_id, int(cantidad or 0), dest_path)
            )
        else:
            if cantidad is not None:
                cur.execute(
                    "UPDATE matriz_detalle SET archivo_ruta=%s, cantidad=%s WHERE matriz_id=%s AND tema_id=%s",
                    (dest_path, int(cantidad or 0), matriz_id, tema_id)
                )
            else:
                cur.execute(
                    "UPDATE matriz_detalle SET archivo_ruta=%s WHERE matriz_id=%s AND tema_id=%s",
                    (dest_path, matriz_id, tema_id)
                )

        conn.commit()
        cur.close(); conn.close()
        return jsonify({"ok": True, "ruta": dest_path})
    except Exception as e:
        try: conn.rollback()
        except: pass
        return jsonify({"error": str(e)}), 500


# POST /api/matriz/<id>/generar  -> une por tema (título + contenido), respetando 'cantidad'
# POST /api/matriz/<id>/generar  -> une por tema (título + contenido), respetando 'cantidad'
@app.route("/api/matriz/<int:matriz_id>/generar", methods=["POST"])
def matriz_generar_db(matriz_id:int):
    try:
        # --- Lee cabecera + detalle ---
        conn = get_connection(); cur = conn.cursor(dictionary=True)
        cur.execute("SELECT id, nombre FROM matriz WHERE id=%s", (matriz_id,))
        head = cur.fetchone()
        if not head:
            cur.close(); conn.close()
            return jsonify({"error": "Matriz no existe"}), 404

        cur.execute("""
            SELECT md.tema_id, t.nombre AS tema_nombre, md.cantidad, md.archivo_ruta
            FROM matriz_detalle md
            JOIN temas t ON t.id = md.tema_id
            WHERE md.matriz_id=%s
            ORDER BY t.nombre
        """, (matriz_id,))
        dets = cur.fetchall()
        cur.close(); conn.close()
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    # --- Archivos faltantes ---
    faltan_arch = [d for d in dets if not d["archivo_ruta"] or not os.path.isfile(d["archivo_ruta"])]
    if faltan_arch:
        nombres = ", ".join(d["tema_nombre"] for d in faltan_arch)
        return jsonify({"error": f"Falta subir DOCX para: {nombres}"}), 400

    # ========== VALIDACIÓN DE CUPO (ANTES DE RECORTAR/UNIR) ==========
    faltantes = []
    for d in dets:
        pedidas = int(d["cantidad"] or 0)
        if pedidas <= 0:
            continue
        detectadas = _contar_preguntas_docx(d["archivo_ruta"])
        if detectadas < pedidas:
            faltantes.append({
                "tema_id": d["tema_id"],
                "tema": d["tema_nombre"],
                "pedidas": pedidas,
                "detectadas": detectadas
            })

    if faltantes:
        return jsonify({
            "ok": False,
            "error": "No hay suficientes preguntas numeradas (nivel 0, decimal) en algunos DOCX.",
            "faltantes": faltantes
        }), 409
    # ================================================================

    # --- Nombre destino ---
    ts = time.strftime("%Y-%m-%d %H-%M")
    base_name = f"Matriz {head['nombre']} - {ts}.docx"
    out_path  = os.path.join(app.config['DESCARGAS_FOLDER'], base_name)

    # --- Recorte y merge ---
    temp_files = []
    grouped = []  # [(titulo, [paths])]
    try:
        for d in dets:
            tema = d["tema_nombre"]
            src  = os.path.abspath(d["archivo_ruta"])
            cant = int(d.get("cantidad") or 0)

            if cant > 0:
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                tmp.close()
                _cut_docx_first_n_questions(src, cant, tmp.name)
                temp_files.append(tmp.name)
                use_path = tmp.name
            else:
                use_path = src

            grouped.append((tema, [use_path]))

        # Une con título por tema
        out, _, _ = _merge_grouped_with_headings(grouped, out_path)

        # Normalizaciones opcionales
        try: _post_merge_fix_numbering(out)
        except Exception: pass
        try: bullets_to_numbers_docx(out)
        except Exception: pass
        try: reparar_docx_inplace(out)
        except Exception: pass

        return send_file(out, as_attachment=True, download_name=os.path.basename(out))

    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        for p in temp_files:
            try: os.remove(p)
            except Exception: pass


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

def _find_question_spans(doc_xml_path: str, numbering_xml_path: str, max_q: int):
    """
    Devuelve spans (start_idx, end_idx) de párrafos que forman las primeras
    'max_q' preguntas numeradas a nivel 0 con formato decimal.
    """
    # Mapa (numId, ilvl) -> numFmt
    num_fmt_map = {}
    if os.path.exists(numbering_xml_path):
        tree_num = ET.parse(numbering_xml_path)
        root_num = tree_num.getroot()
        for num in root_num.findall(f".//{W_NS}num"):
            numId = num.get(f"{W_NS}numId")
            absN = num.find(f"./{W_NS}abstractNumId")
            if absN is None:
                continue
            absId = absN.get(f"{W_NS}val")
            absNode = root_num.find(f".//{W_NS}abstractNum[@{W_NS}abstractNumId='{absId}']")
            if absNode is None:
                continue
            for lvl in absNode.findall(f"./{W_NS}lvl"):
                ilvl = lvl.get(f"{W_NS}ilvl")
                nfmt = lvl.find(f"./{W_NS}numFmt")
                if nfmt is not None:
                    num_fmt_map[(numId, ilvl)] = nfmt.get(f"{W_NS}val")

    # document.xml
    tree_doc = ET.parse(doc_xml_path)
    root_doc = tree_doc.getroot()
    paras = root_doc.findall(f".//{W_NS}p")

    def is_q_start(p):
        numPr = p.find(f"./{W_NS}pPr/{W_NS}numPr")
        if numPr is None:
            return False
        numId_el = numPr.find(f"./{W_NS}numId")
        ilvl_el  = numPr.find(f"./{W_NS}ilvl")
        numId = (numId_el.get(f"{W_NS}val") if numId_el is not None else None)
        ilvl  = (ilvl_el.get(f"{W_NS}val") if ilvl_el is not None else None)
        fmt = num_fmt_map.get((numId, ilvl), None)
        return (fmt == "decimal" and (ilvl or "0") == "0")

    starts = [i for i, p in enumerate(paras) if is_q_start(p)]
    if not starts:
        return []

    spans = []
    for idx, s in enumerate(starts):
        e = (starts[idx+1] if idx+1 < len(starts) else len(paras))
        spans.append((s, e))

    if max_q > 0:
        spans = spans[:max_q]
    return spans

def _cut_docx_first_n_questions(src_docx: str, n: int, out_docx: str):
    """
    Crea un DOCX con solo las primeras n preguntas del src.
    Si no detecta preguntas, copia el original.
    """
    if n <= 0:
        d = DocxDocument()
        d.save(out_docx)
        return

    with tempfile.TemporaryDirectory() as td:
        with zipfile.ZipFile(src_docx, "r") as z:
            z.extractall(td)

        doc_xml = os.path.join(td, "word", "document.xml")
        num_xml = os.path.join(td, "word", "numbering.xml")

        spans = _find_question_spans(doc_xml, num_xml, n)
        if not spans:
            shutil.copyfile(src_docx, out_docx)
            return

        tree_doc = ET.parse(doc_xml)
        root_doc = tree_doc.getroot()
        body = root_doc.find(f"{W_NS}body")
        all_p = list(body.findall(f"./{W_NS}p"))

        keep = set()
        for s, e in spans:
            keep.update(range(s, e))

        for idx, p in enumerate(all_p):
            if idx not in keep:
                body.remove(p)

        tree_doc.write(doc_xml, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(out_docx, "w", zipfile.ZIP_DEFLATED) as outz:
            for root, _dirs, files in os.walk(td):
                for f in files:
                    full = os.path.join(root, f)
                    rel  = os.path.relpath(full, td)
                    outz.write(full, rel)


_WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

def _contar_preguntas_docx(docx_path: str) -> int:
    """
    Devuelve cuántas 'preguntas' detecta (párrafos con numPr decimal ilvl=0).
    Si el DOCX está corrupto o no tiene numbering/document.xml -> 0.
    """
    try:
        if not zipfile.is_zipfile(docx_path):
            return 0

        with zipfile.ZipFile(docx_path, "r") as zin:
            names = set(zin.namelist())
            if "word/document.xml" not in names:
                return 0
            doc_xml = zin.read("word/document.xml")
            numbering_xml = zin.read("word/numbering.xml") if "word/numbering.xml" in names else None

        num_fmt_map = {}
        if numbering_xml:
            try:
                rnum = ET.fromstring(numbering_xml)
                for num in rnum.findall(f".//{_WNS}num"):
                    numId = num.get(f"{_WNS}numId")
                    absN = num.find(f"./{_WNS}abstractNumId")
                    if absN is None: 
                        continue
                    absId = absN.get(f"{_WNS}val")
                    absNode = rnum.find(f".//{_WNS}abstractNum[@{_WNS}abstractNumId='{absId}']")
                    if absNode is None:
                        continue
                    for lvl in absNode.findall(f"./{_WNS}lvl"):
                        ilvl = lvl.get(f"{_WNS}ilvl")
                        nfmt = lvl.find(f"./{_WNS}numFmt")
                        if nfmt is not None:
                            num_fmt_map[(numId, ilvl)] = nfmt.get(f"{_WNS}val")
            except Exception:
                return 0

        try:
            rdoc = ET.fromstring(doc_xml)
        except Exception:
            return 0

        paras = rdoc.findall(f".//{_WNS}p")

        def is_q_start(p):
            numPr = p.find(f"./{_WNS}pPr/{_WNS}numPr")
            if numPr is None:
                return False
            numId_el = numPr.find(f"./{_WNS}numId")
            ilvl_el  = numPr.find(f"./{_WNS}ilvl")
            numId = numId_el.get(f"{_WNS}val") if numId_el is not None else None
            ilvl  = ilvl_el.get(f"{_WNS}val")  if ilvl_el is not None else None
            fmt = num_fmt_map.get((numId, ilvl), None)
            return (fmt == "decimal" and (ilvl or "0") == "0")

        return sum(1 for p in paras if is_q_start(p))
    except Exception:
        return 0

# ---------------------------
# CRUD TEMAS PARA CUADERNILLOS (sin conteo de preguntas)
# ---------------------------

@app.route("/api/temas_cuad", methods=["GET"])
def temas_listar_cuad():
    """
    Versión básica para CUADERNILLOS:
    - Devuelve solo id, nombre, activo
    - ?all=1 para incluir inactivos, por defecto solo activos
    """
    include_all = request.args.get("all") == "1"
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)

        if include_all:
            cur.execute("""
                SELECT t.id, t.nombre, t.activo
                FROM temas t
                ORDER BY t.nombre
            """)
        else:
            cur.execute("""
                SELECT t.id, t.nombre, t.activo
                FROM temas t
                WHERE t.activo = 1
                ORDER BY t.nombre
            """)

        rows = cur.fetchall()
        cur.close()
        conn.close()
        return jsonify(rows)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------------------------
# GRUPOS – generar por matriz
# ---------------------------

# Listar matrices para "Importar matriz"
@app.route("/api/matrices", methods=["GET"])
def api_list_matrices():
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("""
        SELECT id, nombre, fecha_creacion
        FROM matriz
        ORDER BY id DESC
    """)
    data = _row_to_dict_list(cur)
    cur.close()
    conn.close()
    return jsonify(data), 200


def _leer_matriz_detalle(cur, matriz_id: int):
    """
    Lee el detalle de la matriz desde BD.
    Devuelve: [{'tema_id', 'tema_nombre', 'cantidad', 'archivo_ruta'}, ...]
    """
    cur.execute(
        """
        SELECT d.tema_id,
               t.nombre AS tema_nombre,
               d.cantidad,
               d.archivo_ruta
        FROM matriz_detalle d
        JOIN temas t ON t.id = d.tema_id
        WHERE d.matriz_id = %s
        ORDER BY t.nombre          -- orden alfabético de la matriz
        """,
        (matriz_id,),
    )
    return _row_to_dict_list(cur)


def _leer_config_grupos(cur):
    """
    Devuelve:
      {
        idgrupo: {
          idgrupo, clave, nombre,
          temas: [ {tema_id, tema_nombre, cantidad}, ... ]
                 # en el MISMO orden que ves en el modal
        },
        ...
      }
    """
    cur.execute("""
        SELECT g.idgrupo,
               g.clave,
               g.nombre,
               g.activo,
               gt.tema_id,
               IFNULL(gt.cantidad,0) AS cant_tema,
               t.nombre AS tema_nombre
        FROM grupos g
        LEFT JOIN grupo_tema gt
               ON gt.grupos_idgrupo = g.idgrupo
        LEFT JOIN temas t
               ON t.id = gt.tema_id
        WHERE g.activo = 1
        ORDER BY g.idgrupo, t.nombre   -- 👈 mismo orden que el modal (por nombre de tema)
    """)
    rows = _row_to_dict_list(cur)

    grupos = {}
    for r in rows:
        gid = r["idgrupo"]
        if gid not in grupos:
            grupos[gid] = {
                "idgrupo": gid,
                "clave": r["clave"],
                "nombre": r["nombre"],
                "temas": []   # se llenan ya ordenados por t.nombre
            }
        if r["tema_id"] is not None:
            grupos[gid]["temas"].append({
                "tema_id": r["tema_id"],
                "tema_nombre": r["tema_nombre"],
                "cantidad": int(r["cant_tema"] or 0)
            })
    return grupos


def _preguntas_por_tema(cur):
    cur.execute(
        """
        SELECT idpreguntas, tema_id, enunciado,
               COALESCE(alternativa_a,'') AS alternativa_a,
               COALESCE(alternativa_b,'') AS alternativa_b,
               COALESCE(alternativa_c,'') AS alternativa_c,
               COALESCE(alternativa_d,'') AS alternativa_d
        FROM preguntas
        """
    )
    rows = _row_to_dict_list(cur)
    by_tema = {}
    for r in rows:
        by_tema.setdefault(r["tema_id"], []).append(r)
    return by_tema


def _armar_docx_grupo(nombre_grupo: str, clave: str, bloques: list, ruta: str):
    """
    (Si luego no usas esto, lo puedes borrar, pero lo dejo por compatibilidad)
    bloques: [ { 'titulo': 'ARITMÉTICA', 'pregs': [ {...}, ... ] }, ... ]
    """
    doc = Document()
    doc.add_heading(f"Examen - Grupo {clave} ({nombre_grupo})", level=1)
    orden = 1
    for b in bloques:
        doc.add_heading(b["titulo"], level=2)
        for p in b["pregs"]:
            doc.add_paragraph(f"{orden}. {p['enunciado']}")
            a = p.get("alternativa_a")
            b1 = p.get("alternativa_b")
            c = p.get("alternativa_c")
            d = p.get("alternativa_d")
            if any([a, b1, c, d]):
                if a:
                    doc.add_paragraph(f"A) {a}")
                if b1:
                    doc.add_paragraph(f"B) {b1}")
                if c:
                    doc.add_paragraph(f"C) {c}")
                if d:
                    doc.add_paragraph(f"D) {d}")
            doc.add_paragraph("")  # espacio
            orden += 1
    os.makedirs(os.path.dirname(ruta), exist_ok=True)
    doc.save(ruta)


@app.route("/api/grupos/generar", methods=["POST"])
def api_generar_por_grupos():
    data = request.get_json(force=True) or {}
    matriz_id = data.get("matriz_id")
    matriz_inline = data.get("matriz")  # opcional, para futuro

    conn = get_connection()
    cur = conn.cursor()   # cursor normal; _row_to_dict_list lo convierte

    try:
        # -------------------------
        # 1) Construir detalle de matriz (md)
        # -------------------------
        if matriz_id:
            md = _leer_matriz_detalle(cur, matriz_id)
            if not md:
                return jsonify({"error": "La matriz no tiene detalle."}), 400
            nombre_matriz = f"Matriz {matriz_id}"
            matriz_id_for_lote = matriz_id

        elif matriz_inline:
            items = matriz_inline.get("items") or []
            if (
                not isinstance(items, list)
                or not items
                or not all(isinstance(it, dict) and "tema_id" in it for it in items)
            ):
                return jsonify(
                    {"error": "Formato de matriz inválido: items[] incorrectos"}
                ), 400

            tema_ids = tuple(
                {int(it.get("tema_id", 0)) for it in items if it.get("tema_id")}
            )
            nombres = {}
            if tema_ids:
                cur.execute(
                    f"SELECT id, nombre FROM temas WHERE id IN ({','.join(['%s']*len(tema_ids))})",
                    tema_ids,
                )
                for tid, nom in cur.fetchall():
                    nombres[int(tid)] = nom

            md = []
            for it in items:
                tema_id = int(it.get("tema_id") or 0)
                cant = int(it.get("cantidad") or 0)
                if tema_id <= 0:
                    return jsonify(
                        {"error": "Cada item debe incluir tema_id válido"}
                    ), 400
                md.append(
                    {
                        "tema_id": tema_id,
                        "tema_nombre": nombres.get(tema_id, f"Tema {tema_id}"),
                        "cantidad": cant,
                        "archivo_ruta": it.get("archivo_ruta"),
                    }
                )

            nombre_matriz = (matriz_inline.get("nombre") or "Matriz importada").strip()
            matriz_id_for_lote = None
        else:
            return jsonify({"error": "Envía matriz_id o matriz (JSON)"}), 400

        # índice rápido por tema
        md_por_tema = {int(r["tema_id"]): r for r in md}

        # -------------------------
        # 2) Config de grupos
        # -------------------------
        grupos_cfg = _leer_config_grupos(cur)
        if not grupos_cfg:
            return jsonify({"error": "No hay grupos activos configurados."}), 400

        cuotas_por_tema = {}
        for g in grupos_cfg.values():
            for rel in g["temas"]:
                t_id = int(rel["tema_id"])
                q = int(rel["cantidad"] or 0)
                if q <= 0:
                    continue
                if t_id not in md_por_tema:
                    return jsonify(
                        {"error": f"El tema ID {t_id} está en un grupo pero no está en la matriz."}
                    ), 400
                cuotas_por_tema.setdefault(t_id, set()).add(q)

        if not cuotas_por_tema:
            return jsonify({"error": "No hay cuotas definidas para los grupos."}), 400

        # Validar cuotas y archivos
        for t_id, qs in cuotas_por_tema.items():
            if len(qs) > 1:
                tema_nombre = md_por_tema[t_id]["tema_nombre"]
                return jsonify(
                    {
                        "error": f"Las cuotas del tema '{tema_nombre}' no son iguales en todos los grupos. "
                                 f"Para reutilizar las mismas preguntas, la cantidad debe ser la misma."
                    }
                ), 400

            cuota = next(iter(qs))
            cant_matriz = int(md_por_tema[t_id].get("cantidad") or 0)
            if cuota > cant_matriz:
                tema_nombre = md_por_tema[t_id]["tema_nombre"]
                return jsonify(
                    {
                        "error": f"El tema '{tema_nombre}' pide {cuota} preguntas por grupo, "
                                 f"pero la matriz solo tiene {cant_matriz}."
                    }
                ), 400

            if not md_por_tema[t_id].get("archivo_ruta"):
                tema_nombre = md_por_tema[t_id]["tema_nombre"]
                return jsonify(
                    {"error": f"El tema '{tema_nombre}' no tiene DOCX asociado en la matriz."}
                ), 400

        # -------------------------
        # 3) Crear lote
        # -------------------------
        cur.execute(
            """
            INSERT INTO gen_lote(matriz_id, nombre, usuario)
            VALUES (%s, %s, %s)
            """,
            (
                matriz_id_for_lote,
                f"Lote {dt.datetime.now():%Y-%m-%d %H:%M} - {nombre_matriz}",
                "system",
            ),
        )
        conn.commit()
        lote_id = cur.lastrowid

        # -------------------------
        # 4) Recortes DOCX por tema
        # -------------------------
        temp_files = []
        docx_por_tema = {}

        for t_id, qs in cuotas_por_tema.items():
            cuota = next(iter(qs))
            info = md_por_tema[t_id]
            src = info["archivo_ruta"]
            titulo = info["tema_nombre"]

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp.close()
            _cut_docx_first_n_questions(src, cuota, tmp.name)
            temp_files.append(tmp.name)
            docx_por_tema[t_id] = {"path": tmp.name, "titulo": titulo, "n": cuota}

        # -------------------------
        # 5) Generar DOCX por grupo
        # -------------------------
        lote_dir = os.path.join(GRUPOS_OUT_DIR, f"lote_{lote_id}")
        os.makedirs(lote_dir, exist_ok=True)

        docxs_generados = []

        for gid, ginfo in grupos_cfg.items():
            # IMPORTANTE: aquí ya vienen en el mismo orden que el modal,
            # porque _leer_config_grupos ordena por t.nombre
            bloques = []
            for rel in ginfo["temas"]:
                t_id = int(rel["tema_id"])
                cant = int(rel["cantidad"] or 0)
                if cant <= 0:
                    continue
                tema_doc = docx_por_tema.get(t_id)
                if not tema_doc:
                    continue
                bloques.append((tema_doc["titulo"], [tema_doc["path"]]))

            if not bloques:
                continue

            clave = (ginfo.get("clave") or f"G{gid}").strip()
            nom_grupo = (ginfo.get("nombre") or "").strip() or f"Grupo {clave}"
            out_path = os.path.join(lote_dir, f"grupo_{clave}.docx")

            out_docx, _, _ = _merge_grouped_with_headings(bloques, out_path)

            try:
                _post_merge_fix_numbering(out_docx)
            except Exception:
                pass
            try:
                bullets_to_numbers_docx(out_docx)
            except Exception:
                pass
            try:
                reparar_docx_inplace(out_docx)
            except Exception:
                pass

            docxs_generados.append(out_docx)

        if not docxs_generados:
            return jsonify({"error": "No se generó ningún examen de grupo."}), 400

        # -------------------------
        # 6) ZIP
        # -------------------------
        zip_path = os.path.join(lote_dir, f"lote_{lote_id}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for p in docxs_generados:
                zf.write(p, os.path.basename(p))

        return jsonify(
            {"ok": True, "lote_id": lote_id, "zip_url": f"/api/grupos/lote/{lote_id}/zip"}
        )

    except Exception as e:
        try:
            conn.rollback()
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            cur.close()
            conn.close()
        except Exception:
            pass
        for p in locals().get("temp_files", []):
            try:
                os.remove(p)
            except Exception:
                pass


@app.route("/api/grupos/lote/<int:lote_id>/zip", methods=["GET"])
def api_descargar_zip_lote(lote_id):
    lote_dir = os.path.join(GRUPOS_OUT_DIR, f"lote_{lote_id}")
    zip_path = os.path.join(lote_dir, f"lote_{lote_id}.zip")
    if not os.path.isfile(zip_path):
        return jsonify({"error": "ZIP no encontrado"}), 404
    return send_file(zip_path, as_attachment=True,
                     download_name=f"lote_{lote_id}.zip")

# ----------------- Utils -----------------
def sha256sum(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""): h.update(chunk)
    return h.hexdigest()

LTRS = ["A", "B", "C", "D", "E"]

def pick_not_in(forbidden):
    forbidden = set((x or "").upper() for x in forbidden)
    pool = [l for l in LTRS if l not in forbidden]
    return random.choice(pool) if pool else "A"

ALT_RX = re.compile(r"^\s*\(?([A-Ea-e])\s*[\)\.\-]\s+(.*)$")
QSTART_RX = re.compile(r"^\s*\(?\d{1,3}\)?[.)]\s+")


# ----------------- API -------------------

# Subir exámenes (.doc | .docx | .pdf)
# ========= Helpers para matriz DOCX (grupos) =========

def _norm_name(s: str) -> str:
    """Normaliza nombres de temas: mayúsculas, sin tildes, colapsando espacios."""
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = " ".join(s.upper().split())
    return s








def _build_heading_ranges_for_temas(docx_path: str, nombres_objetivo_norm: set):
    """
    Busca en el DOCX los párrafos cuyo texto coincide con los nombres de temas
    (normalizados) y devuelve rangos de párrafos por tema:

      {nombre_norm: (start_idx, end_idx), ...}

    donde start/end son índices de párrafos.
    """
    doc = Document(docx_path)
    indices = []

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue

        # Normalizamos el texto del párrafo
        n = _norm_name(text)

        # Si el texto normalizado coincide con alguno de los nombres de tema,
        # lo tomamos como "título" SIN exigir estilos especiales.
        if n in nombres_objetivo_norm:
            indices.append((n, i))

    # Ordenar por posición
    indices.sort(key=lambda x: x[1])

    ranges = {}
    total = len(doc.paragraphs)
    for idx, (n, start) in enumerate(indices):
        end = indices[idx + 1][1] if idx + 1 < len(indices) else total
        ranges[n] = (start, end)

    return ranges



def _extract_tema_docx_range(src_path: str, start_idx: int, end_idx: int, dest_path: str):
    """
    Crea un DOCX con solo los párrafos del rango [start_idx, end_idx)
    preservando numeración y estilos (se parte del DOCX original y se
    eliminan los párrafos que no interesan).
    """
    doc = Document(src_path)
    # Eliminamos TODO lo que no esté en el rango, recorriendo hacia atrás
    for i in reversed(range(len(doc.paragraphs))):
        if not (start_idx <= i < end_idx):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)
    doc.save(dest_path)


# ========= Endpoint: generar grupos desde matriz DOCX =========

@app.route("/api/grupos/generar_from_docx", methods=["POST"])
def api_generar_grupos_from_docx():
    """
    Recibe un archivo DOCX de 'matriz' (títulos = temas),
    separa por tema y genera exámenes por grupo igual que /api/grupos/generar.
    """
    
    if "file" not in request.files:
        return jsonify({"error": "No se envió ningún archivo DOCX."}), 400

    file = request.files["file"]
    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "El archivo debe ser .docx"}), 400

    # Guardar DOCX temporalmente
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    file.save(tmp_in.name)
    tmp_in.close()

    conn = get_connection()
    cur = conn.cursor()

    temp_files = [tmp_in.name]  # para limpieza al final

    try:
        # 1) Configuración de grupos (ya incluye tema_nombre y cantidades)
        grupos_cfg = _leer_config_grupos(cur)
        if not grupos_cfg:
            return jsonify({"error": "No hay grupos activos configurados."}), 400

        # 2) Construir cuotas por tema y recoger nombres de temas
        cuotas_por_tema = {}         # tema_id -> set(cantidades)
        tema_info = {}               # tema_id -> {"nombre": ..., "norm": ...}

        for g in grupos_cfg.values():
            for rel in g["temas"]:
                t_id = int(rel["tema_id"])
                q = int(rel["cantidad"] or 0)
                if q <= 0:
                    continue
                cuotas_por_tema.setdefault(t_id, set()).add(q)
                if t_id not in tema_info:
                    nom = rel.get("tema_nombre") or ""
                    tema_info[t_id] = {
                        "nombre": nom,
                        "norm": _norm_name(nom),
                    }

        if not cuotas_por_tema:
            return jsonify({"error": "No hay cuotas definidas para los grupos."}), 400

        # Validar que por tema todas las cuotas sean iguales
        for t_id, qs in cuotas_por_tema.items():
            if len(qs) > 1:
                nom = tema_info.get(t_id, {}).get("nombre", f"Tema {t_id}")
                return jsonify({
                    "error": f"Las cuotas del tema '{nom}' no son iguales en todos los grupos."
                }), 400

                # 3) Detectar títulos en el DOCX que correspondan a esos temas
        nombres_norm = {info["norm"] for info in tema_info.values()}
        ranges = _build_heading_ranges_for_temas(tmp_in.name, nombres_norm)

        # Validar que TODOS los temas configurados tienen un rango en el DOCX
        faltan = []
        for t_id, info in tema_info.items():
            nom_norm = info["norm"]
            if nom_norm not in ranges:
                faltan.append(info["nombre"] or f"Tema {t_id}")

        if faltan:
            return jsonify({
                "error": (
                    "En el DOCX de la matriz faltan títulos para algunos temas configurados:\n"
                    + ", ".join(sorted(faltan)) +
                    "\n\nVerifica que en el DOCX exista un párrafo cuyo texto sea exactamente "
                    "el nombre del curso (ignorando mayúsculas, tildes y espacios)."
                )
            }), 400



        # 4) Crear lote (sin matriz_id porque viene de DOCX suelto)
        cur.execute(
            """
            INSERT INTO gen_lote(matriz_id, nombre, usuario)
            VALUES (%s, %s, %s)
            """,
            (
                None,
                f"Lote {dt.datetime.now():%Y-%m-%d %H:%M} - Matriz DOCX",
                "system",
            ),
        )
        conn.commit()
        lote_id = cur.lastrowid

        # 5) Generar DOCX por tema (recortando rango de párrafos)
        docx_por_tema = {}  # tema_id -> {"path":..., "titulo":..., "n": cuota}

        for t_id, qs in cuotas_por_tema.items():
            cuota = next(iter(qs))
            info = tema_info[t_id]
            nom_norm = info["norm"]
            nom_real = info["nombre"] or f"Tema {t_id}"
            start_idx, end_idx = ranges[nom_norm]

            # 👇 Saltamos el párrafo del título para que NO se duplique
            start_preg = min(end_idx, start_idx + 1)

            tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp_out.close()
            _extract_tema_docx_range(tmp_in.name, start_preg, end_idx, tmp_out.name)

            temp_files.append(tmp_out.name)
            docx_por_tema[t_id] = {
                "path": tmp_out.name,
                "titulo": nom_real,
                "n": cuota,
            }


            

        # 6) Generar DOCX por grupo (mismo esquema que /api/grupos/generar)
        lote_dir = os.path.join(GRUPOS_OUT_DIR, f"lote_{lote_id}")
        os.makedirs(lote_dir, exist_ok=True)

        docxs_generados = []

        for gid, ginfo in grupos_cfg.items():
            bloques = []
            for rel in ginfo["temas"]:
                t_id = int(rel["tema_id"])
                cant = int(rel["cantidad"] or 0)
                if cant <= 0:
                    continue
                tema_doc = docx_por_tema.get(t_id)
                if not tema_doc:
                    continue
                bloques.append((tema_doc["titulo"], [tema_doc["path"]]))

            if not bloques:
                continue

            clave = (ginfo.get("clave") or f"G{gid}").strip()
            nom_grupo = (ginfo.get("nombre") or "").strip() or f"Grupo {clave}"
            out_path = os.path.join(lote_dir, f"grupo_{clave}.docx")

            out_docx, _, _ = _merge_grouped_with_headings(bloques, out_path)

            try:
                _post_merge_fix_numbering(out_docx)
            except Exception:
                pass
            try:
                bullets_to_numbers_docx(out_docx)
            except Exception:
                pass
            try:
                reparar_docx_inplace(out_docx)
            except Exception:
                pass

            docxs_generados.append(out_docx)

        if not docxs_generados:
            return jsonify({"error": "No se generó ningún examen de grupo."}), 400

        # 7) ZIP final
        zip_path = os.path.join(lote_dir, f"lote_{lote_id}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for p in docxs_generados:
                zf.write(p, os.path.basename(p))

        return jsonify({
            "ok": True,
            "lote_id": lote_id,
            "zip_url": f"/api/grupos/lote/{lote_id}/zip"
        })

    except Exception as e:
        try:
            conn.rollback()
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        try:
            cur.close()
            conn.close()
        except Exception:
            pass
        for p in temp_files:
            try:
                os.remove(p)
            except Exception:
                pass

#
#examenes de prueba
#

import random
from copy import deepcopy
from docx import Document
VALID_LETTERS = ["A","B","C","D","E"]

def pick_two_distinct_letters():
    """Devuelve (p, q) aleatorios en A–E, p != q."""
    p, q = random.sample(VALID_LETTERS, 2)
    return p, q


# por si algunas preguntas vienen como texto "1." "2)" etc.
QSTART_RX = re.compile(r"^\s*\d{1,3}\s*[\)\.]\s+")

def _get_numid_ilvl(p):
    """Devuelve (numId, ilvl) del párrafo si existe numeración Word."""
    try:
        numPr = p._p.pPr.numPr if p._p.pPr is not None else None
        if numPr is None:
            return (None, None)
        numId = numPr.numId.val if numPr.numId is not None else None
        ilvl  = numPr.ilvl.val if numPr.ilvl is not None else None
        return (str(numId) if numId is not None else None,
                int(ilvl) if ilvl is not None else None)
    except Exception:
        return (None, None)

def _is_question_start_paragraph(p, active_q_numId="1"):
    """
    Inicio de pregunta:
    - numId=1, ilvl=0  (como en tu XML)
    - fallback por regex si no hay numPr
    """
    numId, ilvl = _get_numid_ilvl(p)
    if numId is not None and ilvl == 0 and numId == str(active_q_numId):
        return True

    txt = (p.text or "").strip()
    return QSTART_RX.match(txt) is not None

def _reorder_alt_paragraphs(alt_paras, target_letter):
    """
    alt_paras: lista de párrafos (w:p) de alternativas en orden original.
              Se asume que la alternativa correcta ORIGINAL es la primera.
    target_letter: "A".."E" donde debe caer la correcta.
    Retorna lista reordenada de párrafos.
    """
    letters = ["A","B","C","D","E"]
    target_letter = (target_letter or "A").upper()
    idx_t = letters.index(target_letter)

    correct_p = alt_paras[0]
    others = alt_paras[1:5]

    # ✅ aleatoriza TODAS las incorrectas
    random.shuffle(others)

    new_order = [None]*5
    new_order[idx_t] = correct_p

    j = 0
    for i in range(5):
        if new_order[i] is None:
            new_order[i] = others[j] if j < len(others) else None
            j += 1

    # filtra Nones si hubiera menos de 5
    return [p for p in new_order if p is not None]

def _apply_reorder(doc: Document, claves, modo="P"):
    """
    Reordena alternativas preservando formato (mueve XML).
    claves: lista dict con numero_pregunta, origen, p, q
    modo: "P" o "Q"
    """
    map_claves = {int(r["numero_pregunta"]): r for r in claves}

    paras = doc.paragraphs
    qnum = 0
    i = 0

    while i < len(paras):
        p = paras[i]

        # detectar pregunta
        if _is_question_start_paragraph(p, active_q_numId="1"):
            qnum += 1
            q_numId, _ = _get_numid_ilvl(p)  # normalmente "1"
            i += 1

            # recolectar alternativas por numId distinto al de pregunta
            alt_paras = []
            first_alt_numId = None
            start_alt_index = None

            while i < len(paras):
                p2 = paras[i]

                if _is_question_start_paragraph(p2, active_q_numId="1"):
                    break

                numId2, ilvl2 = _get_numid_ilvl(p2)

                # alternativa = lista nivel 0 con numId != numId de pregunta
                if numId2 is not None and ilvl2 == 0 and numId2 != q_numId:
                    if first_alt_numId is None:
                        first_alt_numId = numId2
                        start_alt_index = i

                    if numId2 == first_alt_numId:
                        alt_paras.append(p2)

                    # si numId cambia, ya terminó bloque de alternativas
                    elif first_alt_numId is not None:
                        break

                i += 1

            # aplicar reorder si hay bloque válido
            if len(alt_paras) >= 2:
                rkey = map_claves.get(qnum, {})
                target = rkey.get("p" if modo=="P" else "q") or "A"

                new_order = _reorder_alt_paragraphs(alt_paras[:5], target)

                # 🔥 mover XML completo preservando estilo
                parent = alt_paras[0]._p.getparent()
                insert_at = parent.index(alt_paras[0]._p)

                # insertar copias en nuevo orden
                for k, pp in enumerate(new_order):
                    parent.insert(insert_at + k, deepcopy(pp._p))

                # eliminar originales
                for pp in alt_paras[:5]:
                    parent.remove(pp._p)

            continue

        i += 1




def generar_docx_pq_para_grupo(ruta_docx, claves, clave_grupo, nombre_base):
    """
    Genera bytes docx P y Q para un grupo, reordenando alternativas
    según claves P/Q.
    """
    docP = Document(ruta_docx)
    docQ = Document(ruta_docx)

    _apply_reorder(docP, claves, modo="P")
    _apply_reorder(docQ, claves, modo="Q")

    bioP = io.BytesIO(); docP.save(bioP)
    bioQ = io.BytesIO(); docQ.save(bioQ)

    return bioP.getvalue(), bioQ.getvalue()


# Listado de grupos (para el select)
@app.route("/api/grupos", methods=["GET"])
def api_grupos():
    conn = get_connection(); cur = conn.cursor(dictionary=True)
    # Ajusta campos: idgrupo, clave, nombre, activo
    cur.execute("SELECT idgrupo AS id, clave, nombre FROM grupos WHERE activo=1 ORDER BY clave")
    data = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(data)



# --- DELETE /api/examenes/importados/<id> ---
@app.route("/api/examenes/importados/<int:id>", methods=["DELETE"])
def api_examen_importado_eliminar(id):
    """
    Elimina un examen importado:
    - Borra registro de BD
    - Opcional: borra archivo físico (si existe)
    """
    borrar_archivo = request.args.get("delete_file", "1") == "1"

    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)

        # 1) Obtener ruta antes de borrar
        cur.execute("SELECT ruta FROM examenes_importados WHERE id=%s", (id,))
        row = cur.fetchone()
        if not row:
            cur.close(); conn.close()
            return jsonify(ok=False, error="Examen no encontrado."), 404

        ruta = row.get("ruta")

        # 2) Borrar de BD
        cur.execute("DELETE FROM examenes_importados WHERE id=%s", (id,))
        conn.commit()

        cur.close(); conn.close()

        # 3) Borrar archivo físico si se pide
        if borrar_archivo and ruta and os.path.exists(ruta):
            try:
                os.remove(ruta)
            except Exception:
                # si falla borrar archivo, no rompemos la respuesta
                pass

        return jsonify(ok=True), 200

    except Exception as e:
        try:
            cur.close(); conn.close()
        except Exception:
            pass
        return jsonify(ok=False, error=str(e)), 500

@app.route("/api/claves/ensure", methods=["POST"])
def api_claves_ensure():
    data = request.get_json(force=True)
    examen_id = int(data.get("examen_id") or 0)
    grupo_id  = int(data.get("grupo_id") or 0)
    if not examen_id or not grupo_id:
        return jsonify(ok=False, error="Faltan examen_id/grupo_id"), 400

    conn = get_connection()
    cur = conn.cursor(dictionary=True)

    # total de preguntas del examen importado
    cur.execute("SELECT total_preguntas FROM examenes_importados WHERE id=%s", (examen_id,))
    ex = cur.fetchone()
    total = int(ex["total_preguntas"] or 0) if ex else 0
    if total <= 0:
        cur.close(); conn.close()
        return jsonify(ok=False, error="Examen sin preguntas detectadas"), 409

    # ¿ya existen claves?
    cur.execute("""
        SELECT COUNT(*) AS n
        FROM claves_respuesta
        WHERE examen_id=%s AND grupo_id=%s
    """, (examen_id, grupo_id))
    n = int(cur.fetchone()["n"])

    if n == 0:
        # crear todas las filas con P/Q válidas
        rows = []
        for i in range(total):
            origen = "A"
            p, q = pick_two_distinct_letters()   # ✅ ahora puede incluir "A"
            rows.append((i+1, origen, p, q, examen_id, grupo_id))


        cur.executemany("""
            INSERT INTO claves_respuesta
                (numero_pregunta, origen, p, q, examen_id, grupo_id, fecha_actualizacion)
            VALUES (%s,%s,%s,%s,%s,%s,NOW())
        """, rows)
        conn.commit()

    else:
        # si ya existen, repara cualquier p/q inválida
        cur.execute("""
            SELECT id, origen, p, q
            FROM claves_respuesta
            WHERE examen_id=%s AND grupo_id=%s
            ORDER BY numero_pregunta
        """, (examen_id, grupo_id))
        filas = cur.fetchall()

        valid = {"A","B","C","D","E"}
        for r in filas:
            origen = (r["origen"] or "A").upper()

            p = (r["p"] or "").upper()
            q = (r["q"] or "").upper()

            if p not in valid:
                p = random.choice(list(valid))
            if q not in valid or q == p:
                q = random.choice([x for x in valid if x != p])

            cur.execute("""
                UPDATE claves_respuesta
                SET p=%s, q=%s, fecha_actualizacion=NOW()
                WHERE id=%s
            """, (p, q, r["id"]))
        conn.commit()

    cur.close(); conn.close()
    return jsonify(ok=True, total=total)


@app.route("/api/claves/guardar", methods=["POST"])
def api_claves_guardar():
    data = request.get_json(force=True)
    examen_id = int(data.get("examen_id") or 0)
    grupo_id  = int(data.get("grupo_id")  or 0)
    filas     = data.get("filas") or []

    if not examen_id or not grupo_id:
        return jsonify(ok=False, error="Faltan examen_id/grupo_id"), 400

    conn = get_connection(); cur = conn.cursor(dictionary=True)

    for r in filas:
        num = int(r.get("numero_pregunta") or 0)
        if num <= 0: 
            continue
        origen = (r.get("origen") or "A").upper()
        valid = {"A","B","C","D","E"}

        p = (r.get("p") or "").upper()
        q = (r.get("q") or "").upper()

        if p not in valid:
            p = random.choice(valid)   # ✅ ahora sí puede salir A
        if q not in valid or q == p:
            q = random.choice([x for x in valid if x != p])



        cur.execute("""
            UPDATE claves_respuesta
            SET origen=%s, p=%s, q=%s, fecha_actualizacion=NOW()
            WHERE examen_id=%s AND grupo_id=%s AND numero_pregunta=%s
        """, (origen, p, q, examen_id, grupo_id, num))

    conn.commit()
    cur.close(); conn.close()
    return jsonify(ok=True)





# Obtener claves Origen/P/Q para examen y grupo
@app.route("/api/claves/origen", methods=["GET"])
def api_claves_origen():
    examen_id = request.args.get("examen_id", type=int)
    grupo_id  = request.args.get("grupo_id",  type=int)
    if not examen_id or not grupo_id:
        return jsonify({"ok": False, "error": "Faltan parámetros"}), 400

    conn = get_connection(); cur = conn.cursor(dictionary=True)
    # Ajusta nombres/keys: numero_pregunta, origen, p, q
    cur.execute("""
        SELECT numero_pregunta, origen, p, q
        FROM claves_respuesta
        WHERE examen_id=%s AND grupo_id=%s
        ORDER BY numero_pregunta
    """, (examen_id, grupo_id))
    filas = cur.fetchall()
    cur.close(); conn.close()

    # Si aún no cargaste claves desde el DOCX, devuelve dummy (evita pantalla vacía)
    if not filas:
        filas = [{"numero_pregunta": i+1, "origen": "A", "p": "B", "q": "C"} for i in range(10)]

    return jsonify(filas)

# (Opcional) Aleatorizar P/Q en servidor y persistir
@app.route("/api/claves/aleatorizar", methods=["POST"])
def api_aleatorizar_pq():
    data = request.get_json(force=True)
    examen_id = int(data.get("examen_id") or 0)
    grupo_id  = int(data.get("grupo_id")  or 0)
    if not examen_id or not grupo_id:
        return jsonify(ok=False, error="Faltan parámetros"), 400

    # asegura claves first
    api_claves_ensure_internal(examen_id, grupo_id)

    conn = get_connection(); cur = conn.cursor(dictionary=True)
    cur.execute("""
        SELECT id, numero_pregunta, origen
        FROM claves_respuesta
        WHERE examen_id=%s AND grupo_id=%s
        ORDER BY numero_pregunta
    """, (examen_id, grupo_id))
    filas = cur.fetchall()

    for r in filas:
        p, q = pick_two_distinct_letters()   # ✅ full random A–E
        cur.execute("""
            UPDATE claves_respuesta
            SET p=%s, q=%s, fecha_actualizacion=NOW()
            WHERE id=%s
        """, (p, q, r["id"]))
    conn.commit()
    cur.close(); conn.close()
    return jsonify(ok=True)

def api_claves_ensure_internal(examen_id, grupo_id):
    """igual que ensure pero interno para no repetir código."""
    conn = get_connection()
    cur = conn.cursor(dictionary=True)

    cur.execute("SELECT total_preguntas FROM examenes_importados WHERE id=%s", (examen_id,))
    ex = cur.fetchone()
    total = int(ex["total_preguntas"] or 0) if ex else 0
    if total <= 0:
        cur.close(); conn.close()
        raise Exception("Examen sin preguntas detectadas")

    cur.execute("""
        SELECT COUNT(*) AS n
        FROM claves_respuesta
        WHERE examen_id=%s AND grupo_id=%s
    """, (examen_id, grupo_id))
    n = int(cur.fetchone()["n"])

    if n == 0:
        rows = []
        for i in range(total):
            origen = "A"
            p, q = pick_two_distinct_letters()   # ✅ ahora puede incluir "A"
            rows.append((i+1, origen, p, q, examen_id, grupo_id))


        cur.executemany("""
            INSERT INTO claves_respuesta
                (numero_pregunta, origen, p, q, examen_id, grupo_id, fecha_actualizacion)
            VALUES (%s,%s,%s,%s,%s,%s,NOW())
        """, rows)
        conn.commit()

    else:
        cur.execute("""
            SELECT id, origen, p, q
            FROM claves_respuesta
            WHERE examen_id=%s AND grupo_id=%s
            ORDER BY numero_pregunta
        """, (examen_id, grupo_id))
        filas = cur.fetchall()

        valid = {"A","B","C","D","E"}
        for r in filas:
            origen = (r["origen"] or "A").upper()
            p = (r["p"] or "").upper()
            q = (r["q"] or "").upper()

            if p not in valid:
                p = random.choice(list(valid))
            if q not in valid or q == p:
                q = random.choice([x for x in valid if x != p])


            cur.execute("""
                UPDATE claves_respuesta
                SET p=%s, q=%s, fecha_actualizacion=NOW()
                WHERE id=%s
            """, (p, q, r["id"]))
        conn.commit()

    cur.close(); conn.close()

# Descargar pruebas (dummy para probar flujo)
@app.route("/api/pruebas/descargar", methods=["POST"])
def api_descargar_pruebas():
    data = request.get_json(force=True)
    grupo_id  = data.get("grupo_id")
    examen_id = data.get("examen_id")
    # Aquí deberías generar tus PDFs/ZIPs reales.
    tmp = os.path.join(tempfile.gettempdir(), f"pruebas_g{grupo_id}_e{examen_id}.txt")
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(f"Pruebas para grupo {grupo_id}, examen {examen_id}\n")
        f.write(f"Generado: {dt.datetime.now()}\n")
    return send_file(tmp, as_attachment=True, download_name=f"pruebas_g{grupo_id}.txt")




from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_docx_claves_all(nombre_base, claves_all):
    doc = Document()

    # Título simple
    doc.add_heading("Claves de respuesta", level=1)

    # Tabla con estilo tipo rejilla
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Encabezados centrados y en negrita
    headers = ["Grupo", "Pregunta", "Origen", "P", "Q"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True

    # Anchos más limpios (ajusta si quieres)
    widths = [Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.7), Inches(0.7)]
    for w, c in zip(widths, hdr_cells):
        c.width = w

    # Filas
    for r in claves_all:
        row = table.add_row().cells
        row[0].text = str(r.get("grupo", ""))
        row[1].text = str(r.get("numero_pregunta", ""))
        row[2].text = str(r.get("origen", ""))
        row[3].text = str(r.get("p", ""))
        row[4].text = str(r.get("q", ""))

        # Centrar contenido
        for c in row:
            if c.paragraphs:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

@app.route("/api/pruebas/descargar_all", methods=["POST"])
def api_descargar_pruebas_all():
    data = request.get_json(force=True)
    examen_id = int(data.get("examen_id") or 0)
    if not examen_id:
        return jsonify(ok=False, error="Falta examen_id"), 400

    conn = get_connection(); cur = conn.cursor(dictionary=True)

    # 1) examen importado (ruta docx)
    cur.execute("SELECT id, nombre, ruta FROM examenes_importados WHERE id=%s", (examen_id,))
    ex = cur.fetchone()
    if not ex or not ex["ruta"] or not os.path.exists(ex["ruta"]):
        cur.close(); conn.close()
        return jsonify(ok=False, error="No se encontró ruta del examen importado"), 404
    ruta_docx = ex["ruta"]
    nombre_base = os.path.splitext(ex["nombre"])[0]

    # 2) grupos activos
    cur.execute("SELECT idgrupo, clave, nombre FROM grupos WHERE activo=1 ORDER BY clave")
    grupos = cur.fetchall()
    if not grupos:
        cur.close(); conn.close()
        return jsonify(ok=False, error="No hay grupos activos"), 409

    # 3) ZIP en memoria
    mem_zip = io.BytesIO()
    zf = zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED)

    # guardaremos también las claves para el docx final
    claves_all = []  # [{grupo, numero_pregunta, origen, p, q}]

    for g in grupos:
        grupo_id = g["idgrupo"]
        clave_grupo = (g["clave"] or f"G{grupo_id}").strip()

        # asegurar claves, aleatorizar si falta p/q
        api_claves_ensure_internal(examen_id, grupo_id)

        cur.execute("""
            SELECT numero_pregunta, origen, p, q
            FROM claves_respuesta
            WHERE examen_id=%s AND grupo_id=%s
            ORDER BY numero_pregunta
        """, (examen_id, grupo_id))
        claves = cur.fetchall()

        # si p/q vienen vacíos, aleatorizamos YA
        # si p/q vienen vacíos, aleatorizamos YA
        if any((not r["p"] or not r["q"]) for r in claves):
            # ✅ FORZAR re-aleatorización siempre
            for r in claves:
                p, q = pick_two_distinct_letters()
                r["p"], r["q"] = p, q

            for r in claves:
                cur.execute("""
                    UPDATE claves_respuesta
                    SET p=%s,q=%s,fecha_actualizacion=NOW()
                    WHERE examen_id=%s AND grupo_id=%s AND numero_pregunta=%s
                """, (r["p"], r["q"], examen_id, grupo_id, r["numero_pregunta"]))
            conn.commit()



        # acumula para docx global
        for r in claves:
            claves_all.append({
                "grupo": clave_grupo,
                "numero_pregunta": r["numero_pregunta"],
                "origen": r["origen"],
                "p": r["p"],
                "q": r["q"],
            })

        # 4) generar DOCX P y Q para este grupo
        docx_p_bytes, docx_q_bytes = generar_docx_pq_para_grupo(
            ruta_docx, claves, clave_grupo, nombre_base
        )

        zf.writestr(f"{clave_grupo}_P.docx", docx_p_bytes)
        zf.writestr(f"{clave_grupo}_Q.docx", docx_q_bytes)

    # 5) crear docx global con claves
    claves_docx_bytes = generar_docx_claves_all(nombre_base, claves_all)
    zf.writestr("Claves de respuesta.docx", claves_docx_bytes)

    zf.close()
    cur.close(); conn.close()

    mem_zip.seek(0)
    return send_file(
        mem_zip,
        as_attachment=True,
        download_name=f"PRUEBAS_{nombre_base}_TODOS.zip",
        mimetype="application/zip",
    )




# ---------- CONFIG ----------
# --- utilidades ---
# --- utilidades ---
ALLOWED_EXTS = {".doc", ".docx", ".pdf"}

# ==== HELPERS DE RUTA CORTA (8.3) ====
def _short83(path: str) -> str:
    try:
        import win32api
        return win32api.GetShortPathName(os.path.abspath(path))
    except Exception:
        return os.path.abspath(path)

def _short_path(path: str) -> str:
    # alias por compatibilidad con tu código existente
    return _short83(path)

# ==== EXTENSIONES Y HASH (una sola vez) ====
ALLOWED_EXTS = {".doc", ".docx", ".pdf"}

def _ext_ok(fname: str) -> bool:
    return os.path.splitext(fname)[1].lower() in ALLOWED_EXTS

def sha256sum(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            h.update(chunk)
    return h.hexdigest()



@app.route("/api/examenes/importar", methods=["POST", "OPTIONS"])
def api_examenes_importar():
    if request.method == "OPTIONS":
        return ("", 204)

    files = request.files.getlist("files") + request.files.getlist("files[]")
    if not files:
        return jsonify(ok=False, error="No se recibieron archivos."), 400

    base = app.config.get("UPLOADS_EXAM_DIR")
    if not base:
        return jsonify(ok=False, error="UPLOADS_EXAM_DIR no configurado."), 500
    os.makedirs(base, exist_ok=True)

    items = []
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
    except Exception:
        conn = cur = None

    for f in files:
        if not f or not f.filename:
            continue
        if not _ext_ok(f.filename):
            return jsonify(ok=False, error=f"Extensión no permitida: {os.path.splitext(f.filename)[1]}"), 415

        safe = secure_filename(f.filename)
        dst = os.path.join(base, safe)
        f.save(dst)

        h = sha256sum(dst)
        ext = os.path.splitext(safe)[1].lower().lstrip(".")

        # ✅ contar preguntas del DOCX recién guardado
        total_p = 0
        try:
            if ext in ("docx", "doc"):
                
                total_p = contar_preguntas_docx(dst)
        except Exception as e:
            print("Error contando preguntas:", e)
            total_p = 0

        if cur:
            cur.execute("""
                INSERT INTO examenes_importados
                    (nombre, ruta, extension, total_preguntas, fuente, hash_archivo)
                VALUES (%s,%s,%s,%s,%s,%s)
                ON DUPLICATE KEY UPDATE
                    ruta=VALUES(ruta),
                    extension=VALUES(extension),
                    total_preguntas=VALUES(total_preguntas)
            """, (safe, dst, ext, total_p, "upload", h))
            conn.commit()

            cur.execute("""
                SELECT id, nombre, total_preguntas, fecha_creacion
                FROM examenes_importados
                WHERE hash_archivo=%s
            """, (h,))
            row = cur.fetchone()

            if not row:
                cur.execute("""
                    SELECT id, nombre, total_preguntas, fecha_creacion
                    FROM examenes_importados
                    WHERE nombre=%s ORDER BY id DESC LIMIT 1
                """, (safe,))
                row = cur.fetchone()

            if row:
                items.append(row)
        else:
            items.append({
                "id": None,
                "nombre": safe,
                "total_preguntas": total_p,
                "fecha_creacion": None
            })

    if cur: cur.close()
    if conn: conn.close()
    return jsonify(ok=True, items=items), 200


@app.route("/api/examenes/importados", methods=["GET"])
def api_examenes_importados():
    try:
        conn = get_connection()
        cur = conn.cursor(dictionary=True)
        cur.execute("""
            SELECT id, nombre, total_preguntas, fecha_creacion
            FROM examenes_importados
            ORDER BY id DESC
        """)
        data = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        print("Error listando importados:", e)
        data = []

    return jsonify(data), 200
def contar_preguntas_docx(ruta_docx: str) -> int:
    """
    Cuenta preguntas sin partir por temas.
    FIX:
      - fija active_q_numId SOLO por numeración real
      - descarta alternativas A)–E)
      - una vez fijado active_q_numId, ignora regex en párrafos SIN numPr
        (evita contar fórmulas/ángulos/unidades sueltas)
    """

    import os, re, tempfile, shutil, zipfile
    import xml.etree.ElementTree as ET

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns_doc = {"w": NS_W}

    QUESTION_FMTS = {"decimal", "decimalZero"}
    NON_QUESTION_TOPLEVEL_FMTS = {
        "upperLetter", "lowerLetter",
        "upperRoman", "lowerRoman",
        "bullet"
    }

    def paragraph_text(p) -> str:
        parts = []
        for node in p.iter():
            if node.tag == W + "t":
                txt = node.text or ""
                if node.get(W + "space") == "preserve":
                    parts.append(txt)
                else:
                    parts.append(txt.strip())
            elif node.tag in (W + "tab", W + "br"):
                parts.append(" ")
        return "".join(parts)

    def is_alt_line(txt_norm: str) -> bool:
        return re.match(r"^\(?[a-e]\s*[\)\.\-]\s+", txt_norm, flags=re.I) is not None

    def regex_qstart(plain_text: str) -> bool:
        return re.match(r"^\s*\(?[1-9]\d{0,2}\)?[.)]?(?:\s+|(?=[^\s]))", plain_text) is not None

    # --- unzip DOCX ---
    tmp = tempfile.mkdtemp(prefix="countq_")
    work = os.path.join(tmp, "w"); os.mkdir(work)
    with zipfile.ZipFile(ruta_docx, "r") as z:
        z.extractall(work)

    # --- numbering map ---
    num_fmt_map = {}
    numbering = os.path.join(work, "word", "numbering.xml")
    if os.path.exists(numbering):
        tn = ET.parse(numbering)
        rn = tn.getroot()
        for num in rn.findall(".//w:num", ns_doc):
            numId = num.attrib.get(W + "numId")
            abs_el = num.find("./w:abstractNumId", ns_doc)
            if abs_el is None:
                continue
            abs_id = abs_el.attrib.get(W + "val")
            abstract = rn.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns_doc)
            if abstract is None:
                continue
            for lvl in abstract.findall("./w:lvl", ns_doc):
                ilvl = lvl.attrib.get(W + "ilvl")
                nf = lvl.find("./w:numFmt", ns_doc)
                if nf is not None:
                    num_fmt_map[(numId, ilvl)] = nf.attrib.get(W + "val")

    # --- document.xml ---
    docxml = os.path.join(work, "word", "document.xml")
    tdoc = ET.parse(docxml)
    root = tdoc.getroot()
    body = root.find(W + "body")

    total = 0
    active_q_numId = None

    for p in body.findall("w:p", ns_doc):
        plain = paragraph_text(p)
        if not plain.strip():
            continue

        txt_norm = re.sub(r"\s+", " ", plain.strip())

        # alternativa => nunca es pregunta
        if is_alt_line(txt_norm):
            continue

        numPr = p.find(".//w:numPr", ns_doc)
        ilvl = numPr.find("./w:ilvl", ns_doc).get(W+"val") if (
            numPr is not None and numPr.find("./w:ilvl", ns_doc) is not None
        ) else ""
        numId = numPr.find("./w:numId", ns_doc).get(W+"val") if (
            numPr is not None and numPr.find("./w:numId", ns_doc) is not None
        ) else ""
        fmt = num_fmt_map.get((numId, ilvl), "")

        by_num = (
            numPr is not None and ilvl == "0"
            and fmt not in NON_QUESTION_TOPLEVEL_FMTS
            and (fmt in QUESTION_FMTS or fmt == "")
        )

        by_regex = regex_qstart(plain)

        # ✅ si ya hay corriente de preguntas, NO aceptar regex sin numPr
        if active_q_numId is not None and numPr is None:
            by_regex = False

        qstart = by_num or by_regex

        # fija numId real SOLO por numeración
        if active_q_numId is None and by_num:
            active_q_numId = numId

        # si ya fijamos numId real, ignora otros numId
        if qstart and active_q_numId and numPr is not None and numId != active_q_numId:
            qstart = False

        if qstart:
            total += 1

    shutil.rmtree(tmp, ignore_errors=True)
    return total



def debug_contar_preguntas_docx(ruta_docx: str):
    """
    DEBUG: imprime cada párrafo con:
      - texto
      - numId, ilvl, fmt
      - si fue detectado como pregunta
      - razón
    """

    import os, re, tempfile, shutil, zipfile
    import xml.etree.ElementTree as ET

    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{%s}" % NS_W
    ns_doc = {"w": NS_W}

    QUESTION_FMTS = {"decimal", "decimalZero"}
    NON_QUESTION_TOPLEVEL_FMTS = {
        "upperLetter", "lowerLetter",
        "upperRoman", "lowerRoman",
        "bullet"
    }

    def paragraph_text(p) -> str:
        parts = []
        for node in p.iter():
            if node.tag == W + "t":
                txt = node.text or ""
                if node.get(W + "space") == "preserve":
                    parts.append(txt)
                else:
                    parts.append(txt.strip())
            elif node.tag in (W + "tab", W + "br"):
                parts.append(" ")
        return "".join(parts)

    def is_alt_line(txt_norm: str) -> bool:
        # detecta alternativas tipo "A) ...." "B. ...." "(C) ...."
        return re.match(r"^\(?[a-e]\s*[\)\.\-]\s+", txt_norm, flags=re.I) is not None

    def regex_qstart(plain_text: str) -> bool:
        # 1) 2. 3) 10) ...
        return re.match(r"^\s*\(?[1-9]\d{0,2}\)?[.)]?(?:\s+|(?=[^\s]))", plain_text) is not None

    # --- unzip DOCX ---
    tmp = tempfile.mkdtemp(prefix="dbg_countq_")
    work = os.path.join(tmp, "w"); os.mkdir(work)
    with zipfile.ZipFile(ruta_docx, "r") as z:
        z.extractall(work)

    # --- map numId+ilvl -> fmt ---
    num_fmt_map = {}
    numbering = os.path.join(work, "word", "numbering.xml")
    if os.path.exists(numbering):
        tn = ET.parse(numbering)
        rn = tn.getroot()
        for num in rn.findall(".//w:num", ns_doc):
            numId = num.attrib.get(W + "numId")
            abs_el = num.find("./w:abstractNumId", ns_doc)
            if abs_el is None:
                continue
            abs_id = abs_el.attrib.get(W + "val")
            abstract = rn.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns_doc)
            if abstract is None:
                continue
            for lvl in abstract.findall("./w:lvl", ns_doc):
                ilvl = lvl.attrib.get(W + "ilvl")
                nf = lvl.find("./w:numFmt", ns_doc)
                if nf is not None:
                    num_fmt_map[(numId, ilvl)] = nf.attrib.get(W + "val")

    # --- leer document.xml ---
    docxml = os.path.join(work, "word", "document.xml")
    tdoc = ET.parse(docxml)
    root = tdoc.getroot()
    body = root.find(W + "body")

    total = 0
    active_q_numId = None

    print("\n===== DEBUG PÁRRAFOS =====\n")

    idxp = 0
    for p in body.findall("w:p", ns_doc):
        idxp += 1
        plain = paragraph_text(p)
        if not plain.strip():
            continue

        txt_norm = re.sub(r"\s+", " ", plain.strip())

        numPr = p.find(".//w:numPr", ns_doc)
        ilvl = numPr.find("./w:ilvl", ns_doc).get(W+"val") if (
            numPr is not None and numPr.find("./w:ilvl", ns_doc) is not None
        ) else ""
        numId = numPr.find("./w:numId", ns_doc).get(W+"val") if (
            numPr is not None and numPr.find("./w:numId", ns_doc) is not None
        ) else ""
        fmt = num_fmt_map.get((numId, ilvl), "")

        by_regex = regex_qstart(plain)
        by_num = (
            numPr is not None and ilvl == "0"
            and fmt not in NON_QUESTION_TOPLEVEL_FMTS
            and (fmt in QUESTION_FMTS or fmt == "")
        )

        reason = []
        qstart = False

        # descartar alternativas A) B) C)
        if is_alt_line(txt_norm):
            reason.append("DESCARTADO: parece alternativa A)/B)/C)")
            qstart = False
        else:
            if by_num:
                qstart = True
                reason.append("QSTART por NUMERACIÓN (ilvl=0, fmt pregunta)")
            elif by_regex:
                qstart = True
                reason.append("QSTART por REGEX (texto empieza con número)")
            else:
                reason.append("NO es inicio de pregunta")

        # fijar active_q_numId SOLO por numeración real (no por regex suelto)
        if active_q_numId is None and by_num:
            active_q_numId = numId
            reason.append(f"FIJA active_q_numId={active_q_numId}")

        # si ya fijamos numId real y este párrafo usa otro numId => no cuenta
        if qstart and active_q_numId and numPr is not None and numId != active_q_numId:
            qstart = False
            reason.append(f"DESCARTADO: numId={numId} != active_q_numId={active_q_numId}")

        if qstart:
            total += 1

        print(f"[{idxp}] text='{txt_norm[:120]}'")
        print(f"     numId={numId or '-'}  ilvl={ilvl or '-'}  fmt={fmt or '-'}")
        print(f"     qstart={qstart}  reasons: {' | '.join(reason)}\n")

    print("===== FIN DEBUG =====")
    print("TOTAL CONTADO =", total)

    shutil.rmtree(tmp, ignore_errors=True)
    return total


# ================== FIN BLOQUE ==================
# ================== FIN BLOQUE ==================

@app.get("/__ping__")
def ping():
    return "ok"

if __name__ == "__main__":
    print("== URL MAP ==")
    for r in app.url_map.iter_rules():
        print(r, r.methods)
    app.run(host="127.0.0.1", port=5050, debug=True)
