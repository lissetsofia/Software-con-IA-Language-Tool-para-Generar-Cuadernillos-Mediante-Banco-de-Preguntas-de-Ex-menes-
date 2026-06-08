# backend/test/test_app_matriz_banco_robusto_cov8.py
import os
import sqlite3
from pathlib import Path

from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov8_matriz_banco.sqlite3"

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    return db_path


def _schema(app_module):
    conn = app_module.get_connection()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS tema_docs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tema_id INTEGER,
            doc_preguntas_nombre TEXT,
            doc_preguntas_ruta TEXT,
            doc_sol_nombre TEXT,
            doc_sol_ruta TEXT
        )
    """)
    conn.commit()
    return conn, cur


def _make_docx(path: Path, text="Pregunta banco"):
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _seed_banco(app_module, tmp_path):
    conn, cur = _schema(app_module)
    cur.execute("INSERT INTO temario (nombre, activo) VALUES ('ÁLGEBRA', 1)")
    tema_id = int(cur.lastrowid)
    preg = _make_docx(tmp_path / "pregunta_banco.docx", "1. Pregunta de banco")
    sol = _make_docx(tmp_path / "sol_banco.docx", "Solución de banco")
    cur.execute(
        """INSERT INTO tema_docs
           (tema_id, doc_preguntas_nombre, doc_preguntas_ruta, doc_sol_nombre, doc_sol_ruta)
           VALUES (?, ?, ?, ?, ?)""",
        (tema_id, preg.name, str(preg), sol.name, str(sol)),
    )
    doc_id = int(cur.lastrowid)
    conn.commit()
    cur.close(); conn.close()
    return tema_id, doc_id, preg, sol


def _patch_merge_ok(app_module, monkeypatch):
    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "_cut_docx_to_individual_question_docs", lambda src, n: [src], raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_post_merge_fix_numbering", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "bullets_to_numbers_docx", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge(grouped, out_path, *args, **kwargs):
        doc = DocxDocument()
        for tema, files in grouped:
            doc.add_paragraph(str(tema))
            for f in files:
                doc.add_paragraph(os.path.basename(str(f)))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)


def test_generar_matriz_banco_robusto_success_y_errores(app_module, tmp_path, monkeypatch):
    _patch_merge_ok(app_module, monkeypatch)
    src = _make_docx(tmp_path / "src_banco.docx", "1. Pregunta banco")
    out = tmp_path / "salida_banco.docx"

    with app_module.app.test_request_context("/"):
        resp = app_module._generar_matriz_banco_docx_robusto(
            [(1, "ÁLGEBRA", [str(src)])],
            str(out),
            log_prefix="[COV8]",
        )
        assert getattr(resp, "status_code", 200) == 200

        missing = app_module._generar_matriz_banco_docx_robusto(
            [(1, "ÁLGEBRA", [str(tmp_path / "no_existe.docx")])],
            str(tmp_path / "missing.docx"),
        )
        assert isinstance(missing, tuple)
        assert missing[1] == 400

        empty = app_module._generar_matriz_banco_docx_robusto([], str(tmp_path / "empty.docx"))
        assert isinstance(empty, tuple)
        assert empty[1] == 400


def test_generar_matriz_banco_robusto_docx_invalido_y_merge_malos(app_module, tmp_path, monkeypatch):
    bad = tmp_path / "bad.docx"
    bad.write_bytes(b"no es docx")

    with app_module.app.test_request_context("/"):
        monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: (_ for _ in ()).throw(ValueError("docx malo")), raising=False)
        invalid = app_module._generar_matriz_banco_docx_robusto([(1, "ÁLGEBRA", [str(bad)])], str(tmp_path / "out.docx"))
        assert isinstance(invalid, tuple)
        assert invalid[1] == 400

    good = _make_docx(tmp_path / "good.docx")
    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "_cut_docx_to_individual_question_docs", lambda src, n: [src], raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge_malo(grouped, out_path, *args, **kwargs):
        _make_docx(Path(out_path), "parcial")
        return out_path, [], [(grouped[0][1][0], "fallo insertando")]

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge_malo, raising=False)
    with app_module.app.test_request_context("/"):
        malo = app_module._generar_matriz_banco_docx_robusto([(1, "ÁLGEBRA", [str(good)])], str(tmp_path / "merge_malo.docx"))
        assert isinstance(malo, tuple)
        assert malo[1] == 409


def test_rutas_matriz_banco_success_y_solucionario(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    tema_id, doc_id, preg, sol = _seed_banco(app_module, tmp_path)

    llamadas = []

    def fake_generar(grouped_data, out_path, log_prefix=""):
        llamadas.append((grouped_data, out_path, log_prefix))
        return app_module.jsonify(ok=True, temas=len(grouped_data), log_prefix=log_prefix)

    monkeypatch.setattr(app_module, "_generar_matriz_banco_docx_robusto", fake_generar, raising=False)

    r_ok = client.post(
        "/api/matriz/generar_desde_banco",
        json={"nombre": "Cov8", "items": [{"tema_id": tema_id, "doc_ids": [doc_id]}]},
    )
    assert r_ok.status_code == 200
    assert r_ok.get_json()["log_prefix"] == "[MATRIZ_BANCO]"
    assert llamadas[-1][0][0][2] == [os.path.abspath(str(preg))]

    r_empty_topic = client.post(
        "/api/matriz/generar_desde_banco",
        json={"items": [{"tema_id": tema_id, "doc_ids": []}]},
    )
    assert r_empty_topic.status_code == 200
    assert llamadas[-1][0][0][2] == []

    r_missing = client.post(
        "/api/matriz/generar_desde_banco",
        json={"items": [{"tema_id": tema_id, "doc_ids": [999999]}]},
    )
    assert r_missing.status_code in (200, 400)

    r_sol = client.post(
        "/api/matriz/generar_desde_banco/solucionario",
        json={"nombre": "Cov8Sol", "items": [{"tema_id": tema_id, "doc_ids": [doc_id]}]},
    )
    assert r_sol.status_code == 200
    assert r_sol.get_json()["log_prefix"] == "[MATRIZ_BANCO_SOL]"
    assert llamadas[-1][0][0][2] == [os.path.abspath(str(sol))]

    # Sin solucionario debe activar la rama 409.
    conn = app_module.get_connection(); cur = conn.cursor()
    cur.execute(
        """INSERT INTO tema_docs (tema_id, doc_preguntas_nombre, doc_preguntas_ruta, doc_sol_nombre, doc_sol_ruta)
           VALUES (?, 'sin_sol.docx', ?, NULL, NULL)""",
        (tema_id, str(preg)),
    )
    sin_sol_id = int(cur.lastrowid)
    conn.commit(); cur.close(); conn.close()
    r_no_sol = client.post(
        "/api/matriz/generar_desde_banco/solucionario",
        json={"items": [{"tema_id": tema_id, "doc_ids": [sin_sol_id]}]},
    )
    assert r_no_sol.status_code == 409
