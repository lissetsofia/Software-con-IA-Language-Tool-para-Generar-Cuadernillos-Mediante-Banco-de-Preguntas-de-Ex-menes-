import os
from docx import Document


class _FakeWebOptions:
    def __setattr__(self, name, value):
        object.__setattr__(self, name, value)


class _FakeRange:
    def __init__(self):
        self.FormattedText = ""
    def InsertFile(self, *args, **kwargs):
        return None
    def Collapse(self, *args, **kwargs):
        return None


class _FakeContent(_FakeRange):
    def __init__(self):
        super().__init__()
        self.Text = ""
        self.End = 1


class _FakeDoc:
    def __init__(self):
        self.Content = _FakeContent()
        self.WebOptions = _FakeWebOptions()
    def ExportAsFixedFormat(self, **kwargs):
        out = kwargs.get("OutputFileName")
        with open(out, "wb") as f:
            f.write(b"%PDF-1.4\n%%EOF")
    def SaveAs(self, *args, **kwargs):
        path = kwargs.get("FileName") or (args[0] if args else None)
        if path:
            if str(path).lower().endswith(".htm"):
                with open(path, "w", encoding="utf-8") as f:
                    f.write("<html><head><meta charset=windows-1252></head><body>ok</body></html>")
            else:
                Document().save(path)
    def SaveAs2(self, *args, **kwargs):
        path = kwargs.get("FileName") or (args[0] if args else None)
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write("<html><head></head><body>ok</body></html>")
    def Close(self, *args, **kwargs):
        return None
    def Range(self, *args, **kwargs):
        return _FakeRange()


class _FakeDocuments:
    def Open(self, *args, **kwargs):
        return _FakeDoc()
    def Add(self, *args, **kwargs):
        return _FakeDoc()


class _FakeWord:
    def __init__(self):
        self.Documents = _FakeDocuments()
        self.Visible = False
        self.DisplayAlerts = 0
        self.AutomationSecurity = 0
    def Quit(self):
        return None


def _patch_fake_word(app_module, monkeypatch):
    monkeypatch.setattr(app_module.pythoncom, "CoInitialize", lambda: None)
    monkeypatch.setattr(app_module.pythoncom, "CoUninitialize", lambda: None)
    monkeypatch.setattr(app_module.win32, "DispatchEx", lambda *a, **k: _FakeWord())


def test_generar_pdf_lt_y_generar_pdf_con_word_fake(app_module, tmp_path, monkeypatch):
    _patch_fake_word(app_module, monkeypatch)
    src = tmp_path / "entrada.docx"
    Document().save(src)

    pdf1 = app_module.generar_pdf_lt(str(src))
    assert os.path.exists(pdf1)
    assert open(pdf1, "rb").read().startswith(b"%PDF")

    pdf2 = app_module.generar_pdf(str(src))
    assert os.path.exists(pdf2)
    assert open(pdf2, "rb").read().startswith(b"%PDF")


def test_generar_html_desde_docx_lt_y_generar_html_word_success(app_module, tmp_path, monkeypatch):
    _patch_fake_word(app_module, monkeypatch)
    desc = tmp_path / "desc"
    desc.mkdir()
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(desc))
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(desc))

    src = tmp_path / "entrada.docx"
    Document().save(src)

    html_lt = app_module.generar_html_desde_docx_lt(str(src), "vista_lt")
    assert os.path.exists(html_lt)
    assert "utf-8" in open(html_lt, encoding="utf-8").read().lower()

    html = app_module.generar_html_desde_docx(str(src), "vista_word")
    assert os.path.exists(html)
    txt = open(html, encoding="utf-8").read().lower()
    assert "body" in txt
