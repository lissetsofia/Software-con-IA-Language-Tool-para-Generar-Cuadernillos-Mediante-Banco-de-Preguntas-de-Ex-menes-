import io
import os
import sqlite3
from pathlib import Path

from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / 'cov12_claves.sqlite3'

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, 'get_connection', get_connection)
    conn = _connect(db_path)
    cur = conn.cursor()
    cur.executescript('''
        CREATE TABLE examenes_importados (
            id INTEGER PRIMARY KEY,
            nombre TEXT,
            ruta TEXT,
            total_preguntas INTEGER DEFAULT 0
        );
        CREATE TABLE grupos (
            idgrupo INTEGER PRIMARY KEY,
            clave TEXT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        );
        CREATE TABLE claves_tipo (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            examen_id INTEGER,
            grupo_id INTEGER,
            codigo TEXT,
            orden INTEGER DEFAULT 0,
            activo INTEGER DEFAULT 1,
            UNIQUE(examen_id, grupo_id, codigo)
        );
        CREATE TABLE claves_respuesta (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            examen_id INTEGER,
            grupo_id INTEGER,
            numero_pregunta INTEGER,
            origen TEXT,
            fecha_actualizacion TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(examen_id, grupo_id, numero_pregunta)
        );
        CREATE TABLE claves_respuesta_detalle (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            claves_respuesta_id INTEGER,
            tipo_id INTEGER,
            clave TEXT,
            fecha_actualizacion TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(claves_respuesta_id, tipo_id)
        );
    ''')
    cur.execute("INSERT INTO grupos(idgrupo, clave, nombre, activo) VALUES (1, 'A', 'Grupo A', 1)")
    conn.commit(); cur.close(); conn.close()
    return db_path


def _make_simple_docx(path: Path):
    doc = DocxDocument()
    doc.add_paragraph('Pregunta simple')
    doc.add_paragraph('Correcta')
    doc.add_paragraph('Incorrecta')
    doc.save(path)
    return path


def test_claves_helpers_guardar_origen_aleatorizar_y_build(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    monkeypatch.setattr(app_module.random, 'sample', lambda pool, n: list(pool)[:n])

    assert len(set(app_module.pick_two_distinct_letters())) == 2
    assert app_module.pick_n_distinct_letters(2, exclude='A') == ['B', 'C']
    try:
        app_module.pick_n_distinct_letters(10)
        assert False, 'debió fallar por letras insuficientes'
    except ValueError:
        pass
    assert app_module.pick_distinct_for_tipos(2, exclude_origen='A') == ['B', 'C']
    assert app_module._norm_code(' p_1 ') == 'P_1'
    assert app_module._norm_code('mal codigo') is None
    assert app_module._infer_tipos_from_filas([{'numero_pregunta': 1, 'origen': 'A', 'p': 'B', 'q': 'C', 'R': 'D'}]) == ['P', 'Q', 'R']
    assert app_module.inferir_clave_grupo_desde_nombre('Examen grupo-B.docx') == 'B'
    assert app_module.inferir_clave_grupo_desde_nombre('sin_clave.docx') is None

    conn = app_module.get_connection(); cur = conn.cursor()
    docx_path = _make_simple_docx(tmp_path / 'grupo_A.docx')
    cur.execute("INSERT INTO examenes_importados(id, nombre, ruta, total_preguntas) VALUES (1, 'grupo_A.docx', ?, 2)", (str(docx_path),))
    conn.commit()

    tipo_map = app_module.ensure_tipos(conn, 1, 1, codigos=('P', 'Q'))
    assert set(tipo_map) == {'P', 'Q'}
    cur.execute("INSERT INTO claves_respuesta(examen_id, grupo_id, numero_pregunta, origen) VALUES (1,1,1,'A')")
    cr_id = cur.lastrowid
    app_module.upsert_detalle(cur, cr_id, tipo_map['P'], 'B')
    app_module.upsert_detalle(cur, cr_id, tipo_map['Q'], 'C')
    conn.commit()

    filas, tipos = app_module.fetch_claves_pivot(conn, 1, 1)
    assert tipos == ['P', 'Q']
    assert filas[0]['p'] == 'B'

    salidas, claves_all, tipos_global, err = app_module._build_salidas_y_claves(conn, [1], todos_los_grupos=False, grupo_id_fijo=1)
    assert err is None
    assert salidas and claves_all and tipos_global
    cur.close(); conn.close()

    bad = client.post('/api/claves/guardar', json={'examen_id': 0, 'grupo_id': 1})
    assert bad.status_code == 400

    saved = client.post('/api/claves/guardar', json={
        'examen_id': 1,
        'grupo_id': 1,
        'filas': [
            {'numero_pregunta': 1, 'origen': 'Z', 'p': 'D', 'q': '', 'R': 'E'},
            {'numero_pregunta': 2, 'origen': 'B', 'P': 'C', 'Q': 'D', 'R': 'A'},
        ],
    })
    assert saved.status_code == 200, saved.get_data(as_text=True)
    assert 'R' in saved.get_json()['tipos']

    origin_missing = client.get('/api/claves/origen')
    assert origin_missing.status_code == 400
    origin = client.get('/api/claves/origen?examen_id=1&grupo_id=1')
    assert origin.status_code == 200
    assert origin.get_json()['ok'] is True
    assert origin.get_json()['filas']

    ens = client.post('/api/claves/ensure', json={'examen_ids': [1], 'grupo_id': 1, 'tipos': ['P', 'Q', 'R']})
    assert ens.status_code == 200
    assert ens.get_json()['items'][0]['total'] == 2

    alea_bad = client.post('/api/claves/aleatorizar', json={'examen_id': 0, 'grupo_id': 1})
    assert alea_bad.status_code == 400
    alea = client.post('/api/claves/aleatorizar', json={'examen_id': 1, 'grupo_id': 1, 'tipos': ['P', 'Q']})
    assert alea.status_code == 200, alea.get_data(as_text=True)
    assert alea.get_json()['tipos'] == ['P', 'Q']

    # Generadores DOCX de claves y por tipo.
    bytes_all = app_module.generar_docx_claves_all('base', [{'grupo': 'A', 'numero_pregunta': 1, 'origen': 'A', 'p': 'B', 'q': 'C'}])
    assert bytes_all.startswith(b'PK')
    bytes_dyn = app_module.generar_docx_claves_all_dinamico(['P', 'Q'], [{'grupo': 'A', 'numero_pregunta': 1, 'P': 'B', 'Q': 'C'}])
    assert bytes_dyn.startswith(b'PK')

    gen_tipo = app_module.generar_docx_tipo_para_grupo(str(docx_path), [{'numero_pregunta': 1, 'P': 'A'}], 'P')
    assert gen_tipo.startswith(b'PK')

    download = client.post('/api/pruebas/descargar', json={'grupo_id': 1, 'examen_id': 1})
    assert download.status_code == 200
