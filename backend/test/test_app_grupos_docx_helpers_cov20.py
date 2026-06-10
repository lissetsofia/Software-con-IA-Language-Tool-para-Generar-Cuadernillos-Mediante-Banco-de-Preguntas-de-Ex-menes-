import io
import os
import sqlite3
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _make_docx(path, paragraphs):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)
    return path


def _docx_upload(paragraphs):
    bio = io.BytesIO()
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(bio)
    bio.seek(0)
    return bio


def _patch_min_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov20_grupos_docx.sqlite3"
    if db_path.exists():
        db_path.unlink()
    conn = _connect(db_path)
    conn.executescript("""
        CREATE TABLE gen_lote(id INTEGER PRIMARY KEY AUTOINCREMENT, matriz_id INTEGER, nombre TEXT, usuario TEXT);
    """)
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    out = tmp_path / "grupos_out"
    out.mkdir()
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out), raising=False)
    return db_path


def test_norm_ranges_y_extract_tema_docx(app_module, tmp_path):
    src = _make_docx(tmp_path / "matriz.docx", ["ÁLGEBRA", "1. Pregunta A", "GEOMETRÍA", "1. Pregunta G"])
    assert app_module._norm_name("  álgebra  ") == "ALGEBRA"
    ranges = app_module._build_heading_ranges_for_temas(str(src), {"ALGEBRA", "GEOMETRIA"})
    assert ranges["ALGEBRA"] == (0, 2)
    out = tmp_path / "algebra.docx"
    app_module._extract_tema_docx_range(str(src), 1, 2, str(out))
    textos = [p.text for p in Document(out).paragraphs]
    assert textos == ["1. Pregunta A"]


def test_generar_from_docx_cuotas_desiguales_y_extract_error(client, app_module, tmp_path, monkeypatch):
    _patch_min_db(app_module, tmp_path, monkeypatch)

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {
        1: {"idgrupo": 1, "clave": "A", "nombre": "A", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1}]},
        2: {"idgrupo": 2, "clave": "B", "nombre": "B", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 2}]},
    }, raising=False)
    r_bad = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_upload(["ÁLGEBRA", "1. pregunta"]), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_bad.status_code == 400
    assert "cuotas" in r_bad.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {
        1: {"idgrupo": 1, "clave": "A", "nombre": "A", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1}]},
    }, raising=False)
    monkeypatch.setattr(app_module, "_extract_tema_docx_range", lambda *_a, **_k: (_ for _ in ()).throw(RuntimeError("extract falla")), raising=False)
    r_ext = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_upload(["ÁLGEBRA", "1. pregunta"]), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_ext.status_code == 500


def test_generar_from_docx_success_merge_sin_com(client, app_module, tmp_path, monkeypatch):
    _patch_min_db(app_module, tmp_path, monkeypatch)
    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {
        1: {"idgrupo": 1, "clave": "A", "nombre": "Grupo A", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1}]},
    }, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge(bloques, out_path, *args, **kwargs):
        doc = Document()
        for titulo, files in bloques:
            doc.add_paragraph(titulo)
            doc.add_paragraph(os.path.basename(files[0]))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    r_ok = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_upload(["ÁLGEBRA", "1. pregunta"]), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    body = r_ok.get_json()
    assert body["ok"] is True
    assert body["zip_url"].endswith("/zip")
