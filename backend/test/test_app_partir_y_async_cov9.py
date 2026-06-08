# backend/test/test_app_partir_y_async_cov9.py
import io
import os
import sqlite3
from pathlib import Path

from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov9_partir.sqlite3"

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    conn = _connect(db_path)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS examenes (
            idexamenes INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            numero TEXT,
            institucion TEXT,
            anio INTEGER,
            archivo_nombre TEXT,
            archivo_ruta TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS preguntas (
            idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT,
            examenes_idexamenes INTEGER,
            tema_id INTEGER,
            numero_p INTEGER,
            archivo_nombre TEXT,
            archivo_ruta TEXT
        )
    """)
    conn.commit()
    cur.close()
    conn.close()
    return db_path


def _make_exam_docx(path: Path):
    doc = DocxDocument()
    doc.add_paragraph("ÁLGEBRA")
    doc.add_paragraph("1) Primera pregunta de álgebra")
    doc.add_paragraph("A) alternativa uno")
    doc.add_paragraph("B) alternativa dos")
    doc.add_paragraph("2) Segunda pregunta de álgebra")
    doc.add_paragraph("A) alternativa uno")
    doc.save(path)
    return path


def _seed_exam(app_module, db_path: Path, docx_path: Path):
    conn = _connect(db_path)
    cur = conn.cursor()
    cur.execute("INSERT INTO temario (nombre, activo) VALUES ('ÁLGEBRA', 1)")
    cur.execute(
        """INSERT INTO examenes
           (idexamenes, nombre, numero, institucion, anio, archivo_nombre, archivo_ruta)
           VALUES (1, 'Examen prueba', 'I', 'UNAMBA', 2025, 'fuente.docx', ?)""",
        (str(docx_path),),
    )
    conn.commit()
    cur.close()
    conn.close()


def test_partir_y_guardar_success_overwrite_y_persistencia(client, app_module, tmp_path, monkeypatch):
    db_path = _patch_db(app_module, tmp_path, monkeypatch)
    src = _make_exam_docx(tmp_path / "fuente.docx")
    _seed_exam(app_module, db_path, src)

    preguntas_dir = tmp_path / "preguntas_partidas"
    preguntas_dir.mkdir()
    monkeypatch.setitem(app_module.app.config, "PREGUNTAS_DIR", str(preguntas_dir))

    # Evita Word COM. La ruta seguirá leyendo y reempacando el DOCX real.
    monkeypatch.setattr(app_module, "normalizar_docx_fuente", lambda path: (None, "skip cov9"), raising=False)

    r = client.post("/api/examenes/1/partir_y_guardar?overwrite=1")
    assert r.status_code == 200
    data = r.get_json()
    assert data["ok"] is True
    assert data["preguntas_insertadas"] == 2
    assert data["por_tema"]["ÁLGEBRA"] == 2

    conn = _connect(db_path)
    cur = conn.cursor()
    rows = cur.execute(
        "SELECT numero_p, archivo_nombre, archivo_ruta FROM preguntas ORDER BY numero_p"
    ).fetchall()
    cur.close(); conn.close()

    assert [r["numero_p"] for r in rows] == [1, 2]
    assert all(Path(r["archivo_ruta"]).exists() for r in rows)


def test_partir_y_guardar_validaciones_y_error_bd(client, app_module, tmp_path, monkeypatch):
    db_path = _patch_db(app_module, tmp_path, monkeypatch)

    # Examen inexistente
    assert client.post("/api/examenes/999/partir_y_guardar").status_code == 404

    # Examen con ruta que no existe
    conn = _connect(db_path); cur = conn.cursor()
    cur.execute("INSERT INTO temario (nombre, activo) VALUES ('ÁLGEBRA', 1)")
    cur.execute(
        """INSERT INTO examenes
           (idexamenes, nombre, archivo_nombre, archivo_ruta)
           VALUES (2, 'sin archivo', 'no.docx', ?)""",
        (str(tmp_path / "no_existe.docx"),),
    )
    conn.commit(); cur.close(); conn.close()
    assert client.post("/api/examenes/2/partir_y_guardar").status_code == 500

    # Error de conexión
    monkeypatch.setattr(app_module, "get_connection", lambda: (_ for _ in ()).throw(RuntimeError("db rota")))
    r = client.post("/api/examenes/1/partir_y_guardar")
    assert r.status_code == 500
    assert "DB error" in r.get_json()["error"]


def test_partir_y_guardar_async_worker_status_y_events(client, app_module, monkeypatch):
    # Worker OK
    job_ok = "cov9_partir_ok"
    with app_module._partir_guardar_jobs_lock:
        app_module._partir_guardar_jobs[job_ok] = {
            "status": "queued", "done": 0, "total": 100, "message": "",
            "result": None, "error": None, "http_status": None,
        }

    monkeypatch.setattr(app_module, "partir_y_guardar", lambda _id: app_module.jsonify(ok=True, hecho=True), raising=False)
    app_module._partir_guardar_async_worker(job_ok, 1, overwrite=True)

    st = client.get(f"/api/examenes/partir_y_guardar/jobs/{job_ok}")
    assert st.status_code == 200
    assert st.get_json()["status"] == "done"
    assert st.get_json()["result"]["hecho"] is True

    ev = client.get(f"/api/examenes/partir_y_guardar/jobs/{job_ok}/events")
    assert ev.status_code == 200
    assert b"event: progress" in ev.data

    # Worker con respuesta HTTP de error
    job_bad = "cov9_partir_bad"
    with app_module._partir_guardar_jobs_lock:
        app_module._partir_guardar_jobs[job_bad] = {
            "status": "queued", "done": 0, "total": 100, "message": "",
            "result": None, "error": None, "http_status": None,
        }
    monkeypatch.setattr(app_module, "partir_y_guardar", lambda _id: (app_module.jsonify(ok=False, error="mal"), 409), raising=False)
    app_module._partir_guardar_async_worker(job_bad, 2, overwrite=False)

    st_bad = client.get(f"/api/examenes/partir_y_guardar/jobs/{job_bad}")
    assert st_bad.status_code == 200
    assert st_bad.get_json()["ok"] is False
    assert st_bad.get_json()["http_status"] == 409

    # Job inexistente en status y events.
    assert client.get("/api/examenes/partir_y_guardar/jobs/nope").status_code == 404
    assert client.get("/api/examenes/partir_y_guardar/jobs/nope/events").status_code == 404
