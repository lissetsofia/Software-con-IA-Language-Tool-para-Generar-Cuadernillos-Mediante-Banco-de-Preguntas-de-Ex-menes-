import io
import os
import sqlite3
from pathlib import Path

from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _docx_bytes(text="Pregunta"):
    bio = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(bio)
    return bio.getvalue()


def _make_docx(path: Path, text="Pregunta"):
    path.write_bytes(_docx_bytes(text))
    return path


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov10_matriz.sqlite3"

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    upload_dir = tmp_path / "uploads"
    down_dir = tmp_path / "descargas"
    upload_dir.mkdir(parents=True, exist_ok=True); down_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setitem(app_module.app.config, "UPLOAD_FOLDER", str(upload_dir))
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(down_dir))

    conn = _connect(db_path)
    cur = conn.cursor()
    cur.executescript(
        """
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            activo INTEGER DEFAULT 1
        );
        CREATE TABLE IF NOT EXISTS matriz (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS matriz_detalle (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0,
            archivo_ruta TEXT,
            UNIQUE(matriz_id, tema_id)
        );
        """
    )
    cur.execute("INSERT INTO temario(nombre, activo) VALUES ('Álgebra', 1)")
    tema_id = int(cur.lastrowid)
    conn.commit(); cur.close(); conn.close()
    return db_path, tema_id, upload_dir, down_dir


def test_matriz_crear_listar_upload_y_generar(client, app_module, tmp_path, monkeypatch):
    _db_path, tema_id, _upload_dir, down_dir = _patch_db(app_module, tmp_path, monkeypatch)

    assert client.post("/api/matriz", json={"items": []}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": "x", "cantidad": 1}]}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": 0, "cantidad": 1}]}).status_code == 400

    r_create = client.post(
        "/api/matriz",
        json={"nombre": "Matriz Cov10", "items": [{"tema_id": tema_id, "cantidad": 2, "orden": 1}]},
    )
    assert r_create.status_code == 200
    matriz_id = r_create.get_json()["matriz_id"]

    r_list = client.get("/api/matriz?detail=1&search=Cov10")
    assert r_list.status_code == 200
    assert r_list.get_json()[0]["items"][0]["tema_id"] == tema_id

    assert client.get("/api/matriz/999999").status_code == 404
    assert client.post(f"/api/matriz/{matriz_id}/upload", data={}, content_type="multipart/form-data").status_code == 400
    assert client.post(
        f"/api/matriz/{matriz_id}/upload",
        data={"tema_id": str(tema_id), "file": (io.BytesIO(b"txt"), "mal.txt")},
        content_type="multipart/form-data",
    ).status_code == 400

    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: None, raising=False)
    r_upload = client.post(
        f"/api/matriz/{matriz_id}/upload",
        data={"tema_id": str(tema_id), "cantidad": "2", "file": (io.BytesIO(_docx_bytes("1. Pregunta")), "tema.docx")},
        content_type="multipart/form-data",
    )
    assert r_upload.status_code == 200
    ruta_subida = r_upload.get_json()["ruta"]
    assert os.path.exists(ruta_subida)

    assert client.post("/api/matriz/999999/generar").status_code == 404

    # Generación exitosa sin Word real: recorte y merge simulados.
    monkeypatch.setattr(app_module, "_contar_preguntas_docx", lambda *_a, **_k: 5, raising=False)
    monkeypatch.setattr(app_module, "_cut_docx_to_individual_question_docs", lambda src, n: [src], raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_post_merge_fix_numbering", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "bullets_to_numbers_docx", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge(grouped, out_path, *args, **kwargs):
        doc = DocxDocument()
        for tema, files in grouped:
            doc.add_paragraph(tema)
            for f in files:
                doc.add_paragraph(os.path.basename(str(f)))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    r_gen = client.post(f"/api/matriz/{matriz_id}/generar")
    assert r_gen.status_code == 200
    assert "application/vnd.openxmlformats" in r_gen.headers.get("Content-Type", "")


def test_matriz_generar_errores_de_archivo_cupo_y_merge(client, app_module, tmp_path, monkeypatch):
    db_path, tema_id, _upload_dir, _down_dir = _patch_db(app_module, tmp_path, monkeypatch)
    conn = _connect(db_path); cur = conn.cursor()
    cur.execute("INSERT INTO matriz(nombre) VALUES ('Matriz errores')")
    matriz_id = int(cur.lastrowid)
    missing = tmp_path / "faltante.docx"
    cur.execute(
        "INSERT INTO matriz_detalle(matriz_id, tema_id, cantidad, orden, archivo_ruta) VALUES (?, ?, 3, 1, ?)",
        (matriz_id, tema_id, str(missing)),
    )
    conn.commit(); cur.close(); conn.close()

    r_missing = client.post(f"/api/matriz/{matriz_id}/generar")
    assert r_missing.status_code == 400

    docx = _make_docx(tmp_path / "tema_ok.docx", "1. Pregunta")
    conn = _connect(db_path); cur = conn.cursor()
    cur.execute("UPDATE matriz_detalle SET archivo_ruta=? WHERE matriz_id=?", (str(docx), matriz_id))
    conn.commit(); cur.close(); conn.close()

    monkeypatch.setattr(app_module, "_contar_preguntas_docx", lambda *_a, **_k: 1, raising=False)
    r_cupo = client.post(f"/api/matriz/{matriz_id}/generar")
    assert r_cupo.status_code == 409

    monkeypatch.setattr(app_module, "_contar_preguntas_docx", lambda *_a, **_k: 5, raising=False)
    monkeypatch.setattr(app_module, "_cut_docx_to_individual_question_docs", lambda src, n: [src], raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_post_merge_fix_numbering", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "bullets_to_numbers_docx", lambda *_a, **_k: None, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge_bad(grouped, out_path, *args, **kwargs):
        _make_docx(Path(out_path), "parcial")
        return out_path, [], [(grouped[0][1][0], "falló")]

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge_bad, raising=False)
    r_malos = client.post(f"/api/matriz/{matriz_id}/generar")
    assert r_malos.status_code == 409
