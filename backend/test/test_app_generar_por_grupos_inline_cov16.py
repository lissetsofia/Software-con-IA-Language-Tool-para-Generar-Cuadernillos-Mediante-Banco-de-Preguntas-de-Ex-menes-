import os
import sqlite3
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _make_docx(path, text="Pregunta"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return str(path)


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov16_grupos_inline.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE gen_lote(id INTEGER PRIMARY KEY AUTOINCREMENT, matriz_id INTEGER, nombre TEXT, usuario TEXT);
        """
    )
    conn.execute("INSERT INTO temario(id,nombre,activo) VALUES(10,'Álgebra',1)")
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    out = tmp_path / "grupos_out"; out.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out), raising=False)
    return db_path, out


def _cfg(cantidad_a=1, cantidad_b=1, tema_id=10):
    return {
        1: {"idgrupo": 1, "clave": "A", "nombre": "Grupo A", "temas": [
            {"tema_id": tema_id, "tema_nombre": "Álgebra", "cantidad": cantidad_a, "orden": 1}
        ]},
        2: {"idgrupo": 2, "clave": "B", "nombre": "Grupo B", "temas": [
            {"tema_id": tema_id, "tema_nombre": "Álgebra", "cantidad": cantidad_b, "orden": 1}
        ]},
    }


def test_api_generar_por_grupos_inline_validaciones(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)

    assert client.post("/api/grupos/generar", json={}).status_code == 400
    assert client.post("/api/grupos/generar", json={"matriz": {"items": []}}).status_code == 400
    assert client.post("/api/grupos/generar", json={"matriz": {"items": [{"cantidad": 1}]}}).status_code == 400

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {}, raising=False)
    r_no_groups = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 1, "archivo_ruta": "x.docx"}]}})
    assert r_no_groups.status_code == 400
    assert "grupos" in r_no_groups.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _cfg(tema_id=99), raising=False)
    r_not_in_matrix = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 1, "archivo_ruta": "x.docx"}]}})
    assert r_not_in_matrix.status_code == 400
    assert "matriz" in r_not_in_matrix.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _cfg(1, 2), raising=False)
    r_diff = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 2, "archivo_ruta": "x.docx"}]}})
    assert r_diff.status_code == 400
    assert "cuotas" in r_diff.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _cfg(3, 3), raising=False)
    r_quota = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 2, "archivo_ruta": "x.docx"}]}})
    assert r_quota.status_code == 400
    assert "matriz solo tiene" in r_quota.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _cfg(1, 1), raising=False)
    r_no_file = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 1}]}})
    assert r_no_file.status_code == 400
    assert "docx" in r_no_file.get_json()["error"].lower()


def test_api_generar_por_grupos_inline_success_y_zip(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    src = _make_docx(tmp_path / "tema_algebra.docx", "1. Pregunta base")
    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _cfg(1, 1), raising=False)

    def fake_cut(src_docx, n, dest_docx):
        _make_docx(Path(dest_docx), f"recorte {n} de {os.path.basename(src_docx)}")

    def fake_merge(bloques, out_path):
        doc = Document()
        for tema, files in bloques:
            doc.add_paragraph(f"TEMA: {tema}")
            for f in files:
                doc.add_paragraph(os.path.basename(f))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_cut_docx_first_n_questions", fake_cut, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    r = client.post("/api/grupos/generar", json={"matriz": {"nombre": "Matriz cov16", "items": [{"tema_id": 10, "cantidad": 1, "archivo_ruta": src}]}})
    assert r.status_code == 200
    data = r.get_json()
    assert data["ok"] is True
    assert data["zip_url"].startswith("/api/grupos/lote/")
    lote_id = data["lote_id"]
    assert (tmp_path / "grupos_out" / f"lote_{lote_id}" / f"grupos_{lote_id}.zip").exists()

    # Rama de error después del recorte: merge reporta archivos malos.
    def merge_malos(bloques, out_path):
        _make_docx(Path(out_path), "doc incompleto")
        return out_path, [], [("x.docx", "falló insert")]

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", merge_malos, raising=False)
    r_bad_merge = client.post("/api/grupos/generar", json={"matriz": {"items": [{"tema_id": 10, "cantidad": 1, "archivo_ruta": src}]}})
    assert r_bad_merge.status_code == 409
    assert r_bad_merge.get_json()["ok"] is False
