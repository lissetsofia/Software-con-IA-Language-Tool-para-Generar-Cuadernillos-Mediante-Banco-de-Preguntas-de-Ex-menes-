from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


def _make_docx_with_math(path):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("Haber, revisa el numero entera")
    doc.add_paragraph("A) 2 alternativa correcta")
    doc.add_paragraph("1. pregunta antigua")
    p_math = doc.add_paragraph()
    run = p_math.add_run("valor viejo")
    # Inserta un nodo OMML para forzar la rama de párrafo con objetos.
    try:
        run._r.append(OxmlElement("m:oMath"))
    except Exception:
        pass
    doc.save(path)
    return path


def test_generar_docx_corregido_highlight_objetos_y_opciones(app_module, tmp_path, monkeypatch):
    src = _make_docx_with_math(tmp_path / "original.docx")
    out = tmp_path / "corregido.docx"

    texto_corregido = "\n".join([
        "A ver, revisa el número entero",
        "A) 2 alternativa correcta",
        "1. pregunta nueva",
        "valor nuevo",
    ])

    app_module.generar_docx_corregido(str(src), texto_corregido, str(out), highlight=True)
    assert out.exists() and out.stat().st_size > 0

    doc = Document(out)
    textos = [p.text for p in doc.paragraphs]
    assert "A ver" in textos[0]
    # La alternativa conserva la etiqueta y la firma numérica.
    assert textos[1].startswith("A) 2")


def test_generar_docx_corregido_no_toca_alternativa_insegura(app_module, tmp_path):
    src = tmp_path / "alt.docx"
    doc = Document()
    doc.add_paragraph("A) 123 alternativa original")
    doc.save(src)

    out = tmp_path / "alt_out.docx"
    # Cambia la firma numérica; por seguridad no debería reescribir la alternativa.
    app_module.generar_docx_corregido(str(src), "A) 999 alternativa cambiada", str(out), highlight=False)
    texto = Document(out).paragraphs[0].text
    assert "123" in texto


def test_reemplazos_runs_parciales_flexible_y_diff_helpers(app_module):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("geo")
    p.add_run("metria")
    assert app_module.reemplazo_en_runs_parciales(p, "geo", "geo") is True
    assert app_module.reemplazo_en_runs_flexible(p, "geometria", "geometría") is True
    pares = app_module.pares_reemplazo_palabra_a_palabra("numero entero", "número entero")
    assert ("numero", "número") in pares
    assert app_module.similitud_suave("casa", "casa") == 1.0
