import os
import shutil
import zipfile
from pathlib import Path
import xml.etree.ElementTree as ET

from docx import Document

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{%s}" % NS_W


def _make_docx(path, paragraphs=("1. Pregunta", "A) alt", "2. Pregunta")):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)
    return path


def _write_xml(path, text):
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(text, encoding="utf-8")


def _document_xml():
    return f'''<w:document xmlns:w="{NS_W}"><w:body>
<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>P1</w:t></w:r></w:p>
<w:p><w:r><w:t>detalle p1</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="2"/></w:numPr></w:pPr><w:r><w:t>Alternativa</w:t></w:r></w:p>
<w:p><w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr></w:pPr><w:r><w:t>P2</w:t></w:r></w:p>
<w:sectPr/>
</w:body></w:document>'''


def _numbering_xml():
    return f'''<w:numbering xmlns:w="{NS_W}">
<w:abstractNum w:abstractNumId="1"><w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/></w:lvl></w:abstractNum>
<w:abstractNum w:abstractNumId="2"><w:lvl w:ilvl="0"><w:numFmt w:val="upperLetter"/></w:lvl></w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="1"/></w:num>
<w:num w:numId="2"><w:abstractNumId w:val="2"/></w:num>
</w:numbering>'''


def test_find_question_spans_y_contar_preguntas_docx(app_module, tmp_path):
    doc_xml = tmp_path / "document.xml"
    num_xml = tmp_path / "numbering.xml"
    _write_xml(doc_xml, _document_xml())
    _write_xml(num_xml, _numbering_xml())

    spans = app_module._find_question_spans(str(doc_xml), str(num_xml), 5)
    assert spans == [(0, 3), (3, 4)]
    assert app_module._find_question_spans(str(doc_xml), str(num_xml), 1) == [(0, 3)]

    no_num = tmp_path / "no_numbering.xml"
    _write_xml(no_num, "<w:numbering xmlns:w='%s'/>" % NS_W)
    assert app_module._find_question_spans(str(doc_xml), str(no_num), 5) == []

    docx = _make_docx(tmp_path / "count.docx")
    with zipfile.ZipFile(docx, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}
    files["word/document.xml"] = _document_xml().encode("utf-8")
    files["word/numbering.xml"] = _numbering_xml().encode("utf-8")
    tmp = str(docx) + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, b in files.items():
            zout.writestr(n, b)
    os.replace(tmp, docx)

    assert app_module._contar_preguntas_docx(str(docx)) == 2
    assert app_module.debug_contar_preguntas_docx(str(docx)) == 2


def test_cut_docx_first_n_y_individuales_con_fallbacks(app_module, tmp_path, monkeypatch):
    src = _make_docx(tmp_path / "src.docx", ("P1", "detalle", "P2"))

    # n<=0 crea un DOCX vacío válido.
    out_empty = tmp_path / "empty.docx"
    app_module._cut_docx_first_n_questions(str(src), 0, str(out_empty))
    assert out_empty.exists()

    # Si no detecta spans, copia el documento original.
    monkeypatch.setattr(app_module, "_find_question_spans", lambda *_a, **_k: [], raising=False)
    out_copy = tmp_path / "copy.docx"
    app_module._cut_docx_first_n_questions(str(src), 2, str(out_copy))
    assert out_copy.exists() and out_copy.stat().st_size > 0

    # Spans individuales: evitamos depender del reempaque real y comprobamos que devuelve 2 archivos.
    monkeypatch.setattr(app_module, "_find_question_spans", lambda *_a, **_k: [(0, 2), (2, 3)], raising=False)

    def fake_reempacar(work_dir, seleccion, destino):
        doc = Document()
        doc.add_paragraph(f"seleccion={len(seleccion)}")
        Path(destino).parent.mkdir(parents=True, exist_ok=True)
        doc.save(destino)

    monkeypatch.setattr(app_module, "_reempacar_docx", fake_reempacar, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)
    outs = app_module._cut_docx_to_individual_question_docs(str(src), 2)
    assert len(outs) == 2
    assert all(os.path.exists(p) for p in outs)
    for p in outs:
        try:
            os.remove(p)
        except Exception:
            pass


def test_cut_docx_first_n_error_controlado(app_module, tmp_path, monkeypatch):
    bad = tmp_path / "bad.docx"
    bad.write_text("no zip", encoding="utf-8")
    try:
        app_module._cut_docx_first_n_questions(str(bad), 1, str(tmp_path / "out.docx"))
        assert False, "Debió encapsular el error de DOCX"
    except RuntimeError as e:
        assert "error procesando docx" in str(e).lower()
