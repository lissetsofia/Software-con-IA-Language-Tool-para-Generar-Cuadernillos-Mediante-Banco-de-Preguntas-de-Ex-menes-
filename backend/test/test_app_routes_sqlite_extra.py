# backend/test/test_app_routes_sqlite_extra.py
import io
import os
import sqlite3
from pathlib import Path

import pytest
from docx import Document


SCHEMA = """
CREATE TABLE IF NOT EXISTS usuarios (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT NOT NULL,
    password TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS sesiones_app (
    token TEXT PRIMARY KEY,
    username TEXT NOT NULL
);
CREATE TABLE IF NOT EXISTS examenes (
    idexamenes INTEGER PRIMARY KEY AUTOINCREMENT,
    nombre TEXT,
    numero TEXT,
    institucion TEXT,
    anio INTEGER,
    archivo_nombre TEXT,
    archivo_ruta TEXT
);
CREATE TABLE IF NOT EXISTS temario (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nombre TEXT NOT NULL,
    activo INTEGER DEFAULT 1
);
CREATE TABLE IF NOT EXISTS preguntas (
    idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT,
    examenes_idexamenes INTEGER,
    tema_id INTEGER,
    numero_p INTEGER,
    archivo_nombre TEXT,
    archivo_ruta TEXT
);
CREATE TABLE IF NOT EXISTS grupos (
    idgrupo INTEGER PRIMARY KEY AUTOINCREMENT,
    clave TEXT NOT NULL UNIQUE,
    nombre TEXT,
    activo INTEGER DEFAULT 1,
    fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
);
CREATE TABLE IF NOT EXISTS grupo_tema (
    idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT,
    grupos_idgrupo INTEGER,
    tema_id INTEGER,
    cantidad INTEGER,
    orden INTEGER DEFAULT 0,
    UNIQUE(grupos_idgrupo, tema_id)
);
CREATE TABLE IF NOT EXISTS matriz (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nombre TEXT,
    fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
);
CREATE TABLE IF NOT EXISTS matriz_detalle (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    matriz_id INTEGER,
    tema_id INTEGER,
    cantidad INTEGER,
    orden INTEGER DEFAULT 0,
    archivo_ruta TEXT,
    UNIQUE(matriz_id, tema_id)
);
CREATE TABLE IF NOT EXISTS tema_docs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    tema_id INTEGER,
    doc_preguntas_nombre TEXT,
    doc_preguntas_ruta TEXT,
    doc_sol_ruta TEXT
);
"""


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


@pytest.fixture()
def sqlite_db(tmp_path, app_module, monkeypatch):
    db_path = tmp_path / "evalunia_test.sqlite3"
    conn = _connect(db_path)
    conn.executescript(SCHEMA)
    conn.commit()
    conn.close()

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    return db_path


def _insert_seed_data(db_path: Path, docx_path: str | None = None):
    conn = _connect(db_path)
    cur = conn.cursor()
    cur.execute("INSERT INTO usuarios(username, password) VALUES (?,?)", ("admin", "1234"))
    cur.execute("INSERT INTO temario(nombre, activo) VALUES (?,1)", ("Álgebra",))
    cur.execute("INSERT INTO temario(nombre, activo) VALUES (?,0)", ("Historia",))
    cur.execute("INSERT INTO grupos(clave, nombre, activo) VALUES ('A','Grupo A',1)")
    cur.execute("INSERT INTO grupo_tema(grupos_idgrupo, tema_id, cantidad, orden) VALUES (1,1,10,1)")
    if docx_path:
        cur.execute(
            "INSERT INTO examenes(nombre, numero, institucion, anio, archivo_nombre, archivo_ruta) VALUES (?,?,?,?,?,?)",
            ("Examen Ordinario", "I", "UNAMBA", 2023, os.path.basename(docx_path), docx_path),
        )
        cur.execute(
            "INSERT INTO preguntas(examenes_idexamenes, tema_id, numero_p, archivo_nombre, archivo_ruta) VALUES (1,1,1,?,?)",
            (os.path.basename(docx_path), docx_path),
        )
    conn.commit()
    conn.close()


def _make_docx(path: Path, paragraphs=("Pregunta uno", "A) Uno", "B) Dos")):
    doc = Document()
    for txt in paragraphs:
        doc.add_paragraph(txt)
    doc.save(path)
    return path


# -------------------------
# Auth y sesión reales
# -------------------------
def test_login_session_logout_con_sqlite(client, app_module, sqlite_db, monkeypatch):
    _insert_seed_data(sqlite_db)
    monkeypatch.setattr(app_module.secrets, "token_urlsafe", lambda n: "tok-prueba")

    r = client.post("/login", json={"usuario": "admin", "clave": "1234"})
    assert r.status_code == 200
    assert r.get_json()["token"] == "tok-prueba"

    r_sess = client.get("/api/session", headers={"Authorization": "Bearer tok-prueba"})
    assert r_sess.status_code == 200
    assert r_sess.get_json()["usuario"] == "admin"

    r_logout = client.post("/logout", headers={"Authorization": "Bearer tok-prueba"})
    assert r_logout.status_code == 200

    r_sess2 = client.get("/api/session", headers={"Authorization": "Bearer tok-prueba"})
    assert r_sess2.status_code == 401


def test_login_invalido_y_probar_conexion(client, sqlite_db):
    _insert_seed_data(sqlite_db)
    r = client.post("/login", json={"usuario": "admin", "clave": "mala"})
    assert r.status_code == 401

    r2 = client.get("/probar-conexion")
    assert r2.status_code == 200
    assert r2.get_json()["conexion"] == "ok"


# -------------------------
# Exámenes y preguntas
# -------------------------
def test_examenes_listar_nombre_exportar_word_y_preguntas(client, sqlite_db, tmp_path):
    docx = _make_docx(tmp_path / "examen.docx")
    _insert_seed_data(sqlite_db, str(docx))

    r = client.get("/api/examenes")
    assert r.status_code == 200
    assert r.get_json()[0]["institucion"] == "UNAMBA"

    r_nombre = client.get("/api/examen_nombre/1")
    assert r_nombre.status_code == 200
    assert r_nombre.get_json()["archivo_nombre"] == "examen.docx"

    r_pregs = client.get("/api/preguntas?examen=1&tema=1")
    assert r_pregs.status_code == 200
    assert r_pregs.get_json()[0]["numero_p"] == 1

    r_export = client.get("/api/exportar_examen/1?formato=word")
    assert r_export.status_code == 200
    assert r_export.headers["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_importar_examen_validaciones_y_ok(client, sqlite_db):
    r_bad = client.post("/api/importar_examen", data={}, content_type="multipart/form-data")
    assert r_bad.status_code == 400

    data = {
        "archivo": (
            io.BytesIO(b"contenido simulado"),
            "examen ordinario I UNAMBA 2023.docx",
        )
    }
    r = client.post("/api/importar_examen", data=data, content_type="multipart/form-data")
    assert r.status_code == 200
    assert r.get_json()["exito"] is True

    # Mismo archivo: debe detectar duplicado por archivo_nombre.
    data_dup = {
        "archivo": (
            io.BytesIO(b"contenido simulado"),
            "examen ordinario I UNAMBA 2023.docx",
        )
    }
    r_dup = client.post("/api/importar_examen", data=data_dup, content_type="multipart/form-data")
    assert r_dup.status_code == 400


def test_eliminar_examen_borra_bd_y_archivo(client, sqlite_db, tmp_path):
    docx = _make_docx(tmp_path / "a_borrar.docx")
    _insert_seed_data(sqlite_db, str(docx))

    r = client.delete("/api/examenes/1")
    assert r.status_code == 200
    assert "eliminados" in r.get_json()["mensaje"].lower()
    assert not docx.exists()

    r404 = client.delete("/api/examenes/999")
    assert r404.status_code == 404


# -------------------------
# Temas y grupos reales
# -------------------------
def test_temas_crud_con_sqlite(client, sqlite_db):
    _insert_seed_data(sqlite_db)

    r_list = client.get("/api/temas")
    assert r_list.status_code == 200
    assert all(x["activo"] == 1 for x in r_list.get_json())

    r_all = client.get("/api/temas?all=1")
    assert r_all.status_code == 200
    assert len(r_all.get_json()) >= 2

    assert client.post("/api/temas", json={"nombre": ""}).status_code == 400
    assert client.post("/api/temas", json={"nombre": "Álgebra"}).status_code == 409

    r_new = client.post("/api/temas", json={"nombre": "Geometría"})
    assert r_new.status_code == 201
    new_id = r_new.get_json()["id"]

    r_edit = client.put(f"/api/temas/{new_id}", json={"nombre": "Geometría Analítica"})
    assert r_edit.status_code == 200

    r_toggle = client.patch(f"/api/temas/{new_id}/toggle")
    assert r_toggle.status_code == 200

    # Tema 1 tiene preguntas/cupos: sin force bloquea; con force elimina.
    conn = _connect(sqlite_db)
    conn.execute("INSERT INTO preguntas(examenes_idexamenes, tema_id, numero_p) VALUES (1,1,1)")
    conn.commit(); conn.close()
    assert client.delete("/api/temas/1").status_code == 409
    assert client.delete("/api/temas/1?force=1").status_code == 200


def test_grupos_crud_y_cuotas_con_sqlite(client, sqlite_db):
    _insert_seed_data(sqlite_db)

    r_get = client.get("/api/grupos")
    assert r_get.status_code == 200
    assert len(r_get.get_json()) >= 1

    assert client.post("/api/grupos", json={}).status_code == 400
    assert client.post("/api/grupos", json={"clave": "A", "nombre": "Duplicado"}).status_code == 409

    r_new = client.post(
        "/api/grupos",
        json={"clave": " b123456 ", "nombre": "Grupo B", "cuotas": [{"tema_id": 1, "cantidad": 5}]},
    )
    assert r_new.status_code == 200
    idgrupo = r_new.get_json()["idgrupo"]

    r_cuotas = client.get(f"/api/grupos/{idgrupo}/cuotas")
    assert r_cuotas.status_code == 200
    assert r_cuotas.get_json()[0]["cantidad"] == 5

    r_put = client.put(
        f"/api/grupos/{idgrupo}/cuotas",
        json={"cuotas": [{"tema_id": 1, "cantidad": 9, "orden": 3}]},
    )
    assert r_put.status_code == 200
    assert r_put.get_json()["status"] == "ok"

    r_edit = client.put(f"/api/grupos/{idgrupo}", json={"nombre": "Grupo B actualizado", "activo": 0})
    assert r_edit.status_code == 200

    assert client.patch(f"/api/grupos/{idgrupo}/toggle").status_code == 200
    assert client.delete(f"/api/grupos/{idgrupo}").status_code == 409
    assert client.delete(f"/api/grupos/{idgrupo}?force=1").status_code == 200


# -------------------------
# Matriz: cabecera, detalle, upload y generación con error controlado
# -------------------------
def test_matriz_crear_listar_get_upload_y_generar_faltante(client, app_module, sqlite_db, monkeypatch, tmp_path):
    _insert_seed_data(sqlite_db)
    monkeypatch.setattr(app_module, "_validar_docx_real", lambda path: None, raising=False)

    assert client.post("/api/matriz", json={"nombre": "M1", "items": []}).status_code == 400

    r_create = client.post(
        "/api/matriz",
        json={"nombre": "M1", "items": [{"tema_id": 1, "cantidad": 2, "orden": 1}]},
    )
    assert r_create.status_code == 200
    matriz_id = r_create.get_json()["matriz_id"]

    r_list = client.get("/api/matriz?detail=1")
    assert r_list.status_code == 200
    assert r_list.get_json()[0]["items"][0]["tema_id"] == 1

    r_get = client.get(f"/api/matriz/{matriz_id}")
    assert r_get.status_code == 200
    assert r_get.get_json()["n_items"] == 1

    r_upload_bad = client.post(f"/api/matriz/{matriz_id}/upload", data={})
    assert r_upload_bad.status_code == 400

    r_upload = client.post(
        f"/api/matriz/{matriz_id}/upload",
        data={"tema_id": "1", "cantidad": "2", "file": (io.BytesIO(b"x"), "tema.docx")},
        content_type="multipart/form-data",
    )
    assert r_upload.status_code == 200
    assert r_upload.get_json()["ok"] is True

    # Crear otra matriz sin archivo para cubrir la rama de error 'Falta subir DOCX'.
    r_create2 = client.post(
        "/api/matriz",
        json={"nombre": "M2", "items": [{"tema_id": 1, "cantidad": 1, "orden": 1}]},
    )
    matriz2 = r_create2.get_json()["matriz_id"]
    r_gen = client.post(f"/api/matriz/{matriz2}/generar")
    assert r_gen.status_code == 400
    assert "Falta subir DOCX" in r_gen.get_json()["error"]


# -------------------------
# Corrector: validaciones, preview y corrección simulada sin LanguageTool real
# -------------------------
def test_corregir_archivo_validaciones_preview_y_corregir_simulado(client, app_module, sqlite_db, monkeypatch, tmp_path):
    _insert_seed_data(sqlite_db)

    assert client.post("/api/corregir_archivo", data={}, content_type="multipart/form-data").status_code == 400

    bad = client.post(
        "/api/corregir_archivo",
        data={"archivo": (io.BytesIO(b"txt"), "archivo.txt")},
        content_type="multipart/form-data",
    )
    assert bad.status_code == 400

    doc_path = tmp_path / "corrector.docx"
    _make_docx(doc_path, paragraphs=("numero entero", "A) numero sin tocar"))
    raw = doc_path.read_bytes()

    monkeypatch.setattr(app_module, "detectar_indices_alternativas_por_numid", lambda doc, active_q_numId=None: ({1}, "1"), raising=False)
    monkeypatch.setattr(app_module, "lt_check_smart", lambda texto, lang="es": {"matches": []}, raising=False)

    def fake_generar_docx_corregido(path_in, texto_corregido, path_salida_docx, highlight=False, texto_original_para_highlight=None):
        out = Document()
        out.add_paragraph(texto_corregido or "ok")
        out.save(path_salida_docx)

    monkeypatch.setattr(app_module, "generar_docx_corregido", fake_generar_docx_corregido, raising=False)

    preview = client.post(
        "/api/corregir_archivo",
        data={"modo": "preview", "archivo": (io.BytesIO(raw), "corrector.docx")},
        content_type="multipart/form-data",
    )
    assert preview.status_code == 200
    assert preview.get_json()["ok"] is True
    assert preview.get_json()["total_alertas"] == 0

    corregir = client.post(
        "/api/corregir_archivo",
        data={"modo": "corregir", "archivo": (io.BytesIO(raw), "corrector.docx")},
        content_type="multipart/form-data",
    )
    assert corregir.status_code == 200
    payload = corregir.get_json()
    assert payload["ok"] is True
    assert payload["descargas"]["docx"].endswith("_corregido_limpio.docx")
