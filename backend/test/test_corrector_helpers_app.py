# backend/test/test_corrector_helpers_app.py
import zipfile
from docx import Document


def test_normalize_ocr_noise_limpia_espacios_y_guiones(app_module):
    texto = "núme-\nro\u00a0\u200b  entero"
    assert app_module.normalize_ocr_noise(texto) == "número entero"


def test_utilidades_de_texto_basicas(app_module):
    assert app_module._strip_accents("Árbol") == "Arbol"
    assert app_module._same_casing("HOLA", "mundo") == "MUNDO"
    assert app_module._same_casing("Hola", "mundo") == "Mundo"
    assert app_module._prev_token("El número", 3) == "El"
    assert app_module.is_upper_acronym("UNAMBA")
    assert app_module.is_upper_acronym("(DNI)")
    assert not app_module.is_upper_acronym("Casa")
    assert app_module._has_any_digit("x²") is True
    assert app_module._has_any_digit("x2") is True
    assert app_module._edit_distance("casa", "caso") == 1


def test_choose_best_suggestion_prefiere_acentos(app_module):
    sug = [{"value": "número"}, {"value": "numero"}]
    assert app_module._choose_best_suggestion("numero", sug, "") == "número"


def test_looks_reasonable_replacement(app_module):
    assert app_module._looks_reasonable_replacement("numero", "número")
    # La función es estricta: solo acepta cambios muy seguros, como acentos.
    assert not app_module._looks_reasonable_replacement("casa", "caso")
    assert not app_module._looks_reasonable_replacement("abc", "universidad")


def test_spans_de_alternativas_y_restauracion(app_module):
    texto = "Pregunta\nA) Uno\nB) Dos\nTexto final"
    spans = app_module.detectar_spans_alternativas(texto)
    assert len(spans) == 2
    assert app_module.intersecta_spans(spans[0][0], 3, spans)

    corregido = "Pregunta\nA) UNO\nB) DOS\nTexto final"
    restaurado = app_module.restaurar_segmentos_protegidos(texto, corregido, spans)
    assert "A) Uno" in restaurado
    assert "B) Dos" in restaurado


def test_apply_lt_corrections_classic_corrige_y_respeta_protegidos(app_module):
    texto = "El numero entero\nA) numero incorrecto"
    matches = [
        {"offset": 3, "length": 6, "replacements": [{"value": "número"}]},
        {
            "offset": texto.index("numero incorrecto"),
            "length": 6,
            "replacements": [{"value": "número"}],
        },
    ]
    spans = app_module.detectar_spans_alternativas(texto)
    out = app_module.apply_lt_corrections_classic(texto, matches, protected_spans=spans)
    assert out.startswith("El número")
    assert "A) numero incorrecto" in out


def test_apply_lt_corrections_omite_digitos_siglas_y_variables(app_module):
    texto = "UNAMBA x X2"
    matches = [
        {"offset": 0, "length": 6, "replacements": [{"value": "Unamba"}]},
        {"offset": 7, "length": 1, "replacements": [{"value": "equis"}]},
        {"offset": 9, "length": 2, "replacements": [{"value": "Xdos"}]},
    ]
    assert app_module.apply_lt_corrections(texto, matches) == texto


def test_post_correcciones_reglas_contextuales(app_module):
    out = app_module.post_correcciones("de el numero entera. haber si funciona?")
    assert "del número entero" in out.lower()
    assert "a ver si" in out.lower()
    assert out.endswith("?")


def test_tokenizacion_y_pares_de_reemplazo(app_module):
    toks = app_module.tokenize_preservando("hola, mundo")
    assert "hola" in toks
    assert any("," in tok for tok in toks)
    assert app_module.es_token_palabra("hola")

    pares = app_module.pares_reemplazo_palabra_a_palabra("numero entero", "número entero")
    assert ("numero", "número") in pares
    assert app_module.similitud_suave("casa", "casa") == 1


def test_generar_diff_html_contiene_columnas(app_module):
    html = app_module.generar_diff_html("hola", "hola mundo")
    assert "Original" in html
    assert "Corregido" in html


def test_extraer_y_crear_docx_desde_texto(app_module, tmp_path):
    doc = app_module.crear_docx_desde_texto("uno\ndos")
    path = tmp_path / "texto.docx"
    doc.save(path)
    assert app_module.extraer_texto_docx(str(path)) == "uno\ndos"


def test_detectar_indices_alternativas_por_numid_fallback_docx(app_module, tmp_path):
    doc = Document()
    doc.add_paragraph("Pregunta extensa con suficiente texto para detectar como pregunta 1")
    doc.add_paragraph("A) Alternativa uno")
    doc.add_paragraph("B) Alternativa dos")
    doc.add_paragraph("Otra explicación")
    path = tmp_path / "alts.docx"
    doc.save(path)

    spans = app_module.detectar_spans_alternativas_docx(str(path), active_q_numId="1")
    assert len(spans) == 2
