# backend/test/test_app_generar_from_docx_cov7.py
import io
import os
import sqlite3
import zipfile
from pathlib import Path

from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov7_from_docx.sqlite3"

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    return db_path


def _schema(app_module):
    conn = app_module.get_connection()
    cur = conn.cursor()
    cur.execute("""CREATE TABLE IF NOT EXISTS grupos (
        idgrupo INTEGER PRIMARY KEY AUTOINCREMENT, clave TEXT, nombre TEXT, activo INTEGER DEFAULT 1
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS temario (
        id INTEGER PRIMARY KEY AUTOINCREMENT, nombre TEXT, activo INTEGER DEFAULT 1
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS grupo_tema (
        idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT,
        grupos_idgrupo INTEGER, tema_id INTEGER, cantidad INTEGER DEFAULT 0, orden INTEGER DEFAULT 0
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS gen_lote (
        id INTEGER PRIMARY KEY AUTOINCREMENT, matriz_id INTEGER, nombre TEXT, usuario TEXT
    )""")
    conn.commit()
    return conn, cur


def _insert_config(app_module, *, tema="ÁLGEBRA", cantidad=1, clave="A"):
    conn, cur = _schema(app_module)
    cur.execute("INSERT INTO temario (nombre, activo) VALUES (?, 1)", (tema,))
    tid = int(cur.lastrowid)
    cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES (?, ?, 1)", (clave, f"Grupo {clave}"))
    gid = int(cur.lastrowid)
    cur.execute("INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad, orden) VALUES (?, ?, ?, 1)", (gid, tid, cantidad))
    conn.commit(); cur.close(); conn.close()
    return gid, tid


def _make_matrix_docx(path: Path, *temas):
    doc = DocxDocument()
    for tema in temas:
        doc.add_paragraph(tema)
        doc.add_paragraph(f"1) Pregunta de {tema}")
        doc.add_paragraph("A) alternativa")
    doc.save(path)
    return path


def _patch_generation(app_module, tmp_path, monkeypatch):
    out_root = tmp_path / "grupos_docx_out"
    out_root.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out_root), raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None)

    def fake_merge(grouped, out_path, merge_step_cb=None, merge_ops=None):
        doc = DocxDocument()
        for titulo, files in grouped:
            doc.add_paragraph(titulo)
            for f in files:
                doc.add_paragraph(os.path.basename(str(f)))
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge)
    return out_root


def test_generar_from_docx_success_y_descargar_zip(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)
    _patch_generation(app_module, tmp_path, monkeypatch)
    _insert_config(app_module, tema="ÁLGEBRA", cantidad=1, clave="A")

    matrix = _make_matrix_docx(tmp_path / "matriz.docx", "ÁLGEBRA", "GEOMETRÍA")
    data = {"file": (io.BytesIO(matrix.read_bytes()), "matriz.docx")}
    r = client.post("/api/grupos/generar_from_docx", data=data, content_type="multipart/form-data")
    assert r.status_code == 200
    js = r.get_json()
    assert js["ok"] is True
    assert js["zip_url"].startswith("/api/grupos/lote/")

    rz = client.get(js["zip_url"])
    assert rz.status_code == 200
    assert rz.data.startswith(b"PK")
    with zipfile.ZipFile(io.BytesIO(rz.data)) as z:
        assert z.namelist()


def test_generar_from_docx_validaciones(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)

    assert client.post("/api/grupos/generar_from_docx", data={}, content_type="multipart/form-data").status_code == 400
    r_txt = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(b"hola"), "matriz.txt")},
        content_type="multipart/form-data",
    )
    assert r_txt.status_code == 400

    matrix = _make_matrix_docx(tmp_path / "matriz_sin_grupos.docx", "ÁLGEBRA")
    r_no_group = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(matrix.read_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_no_group.status_code == 400
    assert "grupos activos" in r_no_group.get_json()["error"].lower()


def test_generar_from_docx_faltan_titulos(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)
    _insert_config(app_module, tema="TRIGONOMETRÍA", cantidad=1, clave="A")

    matrix = _make_matrix_docx(tmp_path / "matriz_faltante.docx", "ÁLGEBRA")
    r = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(matrix.read_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r.status_code == 400
    assert "faltan títulos" in r.get_json()["error"].lower()


def test_helpers_heading_ranges_y_extract_range(app_module, tmp_path):
    src = _make_matrix_docx(tmp_path / "rangos.docx", "ÁLGEBRA", "GEOMETRÍA")
    ranges = app_module._build_heading_ranges_for_temas(str(src), {"ALGEBRA", "GEOMETRIA"})
    assert "ALGEBRA" in ranges and "GEOMETRIA" in ranges

    out = tmp_path / "solo_algebra.docx"
    start, end = ranges["ALGEBRA"]
    app_module._extract_tema_docx_range(str(src), start + 1, end, str(out))
    doc = DocxDocument(out)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Pregunta de ÁLGEBRA" in texto
    assert "GEOMETRÍA" not in texto
