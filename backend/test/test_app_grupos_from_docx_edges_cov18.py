import io
import os
import sqlite3
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov18_from_docx.sqlite3"
    conn = _connect(db_path)
    conn.executescript("""
        CREATE TABLE gen_lote(id INTEGER PRIMARY KEY AUTOINCREMENT, matriz_id INTEGER, nombre TEXT, usuario TEXT);
    """)
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    out = tmp_path / "grupos_from_docx_out"
    out.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out), raising=False)
    return db_path


def _docx_bytes(*paragraphs):
    bio = io.BytesIO()
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(bio)
    bio.seek(0)
    return bio


def _cfg(*rels):
    return {1: {"idgrupo": 1, "clave": "A", "nombre": "Grupo A", "temas": list(rels)}}


def test_generar_from_docx_validaciones_archivo_y_config(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)

    assert client.post("/api/grupos/generar_from_docx", data={}, content_type="multipart/form-data").status_code == 400

    r_ext = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(b"x"), "matriz.txt")},
        content_type="multipart/form-data",
    )
    assert r_ext.status_code == 400

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {}, raising=False)
    r_no_groups = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_no_groups.status_code == 400
    assert "grupos" in r_no_groups.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _cfg({"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 0, "orden": 1}),
        raising=False,
    )
    r_no_cuotas = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_no_cuotas.status_code == 400
    assert "cuotas" in r_no_cuotas.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: {
            1: {"idgrupo": 1, "clave": "A", "nombre": "A", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1}]},
            2: {"idgrupo": 2, "clave": "B", "nombre": "B", "temas": [{"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 2}]},
        },
        raising=False,
    )
    r_inconsistente = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_inconsistente.status_code == 400
    assert "no son iguales" in r_inconsistente.get_json()["error"].lower()


def test_generar_from_docx_missing_extract_merge_y_success(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _cfg({"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}),
        raising=False,
    )

    r_missing = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("GEOMETRÍA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_missing.status_code == 400
    assert "faltan" in r_missing.get_json()["error"].lower()

    monkeypatch.setattr(app_module, "_extract_tema_docx_range", lambda *_a, **_k: (_ for _ in ()).throw(RuntimeError("extract falla")), raising=False)
    r_extract = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_extract.status_code == 500
    assert "extracción" in r_extract.get_json()["error"].lower()

    def fake_extract(src, start, end, dest):
        doc = Document()
        doc.add_paragraph("pregunta extraída")
        Path(dest).parent.mkdir(parents=True, exist_ok=True)
        doc.save(dest)

    def merge_malos(bloques, out_path):
        doc = Document(); doc.add_paragraph("parcial"); Path(out_path).parent.mkdir(parents=True, exist_ok=True); doc.save(out_path)
        return out_path, [], [("x.docx", "fallo merge")]

    monkeypatch.setattr(app_module, "_extract_tema_docx_range", fake_extract, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", merge_malos, raising=False)
    r_merge = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_merge.status_code == 409

    def merge_ok(bloques, out_path):
        doc = Document()
        for tema, files in bloques:
            doc.add_paragraph(tema)
            doc.add_paragraph(os.path.basename(files[0]))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", merge_ok, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)
    r_ok = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (_docx_bytes("ÁLGEBRA", "1. Pregunta"), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    assert r_ok.get_json()["ok"] is True
