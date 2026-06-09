import io
import os
import sqlite3
import zipfile
from pathlib import Path
from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _docx_bytes(text="1. Pregunta"):
    bio = io.BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(bio)
    bio.seek(0)
    return bio


def _patch_matriz(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov14_matriz.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE matriz(id INTEGER PRIMARY KEY AUTOINCREMENT, nombre TEXT, fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE matriz_detalle(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0,
            archivo_ruta TEXT,
            UNIQUE(matriz_id, tema_id)
        );
        CREATE TABLE grupos(idgrupo INTEGER PRIMARY KEY, clave TEXT, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE grupo_tema(idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT, grupos_idgrupo INTEGER, tema_id INTEGER, cantidad INTEGER, orden INTEGER DEFAULT 0);
        CREATE TABLE preguntas(idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT, tema_id INTEGER, enunciado TEXT, alternativa_a TEXT, alternativa_b TEXT, alternativa_c TEXT, alternativa_d TEXT);
        """
    )
    conn.execute("INSERT INTO temario(id,nombre,activo) VALUES(1,'Álgebra',1)")
    conn.execute("INSERT INTO temario(id,nombre,activo) VALUES(2,'Física',1)")
    conn.execute("INSERT INTO grupos(idgrupo,clave,nombre,activo) VALUES(1,'A','Ingenierías',1)")
    conn.execute("INSERT INTO grupo_tema(grupos_idgrupo,tema_id,cantidad,orden) VALUES(1,1,1,2)")
    conn.execute("INSERT INTO grupo_tema(grupos_idgrupo,tema_id,cantidad,orden) VALUES(1,2,1,1)")
    conn.execute("INSERT INTO preguntas(tema_id,enunciado,alternativa_a,alternativa_b,alternativa_c,alternativa_d) VALUES(1,'Enunciado','a','b','c','d')")
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    uploads = tmp_path / "uploads"; uploads.mkdir(exist_ok=True)
    desc = tmp_path / "descargas"; desc.mkdir(exist_ok=True)
    grupos_out = tmp_path / "grupos"; grupos_out.mkdir(exist_ok=True)
    monkeypatch.setitem(app_module.app.config, "UPLOAD_FOLDER", str(uploads))
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(desc))
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(grupos_out), raising=False)
    monkeypatch.setattr(app_module, "_validar_docx_real", lambda _p: None, raising=False)
    return db_path, uploads, grupos_out


def test_matriz_crud_upload_helpers_y_zip(client, app_module, tmp_path, monkeypatch):
    db_path, uploads, grupos_out = _patch_matriz(app_module, tmp_path, monkeypatch)

    assert client.post("/api/matriz", json={"nombre": "Vacía", "items": []}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": "x"}]}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": 0, "cantidad": 1}]}).status_code == 400

    created = client.post("/api/matriz", json={
        "nombre": "Matriz Cov14",
        "items": [
            {"tema_id": 1, "cantidad": 2, "orden": 2},
            {"tema_id": 2, "cantidad": -5, "orden": 1},
        ],
    })
    assert created.status_code == 200
    matriz_id = created.get_json()["matriz_id"]

    assert client.get("/api/matriz?detail=1&search=Cov14").status_code == 200
    detalle = client.get(f"/api/matriz/{matriz_id}")
    assert detalle.status_code == 200
    assert detalle.get_json()["n_items"] == 2
    assert client.get("/api/matriz/9999").status_code == 404

    assert client.post(f"/api/matriz/{matriz_id}/upload", data={}).status_code == 400
    bad_ext = client.post(f"/api/matriz/{matriz_id}/upload", data={
        "tema_id": "1",
        "file": (io.BytesIO(b"x"), "mal.txt"),
    }, content_type="multipart/form-data")
    assert bad_ext.status_code == 400

    up = client.post(f"/api/matriz/{matriz_id}/upload", data={
        "tema_id": "1",
        "cantidad": "3",
        "file": (_docx_bytes("1. Matriz"), "tema.docx"),
    }, content_type="multipart/form-data")
    assert up.status_code == 200
    assert up.get_json()["ok"] is True

    # Helpers de grupos/matriz que respetan orden y agrupan por tema.
    conn = _connect(db_path); cur = conn.cursor()
    md = app_module._leer_matriz_detalle(cur, matriz_id)
    grupos = app_module._leer_config_grupos(cur)
    por_tema = app_module._preguntas_por_tema(cur)
    cur.close(); conn.close()
    assert [x["tema_id"] for x in md] == [2, 1]
    assert [x["tema_id"] for x in grupos[1]["temas"]] == [2, 1]
    assert 1 in por_tema

    # Arma un DOCX simple por grupo y lo valida.
    out_docx = tmp_path / "grupo_A.docx"
    app_module._armar_docx_grupo("Ingenierías", "A", [{"titulo": "Álgebra", "pregs": por_tema[1]}], str(out_docx))
    assert out_docx.exists()
    assert Document(str(out_docx)).paragraphs[0].text.startswith("Examen - Grupo A")

    # Ruta de descarga de lote: 404 y éxito.
    assert client.get("/api/grupos/lote/77/zip").status_code == 404
    lote_dir = grupos_out / "lote_77"; lote_dir.mkdir(exist_ok=True)
    zip_path = lote_dir / "grupos_77.zip"
    with zipfile.ZipFile(zip_path, "w") as z:
        z.writestr("grupo_A.txt", "ok")
    assert client.get("/api/grupos/lote/77/zip").status_code == 200
