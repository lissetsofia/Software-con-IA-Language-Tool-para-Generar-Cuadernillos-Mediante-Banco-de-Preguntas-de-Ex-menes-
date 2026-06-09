import os
import zipfile
import tempfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{%s}" % WNS


def _paragraph_texts(path):
    with zipfile.ZipFile(path, "r") as z:
        root = ET.fromstring(z.read("word/document.xml"))
    out = []
    for p in root.findall(f".//{W}p"):
        txt = "".join((t.text or "") for t in p.findall(f".//{W}t"))
        if txt:
            out.append(txt)
    return out


def _rewrite_docx(path, files):
    tmp = str(path) + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    os.replace(tmp, path)


def _make_numbered_docx(path: Path, n=3, num_id="77", fmt="decimal"):
    doc = Document()
    for i in range(1, n + 1):
        doc.add_paragraph(f"Pregunta {i}")
        doc.add_paragraph(f"Detalle {i}")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    root = ET.fromstring(files["word/document.xml"])
    body = root.find(f"{W}body")
    paras = [ch for ch in list(body) if ch.tag == f"{W}p"]
    for idx in range(0, min(len(paras), n * 2), 2):
        p = paras[idx]
        pPr = p.find(f"{W}pPr")
        if pPr is None:
            pPr = ET.Element(f"{W}pPr")
            p.insert(0, pPr)
        old = pPr.find(f"{W}numPr")
        if old is not None:
            pPr.remove(old)
        numPr = ET.Element(f"{W}numPr")
        ilvl = ET.SubElement(numPr, f"{W}ilvl")
        ilvl.set(f"{W}val", "0")
        numId = ET.SubElement(numPr, f"{W}numId")
        numId.set(f"{W}val", str(num_id))
        pPr.insert(0, numPr)

    numbering = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{WNS}">
  <w:abstractNum w:abstractNumId="{num_id}">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="{fmt}"/>
      <w:lvlText w:val="%1."/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="{num_id}"><w:abstractNumId w:val="{num_id}"/></w:num>
</w:numbering>'''.encode("utf-8")

    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    files["word/numbering.xml"] = numbering
    _rewrite_docx(path, files)
    return path


def test_ooxml_question_spans_cut_count_and_numbering_helpers(app_module, tmp_path, monkeypatch):
    src = _make_numbered_docx(tmp_path / "preguntas.docx", n=3)
    assert app_module._contar_preguntas_docx(str(src)) == 3
    assert app_module.contar_preguntas_docx(str(src)) == 3
    assert app_module.debug_contar_preguntas_docx(str(src)) == 3

    with tempfile.TemporaryDirectory() as td:
        with zipfile.ZipFile(src, "r") as z:
            z.extractall(td)
        spans = app_module._find_question_spans(
            os.path.join(td, "word", "document.xml"),
            os.path.join(td, "word", "numbering.xml"),
            2,
        )
    assert spans == [(0, 2), (2, 4)]

    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: None, raising=False)
    out_first = tmp_path / "primeras.docx"
    app_module._cut_docx_first_n_questions(str(src), 2, str(out_first))
    assert zipfile.is_zipfile(out_first)
    assert app_module._contar_preguntas_docx(str(out_first)) == 2

    individuales = app_module._cut_docx_to_individual_question_docs(str(src), 2)
    try:
        assert len(individuales) == 2
        assert all(zipfile.is_zipfile(p) for p in individuales)
    finally:
        for p in individuales:
            try:
                os.remove(p)
            except Exception:
                pass

    # Repara la numeración tras un merge: todo numId decimal pasa a 1 / ilvl 0.
    fix_num = _make_numbered_docx(tmp_path / "fix_num.docx", n=1, num_id="77")
    app_module._post_merge_fix_numbering(str(fix_num))
    with zipfile.ZipFile(fix_num, "r") as z:
        root = ET.fromstring(z.read("word/document.xml"))
    num_ids = [n.get(f"{W}val") for n in root.findall(f".//{W}numId")]
    assert "1" in num_ids

    # Convierte listas bullet en decimal dentro de numbering.xml.
    bullet = _make_numbered_docx(tmp_path / "bullet.docx", n=1, num_id="88", fmt="bullet")
    app_module.bullets_to_numbers_docx(str(bullet))
    with zipfile.ZipFile(bullet, "r") as z:
        numbering = z.read("word/numbering.xml").decode("utf-8")
    assert "bullet" not in numbering
    assert "decimal" in numbering

    # Aplana numeración como texto y elimina numbering.xml.
    flat = _make_numbered_docx(tmp_path / "flat.docx", n=2, num_id="77")
    app_module.aplanar_listas_a_texto(str(flat))
    with zipfile.ZipFile(flat, "r") as z:
        names = set(z.namelist())
        document_xml = z.read("word/document.xml").decode("utf-8")
    assert "word/numbering.xml" not in names
    assert "numPr" not in document_xml
    assert _paragraph_texts(flat)[0].startswith("1. ")


def test_sanear_fragmento_unwrap_sdt_y_limpieza_ooxml(app_module):
    xml = f'''
    <w:p xmlns:w="{WNS}" w:rsidR="00AA">
      <w:pPr><w:sectPr/><w:spacing/></w:pPr>
      <w:sdt><w:sdtPr/><w:sdtContent><w:r><w:t>Dentro</w:t></w:r></w:sdtContent></w:sdt>
      <w:ins><w:r><w:t> Insertado</w:t></w:r></w:ins>
      <w:hyperlink><w:r><w:t> Link</w:t></w:r></w:hyperlink>
      <w:del><w:r><w:delText>Borrado</w:delText></w:r></w:del>
      <w:bookmarkStart w:id="1" w:name="b"/><w:bookmarkEnd w:id="1"/>
    </w:p>
    '''
    node = ET.fromstring(xml)
    app_module._sanear_fragmento(node)
    final_xml = ET.tostring(node, encoding="unicode")
    text = "".join((t.text or "") for t in node.findall(f".//{W}t"))
    assert "Dentro" in text
    assert "Insertado" in text
    assert "Link" in text
    assert "Borrado" not in final_xml
    assert "sdt" not in final_xml
    assert "bookmark" not in final_xml
    assert "rsid" not in final_xml
