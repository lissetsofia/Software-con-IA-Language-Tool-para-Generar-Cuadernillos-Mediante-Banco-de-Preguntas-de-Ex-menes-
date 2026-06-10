import os
import sqlite3
import zipfile
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _make_docx(path, text="Pregunta de prueba"):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov17_grupos_generar.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE gen_lote(id INTEGER PRIMARY KEY AUTOINCREMENT, matriz_id INTEGER, nombre TEXT, usuario TEXT);
        """
    )
    conn.executemany(
        "INSERT INTO temario(id,nombre,activo) VALUES(?,?,1)",
        [(1, "Álgebra"), (2, "Geometría"), (3, "Física")],
    )
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    out = tmp_path / "grupos_out"
    out.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out), raising=False)
    return db_path, out


def _grupo_con_temas(*rels):
    return {
        1: {
            "idgrupo": 1,
            "clave": "A",
            "nombre": "Grupo A",
            "temas": list(rels),
        }
    }


def test_api_grupos_generar_validaciones_matriz_inline(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)

    r0 = client.post("/api/grupos/generar", json={})
    assert r0.status_code == 400
    assert "matriz_id" in r0.get_json()["error"]

    r_bad_items = client.post("/api/grupos/generar", json={"matriz": {"items": []}})
    assert r_bad_items.status_code == 400

    r_bad_item = client.post("/api/grupos/generar", json={"matriz": {"items": [{"cantidad": 1}]}})
    assert r_bad_item.status_code == 400

    src = _make_docx(tmp_path / "tema1.docx")

    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {}, raising=False)
    r_no_groups = client.post(
        "/api/grupos/generar",
        json={"matriz": {"items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_no_groups.status_code == 400
    assert "grupos" in r_no_groups.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _grupo_con_temas({"tema_id": 2, "tema_nombre": "Geometría", "cantidad": 1, "orden": 1}),
        raising=False,
    )
    r_tema_fuera = client.post(
        "/api/grupos/generar",
        json={"matriz": {"items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_tema_fuera.status_code == 400
    assert "matriz" in r_tema_fuera.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _grupo_con_temas({"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 0, "orden": 1}),
        raising=False,
    )
    r_sin_cuotas = client.post(
        "/api/grupos/generar",
        json={"matriz": {"items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_sin_cuotas.status_code == 400
    assert "cuotas" in r_sin_cuotas.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _grupo_con_temas({"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 3, "orden": 1}),
        raising=False,
    )
    r_cuota_mayor = client.post(
        "/api/grupos/generar",
        json={"matriz": {"items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_cuota_mayor.status_code == 400
    assert "matriz solo tiene" in r_cuota_mayor.get_json()["error"].lower()

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: _grupo_con_temas({"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}),
        raising=False,
    )
    r_sin_archivo = client.post(
        "/api/grupos/generar",
        json={"matriz": {"items": [{"tema_id": 1, "cantidad": 1}]}},
    )
    assert r_sin_archivo.status_code == 400
    assert "docx" in r_sin_archivo.get_json()["error"].lower()


def test_api_grupos_generar_success_zip_y_merge_malos(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    src = _make_docx(tmp_path / "tema1.docx", "1. Pregunta de álgebra")

    monkeypatch.setattr(
        app_module,
        "_leer_config_grupos",
        lambda cur: {
            1: {"idgrupo": 1, "clave": "A", "nombre": "Grupo A", "temas": [
                {"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}
            ]},
            2: {"idgrupo": 2, "clave": "B", "nombre": "Grupo B", "temas": [
                {"tema_id": 1, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}
            ]},
        },
        raising=False,
    )

    def fake_cut(src_path, n, out_docx):
        _make_docx(out_docx, f"recortado {n}")

    def fake_merge(bloques, out_path):
        doc = Document()
        for tema, files in bloques:
            doc.add_paragraph(f"TEMA {tema}")
            for f in files:
                doc.add_paragraph(os.path.basename(f))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_cut_docx_first_n_questions", fake_cut, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    r_ok = client.post(
        "/api/grupos/generar",
        json={"matriz": {"nombre": "Matriz COV17", "items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_ok.status_code == 200
    body = r_ok.get_json()
    assert body["ok"] is True
    zip_resp = client.get(body["zip_url"])
    assert zip_resp.status_code == 200
    assert "zip" in zip_resp.headers["Content-Type"].lower()

    lote_zip = tmp_path / "grupos_out" / f"lote_{body['lote_id']}" / f"grupos_{body['lote_id']}.zip"
    assert zipfile.is_zipfile(lote_zip)

    def merge_con_malos(bloques, out_path):
        _make_docx(out_path, "parcial")
        return out_path, [], [("fallo.docx", "no insertó")]

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", merge_con_malos, raising=False)
    r_malos = client.post(
        "/api/grupos/generar",
        json={"matriz": {"nombre": "Matriz COV17", "items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": str(src)}]}},
    )
    assert r_malos.status_code == 409
    assert r_malos.get_json()["ok"] is False

    assert client.get("/api/grupos/lote/999999/zip").status_code == 404
