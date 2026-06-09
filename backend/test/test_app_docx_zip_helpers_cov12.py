import os
import shutil
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{%s}' % W_NS
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'


def _set_numpr_xml(p, num_id='1', ilvl='0'):
    ppr = p.find(W + 'pPr')
    if ppr is None:
        ppr = ET.Element(W + 'pPr')
        p.insert(0, ppr)
    numpr = ET.SubElement(ppr, W + 'numPr')
    ET.SubElement(numpr, W + 'ilvl', {W + 'val': str(ilvl)})
    ET.SubElement(numpr, W + 'numId', {W + 'val': str(num_id)})


def _make_numbered_docx(path: Path):
    doc = DocxDocument()
    for txt in ['Pregunta uno', 'A) alternativa', 'B) alternativa', 'Pregunta dos', 'A) otra']:
        doc.add_paragraph(txt)
    doc.save(path)

    tmp = tempfile.mkdtemp(prefix='numbered_docx_')
    try:
        with zipfile.ZipFile(path, 'r') as z:
            z.extractall(tmp)

        doc_xml = Path(tmp) / 'word' / 'document.xml'
        tree = ET.parse(doc_xml)
        root = tree.getroot()
        paras = root.findall('.//' + W + 'p')
        _set_numpr_xml(paras[0], '1', '0')
        _set_numpr_xml(paras[3], '1', '0')
        tree.write(doc_xml, encoding='utf-8', xml_declaration=True)

        numbering = ET.Element(W + 'numbering')
        absn = ET.SubElement(numbering, W + 'abstractNum', {W + 'abstractNumId': '0'})
        lvl = ET.SubElement(absn, W + 'lvl', {W + 'ilvl': '0'})
        ET.SubElement(lvl, W + 'numFmt', {W + 'val': 'decimal'})
        ET.SubElement(lvl, W + 'lvlText', {W + 'val': '%1.'})
        num = ET.SubElement(numbering, W + 'num', {W + 'numId': '1'})
        ET.SubElement(num, W + 'abstractNumId', {W + 'val': '0'})
        ET.ElementTree(numbering).write(Path(tmp) / 'word' / 'numbering.xml', encoding='utf-8', xml_declaration=True)

        rels_path = Path(tmp) / 'word' / '_rels' / 'document.xml.rels'
        rels = ET.parse(rels_path).getroot()
        rel_tag = '{%s}Relationship' % REL_NS
        if not any(r.get('Target') == 'numbering.xml' for r in rels.findall(rel_tag)):
            ET.SubElement(rels, rel_tag, {
                'Id': 'rId999',
                'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering',
                'Target': 'numbering.xml',
            })
        ET.ElementTree(rels).write(rels_path, encoding='utf-8', xml_declaration=True)

        ct_path = Path(tmp) / '[Content_Types].xml'
        ct = ET.parse(ct_path).getroot()
        override = '{%s}Override' % CT_NS
        if not any(x.get('PartName') == '/word/numbering.xml' for x in ct.findall(override)):
            ET.SubElement(ct, override, {
                'PartName': '/word/numbering.xml',
                'ContentType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml',
            })
        ET.ElementTree(ct).write(ct_path, encoding='utf-8', xml_declaration=True)

        with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as out:
            for base, _dirs, files in os.walk(tmp):
                for f in files:
                    p = Path(base) / f
                    out.write(p, p.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
    return path


def test_docx_zip_helpers_spans_cut_y_rezip(app_module, tmp_path, monkeypatch):
    src = _make_numbered_docx(tmp_path / 'numbered.docx')
    app_module._validar_docx_real(str(src))
    assert app_module._tiene_texto_o_contenido(str(src)) is True
    assert app_module._docx_requiere_resave_para_matriz('TRIGONOMETRÍA', str(src)) is True
    assert app_module._docx_requiere_resave_para_matriz('ÁLGEBRA', str(src)) is False

    with tempfile.TemporaryDirectory() as td:
        with zipfile.ZipFile(src, 'r') as z:
            z.extractall(td)
        spans = app_module._find_question_spans(
            os.path.join(td, 'word', 'document.xml'),
            os.path.join(td, 'word', 'numbering.xml'),
            2,
        )
    assert spans == [(0, 3), (3, 5)]
    assert app_module._contar_preguntas_docx(str(src)) == 2

    out_one = tmp_path / 'first_question.docx'
    app_module._cut_docx_first_n_questions(str(src), 1, str(out_one))
    assert out_one.exists()
    app_module._validar_docx_real(str(out_one))

    individuales = app_module._cut_docx_to_individual_question_docs(str(src), 2)
    try:
        assert len(individuales) == 2
        for p in individuales:
            assert os.path.exists(p)
            app_module._validar_docx_real(p)
    finally:
        for p in individuales:
            try:
                os.remove(p)
            except Exception:
                pass

    # Rutas sin numeración también deben salir sin romper.
    plain = tmp_path / 'plain.docx'
    DocxDocument().save(plain)
    app_module._post_merge_fix_numbering(str(plain))
    app_module._hacer_secciones_continuas(str(plain))
    app_module.aplanar_listas_a_texto(str(plain))
    app_module.bullets_to_numbers_docx(str(plain))

    # _safe_rezip acepta bytes y strings.
    files = {'a.txt': b'A', 'b.txt': 'B'}
    zpath = tmp_path / 'simple.zip'
    app_module._safe_rezip(str(zpath), files)
    with zipfile.ZipFile(zpath) as z:
        assert z.read('a.txt') == b'A'
        assert z.read('b.txt') == b'B'


def test_sanear_fragmento_y_parent_helpers(app_module):
    root = ET.Element(W + 'p')
    ins = ET.SubElement(root, W + 'ins')
    run = ET.SubElement(ins, W + 'r')
    t = ET.SubElement(run, W + 't')
    t.text = 'texto'
    ET.SubElement(root, W + 'bookmarkStart')

    app_module._sanear_fragmento(root)
    xml = ET.tostring(root, encoding='unicode')
    assert 'bookmarkStart' not in xml
    assert 'texto' in xml

    parent_map = app_module._build_parent_map(root)
    text_node = root.find('.//' + W + 't')
    assert parent_map[text_node].tag == W + 'r'
