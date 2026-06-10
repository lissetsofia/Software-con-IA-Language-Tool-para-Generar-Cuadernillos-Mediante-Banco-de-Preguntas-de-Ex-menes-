import io
import os
import sqlite3
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch, with_doc=True):
    db_path = tmp_path / "cov19_banco.sqlite3"
    if db_path.exists():
        db_path.unlink()
    preg_dir = tmp_path / "banco_preg"
    sol_dir = tmp_path / "banco_sol"
    preg_dir.mkdir(exist_ok=True)
    sol_dir.mkdir(exist_ok=True)

    doc_path = preg_dir / "p_actual.docx"
    sol_path = sol_dir / "s_actual.docx"
    if with_doc:
        _make_docx(doc_path, "pregunta actual")
        _make_docx(sol_path, "sol actual")

    conn = _connect(db_path)
    conn.executescript("""
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT);
        CREATE TABLE tema_docs(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tema_id INTEGER,
            doc_preguntas_nombre TEXT,
            doc_preguntas_ruta TEXT,
            doc_sol_nombre TEXT,
            doc_sol_ruta TEXT,
            fecha_creacion TEXT DEFAULT '2025-01-01'
        );
        CREATE TABLE preguntas(
            idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT,
            examenes_idexamenes INTEGER,
            tema_id INTEGER,
            numero_p INTEGER,
            archivo_nombre TEXT,
            archivo_ruta TEXT
        );
    """)
    conn.execute("INSERT INTO temario(id,nombre) VALUES(1,'Álgebra')")
    conn.execute(
        "INSERT INTO tema_docs(id,tema_id,doc_preguntas_nombre,doc_preguntas_ruta,doc_sol_nombre,doc_sol_ruta) VALUES(1,1,?,?,?,?)",
        ("p_actual.docx", str(doc_path), "s_actual.docx", str(sol_path) if with_doc else None),
    )
    conn.execute("INSERT INTO preguntas(examenes_idexamenes,tema_id,numero_p,archivo_nombre,archivo_ruta) VALUES(NULL,1,1,?,?)", ("p_actual.docx", str(doc_path)))
    conn.commit(); conn.close()

    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    monkeypatch.setattr(app_module, "BANCO_PREG_DIR", str(preg_dir), raising=False)
    monkeypatch.setattr(app_module, "BANCO_SOL_DIR", str(sol_dir), raising=False)
    return db_path, doc_path, sol_path


def _make_docx(path, text="1. Pregunta"):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(); doc.add_paragraph(text); doc.save(path)
    return path


def _docx_upload(name="x.docx"):
    bio = io.BytesIO()
    doc = Document(); doc.add_paragraph("1. Pregunta"); doc.save(bio); bio.seek(0)
    return bio, name


def test_banco_importar_y_reemplazar_validaciones(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)

    assert client.post("/api/banco_preguntas", data={}, content_type="multipart/form-data").status_code == 400
    assert client.post("/api/banco_preguntas/999/reemplazar/preguntas", data={}, content_type="multipart/form-data").status_code == 400

    # Archivo con más de una pregunta: se rechaza y no toca BD.
    monkeypatch.setattr(app_module, "contar_preguntas_docx", lambda *_a, **_k: 2, raising=False)
    r_bad = client.post(
        "/api/banco_preguntas",
        data={"tema_id": "1", "doc_preguntas": _docx_upload("dos.docx")},
        content_type="multipart/form-data",
    )
    assert r_bad.status_code == 400
    assert r_bad.get_json()["n_preguntas"] == 2

    # Reemplazo con registro inexistente.
    monkeypatch.setattr(app_module, "contar_preguntas_docx", lambda *_a, **_k: 1, raising=False)
    r_missing = client.post(
        "/api/banco_preguntas/999/reemplazar/preguntas",
        data={"doc_preguntas": _docx_upload("nuevo.docx")},
        content_type="multipart/form-data",
    )
    assert r_missing.status_code == 404

    # Reemplazo correcto donde no coincide archivo viejo en preguntas: inserta nueva fila.
    r_ok = client.post(
        "/api/banco_preguntas/1/reemplazar/preguntas",
        data={"doc_preguntas": _docx_upload("nuevo_ok.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    assert r_ok.get_json()["ok"] is True


def test_banco_solucionario_y_descargas_errores(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch, with_doc=True)

    assert client.post("/api/banco_preguntas/solucionario", data={}, content_type="multipart/form-data").status_code == 400
    assert client.post("/api/banco_preguntas/1/reemplazar/solucionario", data={}, content_type="multipart/form-data").status_code == 400

    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: (_ for _ in ()).throw(ValueError("docx inválido")), raising=False)
    r_sol_bad = client.post(
        "/api/banco_preguntas/solucionario",
        data={"tema_id": "1", "doc_solucionario": _docx_upload("sol_bad.docx")},
        content_type="multipart/form-data",
    )
    assert r_sol_bad.status_code == 400
    assert "inválido" in r_sol_bad.get_json()["error"].lower()

    r_repl_bad = client.post(
        "/api/banco_preguntas/1/reemplazar/solucionario",
        data={"doc_solucionario": _docx_upload("sol_repl_bad.docx")},
        content_type="multipart/form-data",
    )
    assert r_repl_bad.status_code == 400

    # Descargas no disponibles por id inexistente.
    assert client.get("/api/banco_preguntas/999/download/preguntas").status_code == 404
    assert client.get("/api/banco_preguntas/999/download/solucionario").status_code == 404
    assert client.get("/api/banco_preguntas/999/download").status_code == 404


def test_banco_preview_y_download_full_sin_solucionario(client, app_module, tmp_path, monkeypatch):
    db_path, doc_path, sol_path = _patch_db(app_module, tmp_path, monkeypatch, with_doc=True)
    # Quitamos el solucionario físico para cubrir rama sin solucionario en preview/full.
    try:
        os.remove(sol_path)
    except Exception:
        pass

    out_pdf = tmp_path / "preview.pdf"
    out_pdf.write_bytes(b"%PDF")
    monkeypatch.setattr(app_module, "generar_pdf_preview", lambda *_a, **_k: str(out_pdf), raising=False)
    monkeypatch.setattr(app_module.time, "time", lambda: 1700000000, raising=False)

    r_prev = client.get("/api/banco_preguntas/1/preview")
    assert r_prev.status_code == 200
    assert r_prev.get_json()["tiene_solucionario"] is False

    r_zip = client.get("/api/banco_preguntas/1/download")
    assert r_zip.status_code == 200
    assert "zip" in r_zip.headers.get("Content-Type", "").lower()

    # Delete borra el registro y limpia archivos si existen.
    r_del = client.delete("/api/banco_preguntas/1")
    assert r_del.status_code == 200
    assert r_del.get_json()["ok"] is True
    assert client.delete("/api/banco_preguntas/1").status_code == 404
