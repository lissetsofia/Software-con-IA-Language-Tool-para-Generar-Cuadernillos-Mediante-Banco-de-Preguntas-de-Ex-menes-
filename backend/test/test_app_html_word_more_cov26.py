import os
import sys
import types
from docx import Document


class _FakeWebOptions:
    def __init__(self):
        self.AllowPNG = False
        self.OptimizeForBrowser = False
        self.RelyOnCSS = False


class _FakeRange:
    def __init__(self):
        self.FormattedText = "FORMATEADO"
    def InsertFile(self, *args, **kwargs):
        return None
    def Collapse(self, *args, **kwargs):
        return None


class _FakeDocHtml:
    def __init__(self, path=None):
        self.path = path
        self.WebOptions = _FakeWebOptions()
        self.Content = types.SimpleNamespace(FormattedText="FORMATEADO")
        self.Fields = types.SimpleNamespace(Update=lambda: None)
    def Range(self, *args, **kwargs):
        return _FakeRange()
    def Repaginate(self):
        return None
    def SaveAs(self, path=None, FileName=None, FileFormat=None, **kwargs):
        target = FileName or path
        self._write_target(target)
    def SaveAs2(self, FileName=None, path=None, FileFormat=None, Encoding=None, **kwargs):
        target = FileName or path
        self._write_target(target, make_assets=True)
    def ExportAsFixedFormat(self, OutputFileName=None, **kwargs):
        with open(OutputFileName, "wb") as f:
            f.write(b"%PDF-1.4 fake")
    def Close(self, *args, **kwargs):
        return None
    def _write_target(self, target, make_assets=False):
        target = os.path.abspath(str(target))
        os.makedirs(os.path.dirname(target), exist_ok=True)
        low = target.lower()
        if low.endswith(('.htm', '.html')):
            base = os.path.splitext(os.path.basename(target))[0]
            with open(target, "w", encoding="utf-8") as f:
                f.write(f"<html><head><meta charset=windows-1252></head><body><img src='{base}_archivos/a.png'>Hola</body></html>")
            if make_assets:
                os.makedirs(os.path.join(os.path.dirname(target), f"{base}_archivos"), exist_ok=True)
        elif low.endswith('.docx'):
            d = Document()
            d.add_paragraph("doc fake")
            d.save(target)
        else:
            with open(target, "wb") as f:
                f.write(b"x")


class _FakeDocumentsHtml:
    def Open(self, path, *args, **kwargs):
        return _FakeDocHtml(path)
    def Add(self, *args, **kwargs):
        return _FakeDocHtml()


class _FakeWordHtml:
    def __init__(self):
        self.Documents = _FakeDocumentsHtml()
        self.Visible = False
        self.DisplayAlerts = 0
        self.AutomationSecurity = None
    def Quit(self):
        return None


def _patch_word_html(app_module, monkeypatch):
    monkeypatch.setattr(app_module.pythoncom, "CoInitialize", lambda: None)
    monkeypatch.setattr(app_module.pythoncom, "CoUninitialize", lambda: None)
    fake_client = types.SimpleNamespace(DispatchEx=lambda *a, **k: _FakeWordHtml())
    monkeypatch.setattr(app_module.win32, "DispatchEx", fake_client.DispatchEx, raising=False)
    if "win32com.client" in sys.modules:
        monkeypatch.setattr(sys.modules["win32com.client"], "DispatchEx", fake_client.DispatchEx, raising=False)


def test_html_word_export_y_preview_word_fake(app_module, tmp_path, monkeypatch):
    _patch_word_html(app_module, monkeypatch)
    desc = tmp_path / "desc"; desc.mkdir()
    prev = tmp_path / "prev"; prev.mkdir()
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(desc))
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(desc), raising=False)
    monkeypatch.setattr(app_module, "PREVIEWS_DIR", str(prev), raising=False)
    monkeypatch.setattr(app_module, "_short_path", lambda p: p, raising=False)
    monkeypatch.setattr(app_module, "_wait_exists_nonzero", lambda p, *a, **k: os.path.exists(p) and os.path.getsize(p) > 0, raising=False)

    src = tmp_path / "fuente.docx"
    Document().save(src)

    # Cubre generar_html_desde_docx por la ruta Word exitosa.
    html = app_module.generar_html_desde_docx(str(src), "vista_word")
    assert os.path.exists(html)
    assert "utf-8" in open(html, encoding="utf-8").read().lower()

    # Cubre exportación HTML Word directa y normalización de recursos.
    out_html, warnings = app_module._exportar_html_con_word(str(src), "base con espacios")
    assert os.path.exists(out_html)
    assert isinstance(warnings, list)

    # Cubre docx_a_html_filtrado con carpeta *_archivos y devolución de URL relativa.
    html2, rel = app_module.docx_a_html_filtrado(str(src), str(prev / "sub"))
    assert html2 and os.path.exists(html2)
    assert rel.startswith("/static/previews/")


def test_guardar_pdf_y_pdf_from_docx_con_word_fake(client, app_module, tmp_path, monkeypatch):
    _patch_word_html(app_module, monkeypatch)
    desc = tmp_path / "desc"; desc.mkdir()
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(desc))
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(desc), raising=False)
    src = desc / "doc.docx"
    Document().save(src)

    pdf = app_module.guardar_pdf(str(src))
    assert os.path.exists(pdf) and os.path.getsize(pdf) > 0

    # Para la ruta /api/pdf_from_docx, evita depender de Word real.
    monkeypatch.setattr(app_module, "resave_docx_formatted", lambda a, b: Document(a).save(b))
    monkeypatch.setattr(app_module, "docx_a_pdf", lambda a, b: (open(b, "wb").write(b"%PDF fake") and b))
    r = client.post("/api/pdf_from_docx", json={"docx": "doc.docx"})
    assert r.status_code == 200
    assert r.json["ok"] is True
