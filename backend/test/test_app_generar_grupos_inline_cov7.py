# backend/test/test_app_generar_grupos_inline_cov7.py
import io
import os
import sqlite3
import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_db(app_module, tmp_path, monkeypatch, name="cov7_grupos.sqlite3"):
    db_path = tmp_path / name

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    return db_path


def _schema(app_module):
    conn = app_module.get_connection()
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS grupos (
            idgrupo INTEGER PRIMARY KEY AUTOINCREMENT,
            clave TEXT,
            nombre TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS grupo_tema (
            idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT,
            grupos_idgrupo INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS gen_lote (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            nombre TEXT,
            usuario TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS matriz_detalle (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0,
            archivo_ruta TEXT
        )
    """)
    conn.commit()
    return conn, cur


def _insert_topic_group(app_module, *, cantidad=1, clave="A", tema="ÁLGEBRA"):
    conn, cur = _schema(app_module)
    cur.execute("INSERT INTO temario (nombre, activo) VALUES (?, 1)", (tema,))
    tid = int(cur.lastrowid)
    cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES (?, ?, 1)", (clave, f"Grupo {clave}"))
    gid = int(cur.lastrowid)
    cur.execute(
        "INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad, orden) VALUES (?, ?, ?, 1)",
        (gid, tid, cantidad),
    )
    conn.commit()
    cur.close()
    conn.close()
    return gid, tid


def _make_docx(path: Path, text="Pregunta de prueba"):
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _patch_docx_generation(app_module, tmp_path, monkeypatch):
    out_root = tmp_path / "grupos_out"
    out_root.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out_root), raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None)

    def fake_cut(src, n, dst):
        doc = DocxDocument()
        doc.add_paragraph(f"recorte {n}")
        doc.save(dst)

    def fake_merge(grouped, out_path, merge_step_cb=None, merge_ops=None):
        assert grouped
        doc = DocxDocument()
        for titulo, files in grouped:
            doc.add_paragraph(str(titulo))
            for f in files:
                doc.add_paragraph(os.path.basename(str(f)))
        doc.save(out_path)
        if merge_step_cb:
            merge_step_cb(1, "merge")
        return out_path, [], []

    monkeypatch.setattr(app_module, "_cut_docx_first_n_questions", fake_cut)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge)
    return out_root


def test_generar_por_grupos_inline_success_y_descargar_zip(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)
    _patch_docx_generation(app_module, tmp_path, monkeypatch)

    # Dos grupos con la misma cuota para cubrir el loop de generación por grupo.
    gid_a, tid = _insert_topic_group(app_module, cantidad=1, clave="A")
    conn = app_module.get_connection()
    cur = conn.cursor()
    cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES ('B', 'Grupo B', 1)")
    gid_b = int(cur.lastrowid)
    cur.execute("INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad, orden) VALUES (?, ?, 1, 1)", (gid_b, tid))
    conn.commit(); cur.close(); conn.close()

    src = _make_docx(tmp_path / "fuente.docx")
    r = client.post("/api/grupos/generar", json={
        "matriz": {
            "nombre": "Matriz inline cov7",
            "items": [{"tema_id": tid, "cantidad": 2, "archivo_ruta": str(src)}],
        }
    })
    assert r.status_code == 200
    data = r.get_json()
    assert data["ok"] is True
    assert data["zip_url"].startswith("/api/grupos/lote/")

    rz = client.get(data["zip_url"])
    assert rz.status_code == 200
    assert rz.data.startswith(b"PK")
    with zipfile.ZipFile(io.BytesIO(rz.data)) as z:
        assert any(name.endswith(".docx") for name in z.namelist())


def test_generar_por_grupos_validaciones_basicas(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)

    assert client.post("/api/grupos/generar", json={}).status_code == 400
    assert client.post("/api/grupos/generar", json={"matriz": {"items": []}}).status_code == 400
    assert client.post("/api/grupos/generar", json={"matriz": {"items": [{"cantidad": 1}]}}).status_code == 400

    # Matriz válida, pero sin grupos activos configurados.
    r_no_groups = client.post("/api/grupos/generar", json={
        "matriz": {"items": [{"tema_id": 1, "cantidad": 1, "archivo_ruta": "x.docx"}]}
    })
    assert r_no_groups.status_code == 400
    assert "grupos activos" in r_no_groups.get_json()["error"].lower()


def test_generar_por_grupos_errores_de_cuotas_y_archivo(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)
    _, tid = _insert_topic_group(app_module, cantidad=3, clave="A")

    # La cuota del grupo supera la cantidad de la matriz.
    r_cuota = client.post("/api/grupos/generar", json={
        "matriz": {"items": [{"tema_id": tid, "cantidad": 1, "archivo_ruta": str(tmp_path / "x.docx")}]}
    })
    assert r_cuota.status_code == 400
    assert "matriz solo tiene" in r_cuota.get_json()["error"].lower()

    # Reinicia DB: cuota válida pero sin DOCX asociado.
    _patch_db(app_module, tmp_path, monkeypatch, name="cov7_grupos_archivo.sqlite3")
    _schema(app_module)
    _, tid2 = _insert_topic_group(app_module, cantidad=1, clave="A")
    r_archivo = client.post("/api/grupos/generar", json={
        "matriz": {"items": [{"tema_id": tid2, "cantidad": 1, "archivo_ruta": None}]}
    })
    assert r_archivo.status_code == 400
    assert "no tiene docx" in r_archivo.get_json()["error"].lower()


def test_generar_por_grupos_cuotas_distintas(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    _schema(app_module)
    _, tid = _insert_topic_group(app_module, cantidad=1, clave="A")
    conn = app_module.get_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES ('B', 'Grupo B', 1)")
    gid_b = int(cur.lastrowid)
    cur.execute("INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad, orden) VALUES (?, ?, 2, 1)", (gid_b, tid))
    conn.commit(); cur.close(); conn.close()

    r = client.post("/api/grupos/generar", json={
        "matriz": {"items": [{"tema_id": tid, "cantidad": 2, "archivo_ruta": str(tmp_path / "x.docx")}]}
    })
    assert r.status_code == 400
    assert "no son iguales" in r.get_json()["error"].lower()
