import os
import sqlite3
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _make_docx(path: Path, text="Pregunta"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov15_generar_doc.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE grupos(idgrupo INTEGER PRIMARY KEY, clave TEXT, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE grupo_tema(idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT, grupos_idgrupo INTEGER, tema_id INTEGER, cantidad INTEGER, orden INTEGER DEFAULT 0);
        CREATE TABLE preguntas(idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT, tema_id INTEGER, archivo_ruta TEXT);
        """
    )
    conn.execute("INSERT INTO grupos(idgrupo, clave, nombre, activo) VALUES(1,'A','Ingenierías',1)")
    conn.execute("INSERT INTO grupos(idgrupo, clave, nombre, activo) VALUES(2,'B','Sin cuotas',1)")
    conn.execute("INSERT INTO temario(id,nombre,activo) VALUES(1,'Álgebra',1)")
    conn.execute("INSERT INTO temario(id,nombre,activo) VALUES(2,'Física',1)")
    conn.execute("INSERT INTO grupo_tema(grupos_idgrupo, tema_id, cantidad, orden) VALUES(1,1,1,2)")
    conn.execute("INSERT INTO grupo_tema(grupos_idgrupo, tema_id, cantidad, orden) VALUES(1,2,1,1)")
    p1 = _make_docx(tmp_path / "p1.docx", "Pregunta álgebra")
    p2 = _make_docx(tmp_path / "p2.docx", "Pregunta física")
    conn.execute("INSERT INTO preguntas(tema_id, archivo_ruta) VALUES(1, ?)", (str(p1),))
    conn.execute("INSERT INTO preguntas(tema_id, archivo_ruta) VALUES(2, ?)", (str(p2),))
    conn.commit(); conn.close()

    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    desc = tmp_path / "descargas"; desc.mkdir(exist_ok=True)
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(desc))
    return db_path, desc


def test_grupos_generar_doc_run_success_debug_and_errors(app_module, tmp_path, monkeypatch):
    _db_path, desc = _patch_db(app_module, tmp_path, monkeypatch)

    assert app_module._grupos_generar_doc_run(1, "rar", {}, None)[0:2] == ("err", 400)
    assert app_module._grupos_generar_doc_run(999, "word", {}, None)[0:2] == ("err", 404)
    assert app_module._grupos_generar_doc_run(2, "word", {}, None)[0:2] == ("err", 400)

    dbg = app_module._grupos_generar_doc_run(1, "word", {"debug": "1"}, None)
    assert dbg[0] == "ok"
    assert dbg[1]["total_requeridas"] == 2
    assert dbg[1]["disponibilidad_por_tema"][1] == 1

    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_fuerte", lambda p: (True, None), raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    def fake_merge(grouped, out_path, merge_step_cb=None, merge_ops=None):
        doc = Document()
        for tema, files in grouped:
            doc.add_paragraph(f"TEMA: {tema}")
            for f in files:
                doc.add_paragraph(os.path.basename(str(f)))
                if merge_step_cb:
                    merge_step_cb(1, "merge")
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    def fake_pdf(src_docx, dst_pdf):
        Path(dst_pdf).write_bytes(b"%PDF-1.4\n%cov15\n")
        return dst_pdf

    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    monkeypatch.setattr(app_module, "docx_a_pdf", fake_pdf, raising=False)

    progress = []
    out = app_module._grupos_generar_doc_run(
        1,
        "pdf",
        {"flat": "1"},
        lambda done, total, msg: progress.append((done, total, msg)),
    )
    assert out[0] == "ok"
    body = out[1]
    assert body["ok"] is True
    assert body["preview_kind"] == "pdf"
    assert body["archivo_docx"].endswith(".docx")
    assert body["archivo_pdf"].endswith(".pdf")
    assert progress
    assert any(p.name.endswith(".docx") for p in desc.iterdir())
    assert any(p.name.endswith(".pdf") for p in desc.iterdir())
