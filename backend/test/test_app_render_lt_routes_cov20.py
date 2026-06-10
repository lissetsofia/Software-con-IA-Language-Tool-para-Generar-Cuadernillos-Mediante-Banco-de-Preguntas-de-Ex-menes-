import io
import os
from pathlib import Path

from docx import Document


def _docx_bytes(text="Documento"):
    bio = io.BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(bio)
    bio.seek(0)
    return bio


def _make_docx(path, text="Documento"):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_render_vista_validaciones_y_success(client, app_module, tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    desc = tmp_path / "descargas"
    uploads.mkdir(exist_ok=True)
    desc.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "UPLOAD_DIR", str(uploads), raising=False)
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(desc), raising=False)

    assert client.post("/api/render_vista", data={}).status_code == 400

    r_ext = client.post(
        "/api/render_vista",
        data={"archivo": (io.BytesIO(b"no es docx"), "mal.txt")},
        content_type="multipart/form-data",
    )
    assert r_ext.status_code == 400

    def fake_sha(path):
        return "abc1234567890000"

    def fake_preview(ruta_docx, nombre_base=None):
        pdf = desc / f"{nombre_base or 'preview'}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n%fake")
        return str(pdf)

    monkeypatch.setattr(app_module, "_sha1_file", fake_sha, raising=False)
    monkeypatch.setattr(app_module, "generar_pdf_preview", fake_preview, raising=False)

    r_ok = client.post(
        "/api/render_vista",
        data={"archivo": (_docx_bytes("Hola"), "entrada.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    assert r_ok.get_json()["ok"] is True
    assert "html_url" in r_ok.get_json()

    def boom_preview(*_a, **_k):
        raise RuntimeError("pdf boom")

    monkeypatch.setattr(app_module, "generar_pdf_preview", boom_preview, raising=False)
    r_err = client.post(
        "/api/render_vista",
        data={"archivo": (_docx_bytes("Hola"), "entrada2.docx")},
        content_type="multipart/form-data",
    )
    assert r_err.status_code == 500


def test_render_docx_guardado_lt_y_lt_status_ensure(client, app_module, tmp_path, monkeypatch):
    desc = tmp_path / "descargas"
    desc.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(desc), raising=False)

    assert client.get("/api/render_docx_guardado_lt/no_existe.docx").status_code == 404

    limpio = _make_docx(desc / "prueba_corregido_limpio.docx", "limpio")
    preview = _make_docx(desc / "prueba_corregido.docx", "preview")

    def fake_preview(ruta_docx, nombre_base=None):
        assert os.path.basename(ruta_docx) == "prueba_corregido.docx"
        pdf = desc / f"{nombre_base}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n")
        return str(pdf)

    monkeypatch.setattr(app_module, "generar_pdf_preview", fake_preview, raising=False)
    r_ok = client.get("/api/render_docx_guardado_lt/prueba_corregido_limpio.docx")
    assert r_ok.status_code == 200
    assert r_ok.get_json()["ok"] is True

    def bad_preview(*_a, **_k):
        raise RuntimeError("no convierte")

    monkeypatch.setattr(app_module, "generar_pdf_preview", bad_preview, raising=False)
    r_500 = client.get("/api/render_docx_guardado_lt/prueba_corregido.docx")
    assert r_500.status_code == 500

    monkeypatch.setattr(app_module, "LT_DIR", "LT_LOCAL", raising=False)
    monkeypatch.setattr(app_module, "LT_PORT", 8010, raising=False)
    monkeypatch.setattr(app_module, "lt_is_running", lambda: True, raising=False)
    st = client.get("/lt/status")
    assert st.status_code == 200
    assert st.get_json()["running"] is True

    monkeypatch.setattr(app_module, "lt_start_server", lambda: None, raising=False)
    ens = client.get("/lt/ensure")
    assert ens.status_code == 200
    assert ens.get_json()["ok"] is True

    def fail_start():
        raise RuntimeError("LT no inicia")

    monkeypatch.setattr(app_module, "lt_start_server", fail_start, raising=False)
    ens_fail = client.get("/lt/ensure")
    assert ens_fail.status_code == 500
