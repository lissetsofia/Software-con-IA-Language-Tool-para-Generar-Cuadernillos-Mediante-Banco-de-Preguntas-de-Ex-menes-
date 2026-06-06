
# backend/test/test_app_grupos_jobs_cov6.py
import os
import sqlite3
from pathlib import Path

import pytest
from docx import Document as DocxDocument


def _connect(db_path: Path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _patch_sqlite_connection(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov6_grupos.sqlite3"

    def get_connection():
        return _connect(db_path)

    monkeypatch.setattr(app_module, "get_connection", get_connection)
    return db_path


def _cols(cur, table):
    try:
        return {r["name"] if hasattr(r, "keys") and "name" in r.keys() else r[1] for r in cur.execute(f"PRAGMA table_info({table})").fetchall()}
    except Exception:
        return set()


def _ensure_generation_schema(app_module):
    conn = app_module.get_connection()
    cur = conn.cursor()

    cur.execute("""
        CREATE TABLE IF NOT EXISTS grupos (
            idgrupo INTEGER PRIMARY KEY AUTOINCREMENT,
            clave TEXT,
            nombre TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS grupo_tema (
            idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT,
            grupos_idgrupo INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS preguntas (
            idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT,
            examenes_idexamenes INTEGER,
            tema_id INTEGER,
            numero_p INTEGER,
            archivo_nombre TEXT,
            archivo_ruta TEXT
        )
    """)

    conn.commit()
    return conn, cur


def _insert_group_with_topic(app_module, cantidad=1, pregunta_path=None):
    conn, cur = _ensure_generation_schema(app_module)
    cur.execute("INSERT INTO grupos (clave, nombre, activo) VALUES (?, ?, 1)", ("T" + os.urandom(2).hex().upper()[:3], "Grupo test"))
    gid = int(cur.lastrowid)
    cur.execute("INSERT INTO temario (nombre, activo) VALUES (?, 1)", ("Tema cobertura " + os.urandom(2).hex(),))
    tid = int(cur.lastrowid)
    cur.execute(
        "INSERT INTO grupo_tema (grupos_idgrupo, tema_id, cantidad, orden) VALUES (?, ?, ?, 1)",
        (gid, tid, cantidad),
    )
    if pregunta_path is not None:
        cur.execute(
            "INSERT INTO preguntas (examenes_idexamenes, tema_id, numero_p, archivo_nombre, archivo_ruta) VALUES (?, ?, ?, ?, ?)",
            (1, tid, 1, os.path.basename(str(pregunta_path)), str(pregunta_path)),
        )
    conn.commit()
    cur.close()
    conn.close()
    return gid, tid


def _make_docx(path: Path, text="Pregunta de prueba"):
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_grupos_generar_doc_run_errores_basicos(app_module, tmp_path, monkeypatch):
    _patch_sqlite_connection(app_module, tmp_path, monkeypatch)
    assert app_module._grupos_generar_doc_run(1, "xlsx", {}, None)[0:2] == ("err", 400)

    _ensure_generation_schema(app_module)
    missing = app_module._grupos_generar_doc_run(987654, "word", {}, None)
    assert missing[0] == "err" and missing[1] == 404

    gid, _tid = _insert_group_with_topic(app_module, cantidad=0)
    zero = app_module._grupos_generar_doc_run(gid, "word", {}, None)
    assert zero[0] == "err"
    assert zero[1] == 400

    gid2, _tid2 = _insert_group_with_topic(app_module, cantidad=2, pregunta_path=None)
    insufficient = app_module._grupos_generar_doc_run(gid2, "word", {}, None)
    assert insufficient[0] == "err"
    assert insufficient[1] == 409
    assert "faltantes" in insufficient[2]


def test_grupos_generar_doc_run_debug_y_archivo_faltante(app_module, tmp_path, monkeypatch):
    _patch_sqlite_connection(app_module, tmp_path, monkeypatch)
    missing_file = tmp_path / "no_existe.docx"
    gid, _tid = _insert_group_with_topic(app_module, cantidad=1, pregunta_path=missing_file)

    debug = app_module._grupos_generar_doc_run(gid, "word", {"debug": "1"}, None)
    assert debug[0] == "ok"
    assert debug[1]["total_requeridas"] == 1

    normal = app_module._grupos_generar_doc_run(gid, "word", {}, None)
    assert normal[0] == "err"
    assert normal[1] == 409
    assert "archivo no existe" in str(normal[2]).lower()


def test_grupos_generar_doc_run_success_y_pdf_failure(app_module, tmp_path, monkeypatch):
    _patch_sqlite_connection(app_module, tmp_path, monkeypatch)
    out_dir = tmp_path / "descargas"
    out_dir.mkdir()
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(out_dir))

    pregunta = _make_docx(tmp_path / "pregunta.docx")
    gid, _tid = _insert_group_with_topic(app_module, cantidad=1, pregunta_path=pregunta)

    def fake_merge(grouped, out_path, merge_step_cb=None, merge_ops=None):
        assert grouped
        doc = DocxDocument()
        doc.add_paragraph("Examen unido")
        doc.save(out_path)
        if merge_step_cb:
            merge_step_cb(1, "merge")
        return out_path, [], []

    def fake_pdf(src, dst):
        Path(dst).write_bytes(b"%PDF-1.4\n%%EOF")
        return dst

    monkeypatch.setattr(app_module, "_com_disponible", lambda: False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge)
    monkeypatch.setattr(app_module, "docx_a_pdf", fake_pdf)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None)
    monkeypatch.setattr(app_module, "aplanar_listas_a_texto", lambda *_a, **_k: None)

    progress = []
    ok = app_module._grupos_generar_doc_run(
        gid,
        "pdf",
        {"flat": "1"},
        lambda done, total, msg: progress.append((done, total, msg)),
    )
    assert ok[0] == "ok"
    assert ok[1]["preview_kind"] == "pdf"
    assert ok[1]["ruta_rel_pdf"].endswith(".pdf")
    assert progress

    # Fuerza la rama donde falla la generación del PDF.
    gid2, _tid2 = _insert_group_with_topic(app_module, cantidad=1, pregunta_path=pregunta)
    monkeypatch.setattr(app_module, "docx_a_pdf", lambda *_a, **_k: (_ for _ in ()).throw(RuntimeError("pdf roto")))
    fail = app_module._grupos_generar_doc_run(gid2, "word", {}, None)
    assert fail[0] == "err"
    assert fail[1] == 500
    assert "PDF" in fail[2]["error"]


def test_generar_doc_async_worker_ok_y_error(app_module, monkeypatch):
    job_ok = "job_ok_cov6"
    with app_module._generar_doc_jobs_lock:
        app_module._generar_doc_jobs[job_ok] = {
            "status": "queued", "done": 0, "total": 1, "message": "",
            "result": None, "error": None, "http_status": None,
        }

    monkeypatch.setattr(app_module, "_grupos_generar_doc_run", lambda *a, **k: ("ok", {"ok": True, "archivo": "x.docx"}))
    app_module._grupos_generar_doc_async_worker(job_ok, 1, "word", {})
    assert app_module._generar_doc_jobs[job_ok]["status"] == "done"
    assert app_module._generar_doc_jobs[job_ok]["result"]["ok"] is True

    job_err = "job_err_cov6"
    with app_module._generar_doc_jobs_lock:
        app_module._generar_doc_jobs[job_err] = {
            "status": "queued", "done": 0, "total": 1, "message": "",
            "result": None, "error": None, "http_status": None,
        }

    monkeypatch.setattr(app_module, "_grupos_generar_doc_run", lambda *a, **k: ("err", 409, {"error": "sin stock"}))
    app_module._grupos_generar_doc_async_worker(job_err, 1, "word", {})
    assert app_module._generar_doc_jobs[job_err]["status"] == "error"
    assert app_module._generar_doc_jobs[job_err]["http_status"] == 409


def test_partir_guardar_async_worker_ok_error_y_events(client, app_module, monkeypatch):
    job_ok = "partir_ok_cov6"
    with app_module._partir_guardar_jobs_lock:
        app_module._partir_guardar_jobs[job_ok] = {
            "status": "queued", "done": 0, "total": 100, "message": "",
            "result": None, "error": None, "http_status": None,
        }

    def fake_partir_ok(idexamen):
        return app_module.jsonify({"ok": True, "examen": idexamen})

    monkeypatch.setattr(app_module, "partir_y_guardar", fake_partir_ok)
    app_module._partir_guardar_async_worker(job_ok, 77, overwrite=True)
    assert app_module._partir_guardar_jobs[job_ok]["status"] == "done"

    # Lee el endpoint normal y el stream de eventos con un job ya terminado.
    r_status = client.get(f"/api/examenes/partir_y_guardar/jobs/{job_ok}")
    assert r_status.status_code == 200
    r_events = client.get(f"/api/examenes/partir_y_guardar/jobs/{job_ok}/events")
    assert r_events.status_code == 200
    assert b"progress" in r_events.data

    job_err = "partir_err_cov6"
    with app_module._partir_guardar_jobs_lock:
        app_module._partir_guardar_jobs[job_err] = {
            "status": "queued", "done": 0, "total": 100, "message": "",
            "result": None, "error": None, "http_status": None,
        }

    def fake_partir_error(idexamen):
        return app_module.jsonify({"ok": False, "error": "docx invalido"}), 409

    monkeypatch.setattr(app_module, "partir_y_guardar", fake_partir_error)
    app_module._partir_guardar_async_worker(job_err, 78, overwrite=False)
    assert app_module._partir_guardar_jobs[job_err]["status"] == "error"
    assert app_module._partir_guardar_jobs[job_err]["http_status"] == 409


def test_generar_doc_job_status_y_events(client, app_module):
    job = "gen_events_cov6"
    with app_module._generar_doc_jobs_lock:
        app_module._generar_doc_jobs[job] = {
            "status": "done", "done": 1, "total": 1, "message": "done",
            "result": {"ok": True}, "error": None, "http_status": None,
        }

    r = client.get(f"/api/grupos/generar_doc/jobs/{job}")
    assert r.status_code == 200
    assert r.get_json()["status"] == "done"

    ev = client.get(f"/api/grupos/generar_doc/jobs/{job}/events")
    assert ev.status_code == 200
    assert b"progress" in ev.data

    assert client.get("/api/grupos/generar_doc/jobs/no_existe_cov6").status_code == 404
    assert client.get("/api/grupos/generar_doc/jobs/no_existe_cov6/events").status_code == 404
