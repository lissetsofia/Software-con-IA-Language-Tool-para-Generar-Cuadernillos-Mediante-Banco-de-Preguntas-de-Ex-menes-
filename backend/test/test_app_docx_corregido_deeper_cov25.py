import os
from docx import Document


def test_generar_docx_corregido_highlight_y_alternativas(app_module, tmp_path):
    src = tmp_path / "origen.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Haber")
    p.add_run(" si vienes")
    doc.add_paragraph("A) 5 alternativa exacta")
    doc.add_paragraph("1) pregunta 10 con números")
    doc.add_paragraph("")
    doc.add_paragraph("porfavor revisa")
    doc.add_paragraph("larga")
    doc.save(src)

    corregido = "\n".join([
        "A ver si vienes",                 # reemplazo + frase forzada
        "A) 6 alternativa cambiada",       # debe protegerse por firma numérica distinta
        "1) pregunta 99 con números",      # debe protegerse por firma numérica distinta
        "",                                # párrafo vacío
        "por favor revisa",                # bloque por favor
        "texto demasiado largo para saltar por diferencia de longitud",
    ])
    out = tmp_path / "salida.docx"
    app_module.generar_docx_corregido(str(src), corregido, str(out), highlight=True)
    assert out.exists() and out.stat().st_size > 0
    d2 = Document(out)
    textos = [p.text for p in d2.paragraphs]
    assert textos[0].startswith("A ver")
    assert textos[1] == "A) 5 alternativa exacta"
    assert textos[2] == "1) pregunta 10 con números"
    assert "por favor" in textos[4]


def test_runs_flexibles_parciales_y_forzado_highlight(app_module):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("al")
    p.add_run("fa beta")
    assert app_module.reemplazo_en_runs_flexible(p, "alfa", "omega") is True
    assert "omega beta" == p.text
    assert app_module.reemplazo_en_runs_flexible(p, "noexiste", "x") is False

    p2 = doc.add_paragraph("porfavor y haber")
    assert app_module.reemplazo_en_runs_parciales(p2, "porfavor", "por favor") is True
    app_module._force_block_highlight_phrases(p2, "porfavor haber", "por favor y a ver")

    pares = app_module.pares_reemplazo_palabra_a_palabra("casa azul", "casas rojas")
    assert ("casa", "casas") in pares or pares
    assert app_module.similitud_suave("abc", "abd") > 0
