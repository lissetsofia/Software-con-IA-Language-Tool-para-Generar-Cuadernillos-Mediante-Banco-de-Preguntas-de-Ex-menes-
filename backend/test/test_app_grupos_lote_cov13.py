import os
import sqlite3
import zipfile
from pathlib import Path

from docx import Document


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov13_grupos.sqlite"

    def connect():
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        return conn

    monkeypatch.setattr(app_module, "get_connection", connect)

    uploads = tmp_path / "uploads"
    descargas = tmp_path / "descargas"
    grupos_out = tmp_path / "grupos_out"
    exam_import = uploads / "examenes"
    for p in (uploads, descargas, grupos_out, exam_import):
        p.mkdir(parents=True, exist_ok=True)

    app_module.app.config["UPLOAD_FOLDER"] = str(uploads)
    app_module.app.config["DESCARGAS_FOLDER"] = str(descargas)
    app_module.app.config["UPLOADS_EXAM_DIR"] = str(exam_import)
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(descargas), raising=False)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(grupos_out), raising=False)

    conn = connect()
    cur = conn.cursor()
    cur.executescript(
        """
        CREATE TABLE IF NOT EXISTS temario (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            activo INTEGER DEFAULT 1
        );
        CREATE TABLE IF NOT EXISTS matriz (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS matriz_detalle (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0,
            archivo_ruta TEXT,
            UNIQUE(matriz_id, tema_id)
        );
        CREATE TABLE IF NOT EXISTS grupos (
            idgrupo INTEGER PRIMARY KEY AUTOINCREMENT,
            clave TEXT,
            nombre TEXT,
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS grupo_tema (
            idgrupo_tema INTEGER PRIMARY KEY AUTOINCREMENT,
            grupos_idgrupo INTEGER,
            tema_id INTEGER,
            cantidad INTEGER DEFAULT 0,
            orden INTEGER DEFAULT 0,
            UNIQUE(grupos_idgrupo, tema_id)
        );
        CREATE TABLE IF NOT EXISTS preguntas (
            idpreguntas INTEGER PRIMARY KEY AUTOINCREMENT,
            examenes_idexamenes INTEGER,
            tema_id INTEGER,
            numero_p INTEGER,
            enunciado TEXT,
            alternativa_a TEXT,
            alternativa_b TEXT,
            alternativa_c TEXT,
            alternativa_d TEXT,
            archivo_nombre TEXT,
            archivo_ruta TEXT
        );
        CREATE TABLE IF NOT EXISTS gen_lote (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            nombre TEXT,
            usuario TEXT
        );
        """
    )
    conn.commit()
    cur.close()
    conn.close()
    return db_path


def _make_docx(path: Path, texto="1. Pregunta de prueba"):
    doc = Document()
    doc.add_paragraph(texto)
    doc.save(path)
    return path


def _seed_matriz_grupo(app_module, tmp_path, monkeypatch, *, cantidad_matriz=2, cantidad_grupo=1, dos_grupos=False):
    db_path = _patch_db(app_module, tmp_path, monkeypatch)
    src = _make_docx(tmp_path / "tema.docx")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("INSERT INTO temario(id, nombre, activo) VALUES (1, 'ARITMÉTICA', 1)")
    cur.execute("INSERT INTO matriz(id, nombre) VALUES (1, 'Matriz Cov13')")
    cur.execute(
        "INSERT INTO matriz_detalle(matriz_id, tema_id, cantidad, orden, archivo_ruta) VALUES (1, 1, ?, 1, ?)",
        (cantidad_matriz, str(src)),
    )
    cur.execute("INSERT INTO grupos(idgrupo, clave, nombre, activo) VALUES (1, 'A', 'Grupo A', 1)")
    cur.execute(
        "INSERT INTO grupo_tema(grupos_idgrupo, tema_id, cantidad, orden) VALUES (1, 1, ?, 1)",
        (cantidad_grupo,),
    )
    if dos_grupos:
        cur.execute("INSERT INTO grupos(idgrupo, clave, nombre, activo) VALUES (2, 'B', 'Grupo B', 1)")
        cur.execute(
            "INSERT INTO grupo_tema(grupos_idgrupo, tema_id, cantidad, orden) VALUES (2, 1, ?, 1)",
            (cantidad_grupo,),
        )
    conn.commit(); cur.close(); conn.close()
    return db_path, src


def test_api_generar_por_grupos_success_y_descarga_zip(client, app_module, tmp_path, monkeypatch):
    _seed_matriz_grupo(app_module, tmp_path, monkeypatch, dos_grupos=True)

    def fake_cut(src, n, dst):
        _make_docx(Path(dst), f"1. Recortada {n}")

    def fake_merge(grouped, out_path, *args, **kwargs):
        doc = Document()
        for tema, files in grouped:
            doc.add_paragraph(tema)
            doc.add_paragraph(os.path.basename(files[0]))
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_cut_docx_first_n_questions", fake_cut, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    resp = client.post("/api/grupos/generar", json={"matriz_id": 1})
    assert resp.status_code == 200, resp.get_data(as_text=True)
    data = resp.get_json()
    assert data["ok"] is True and data["zip_url"].endswith("/zip")

    zip_resp = client.get(data["zip_url"])
    assert zip_resp.status_code == 200
    assert zip_resp.mimetype in ("application/zip", "application/x-zip-compressed")


def test_api_generar_por_grupos_errores_de_validacion(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)

    assert client.post("/api/grupos/generar", json={}).status_code == 400
    assert client.post("/api/grupos/generar", json={"matriz": {"items": []}}).status_code == 400
    assert client.get("/api/grupos/lote/999/zip").status_code == 404

    # matriz con detalle, pero sin grupos activos
    conn = app_module.get_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO temario(id, nombre) VALUES (1, 'BIOLOGÍA')")
    cur.execute("INSERT INTO matriz(id, nombre) VALUES (1, 'M')")
    cur.execute("INSERT INTO matriz_detalle(matriz_id, tema_id, cantidad, orden, archivo_ruta) VALUES (1,1,1,1,'x.docx')")
    conn.commit(); cur.close(); conn.close()
    r = client.post("/api/grupos/generar", json={"matriz_id": 1})
    assert r.status_code == 400
    assert "grupos" in r.get_json().get("error", "").lower()


def test_api_generar_por_grupos_cuotas_y_archivos_invalidos(client, app_module, tmp_path, monkeypatch):
    # cuota mayor que matriz
    _seed_matriz_grupo(app_module, tmp_path, monkeypatch, cantidad_matriz=1, cantidad_grupo=3)
    r = client.post("/api/grupos/generar", json={"matriz_id": 1})
    assert r.status_code == 400
    assert "matriz" in r.get_json().get("error", "").lower()

    # tema del grupo no está en matriz
    _patch_db(app_module, tmp_path / "case2", monkeypatch)
    conn = app_module.get_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO temario(id,nombre) VALUES (1,'A'),(2,'B')")
    cur.execute("INSERT INTO matriz(id,nombre) VALUES (1,'M')")
    cur.execute("INSERT INTO matriz_detalle(matriz_id,tema_id,cantidad,orden,archivo_ruta) VALUES (1,1,2,1,'a.docx')")
    cur.execute("INSERT INTO grupos(idgrupo,clave,nombre,activo) VALUES (1,'A','A',1)")
    cur.execute("INSERT INTO grupo_tema(grupos_idgrupo,tema_id,cantidad,orden) VALUES (1,2,1,1)")
    conn.commit(); cur.close(); conn.close()
    r = client.post("/api/grupos/generar", json={"matriz_id": 1})
    assert r.status_code == 400
    assert "no está en la matriz" in r.get_json().get("error", "")


def test_helpers_de_grupos_y_matriz_directos(app_module, tmp_path, monkeypatch):
    db_path, _src = _seed_matriz_grupo(app_module, tmp_path, monkeypatch)
    conn = sqlite3.connect(db_path); conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    detalle = app_module._leer_matriz_detalle(cur, 1)
    assert detalle and detalle[0]["tema_nombre"] == "ARITMÉTICA"

    grupos = app_module._leer_config_grupos(cur)
    assert 1 in grupos and grupos[1]["temas"][0]["cantidad"] == 1

    cur.execute("INSERT INTO preguntas(tema_id,enunciado,alternativa_a,alternativa_b,alternativa_c,alternativa_d) VALUES (1,'E','A','B','C','D')")
    conn.commit()
    por_tema = app_module._preguntas_por_tema(cur)
    assert 1 in por_tema and por_tema[1][0]["enunciado"] == "E"
    cur.close(); conn.close()

    out_docx = tmp_path / "grupo.docx"
    app_module._armar_docx_grupo("Ingenierías", "A", [{"titulo": "ARITMÉTICA", "pregs": por_tema[1]}], str(out_docx))
    assert out_docx.exists() and out_docx.stat().st_size > 0


def test_matriz_endpoints_basicos_y_upload(client, app_module, tmp_path, monkeypatch):
    _patch_db(app_module, tmp_path, monkeypatch)
    conn = app_module.get_connection(); cur = conn.cursor()
    cur.execute("INSERT INTO temario(id,nombre) VALUES (1,'FÍSICA')")
    conn.commit(); cur.close(); conn.close()

    assert client.post("/api/matriz", json={"items": []}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": "x"}]}).status_code == 400
    assert client.post("/api/matriz", json={"items": [{"tema_id": 0, "cantidad": 1}]}).status_code == 400

    r = client.post("/api/matriz", json={"nombre": "M1", "items": [{"tema_id": 1, "cantidad": 2, "orden": 1}]})
    assert r.status_code == 200, r.get_data(as_text=True)
    mid = r.get_json()["matriz_id"]

    assert client.get("/api/matriz?detail=1&search=M1").status_code == 200
    g = client.get(f"/api/matriz/{mid}")
    assert g.status_code == 200 and g.get_json()["n_items"] == 1
    assert client.get("/api/matriz/9999").status_code == 404

    # upload: parámetros faltantes y extensión incorrecta cubren salidas tempranas.
    assert client.post(f"/api/matriz/{mid}/upload", data={}).status_code == 400
    bad = {"tema_id": "1", "file": (open(__file__, "rb"), "no.txt")}
    assert client.post(f"/api/matriz/{mid}/upload", data=bad, content_type="multipart/form-data").status_code == 400

    monkeypatch.setattr(app_module, "_validar_docx_real", lambda *_a, **_k: None, raising=False)
    docx = _make_docx(tmp_path / "subida.docx")
    with open(docx, "rb") as fh:
        ok = client.post(
            f"/api/matriz/{mid}/upload",
            data={"tema_id": "1", "cantidad": "3", "file": (fh, "subida.docx")},
            content_type="multipart/form-data",
        )
    assert ok.status_code == 200, ok.get_data(as_text=True)
    assert ok.get_json()["ok"] is True
