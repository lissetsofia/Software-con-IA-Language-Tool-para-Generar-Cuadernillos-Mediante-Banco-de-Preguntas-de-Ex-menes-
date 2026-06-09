import io
import os
import sqlite3
import zipfile
from pathlib import Path

from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _docx_bytes(text="Contenido"):
    bio = io.BytesIO()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def _make_docx(path, text="Contenido"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _patch_lote_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov16_from_docx.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE gen_lote(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            matriz_id INTEGER,
            nombre TEXT,
            usuario TEXT
        );
        """
    )
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    out = tmp_path / "grupos_out"; out.mkdir(exist_ok=True)
    monkeypatch.setattr(app_module, "GRUPOS_OUT_DIR", str(out), raising=False)
    return db_path, out


def _groups_same_quota():
    return {
        1: {"idgrupo": 1, "clave": "A", "nombre": "Grupo A", "temas": [
            {"tema_id": 10, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}
        ]},
        2: {"idgrupo": 2, "clave": "B", "nombre": "Grupo B", "temas": [
            {"tema_id": 10, "tema_nombre": "Álgebra", "cantidad": 1, "orden": 1}
        ]},
    }


def test_api_generar_grupos_from_docx_validaciones_y_success(client, app_module, tmp_path, monkeypatch):
    _patch_lote_db(app_module, tmp_path, monkeypatch)

    # Validaciones iniciales antes de entrar a la lógica pesada.
    assert client.post("/api/grupos/generar_from_docx", data={}, content_type="multipart/form-data").status_code == 400
    bad = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(b"x"), "matriz.txt")},
        content_type="multipart/form-data",
    )
    assert bad.status_code == 400

    # Rama: no hay grupos configurados.
    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: {}, raising=False)
    r_no_groups = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_no_groups.status_code == 400
    assert "grupos" in r_no_groups.get_json()["error"].lower()

    # Rama: cuotas distintas entre grupos para el mismo tema.
    def grupos_cuotas_distintas(_cur):
        g = _groups_same_quota()
        g[2]["temas"][0]["cantidad"] = 2
        return g

    monkeypatch.setattr(app_module, "_leer_config_grupos", grupos_cuotas_distintas, raising=False)
    r_diff = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_diff.status_code == 400
    assert "cuotas" in r_diff.get_json()["error"].lower()

    # Rama: faltan títulos en el DOCX para los temas configurados.
    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _groups_same_quota(), raising=False)
    monkeypatch.setattr(app_module, "_build_heading_ranges_for_temas", lambda path, nombres: {}, raising=False)
    r_missing = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_missing.status_code == 400
    assert "faltan" in r_missing.get_json()["error"].lower()

    # Camino correcto: extrae el tema, arma dos grupos y genera el ZIP.
    monkeypatch.setattr(app_module, "_build_heading_ranges_for_temas", lambda path, nombres: {"ALGEBRA": (0, 2)}, raising=False)

    def fake_extract(src, start_idx, end_idx, dest_path):
        _make_docx(dest_path, f"extraido {start_idx}-{end_idx}")

    def fake_merge(bloques, out_path):
        doc = Document()
        for tema, files in bloques:
            doc.add_paragraph(f"TEMA {tema}")
            for f in files:
                doc.add_paragraph(os.path.basename(f))
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        return out_path, [], []

    monkeypatch.setattr(app_module, "_extract_tema_docx_range", fake_extract, raising=False)
    monkeypatch.setattr(app_module, "_com_disponible", lambda: False, raising=False)
    monkeypatch.setattr(app_module, "_merge_grouped_with_headings", fake_merge, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda *_a, **_k: None, raising=False)

    r_ok = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes("Álgebra\n1. Pregunta")), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    body = r_ok.get_json()
    assert body["ok"] is True
    assert body["zip_url"].endswith(f"/api/grupos/lote/{body['lote_id']}/zip")
    zip_path = tmp_path / "grupos_out" / f"lote_{body['lote_id']}" / f"grupos_{body['lote_id']}.zip"
    assert zipfile.is_zipfile(zip_path)
    with zipfile.ZipFile(zip_path) as zf:
        assert {"grupo_A.docx", "grupo_B.docx"}.issubset(set(zf.namelist()))


def test_api_generar_grupos_from_docx_ramas_de_error_post_extraccion(client, app_module, tmp_path, monkeypatch):
    _patch_lote_db(app_module, tmp_path, monkeypatch)
    monkeypatch.setattr(app_module, "_leer_config_grupos", lambda cur: _groups_same_quota(), raising=False)
    monkeypatch.setattr(app_module, "_build_heading_ranges_for_temas", lambda path, nombres: {"ALGEBRA": (0, 2)}, raising=False)

    def extract_falla(src, start_idx, end_idx, dest_path):
        raise RuntimeError("corte fallido")

    monkeypatch.setattr(app_module, "_extract_tema_docx_range", extract_falla, raising=False)
    r_extract = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_extract.status_code == 500
    assert r_extract.get_json()["ok"] is False

    def extract_invalido(src, start_idx, end_idx, dest_path):
        Path(dest_path).write_bytes(b"no es docx")

    monkeypatch.setattr(app_module, "_extract_tema_docx_range", extract_invalido, raising=False)
    monkeypatch.setattr(app_module, "reparar_docx_fuerte", lambda p: (False, "sin word"), raising=False)
    r_invalid = client.post(
        "/api/grupos/generar_from_docx",
        data={"file": (io.BytesIO(_docx_bytes()), "matriz.docx")},
        content_type="multipart/form-data",
    )
    assert r_invalid.status_code == 500
    assert "inválido" in r_invalid.get_json()["error"] or "invalido" in r_invalid.get_json()["error"].lower()
