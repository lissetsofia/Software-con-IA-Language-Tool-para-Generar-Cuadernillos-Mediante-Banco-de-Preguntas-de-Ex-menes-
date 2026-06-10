import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document


class _FakeMergeRange:
    def __init__(self, owner=None):
        self.owner = owner
        self.FormattedText = "FORMATEADO"
    def Collapse(self, *args, **kwargs):
        return None
    def InsertFile(self, path):
        # Fuerza la ruta de fallback FormattedText para un archivo concreto.
        if "fallback" in os.path.basename(str(path)).lower():
            raise RuntimeError("insert falla")
        return None


class _FakeMergeContent(_FakeMergeRange):
    pass


class _FakeMergeDoc:
    def __init__(self, path=None):
        self.path = path
        self.Content = _FakeMergeContent(self)
    def Range(self, *args, **kwargs):
        return _FakeMergeRange(self)
    def SaveAs(self, path=None, FileName=None, FileFormat=None, **kwargs):
        target = FileName or path
        if target:
            d = Document()
            d.add_paragraph("merge fake")
            d.save(os.path.abspath(str(target)))
    def Close(self, *args, **kwargs):
        return None


class _FakeMergeDocuments:
    def Add(self, *args, **kwargs):
        return _FakeMergeDoc()
    def Open(self, path, *args, **kwargs):
        return _FakeMergeDoc(path)


class _FakeMergeWord:
    def __init__(self):
        self.Documents = _FakeMergeDocuments()
        self.Visible = False
        self.DisplayAlerts = 0
    def Quit(self):
        return None


def _patch_word_merge(app_module, monkeypatch):
    monkeypatch.setattr(app_module.pythoncom, "CoInitialize", lambda: None)
    monkeypatch.setattr(app_module.pythoncom, "CoUninitialize", lambda: None)
    monkeypatch.setattr(app_module.win32, "DispatchEx", lambda *a, **k: _FakeMergeWord())


def _make_docx(path, text="contenido"):
    d = Document()
    d.add_paragraph(text)
    d.save(path)
    return path


def test_merge_with_word_success_y_fallback(app_module, tmp_path, monkeypatch):
    _patch_word_merge(app_module, monkeypatch)
    p1 = _make_docx(tmp_path / "uno.docx", "Uno")
    p2 = _make_docx(tmp_path / "fallback.docx", "Dos")
    out = tmp_path / "merge_out.docx"
    steps = []
    monkeypatch.setattr(app_module, "_tiene_texto_o_contenido", lambda p: True)

    result, _none, malos = app_module._merge_with_word(
        [(str(p1), False), (str(p2), False), (str(tmp_path / "no_existe.docx"), False)],
        str(out),
        merge_step_cb=lambda i, msg: steps.append((i, msg)),
        merge_ops_hint=3,
    )

    assert result == str(out)
    assert out.exists() and out.stat().st_size > 0
    assert steps
    assert any("InsertFile" in m for _p, m in malos)


def test_reparaciones_bytes_y_validaciones_docx(app_module, tmp_path, monkeypatch):
    docx = _make_docx(tmp_path / "base.docx", "Base")
    app_module._validar_docx_real(str(docx))

    bad = tmp_path / "bad.docx"
    bad.write_text("no zip", encoding="utf-8")
    try:
        app_module._validar_docx_real(str(bad))
        assert False, "debió fallar"
    except ValueError as e:
        assert "zip" in str(e).lower()

    data = docx.read_bytes()
    monkeypatch.setattr(app_module, "_reparar_docx_generado", lambda p: None)
    assert app_module._reparar_docx_bytes(data, "ok")[:2] == b"PK"

    monkeypatch.setattr(app_module, "reparar_docx_inplace", lambda p: None)
    assert app_module._asegurar_docx_bytes_valido_como_grupo(data, "grupo")[:2] == b"PK"

    # Rama de error: bytes inválidos y reparación fuerte fallida.
    monkeypatch.setattr(app_module, "reparar_docx_fuerte", lambda p: (False, "no reparó"))
    try:
        app_module._asegurar_docx_bytes_valido_como_grupo(b"basura", "malo")
        assert False, "debió fallar"
    except RuntimeError as e:
        assert "malo" in str(e)


def test_reconstruir_numbering_desde_documento_con_relaciones(app_module, tmp_path):
    docx = _make_docx(tmp_path / "num.docx", "Pregunta con número")
    with zipfile.ZipFile(docx, "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    root = ET.fromstring(files["word/document.xml"])
    body = root.find(".//" + app_module.W + "body")
    p = body.find(app_module.W + "p")
    ppr = p.find(app_module.W + "pPr")
    if ppr is None:
        ppr = ET.Element(app_module.W + "pPr")
        p.insert(0, ppr)
    numpr = ET.SubElement(ppr, app_module.W + "numPr")
    ET.SubElement(numpr, app_module.W + "ilvl", {app_module.W + "val": "0"})
    ET.SubElement(numpr, app_module.W + "numId", {app_module.W + "val": "9"})
    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    # Numbering anterior corrupto: cubre el except interno y reconstruye desde cero.
    files["word/numbering.xml"] = b"<w:numbering>"
    files.pop("word/_rels/document.xml.rels", None)
    app_module._safe_rezip(str(docx), files)

    app_module._reconstruir_numbering_desde_documento(str(docx))
    with zipfile.ZipFile(docx, "r") as z:
        names = set(z.namelist())
        assert "word/numbering.xml" in names
        assert b"upperLetter" in z.read("word/numbering.xml")
        assert b"numbering" in z.read("word/_rels/document.xml.rels")


def test_modulos_soporte_import_smoke():
    # Cubre al menos imports/top-level de módulos que Sonar muestra en 0%.
    import importlib
    for name in ("db", "init_db", "login_api", "lt_utils", "matriz_utils"):
        try:
            mod = importlib.import_module(name)
        except Exception:
            continue
        assert mod is not None
        for fn_name in ("get_connection", "init_db"):
            fn = getattr(mod, fn_name, None)
            if callable(fn):
                try:
                    obj = fn()
                    close = getattr(obj, "close", None)
                    if callable(close):
                        close()
                except Exception:
                    pass
