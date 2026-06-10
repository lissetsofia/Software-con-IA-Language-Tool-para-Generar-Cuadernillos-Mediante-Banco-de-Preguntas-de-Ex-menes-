import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document


def _p(app_module, text="", with_page_break=False):
    p = ET.Element(app_module.W + "p")
    if text:
        r = ET.SubElement(p, app_module.W + "r")
        t = ET.SubElement(r, app_module.W + "t")
        t.text = text
    if with_page_break:
        r = ET.SubElement(p, app_module.W + "r")
        ET.SubElement(r, app_module.W + "br", {app_module.W + "type": "page"})
    return p


def test_sanear_fragmentos_y_reempacar_docx(app_module, tmp_path):
    base = tmp_path / "base.docx"
    d = Document()
    d.add_paragraph("base")
    d.save(base)

    work = tmp_path / "work"
    work.mkdir()
    with zipfile.ZipFile(base, "r") as z:
        z.extractall(work)

    # Fragmento con sdt, delete y sectPr interno para cubrir saneamiento.
    p = ET.Element(app_module.W + "p", {app_module.W + "rsidR": "001"})
    ppr = ET.SubElement(p, app_module.W + "pPr")
    ET.SubElement(ppr, app_module.W + "sectPr")
    sdt = ET.SubElement(p, app_module.W + "sdt")
    ET.SubElement(sdt, app_module.W + "sdtPr")
    content = ET.SubElement(sdt, app_module.W + "sdtContent")
    r = ET.SubElement(content, app_module.W + "r")
    t = ET.SubElement(r, app_module.W + "t")
    t.text = "Pregunta saneada"
    ET.SubElement(p, app_module.W + "del")

    app_module._sanear_fragmento(p)
    assert "rsid" not in " ".join(p.attrib.keys())
    assert p.find(".//" + app_module.W + "sdtPr") is None
    assert p.find(".//" + app_module.W + "del") is None

    vacio = _p(app_module, with_page_break=True)
    assert app_module._parrafo_esta_vacio_o_es_solo_salto(vacio) is True
    assert app_module._parrafo_esta_vacio_o_es_solo_salto(_p(app_module, "texto")) is False

    out = tmp_path / "out.docx"
    app_module._reempacar_docx(str(work), [p, ET.Element(app_module.W + "sectPr")], str(out))
    assert out.exists() and out.stat().st_size > 0
    assert "Pregunta saneada" in "\n".join(x.text for x in Document(out).paragraphs)

    assert app_module._norm("Álgebra  II") == "algebra ii"
    assert app_module._norm_tema("Álgebra") == "ALGEBRA"
    assert app_module._slug("Álgebra II !!").rstrip("_") == "algebra_ii"
    assert app_module._texto_visible_de_bloque(_p(app_module, "Visible")) == "Visible"
    tbl = ET.Element(app_module.W + "tbl")
    assert app_module._texto_visible_de_bloque(tbl) == "[TABLA]"


def test_reempacar_docx_sin_elementos_crea_docx_valido(app_module, tmp_path):
    base = tmp_path / "base.docx"
    Document().save(base)
    work = tmp_path / "work2"; work.mkdir()
    with zipfile.ZipFile(base, "r") as z:
        z.extractall(work)
    out = tmp_path / "vacio.docx"
    app_module._reempacar_docx(str(work), [], str(out))
    assert out.exists()
    Document(out)
