import os
import sqlite3
from docx import Document


def _connect(db_path):
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn


def _make_docx(path, text='Pregunta válida'):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def _patch_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / 'cov22_generar_doc.sqlite3'
    if db_path.exists():
        db_path.unlink()
    desc = tmp_path / 'desc'; desc.mkdir(exist_ok=True)
    qdir = tmp_path / 'preguntas'; qdir.mkdir(exist_ok=True)
    good = _make_docx(qdir / 'p1.docx')
    missing = qdir / 'missing.docx'
    conn = _connect(db_path)
    conn.executescript(
        f"""
        CREATE TABLE grupos(idgrupo INTEGER PRIMARY KEY, clave TEXT, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE temario(id INTEGER PRIMARY KEY, nombre TEXT, activo INTEGER DEFAULT 1);
        CREATE TABLE grupo_tema(idgrupo_tema INTEGER PRIMARY KEY, grupos_idgrupo INTEGER, tema_id INTEGER, cantidad INTEGER, orden INTEGER);
        CREATE TABLE preguntas(idpreguntas INTEGER PRIMARY KEY, tema_id INTEGER, archivo_ruta TEXT);
        INSERT INTO temario(id,nombre,activo) VALUES(1,'Álgebra',1);
        INSERT INTO grupos(idgrupo,clave,nombre,activo) VALUES(1,'A','Sin cuotas',1),(2,'B','Cero',1),(3,'C','Faltan',1),(4,'D','Ok',1),(5,'E','Missing',1);
        INSERT INTO grupo_tema(grupos_idgrupo,tema_id,cantidad,orden) VALUES(2,1,0,1),(3,1,2,1),(4,1,1,1),(5,1,1,1);
        INSERT INTO preguntas(tema_id,archivo_ruta) VALUES(1,'{str(good).replace("'", "''")}');
        """
    )
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, 'get_connection', lambda: _connect(db_path))
    monkeypatch.setitem(app_module.app.config, 'DESCARGAS_FOLDER', str(desc))
    monkeypatch.setattr(app_module, '_com_disponible', lambda: False)
    monkeypatch.setattr(app_module, 'reparar_docx_inplace', lambda p: None)
    monkeypatch.setattr(app_module, 'aplanar_listas_a_texto', lambda p: None)
    return db_path, desc, good, missing


def test_grupos_generar_doc_run_errores_y_debug(client, app_module, tmp_path, monkeypatch):
    db_path, desc, good, missing = _patch_db(app_module, tmp_path, monkeypatch)

    assert app_module._grupos_generar_doc_run(1, 'xls', {}, None)[1] == 400
    assert app_module._grupos_generar_doc_run(99, 'word', {}, None)[1] == 404
    assert app_module._grupos_generar_doc_run(1, 'word', {}, None)[1] == 400
    assert app_module._grupos_generar_doc_run(2, 'word', {}, None)[1] == 400
    insufficient = app_module._grupos_generar_doc_run(3, 'word', {}, None)
    assert insufficient[0] == 'err' and insufficient[1] == 409

    dbg = app_module._grupos_generar_doc_run(4, 'word', {'debug': '1'}, None)
    assert dbg[0] == 'ok'
    assert dbg[1]['total_requeridas'] == 1

    conn = _connect(db_path)
    conn.execute('UPDATE preguntas SET archivo_ruta=?', (str(missing),))
    conn.commit(); conn.close()
    miss = app_module._grupos_generar_doc_run(5, 'word', {}, None)
    assert miss[0] == 'err' and miss[1] == 409


def test_grupos_generar_doc_route_success_merge_malos_pdf_error(client, app_module, tmp_path, monkeypatch):
    db_path, desc, good, _missing = _patch_db(app_module, tmp_path, monkeypatch)
    progress = []

    def fake_merge(grouped, out_path, merge_step_cb=None, merge_ops=None):
        progress.append(('merge', len(grouped)))
        doc = Document(); doc.add_paragraph('Final'); doc.save(out_path)
        if merge_step_cb:
            merge_step_cb(1, 'merge')
        return out_path, [], []

    def fake_pdf(src, dst):
        with open(dst, 'wb') as f:
            f.write(b'%PDF-1.4\n%%EOF')

    monkeypatch.setattr(app_module, '_merge_grouped_with_headings', fake_merge)
    monkeypatch.setattr(app_module, 'docx_a_pdf', fake_pdf)

    r = client.get('/api/grupos/4/generar_doc?formato=pdf&flat=1')
    assert r.status_code == 200
    assert r.get_json()['ok'] is True
    assert r.get_json()['preview_kind'] == 'pdf'
    assert progress

    def merge_malos(grouped, out_path, *a, **k):
        doc = Document(); doc.add_paragraph('bad'); doc.save(out_path)
        return out_path, [], [(grouped[0][1][0], 'fallo')]
    monkeypatch.setattr(app_module, '_merge_grouped_with_headings', merge_malos)
    bad = client.get('/api/grupos/4/generar_doc')
    assert bad.status_code == 409

    monkeypatch.setattr(app_module, '_merge_grouped_with_headings', fake_merge)
    monkeypatch.setattr(app_module, 'docx_a_pdf', lambda src, dst: (_ for _ in ()).throw(RuntimeError('sin pdf')))
    no_pdf = client.get('/api/grupos/4/generar_doc')
    assert no_pdf.status_code == 500

    assert client.get('/api/grupos/D/generar_doc?formato=word').status_code == 500
    assert client.get('/api/grupos/NOEXISTE/generar_doc').status_code == 404
