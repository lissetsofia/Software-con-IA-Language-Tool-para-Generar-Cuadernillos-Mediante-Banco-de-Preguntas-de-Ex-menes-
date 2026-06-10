import os
import xml.etree.ElementTree as ET
from docx import Document


class FakeFind:
    def __init__(self, owner):
        self.owner = owner
        self.Text = ""
        self.Forward = True
        self.Wrap = 0
    def ClearFormatting(self):
        return None
    def Execute(self):
        ok = self.Text and self.Text in self.owner.doc_text
        if ok:
            self.owner.Start = 1
            self.owner.End = min(len(self.owner.doc_text), len(self.Text) + 1)
        return bool(ok)


class FakeRange:
    def __init__(self, text="contenido"):
        self.doc_text = text
        self.Start = 0
        self.End = len(text)
        self.FormattedText = "FORMATO"
        self.Text = text
        self.Find = FakeFind(self)
    def Collapse(self, *a, **k):
        return None
    def InsertFile(self, *a, **k):
        return None


class FakeContent(FakeRange):
    pass


class FakeDoc:
    def __init__(self, path=None, text="Inicio visible\nFin visible"):
        self.path = path
        self.Content = FakeContent(text)
        self._text = text
    def Range(self, *args, **kwargs):
        return FakeRange(self._text)
    def SaveAs(self, path=None, FileName=None, FileFormat=None, **kwargs):
        target = FileName or path
        if target:
            Document().save(os.path.abspath(target))
    def Close(self, *a, **k):
        return None


class FakeDocuments:
    def __init__(self, fail_open=False):
        self.fail_open = fail_open
    def Open(self, path, *args, **kwargs):
        if self.fail_open:
            raise RuntimeError("no abre")
        return FakeDoc(path)
    def Add(self, *args, **kwargs):
        return FakeDoc()


class FakeWord:
    def __init__(self, fail_open=False):
        self.Documents = FakeDocuments(fail_open=fail_open)
        self.Visible = False
        self.DisplayAlerts = 0
    def Quit(self):
        return None


def _patch_word(app_module, monkeypatch, fail_open=False):
    monkeypatch.setattr(app_module.pythoncom, "CoInitialize", lambda: None)
    monkeypatch.setattr(app_module.pythoncom, "CoUninitialize", lambda: None)
    monkeypatch.setattr(app_module.win32, "DispatchEx", lambda *a, **k: FakeWord(fail_open=fail_open))


def _p(app_module, text):
    p = ET.Element(app_module.W + "p")
    r = ET.SubElement(p, app_module.W + "r")
    t = ET.SubElement(r, app_module.W + "t")
    t.text = text
    return p


def test_guardar_y_reparar_docx_con_word_fake(app_module, tmp_path, monkeypatch):
    _patch_word(app_module, monkeypatch)
    origen = tmp_path / "origen.docx"
    destino = tmp_path / "destino.docx"
    Document().save(origen)
    elementos = [_p(app_module, "Inicio visible"), _p(app_module, "Fin visible")]

    ok, err = app_module.guardar_pregunta_docx_wordcom(str(origen), str(destino), elementos)
    assert ok and err is None and destino.exists()

    destino2 = tmp_path / "destino2.docx"
    ok, err = app_module.guardar_pregunta_docx_desde_fuente_wordcom(str(origen), str(destino2), elementos)
    assert ok and err is None and destino2.exists()

    reparable = tmp_path / "reparable.docx"
    Document().save(reparable)
    app_module.reparar_docx_inplace(str(reparable))
    assert reparable.exists()

    ok, err = app_module.reparar_docx_fuerte(str(reparable))
    assert ok and err is None and reparable.exists()

    limpio, err = app_module.normalizar_docx_fuente(str(origen))
    assert err is None and limpio and os.path.exists(limpio)

    # Búsqueda por rango y guardado desde rango.
    fake_doc = FakeDoc(text="Inicio visible y luego Fin visible")
    rng = app_module._buscar_rango_por_texto(fake_doc, "Inicio visible", 0)
    assert rng is not None
    assert app_module._buscar_rango_por_texto(fake_doc, "NO ESTA", 0) is None

    destino3 = tmp_path / "rango.docx"
    ok, err = app_module.guardar_pregunta_docx_desde_rango_wordcom(str(origen), str(destino3), elementos)
    assert ok and err is None and destino3.exists()


def test_word_fake_ramas_de_error(app_module, tmp_path, monkeypatch):
    _patch_word(app_module, monkeypatch, fail_open=True)
    origen = tmp_path / "origen.docx"
    Document().save(origen)
    elementos = [_p(app_module, "Texto visible")]

    ok, err = app_module.reparar_docx_fuerte(str(origen))
    assert ok is False and err

    limpio, err = app_module.normalizar_docx_fuente(str(origen))
    assert limpio is None and err

    ok, err = app_module.guardar_pregunta_docx_desde_rango_wordcom(str(origen), str(tmp_path / "x.docx"), [])
    assert ok is False and "texto visible" in err.lower()
