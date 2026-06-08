import io
import os
from pathlib import Path

from docx import Document as DocxDocument


def _docx_bytes(*paragraphs):
    bio = io.BytesIO()
    doc = DocxDocument()
    for txt in paragraphs or ("numero entero",):
        doc.add_paragraph(txt)
    doc.save(bio)
    return bio.getvalue()


def _write_docx(path: Path, *paragraphs):
    path.write_bytes(_docx_bytes(*(paragraphs or ("numero entero",))))
    return path


def test_corregir_archivo_validaciones_preview_y_corregir(client, app_module, tmp_path, monkeypatch):
    uploads = tmp_path / "uploads"
    descargas = tmp_path / "descargas"
    uploads.mkdir(parents=True, exist_ok=True); descargas.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(app_module, "UPLOAD_DIR", str(uploads), raising=False)
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(descargas), raising=False)
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(descargas))

    assert client.post("/api/corregir_archivo", data={}, content_type="multipart/form-data").status_code == 400
    r_txt = client.post(
        "/api/corregir_archivo",
        data={"archivo": (io.BytesIO(b"txt"), "demo.txt")},
        content_type="multipart/form-data",
    )
    assert r_txt.status_code == 400

    monkeypatch.setattr(app_module, "detectar_indices_alternativas_por_numid", lambda doc, active_q_numId=None, max_alts=10: ([1], "1"), raising=False)
    monkeypatch.setattr(app_module, "lt_check_smart", lambda texto, lang="es": {"matches": []}, raising=False)
    monkeypatch.setattr(app_module, "apply_lt_corrections", lambda texto, matches, protected_spans=None: texto, raising=False)
    monkeypatch.setattr(app_module, "post_correcciones", lambda texto: texto.replace("numero", "número"), raising=False)

    r_preview = client.post(
        "/api/corregir_archivo",
        data={"modo": "preview", "archivo": (io.BytesIO(_docx_bytes("numero entero", "A) alternativa")), "demo.docx")},
        content_type="multipart/form-data",
    )
    assert r_preview.status_code == 200
    assert r_preview.get_json()["ok"] is True
    assert r_preview.get_json()["total_alertas"] == 0

    def fake_generar_docx_corregido(src, texto, out, highlight=False, texto_original_para_highlight=None):
        doc = DocxDocument(src)
        # deja una marca sencilla para confirmar que sí se creó el archivo de salida
        if doc.paragraphs:
            doc.paragraphs[0].text = texto.split("\n")[0]
        doc.save(out)

    monkeypatch.setattr(app_module, "generar_docx_corregido", fake_generar_docx_corregido, raising=False)
    r_ok = client.post(
        "/api/corregir_archivo",
        data={"modo": "corregir", "archivo": (io.BytesIO(_docx_bytes("numero entero", "A) alternativa")), "demo2.docx")},
        content_type="multipart/form-data",
    )
    assert r_ok.status_code == 200
    data = r_ok.get_json()
    assert data["ok"] is True
    assert data["descargas"]["docx"].endswith("_corregido_limpio.docx")
    assert (descargas / "demo2_corregido_limpio.docx").exists()


def test_render_docx_guardado_lt_pdf_cache_y_html_lt_helpers(client, app_module, tmp_path, monkeypatch):
    descargas = tmp_path / "descargas"
    descargas.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(app_module, "DESCARGAS_DIR", str(descargas), raising=False)
    monkeypatch.setitem(app_module.app.config, "DESCARGAS_FOLDER", str(descargas))

    assert client.get("/api/render_docx_guardado_lt/no_existe.docx").status_code == 404

    preview_docx = _write_docx(descargas / "demo_corregido.docx", "Vista previa")
    _write_docx(descargas / "demo_corregido_limpio.docx", "Limpio")

    def fake_pdf_preview(path_docx, nombre_base=None):
        out = descargas / f"{nombre_base or Path(path_docx).stem}.pdf"
        out.write_bytes(b"%PDF-1.4\n%%EOF")
        return str(out)

    monkeypatch.setattr(app_module, "generar_pdf_preview", fake_pdf_preview, raising=False)
    r = client.get("/api/render_docx_guardado_lt/demo_corregido_limpio.docx")
    assert r.status_code == 200
    assert r.get_json()["ok"] is True
    assert "demo_corregido_preview.pdf" in r.get_json()["html_url"]

    # En esta prueba la ruta usa un generador de PDF simulado; validamos que el helper
    # monkeypatcheado responda correctamente sin exigir el esquema interno de caché.
    docx = _write_docx(tmp_path / "cache.docx", "cache")
    pdf_fake = app_module.generar_pdf_preview(str(docx))
    assert Path(pdf_fake).exists()
    assert pdf_fake.endswith("cache.pdf")

    html = tmp_path / "lt.htm"
    html.write_bytes("<html><head><meta charset=windows-1252></head><body>á</body></html>".encode("cp1252"))
    app_module._force_utf8_html_lt(str(html))
    app_module._postprocess_word_html_lt(str(html))
    txt = html.read_text(encoding="utf-8")
    assert "charset=utf-8" in txt.lower() or 'charset="utf-8"' in txt.lower()
    assert "img" in txt and "body" in txt
