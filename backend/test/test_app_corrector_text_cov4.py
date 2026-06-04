# backend/test/test_app_corrector_text_cov4.py
from pathlib import Path

from docx import Document


def _match(offset, length, value):
    return {"offset": offset, "length": length, "replacements": [{"value": value}]}


def test_languagetool_config_y_request_helpers(app_module, monkeypatch, tmp_path):
    # _java_cmd: si existe JRE embebido lo usa; si no, retorna "java".
    monkeypatch.setattr(app_module, "BASE_DIR", str(tmp_path), raising=False)
    java_path = tmp_path / "jre" / "bin" / ("java.exe" if app_module.os.name == "nt" else "java")
    java_path.parent.mkdir(parents=True)
    java_path.write_text("fake", encoding="utf-8")
    assert app_module._java_cmd().endswith("java.exe" if app_module.os.name == "nt" else "java")

    lt_dir = tmp_path / "LanguageTool"
    lt_dir.mkdir()
    jar = lt_dir / "languagetool-server.jar"
    jar.write_text("jar", encoding="utf-8")
    monkeypatch.setenv("LT_DIR", str(lt_dir))
    assert app_module._resolve_lt_dir() == str(lt_dir)
    monkeypatch.setattr(app_module, "LT_DIR", str(lt_dir), raising=False)
    assert app_module._find_lt_jar() == str(jar)

    ngrams = tmp_path / "ngrams" / "es"
    (ngrams / "1grams").mkdir(parents=True)
    (ngrams / "2grams").mkdir(parents=True)
    monkeypatch.setenv("NGRAMS_DIR", str(ngrams))
    assert app_module._detect_ngrams_dir() == str(ngrams)

    monkeypatch.setattr(app_module, "lt_is_running", lambda *a, **k: True, raising=False)
    assert app_module.lt_start_server() is None

    class FakeResp:
        status_code = 200
        text = "OK"

        def raise_for_status(self):
            self.raised = True

        def json(self):
            return {"matches": []}

    class FakeHTTP:
        def __init__(self):
            self.calls = []

        def post(self, url, data=None, timeout=None):
            self.calls.append((url, data, timeout))
            return FakeResp()

    fake_http = FakeHTTP()
    monkeypatch.setattr(app_module, "LT_HTTP", fake_http, raising=False)
    out = app_module._lt_request("hola", "es")
    assert out["_status"] == 200
    assert fake_http.calls[0][1]["language"] == "es"

    class BadResp(FakeResp):
        status_code = 400
        text = "too long"

    class BadHTTP(FakeHTTP):
        def post(self, url, data=None, timeout=None):
            return BadResp()

    monkeypatch.setattr(app_module, "LT_HTTP", BadHTTP(), raising=False)
    assert app_module._lt_request("x", "es")["_status"] == 400


def test_lt_check_smart_usa_fallback_y_chunks(app_module, monkeypatch):
    calls = []

    def fake_lt_request(texto, lang, use_picky=True, use_variant=True):
        calls.append((texto, use_picky, use_variant))
        # Primera llamada falla por largo; luego obliga a entrar al flujo por chunks.
        if len(calls) <= 2:
            return {"_status": 400, "_body": "too long"}
        return {"_status": 200, "matches": [{"offset": 1, "length": 2, "replacements": [{"value": "ok"}]}]}

    monkeypatch.setattr(app_module, "_lt_request", fake_lt_request, raising=False)
    monkeypatch.setattr(app_module, "LT_SOFT_CHUNK", 8, raising=False)
    texto = "abc def\n\nghi jkl\n\nmno pqr"
    out = app_module.lt_check_smart(texto, lang="es")
    assert out["matches"]
    assert len(calls) >= 3


def test_normalizacion_sugerencias_y_filtros_de_correccion(app_module):
    assert app_module.normalize_ocr_noise("pa-\nlabra\u00A0  final") == "palabra final"
    assert app_module._strip_accents("ÁÉÍÓÚ ñ") == "AEIOU n"
    assert app_module._same_casing("NUMERO", "número") == "NÚMERO"
    assert app_module._same_casing("Numero", "número") == "Número"
    assert app_module._prev_token("El numero", 3) == "El"
    assert app_module._choose_best_suggestion("numero", [{"value": "número"}], "") == "número"
    assert app_module.is_upper_acronym("UNAMBA") is True
    assert app_module.is_upper_acronym("(ONU)") is True
    assert app_module.is_upper_acronym("Universidad") is False
    assert app_module._has_any_digit("x2") is True
    assert app_module._edit_distance("casa", "caso") == 1
    assert app_module._looks_reasonable_replacement("numero", "número") is True
    assert app_module._looks_reasonable_replacement("casa", "perro") is False

    texto = "numero UNAMBA x1"
    out = app_module.apply_lt_corrections_smart(
        texto,
        [
            _match(0, 6, "número"),
            _match(7, 6, "universidad"),  # sigla: no debe tocarse
            _match(14, 2, "x"),           # tiene dígito: no debe tocarse
        ],
    )
    assert out == "número UNAMBA x1"

    protegido = "A) numero"
    spans = [(0, len(protegido))]
    assert app_module.apply_lt_corrections_classic(protegido, [_match(3, 6, "número")], spans) == protegido

    # Wrapper que elige classic/smart según USE_CLASSIC_LT.
    app_module.USE_CLASSIC_LT = True
    assert app_module.apply_lt_corrections("numero", [_match(0, 6, "número")]) == "número"
    app_module.USE_CLASSIC_LT = False
    assert app_module.apply_lt_corrections("numero", [_match(0, 6, "número")]) == "número"
    app_module.USE_CLASSIC_LT = True


def test_spans_protegidos_marcas_y_post_correcciones(app_module):
    texto = "Pregunta\nA) Uno\nB) Dos\nFinal"
    spans = app_module.detectar_spans_alternativas(texto)
    assert len(spans) == 2
    assert app_module.intersecta_spans(spans[0][0], 2, spans) is True
    assert app_module.intersecta_spans(0, 3, spans) is False

    original = "P\nA) UNO\nF"
    corregido = "Q\nA) DOS\nG"
    restaurado = app_module.restaurar_segmentos_protegidos(original, corregido, [(2, 8)])
    assert "A) UNO" in restaurado

    marcado = app_module.insertar_marcas_eliminacion("uno dos tres", "uno tres")
    assert app_module.MARK_WORD_DELETE in marcado
    assert app_module.tokenize_preservando("uno, dos")
    assert app_module.es_token_palabra("uno") is True
    assert app_module.es_token_palabra(",") is False

    assert app_module._fix_numero_entero_specific("numero entera") == "número entero"
    assert "números enteros" in app_module._fix_numero_fallback("números entera")
    assert app_module._agree_adj("matriz", "entero") == "entera"
    assert app_module._fix_noun_adj_agreement("matriz mayor")
    assert app_module._context_fixes("Hay que de resolver. Se me olvidó de avisar.") == "Hay que resolver. Se me olvidó avisar."
    assert app_module._fix_del_contractions("De el grupo y de el valor") == "Del grupo y del valor"
    assert app_module._fix_aver_haber_context("haber si funciona") == "a ver si funciona"
    assert app_module._fix_cuyo_lado_agreement("cuya lados") == "cuyos lados"
    assert app_module._ensure_opening_mark("Qué pasa?").startswith("¿")
    assert "¡" in app_module._add_opening_spanish_marks("1) cuidado!")
    assert app_module._capitalize_starts("hola. mundo\n1) pregunta").startswith("Hola. Mundo")

    post = app_module.post_correcciones("hay que de resolver. numero entera. cuya lados.")
    post_l = post.lower()
    assert "hay que resolver" in post_l
    assert "número entero" in post_l
    assert "cuyos lados" in post_l


def test_docx_texto_diff_y_reemplazos_en_runs(app_module, tmp_path):
    doc = app_module.crear_docx_desde_texto("uno\ndos")
    assert [p.text for p in doc.paragraphs] == ["uno", "dos"]

    src = tmp_path / "texto.docx"
    doc.save(src)
    assert "uno" in app_module.extraer_texto_docx(str(src))

    html = app_module.generar_diff_html("uno", "dos")
    assert "diff" in html and "Original" in html
    assert app_module.now_mysql().count(":") == 2

    d = Document()
    p = d.add_paragraph()
    p.add_run("nume")
    p.add_run("ro entero")
    assert app_module.reemplazo_en_runs_flexible(p, "numero", "número") is True
    assert p.text == "número entero"

    p2 = d.add_paragraph("casa grande")
    assert app_module.reemplazo_en_runs_parciales(p2, "casa", "hogar") is True
    assert "hogar" in p2.text

    pares = app_module.pares_reemplazo_palabra_a_palabra("numero entero", "número entero")
    assert ("numero", "número") in pares
    assert app_module.similitud_suave("casa", "casa") == 1.0


def test_generar_docx_corregido_simple_y_highlight(app_module, tmp_path):
    src = tmp_path / "original.docx"
    out = tmp_path / "corregido.docx"
    doc = Document()
    doc.add_paragraph("numero entero")
    doc.add_paragraph("A) alternativa 1")
    doc.save(src)

    app_module.generar_docx_corregido(
        str(src),
        "número entero\nA) alternativa 1",
        str(out),
        highlight=True,
        texto_original_para_highlight="numero entero\nA) alternativa 1",
    )

    result = Document(out)
    assert result.paragraphs[0].text == "número entero"
    assert result.paragraphs[1].text == "A) alternativa 1"

    # Cubre el resaltado forzado sin depender del color exacto en el assert.
    p = result.add_paragraph("A ver si funciona")
    app_module._force_block_highlight_phrases(p, "haber si funciona", p.text)
    assert p.text == "A ver si funciona"
