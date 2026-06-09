import io
import os
import sqlite3
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{%s}" % WNS


def _connect(db_path):
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def _rewrite_docx(path, files):
    tmp = str(path) + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in files.items():
            z.writestr(name, data)
    os.replace(tmp, path)


def _numbered_docx_bytes(tmp_path, n=2):
    path = tmp_path / "contador.docx"
    doc = Document()
    for i in range(1, n + 1):
        doc.add_paragraph(f"Pregunta {i}")
        doc.add_paragraph("A) alternativa protegida")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}
    root = ET.fromstring(files["word/document.xml"])
    paras = root.findall(f".//{W}p")
    for idx in range(0, len(paras), 2):
        p = paras[idx]
        pPr = p.find(f"{W}pPr")
        if pPr is None:
            pPr = ET.Element(f"{W}pPr")
            p.insert(0, pPr)
        numPr = ET.Element(f"{W}numPr")
        ilvl = ET.SubElement(numPr, f"{W}ilvl"); ilvl.set(f"{W}val", "0")
        numId = ET.SubElement(numPr, f"{W}numId"); numId.set(f"{W}val", "55")
        pPr.insert(0, numPr)
    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    files["word/numbering.xml"] = f'''<w:numbering xmlns:w="{WNS}">
      <w:abstractNum w:abstractNumId="55"><w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/></w:lvl></w:abstractNum>
      <w:num w:numId="55"><w:abstractNumId w:val="55"/></w:num>
    </w:numbering>'''.encode("utf-8")
    _rewrite_docx(path, files)
    return path.read_bytes()


def _patch_import_db(app_module, tmp_path, monkeypatch):
    db_path = tmp_path / "cov15_importados.sqlite3"
    conn = _connect(db_path)
    conn.executescript(
        """
        CREATE TABLE examenes_importados(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT,
            ruta TEXT,
            extension TEXT,
            total_preguntas INTEGER DEFAULT 0,
            fuente TEXT,
            hash_archivo TEXT UNIQUE,
            fecha_creacion TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """
    )
    conn.commit(); conn.close()
    monkeypatch.setattr(app_module, "get_connection", lambda: _connect(db_path), raising=False)
    up = tmp_path / "uploads_importados"; up.mkdir(exist_ok=True)
    monkeypatch.setitem(app_module.app.config, "UPLOADS_EXAM_DIR", str(up))
    return db_path, up


def test_importar_examenes_importados_y_conteo_debug(client, app_module, tmp_path, monkeypatch):
    db_path, up = _patch_import_db(app_module, tmp_path, monkeypatch)
    docx_data = _numbered_docx_bytes(tmp_path, n=2)

    assert client.open("/api/examenes/importar", method="OPTIONS").status_code == 204
    assert client.post("/api/examenes/importar", data={}, content_type="multipart/form-data").status_code == 400
    bad = client.post("/api/examenes/importar", data={"files": (io.BytesIO(b"x"), "mal.exe")}, content_type="multipart/form-data")
    assert bad.status_code == 415

    ok = client.post(
        "/api/examenes/importar",
        data={"files": (io.BytesIO(docx_data), "examen_grupo_A.docx")},
        content_type="multipart/form-data",
    )
    assert ok.status_code == 200
    item = ok.get_json()["items"][0]
    assert item["total_preguntas"] == 2
    assert os.path.exists(up / "examen_grupo_A.docx")

    listado = client.get("/api/examenes/importados")
    assert listado.status_code == 200
    assert listado.get_json()[0]["total_preguntas"] == 2

    # Mismo archivo: cubre rama ON CONFLICT(hash_archivo) de importación.
    ok2 = client.post(
        "/api/examenes/importar",
        data={"files[]": (io.BytesIO(docx_data), "examen_grupo_A_copia.docx")},
        content_type="multipart/form-data",
    )
    assert ok2.status_code == 200

    conn = _connect(db_path)
    total_rows = conn.execute("SELECT COUNT(*) FROM examenes_importados").fetchone()[0]
    conn.close()
    assert total_rows == 1

    # Helpers pequeños cercanos al flujo de importados.
    assert app_module._ext_ok("archivo.pdf") is True
    assert app_module._ext_ok("archivo.exe") is False
    assert app_module.inferir_clave_grupo_desde_nombre("TEMAS_GRUPO-B.docx") == "B"
    assert app_module.inferir_clave_grupo_desde_nombre("sin_grupo.docx") is None
